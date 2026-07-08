"""Announcement DOCX response import, translation validation, apply, and delivery.

Implementation module split out of utils/announcement_docx_harness.py. Import
these symbols through utils.announcement_docx_harness to keep the public
surface stable.
"""
from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook, load_workbook

from utils.announcement_docx_common import (
    AI_RESPONSE_PREFIX,
    FIXED_COLUMNS,
    QA_SUMMARY_NAME,
    TRANSLATION_WORKBOOK_NAME,
    _CJK_RE,
    _clean_cell,
    _expected_paragraphs,
    _load_manifest,
    _manifest_languages,
    _ordered_expected_rows,
    _parse_json_cell,
    _resolve_language_headers,
    _sha256_file,
    _work_dir,
)

_BRACKET_RE = re.compile(r"[\(\)\[\]\{\}\u3010\u3011\uff08\uff09]")


@dataclass(frozen=True)
class ImportedAnnouncementResponses:
    work_dir: Path
    translation_workbook: Path
    languages: list[str]
    row_count: int


@dataclass(frozen=True)
class AppliedAnnouncementHarness:
    work_dir: Path
    output_dir: Path
    qa_summary_path: Path
    output_docx_paths: list[Path]
    hard_blockers: int


@dataclass(frozen=True)
class DeliveredAnnouncementOutputs:
    delivery_dir: Path
    files: list[Path]


def import_announcement_ai_responses(
    input_dir: str | Path,
    *,
    translation_workbook: str | Path | None = None,
    response_dir: str | Path | None = None,
    languages: list[str] | None = None,
    response_prefix: str = AI_RESPONSE_PREFIX,
) -> ImportedAnnouncementResponses:
    """Import model-authored JSONL responses into the intermediate workbook.

    This adapter only validates and imports response files. It intentionally
    does not call translation APIs, browser translators, or MT packages.
    """
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    manifest = _load_manifest(work_dir)
    target_languages = _manifest_languages(manifest)
    expected = _expected_paragraphs(manifest)
    ordered_rows = _ordered_expected_rows(manifest)
    workbook_path = Path(translation_workbook) if translation_workbook else work_dir / TRANSLATION_WORKBOOK_NAME
    response_root = Path(response_dir) if response_dir else work_dir
    selected_headers = _resolve_language_headers(languages, valid_languages=target_languages)

    translations_by_header: dict[str, dict[str, str]] = {}
    for header, code in target_languages:
        if header not in selected_headers:
            continue
        response_path = response_root / f"{response_prefix}{code}.jsonl"
        response_rows = _read_ai_response_rows(response_path, ordered_rows)
        issues: list[dict[str, Any]] = []
        for expected_row, response_row in zip(ordered_rows, response_rows):
            row_context = {
                "source_file": expected_row["source_file"],
                "para_id": str(expected_row["para_id"]),
                "para_index": expected_row["para_index"],
            }
            issues.extend(
                _validate_translation(
                    str(expected_row["CN"]),
                    response_row["translation"],
                    expected_row.get("protected_tokens", []),
                    expected_row.get("term_hits", []),
                    header,
                    code,
                    row_context,
                )
            )
        if issues:
            raise ValueError(f"AI response QA failed for {code}: {len(issues)} issues")
        translations_by_header[header] = {
            response_row["para_id"]: response_row["translation"]
            for response_row in response_rows
        }

    _write_ai_responses_to_workbook(workbook_path, expected, translations_by_header)
    return ImportedAnnouncementResponses(
        work_dir=work_dir,
        translation_workbook=workbook_path,
        languages=selected_headers,
        row_count=len(ordered_rows),
    )


def apply_announcement_translations(
    input_dir: str | Path,
    translation_workbook: str | Path,
) -> AppliedAnnouncementHarness:
    """Validate a filled announcement workbook and write language DOCX outputs."""
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    manifest = _load_manifest(work_dir)
    target_languages = _manifest_languages(manifest)
    rows = _read_translation_rows(Path(translation_workbook), target_languages)
    expected = _expected_paragraphs(manifest)
    _validate_row_coverage(expected, rows)
    _validate_input_drift(input_dir, manifest)

    issues = _validate_all_translations(expected, rows, target_languages)
    output_dir = work_dir / "output"
    qa_summary_path = work_dir / QA_SUMMARY_NAME
    if issues:
        _write_qa_summary(qa_summary_path, issues, [])
        raise ValueError(f"hard blockers: {len(issues)}")

    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_paths = _write_output_docx(input_dir, manifest, rows, output_dir, target_languages)
    _write_qa_summary(qa_summary_path, [], output_paths)
    return AppliedAnnouncementHarness(
        work_dir=work_dir,
        output_dir=output_dir,
        qa_summary_path=qa_summary_path,
        output_docx_paths=output_paths,
        hard_blockers=0,
    )


def deliver_announcement_outputs(
    input_dir: str | Path,
    *,
    date_stamp: str | None = None,
) -> DeliveredAnnouncementOutputs:
    """Copy passed DOCX outputs and QA summary into a clean delivery directory."""
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    output_dir = work_dir / "output"
    qa_summary = work_dir / QA_SUMMARY_NAME
    if not output_dir.exists():
        raise ValueError(f"output directory not found: {output_dir}")
    if not qa_summary.exists():
        raise ValueError(f"QA summary not found: {qa_summary}")

    date_stamp = date_stamp or datetime.now().strftime("%Y%m%d")
    delivery_dir = input_dir / f"{input_dir.name}_多语言交付_{date_stamp}"
    if delivery_dir.exists():
        shutil.rmtree(delivery_dir)
    delivery_dir.mkdir(parents=True, exist_ok=True)

    files: list[Path] = []
    for path in sorted(output_dir.glob("*.docx")):
        target = delivery_dir / path.name
        shutil.copy2(path, target)
        files.append(target)
    qa_target = delivery_dir / QA_SUMMARY_NAME
    shutil.copy2(qa_summary, qa_target)
    files.append(qa_target)
    return DeliveredAnnouncementOutputs(delivery_dir=delivery_dir, files=files)


def _read_translation_rows(path: Path, target_languages: list[tuple[str, str]]) -> dict[str, dict[str, Any]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb["Translations"]
        headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
        header_index = {str(header): idx + 1 for idx, header in enumerate(headers)}
        required = [*FIXED_COLUMNS, *[header for header, _ in target_languages]]
        missing = [name for name in required if name not in header_index]
        if missing:
            raise ValueError(f"translation workbook missing columns: {missing}")

        rows: dict[str, dict[str, Any]] = {}
        duplicates: list[str] = []
        for row_idx in range(2, ws.max_row + 1):
            para_id = _clean_cell(ws.cell(row_idx, header_index["para_id"]).value)
            if not para_id:
                continue
            if para_id in rows:
                duplicates.append(para_id)
                continue
            rows[para_id] = {
                "source_file": _clean_cell(ws.cell(row_idx, header_index["source_file"]).value),
                "para_id": para_id,
                "para_index": int(ws.cell(row_idx, header_index["para_index"]).value),
                "style": _clean_cell(ws.cell(row_idx, header_index["style"]).value),
                "CN": _clean_cell(ws.cell(row_idx, header_index["CN"]).value),
                "protected_tokens": _parse_json_cell(ws.cell(row_idx, header_index["protected_tokens"]).value, []),
                "term_hits": _parse_json_cell(ws.cell(row_idx, header_index["term_hits_json"]).value, []),
                "translations": {
                    header: _clean_cell(ws.cell(row_idx, header_index[header]).value)
                    for header, _ in target_languages
                },
            }
    finally:
        wb.close()
    if duplicates:
        raise ValueError(f"duplicate paragraph IDs: {duplicates}")
    return rows


def _validate_row_coverage(expected: dict[str, dict[str, Any]], rows: dict[str, dict[str, Any]]) -> None:
    expected_ids = set(expected)
    actual_ids = set(rows)
    missing = sorted(expected_ids - actual_ids)
    extra = sorted(actual_ids - expected_ids)
    if missing:
        raise ValueError(f"missing paragraph IDs: {missing}")
    if extra:
        raise ValueError(f"extra paragraph IDs: {extra}")


def _validate_input_drift(input_dir: Path, manifest: dict[str, Any]) -> None:
    for doc in manifest.get("documents", []):
        path = input_dir / str(doc["source_file"])
        if _sha256_file(path) != doc.get("sha256"):
            raise ValueError(f"input drift detected: {path.name}")


def _validate_all_translations(
    expected: dict[str, dict[str, Any]],
    rows: dict[str, dict[str, Any]],
    target_languages: list[tuple[str, str]],
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for para_id, expected_row in expected.items():
        row = rows[para_id]
        source = str(expected_row["CN"])
        term_hits = expected_row.get("term_hits", [])
        protected_tokens = expected_row.get("protected_tokens", [])
        for header, code in target_languages:
            translation = row["translations"].get(header, "")
            issues.extend(_validate_translation(source, translation, protected_tokens, term_hits, header, code, row))
    return issues


def _read_ai_response_rows(response_path: Path, expected_rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    if not response_path.exists():
        raise ValueError(f"AI response not found: {response_path}")

    rows: list[dict[str, str]] = []
    seen: set[str] = set()
    for line_no, line in enumerate(response_path.read_text(encoding="utf-8-sig").splitlines(), start=1):
        if not line.strip():
            continue
        try:
            row = json.loads(line)
        except json.JSONDecodeError as exc:
            raise ValueError(f"invalid JSON in {response_path.name}:{line_no}: {exc}") from exc
        if set(row) != {"para_id", "translation"}:
            raise ValueError(f"AI response row {line_no} must contain only para_id and translation")
        para_id = _clean_cell(row.get("para_id"))
        translation = _clean_cell(row.get("translation"))
        if not para_id:
            raise ValueError(f"AI response row {line_no} missing para_id")
        if para_id in seen:
            raise ValueError(f"duplicate AI response para_id: {para_id}")
        seen.add(para_id)
        rows.append({"para_id": para_id, "translation": translation})

    expected_ids = [str(row["para_id"]) for row in expected_rows]
    actual_ids = [row["para_id"] for row in rows]
    if actual_ids != expected_ids:
        missing = sorted(set(expected_ids) - set(actual_ids))
        extra = sorted(set(actual_ids) - set(expected_ids))
        if missing:
            raise ValueError(f"AI response missing paragraph IDs: {missing}")
        if extra:
            raise ValueError(f"AI response extra paragraph IDs: {extra}")
        raise ValueError("AI response paragraph order mismatch")
    return rows


def _write_ai_responses_to_workbook(
    workbook_path: Path,
    expected: dict[str, dict[str, Any]],
    translations_by_header: dict[str, dict[str, str]],
) -> None:
    wb = load_workbook(workbook_path)
    try:
        ws = wb["Translations"]
        headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
        header_index = {str(header): idx + 1 for idx, header in enumerate(headers)}
        row_by_para_id = {}
        for row_idx in range(2, ws.max_row + 1):
            para_id = _clean_cell(ws.cell(row_idx, header_index["para_id"]).value)
            if para_id:
                row_by_para_id[para_id] = row_idx
        missing_in_workbook = sorted(set(expected) - set(row_by_para_id))
        if missing_in_workbook:
            raise ValueError(f"translation workbook missing paragraph IDs: {missing_in_workbook}")
        for header, translations in translations_by_header.items():
            if header not in header_index:
                raise ValueError(f"translation workbook missing language column: {header}")
            for para_id, translation in translations.items():
                ws.cell(row_by_para_id[para_id], header_index[header]).value = translation
        wb.save(workbook_path)
    finally:
        wb.close()


def _validate_translation(
    source: str,
    translation: str,
    protected_tokens: list[str],
    term_hits: list[dict[str, Any]],
    lang_header: str,
    lang_code: str,
    row: dict[str, Any],
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    base = {
        "source_file": row["source_file"],
        "para_id": row["para_id"],
        "para_index": row["para_index"],
        "lang": lang_code,
        "source": source,
        "translation": translation,
    }
    if not translation.strip():
        issues.append({**base, "check_type": "empty_translation", "message": "Translation is empty"})
        return issues
    if _CJK_RE.search(translation):
        issues.append({**base, "check_type": "chinese_residue", "message": "Chinese residue found"})

    for token in protected_tokens:
        if token and token not in translation:
            issues.append({**base, "check_type": "protected_token_missing", "message": f"Missing protected token: {token}"})

    if _BRACKET_RE.search(source) and not _BRACKET_RE.search(translation):
        issues.append({**base, "check_type": "bracket_shape_missing", "message": "Source has bracketed content but translation has no brackets"})

    for hit in term_hits:
        target = str(hit.get("targets", {}).get(lang_header, "")).strip()
        if target and not _contains_term(translation, target):
            issues.append(
                {
                    **base,
                    "check_type": "term_missing",
                    "message": f"Missing term target for {hit.get('source')}: {target}",
                }
            )
    return issues


def _write_output_docx(
    input_dir: Path,
    manifest: dict[str, Any],
    rows: dict[str, dict[str, Any]],
    output_dir: Path,
    target_languages: list[tuple[str, str]],
) -> list[Path]:
    output_paths: list[Path] = []
    for doc_record in manifest.get("documents", []):
        source_name = str(doc_record["source_file"])
        source_path = input_dir / source_name
        paragraphs = doc_record.get("paragraphs", [])
        for header, code in target_languages:
            doc = Document(str(source_path))
            for para_record in paragraphs:
                para_id = str(para_record["para_id"])
                index = int(para_record["para_index"])
                translation = rows[para_id]["translations"][header]
                _replace_paragraph_text(doc.paragraphs[index], translation)
            output_path = output_dir / f"{Path(source_name).stem}_{code}.docx"
            doc.save(output_path)
            output_paths.append(output_path)
    return output_paths


def _write_qa_summary(path: Path, issues: list[dict[str, Any]], output_paths: list[Path]) -> None:
    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"
    summary.append(["metric", "value"])
    summary.append(["hard_blockers", len(issues)])
    summary.append(["output_docx_count", len(output_paths)])
    summary.append(["generated_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    details = wb.create_sheet("Issues")
    issue_headers = ["source_file", "para_id", "para_index", "lang", "check_type", "message", "source", "translation"]
    details.append(issue_headers)
    for issue in issues:
        details.append([issue.get(header, "") for header in issue_headers])
    outputs = wb.create_sheet("Outputs")
    outputs.append(["path"])
    for output_path in output_paths:
        outputs.append([str(output_path)])
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _contains_term(translation: str, target: str) -> bool:
    normalized_translation = re.sub(r"\s+", " ", translation).casefold()
    normalized_target = re.sub(r"\s+", " ", target).casefold()
    return normalized_target in normalized_translation
