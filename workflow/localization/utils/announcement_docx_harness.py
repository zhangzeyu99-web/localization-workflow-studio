"""DOCX announcement translation harness.

This harness turns paragraph-only announcement DOCX files into a structured
translation workbook, validates terminology-constrained translations, and writes
language-specific DOCX copies back from stable paragraph IDs.
"""
from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook, load_workbook


TARGET_LANGUAGES: tuple[tuple[str, str], ...] = (
    ("EN", "en"),
    ("KR", "ko"),
    ("JP", "ja"),
    ("FR", "fr"),
    ("DE", "de"),
    ("RU", "ru"),
    ("IT", "it"),
    ("ES", "es"),
    ("PT", "pt"),
    ("TR", "tr"),
    ("IDN", "idn"),
    ("TH", "th"),
    ("AR", "ar"),
)
SUPPORTED_LANGUAGES: tuple[tuple[str, str], ...] = (
    *TARGET_LANGUAGES,
    ("KO", "ko"),
    ("JA", "ja"),
    ("TK", "tr"),
    ("ID", "idn"),
)
CANONICAL_LANGUAGE_HEADER = {code: header for header, code in TARGET_LANGUAGES}
LANGUAGE_CODE_BY_HEADER = {header: code for header, code in SUPPORTED_LANGUAGES}

FIXED_COLUMNS = (
    "source_file",
    "para_id",
    "para_index",
    "style",
    "CN",
    "protected_tokens",
    "term_hits_json",
)
TRANSLATION_WORKBOOK_NAME = "announcement_translation_workbook.xlsx"
MANIFEST_NAME = "manifest.json"
QA_SUMMARY_NAME = "QA摘要.xlsx"
AI_RESPONSE_PREFIX = "ai_response_"
WORK_DIR_NAME = "_work"
HARNESS_DIR_NAME = "announcement_docx"

_CJK_RE = re.compile(r"[\u3400-\u9fff]")
_DATE_RE = re.compile(r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b")
_TIME_RANGE_RE = re.compile(r"\b\d{1,2}:\d{2}\s*(?:[-–~]\s*\d{1,2}:\d{2})?\b")
_NUMBER_RE = re.compile(r"\d+")
_BRACKET_TOKEN_RE = re.compile(r"\[[^\]]+\]|\([^)]+\)|\{[^}]+\}|【[^】]+】|（[^）]+）")
_BRACKET_RE = re.compile(r"[\(\)\[\]\{\}\u3010\u3011\uff08\uff09]")
_ANNOUNCEMENT_PLACEHOLDER_RE = re.compile(r"<@\d+>|\$\{\d+\}|%[sdif]|\{[^{}]+\}|<[^>]+>|\\n")
_ANNOUNCEMENT_SENTENCE_PUNCTUATION = set("，。！？；：、,.?!;:")
_LOW_VALUE_ANNOUNCEMENT_TERMS = {
    "好的",
    "游戏",
    "时间",
    "小时",
    "完成",
    "发放",
    "领取",
    "获得",
    "获取",
    "更新",
    "进入",
    "点击",
    "下载",
    "界面",
    "弹窗",
    "邮件",
    "活动",
    "系统",
    "玩法",
    "服务器",
    "维护",
    "补偿",
    "内容",
    "任务",
    "奖励",
    "开启",
    "开放",
    "修复",
    "新增",
    "部分",
    "所有",
    "普通",
    "使用",
    "问题",
    "确认",
    "取消",
    "返回",
    "查看",
    "选择",
    "购买",
    "前往",
    "本次",
    "当前",
    "公告",
    "进入游戏",
    "开始游戏",
}
_LOW_VALUE_ANNOUNCEMENT_PREFIXES = (
    "使用后",
    "用于",
    "可",
    "将",
    "已",
    "未",
    "请",
    "点击",
    "前往",
    "打开",
    "关闭",
    "获得",
    "完成",
    "通关",
    "达到",
    "活动期间",
    "当前",
    "本期",
    "每日",
    "成功",
    "失败",
    "进入",
    "开始",
    "返回",
    "查看",
    "选择",
)
_LOW_VALUE_ANNOUNCEMENT_SUBSTRINGS = (
    "是否",
    "不足",
    "尚未",
    "暂无",
    "已达",
    "未领取",
    "已领取",
    "补发",
    "异常",
    "排行奖励",
    "奖励邮件",
    "正在",
    "倒计时",
)
_TERM_NOUN_SUFFIXES = (
    "活动",
    "系统",
    "玩法",
    "副本",
    "战场",
    "商店",
    "商城",
    "公会",
    "英雄",
    "角色",
    "皮肤",
    "碎片",
    "宝箱",
    "礼包",
    "月卡",
    "通行证",
    "圣器",
    "遗物",
    "机关",
    "装备",
    "材料",
    "水晶",
    "金币",
    "硬币",
    "积分",
    "图鉴",
    "技能",
)


@dataclass(frozen=True)
class LanguageSpec:
    header: str
    code: str
    column_index: int


@dataclass(frozen=True)
class TermEntry:
    source: str
    target: str
    term_id: str = ""


@dataclass(frozen=True)
class AnnouncementTerms:
    path: Path
    languages: list[LanguageSpec]
    by_language: dict[str, dict[str, TermEntry]]


@dataclass(frozen=True)
class PreparedAnnouncementHarness:
    work_dir: Path
    translation_workbook: Path
    manifest_path: Path
    row_count: int
    doc_count: int


@dataclass(frozen=True)
class ImportedAnnouncementResponses:
    work_dir: Path
    translation_workbook: Path
    languages: list[str]
    row_count: int


@dataclass(frozen=True)
class AppliedAnnouncementHarness:
    work_dir: Path
    output_dir: Path
    qa_summary_path: Path
    output_docx_paths: list[Path]
    hard_blockers: int


@dataclass(frozen=True)
class DeliveredAnnouncementOutputs:
    delivery_dir: Path
    files: list[Path]


@dataclass(frozen=True)
class AnnouncementTaskInspection:
    input_dir: Path
    source_files: list[Path]
    term_files: list[Path]
    reference_files: list[Path]
    languages: list[tuple[str, str]]


@dataclass(frozen=True)
class StagedAnnouncementTask:
    staging_dir: Path
    source_files: list[Path]
    term_files: list[Path]
    languages: list[tuple[str, str]]


def prepare_announcement_docx_harness(
    input_dir: str | Path,
    *,
    languages: list[str] | None = None,
) -> PreparedAnnouncementHarness:
    """Create a structured workbook and per-language workpacks for DOCX announcements."""
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    if work_dir.exists():
        shutil.rmtree(work_dir)
    work_dir.mkdir(parents=True, exist_ok=True)
    pairs = discover_announcement_docx_pairs(input_dir)
    target_languages = (
        _resolve_language_pairs(languages)
        if languages
        else _infer_language_pairs_from_terms(pairs)
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Translations"
    headers = [*FIXED_COLUMNS, *[header for header, _ in target_languages]]
    ws.append(headers)

    manifest: dict[str, Any] = {
        "version": 1,
        "input_dir": str(input_dir),
        "languages": [{"header": header, "code": code} for header, code in target_languages],
        "documents": [],
    }
    workpack_rows: dict[str, list[dict[str, Any]]] = {code: [] for _, code in target_languages}
    row_count = 0

    for docx_path, term_path in pairs:
        terms = load_announcement_terms(term_path)
        doc = Document(str(docx_path))
        unsupported = _inspect_unsupported(docx_path, doc)
        doc_record = {
            "source_file": docx_path.name,
            "term_file": term_path.name,
            "sha256": _sha256_file(docx_path),
            "unsupported": unsupported,
            "paragraphs": [],
        }
        for index, paragraph in enumerate(doc.paragraphs):
            source = paragraph.text
            if not source or not source.strip():
                continue
            para_id = _paragraph_id(docx_path.name, index, source)
            style = paragraph.style.name if paragraph.style else ""
            term_hits = _find_term_hits(source, terms)
            protected_tokens = _protected_tokens(source)
            row = {
                "source_file": docx_path.name,
                "para_id": para_id,
                "para_index": index,
                "style": style,
                "CN": source,
                "protected_tokens": protected_tokens,
                "term_hits": term_hits,
            }
            doc_record["paragraphs"].append(row)
            ws.append(
                [
                    docx_path.name,
                    para_id,
                    index,
                    style,
                    source,
                    json.dumps(protected_tokens, ensure_ascii=False),
                    json.dumps(term_hits, ensure_ascii=False),
                    *["" for _ in target_languages],
                ]
            )
            for header, code in target_languages:
                workpack_rows[code].append(
                    {
                        "source_file": docx_path.name,
                        "para_id": para_id,
                        "para_index": index,
                        "source": source,
                        "term_hits": [
                            {"source": hit["source"], "target": hit["targets"].get(header, "")}
                            for hit in term_hits
                            if hit["targets"].get(header)
                        ],
                        "protected_tokens": protected_tokens,
                    }
                )
            row_count += 1
        manifest["documents"].append(doc_record)

    translation_workbook = work_dir / TRANSLATION_WORKBOOK_NAME
    wb.save(translation_workbook)
    wb.close()

    manifest_path = work_dir / MANIFEST_NAME
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    for _, code in target_languages:
        _write_jsonl(work_dir / f"workpack_{code}.jsonl", workpack_rows[code])

    return PreparedAnnouncementHarness(
        work_dir=work_dir,
        translation_workbook=translation_workbook,
        manifest_path=manifest_path,
        row_count=row_count,
        doc_count=len(pairs),
    )


def import_announcement_ai_responses(
    input_dir: str | Path,
    *,
    translation_workbook: str | Path | None = None,
    response_dir: str | Path | None = None,
    languages: list[str] | None = None,
    response_prefix: str = AI_RESPONSE_PREFIX,
) -> ImportedAnnouncementResponses:
    """Import model-authored JSONL responses into the intermediate workbook.

    This adapter only validates and imports response files. It intentionally
    does not call translation APIs, browser translators, or MT packages.
    """
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    manifest = _load_manifest(work_dir)
    target_languages = _manifest_languages(manifest)
    expected = _expected_paragraphs(manifest)
    ordered_rows = _ordered_expected_rows(manifest)
    workbook_path = Path(translation_workbook) if translation_workbook else work_dir / TRANSLATION_WORKBOOK_NAME
    response_root = Path(response_dir) if response_dir else work_dir
    selected_headers = _resolve_language_headers(languages, valid_languages=target_languages)

    translations_by_header: dict[str, dict[str, str]] = {}
    for header, code in target_languages:
        if header not in selected_headers:
            continue
        response_path = response_root / f"{response_prefix}{code}.jsonl"
        response_rows = _read_ai_response_rows(response_path, ordered_rows)
        issues: list[dict[str, Any]] = []
        for expected_row, response_row in zip(ordered_rows, response_rows):
            row_context = {
                "source_file": expected_row["source_file"],
                "para_id": str(expected_row["para_id"]),
                "para_index": expected_row["para_index"],
            }
            issues.extend(
                _validate_translation(
                    str(expected_row["CN"]),
                    response_row["translation"],
                    expected_row.get("protected_tokens", []),
                    expected_row.get("term_hits", []),
                    header,
                    code,
                    row_context,
                )
            )
        if issues:
            raise ValueError(f"AI response QA failed for {code}: {len(issues)} issues")
        translations_by_header[header] = {
            response_row["para_id"]: response_row["translation"]
            for response_row in response_rows
        }

    _write_ai_responses_to_workbook(workbook_path, expected, translations_by_header)
    return ImportedAnnouncementResponses(
        work_dir=work_dir,
        translation_workbook=workbook_path,
        languages=selected_headers,
        row_count=len(ordered_rows),
    )


def apply_announcement_translations(
    input_dir: str | Path,
    translation_workbook: str | Path,
) -> AppliedAnnouncementHarness:
    """Validate a filled announcement workbook and write language DOCX outputs."""
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    manifest = _load_manifest(work_dir)
    target_languages = _manifest_languages(manifest)
    rows = _read_translation_rows(Path(translation_workbook), target_languages)
    expected = _expected_paragraphs(manifest)
    _validate_row_coverage(expected, rows)
    _validate_input_drift(input_dir, manifest)

    issues = _validate_all_translations(expected, rows, target_languages)
    output_dir = work_dir / "output"
    qa_summary_path = work_dir / QA_SUMMARY_NAME
    if issues:
        _write_qa_summary(qa_summary_path, issues, [])
        raise ValueError(f"hard blockers: {len(issues)}")

    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_paths = _write_output_docx(input_dir, manifest, rows, output_dir, target_languages)
    _write_qa_summary(qa_summary_path, [], output_paths)
    return AppliedAnnouncementHarness(
        work_dir=work_dir,
        output_dir=output_dir,
        qa_summary_path=qa_summary_path,
        output_docx_paths=output_paths,
        hard_blockers=0,
    )


def deliver_announcement_outputs(
    input_dir: str | Path,
    *,
    date_stamp: str | None = None,
) -> DeliveredAnnouncementOutputs:
    """Copy passed DOCX outputs and QA summary into a clean delivery directory."""
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    output_dir = work_dir / "output"
    qa_summary = work_dir / QA_SUMMARY_NAME
    if not output_dir.exists():
        raise ValueError(f"output directory not found: {output_dir}")
    if not qa_summary.exists():
        raise ValueError(f"QA summary not found: {qa_summary}")

    date_stamp = date_stamp or datetime.now().strftime("%Y%m%d")
    delivery_dir = input_dir / f"{input_dir.name}_多语言交付_{date_stamp}"
    if delivery_dir.exists():
        shutil.rmtree(delivery_dir)
    delivery_dir.mkdir(parents=True, exist_ok=True)

    files: list[Path] = []
    for path in sorted(output_dir.glob("*.docx")):
        target = delivery_dir / path.name
        shutil.copy2(path, target)
        files.append(target)
    qa_target = delivery_dir / QA_SUMMARY_NAME
    shutil.copy2(qa_summary, qa_target)
    files.append(qa_target)
    return DeliveredAnnouncementOutputs(delivery_dir=delivery_dir, files=files)


def inspect_announcement_task_dir(input_dir: str | Path) -> AnnouncementTaskInspection:
    """Identify loose announcement sources, term delivery files, references, and target languages."""
    input_dir = Path(input_dir)
    source_files = [
        path
        for path in sorted(input_dir.iterdir(), key=lambda item: item.name)
        if path.is_file()
        and not _is_temp_file(path)
        and path.suffix.lower() in {".docx", ".txt"}
        and not _is_generated_docx(path)
    ]
    term_files: list[Path] = []
    reference_files: list[Path] = []
    language_pairs: list[tuple[str, str]] = []
    seen_languages: set[str] = set()
    for path in sorted(input_dir.glob("*.xlsx"), key=lambda item: item.name):
        if _is_temp_file(path):
            continue
        try:
            language_specs = _read_announcement_language_specs(path)
        except Exception:
            continue
        if not language_specs:
            continue
        if _is_loose_announcement_terms_file(path):
            term_files.append(path)
            for spec in language_specs:
                header = spec.header.strip().upper()
                code = spec.code
                if header not in seen_languages:
                    language_pairs.append((header, code))
                    seen_languages.add(header)
        else:
            reference_files.append(path)

    return AnnouncementTaskInspection(
        input_dir=input_dir,
        source_files=source_files,
        term_files=term_files,
        reference_files=reference_files,
        languages=language_pairs,
    )


def stage_announcement_task_dir(input_dir: str | Path) -> StagedAnnouncementTask:
    """Normalize a loose task folder into the strict DOCX+announcement_terms harness layout."""
    inspection = inspect_announcement_task_dir(input_dir)
    if not inspection.source_files:
        raise ValueError(f"no announcement source .docx/.txt files found in {inspection.input_dir}")
    if not inspection.term_files:
        raise ValueError(f"no announcement term delivery workbook found in {inspection.input_dir}")
    if not inspection.languages:
        raise ValueError(f"no supported target language columns found in {inspection.input_dir}")

    staging_dir = inspection.input_dir / WORK_DIR_NAME / HARNESS_DIR_NAME / "source_input"
    if staging_dir.exists():
        shutil.rmtree(staging_dir)
    staging_dir.mkdir(parents=True, exist_ok=True)

    staged_sources: list[Path] = []
    staged_terms: list[Path] = []
    for source_path in inspection.source_files:
        staged_source = staging_dir / f"{source_path.stem}.docx"
        if source_path.suffix.lower() == ".docx":
            shutil.copy2(source_path, staged_source)
        else:
            _convert_txt_to_docx(source_path, staged_source)
        staged_sources.append(staged_source)

        term_path = _select_term_file_for_source(source_path, inspection.term_files)
        term_date = _extract_date_stamp(term_path)
        staged_term = staging_dir / f"{staged_source.stem}_announcement_terms_{term_date}.xlsx"
        shutil.copy2(term_path, staged_term)
        staged_terms.append(staged_term)

    return StagedAnnouncementTask(
        staging_dir=staging_dir,
        source_files=staged_sources,
        term_files=staged_terms,
        languages=inspection.languages,
    )


def discover_announcement_docx_pairs(input_dir: str | Path) -> list[tuple[Path, Path]]:
    input_dir = Path(input_dir)
    term_files = sorted(
        path
        for path in input_dir.glob("*_announcement_terms_*.xlsx")
        if path.is_file() and not _is_temp_file(path)
    )
    docx_files = sorted(
        path
        for path in input_dir.glob("*.docx")
        if path.is_file() and not _is_temp_file(path) and not _is_generated_docx(path)
    )
    pairs: list[tuple[Path, Path]] = []
    for docx_path in docx_files:
        prefix = f"{docx_path.stem}_announcement_terms_"
        matches = [path for path in term_files if path.name.startswith(prefix)]
        if not matches:
            raise ValueError(f"announcement term workbook not found for {docx_path.name}")
        pairs.append((docx_path, sorted(matches, key=lambda item: item.name)[-1]))
    if not pairs:
        raise ValueError(f"no announcement DOCX files found in {input_dir}")
    return pairs


def load_announcement_terms(path: str | Path) -> AnnouncementTerms:
    path = Path(path)
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb.active
        languages = _language_specs_from_headers(
            [str(ws.cell(1, col).value or "").strip() for col in range(1, ws.max_column + 1)]
        )
        by_language: dict[str, dict[str, TermEntry]] = {}
        for spec in languages:
            by_language[spec.header] = {}

        for row in range(2, ws.max_row + 1):
            term_id = _clean_cell(ws.cell(row, 1).value)
            source = _clean_cell(ws.cell(row, 2).value)
            if not source:
                continue
            for spec in languages:
                target = _clean_cell(ws.cell(row, spec.column_index).value)
                if target:
                    by_language[spec.header][source] = TermEntry(source=source, target=target, term_id=term_id)
    finally:
        wb.close()
    return AnnouncementTerms(path=path, languages=languages, by_language=by_language)


def _read_announcement_language_specs(path: Path) -> list[LanguageSpec]:
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb.active
        headers = [str(ws.cell(1, col).value or "").strip() for col in range(1, ws.max_column + 1)]
        return _language_specs_from_headers(headers)
    finally:
        wb.close()


def _language_specs_from_headers(headers: list[str]) -> list[LanguageSpec]:
    languages: list[LanguageSpec] = []
    seen_codes: set[str] = set()
    for col in range(3, len(headers) + 1):
        input_header = headers[col - 1]
        code = _language_code_for_header(input_header)
        if not code or code in seen_codes:
            continue
        seen_codes.add(code)
        header = CANONICAL_LANGUAGE_HEADER[code]
        languages.append(LanguageSpec(header=header, code=code, column_index=col))
    return languages


def _find_term_hits(source: str, terms: AnnouncementTerms) -> list[dict[str, Any]]:
    all_sources = {
        term
        for lang_terms in terms.by_language.values()
        for term in lang_terms
        if len(term) >= 2 and not _is_low_value_announcement_term(term) and _term_occurs(source, term)
    }
    selected: list[str] = []
    for term in sorted(all_sources, key=len, reverse=True):
        if any(term in longer for longer in selected):
            continue
        selected.append(term)

    hits: list[dict[str, Any]] = []
    for source_term in selected:
        targets = {}
        for spec in terms.languages:
            entry = terms.by_language.get(spec.header, {}).get(source_term)
            if entry:
                targets[spec.header] = entry.target
        if targets:
            hits.append({"source": source_term, "targets": targets})
    return hits


def _is_low_value_announcement_term(source: str) -> bool:
    text = str(source or "").strip()
    if not text:
        return True
    if text in _LOW_VALUE_ANNOUNCEMENT_TERMS:
        return True
    if not _CJK_RE.search(text):
        return True
    if len(text) <= 1:
        return True
    if _ANNOUNCEMENT_PLACEHOLDER_RE.search(text):
        return True
    if any(char.isdigit() for char in text) and not any(text.endswith(suffix) for suffix in _TERM_NOUN_SUFFIXES):
        return True
    cjk_len = sum(1 for char in text if "\u3400" <= char <= "\u9fff")
    if cjk_len <= 2 and text.endswith(("级", "次", "个", "点", "日", "月")):
        return True
    if any(mark in text for mark in _ANNOUNCEMENT_SENTENCE_PUNCTUATION) and cjk_len > 4:
        return True
    if any(text.startswith(prefix) for prefix in _LOW_VALUE_ANNOUNCEMENT_PREFIXES) and cjk_len > 4:
        return True
    if any(part in text for part in _LOW_VALUE_ANNOUNCEMENT_SUBSTRINGS) and cjk_len > 5:
        return True
    if cjk_len > 10 and not any(text.endswith(suffix) for suffix in _TERM_NOUN_SUFFIXES):
        return True
    return False


def _term_occurs(source: str, term: str) -> bool:
    start = 0
    while True:
        index = source.find(term, start)
        if index < 0:
            return False
        before = source[index - 1] if index > 0 else ""
        after_index = index + len(term)
        after = source[after_index] if after_index < len(source) else ""
        if term[0].isdigit() and before.isdigit():
            start = index + 1
            continue
        if term[-1].isdigit() and after.isdigit():
            start = index + 1
            continue
        return True


def _read_translation_rows(path: Path, target_languages: list[tuple[str, str]]) -> dict[str, dict[str, Any]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb["Translations"]
        headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
        header_index = {str(header): idx + 1 for idx, header in enumerate(headers)}
        required = [*FIXED_COLUMNS, *[header for header, _ in target_languages]]
        missing = [name for name in required if name not in header_index]
        if missing:
            raise ValueError(f"translation workbook missing columns: {missing}")

        rows: dict[str, dict[str, Any]] = {}
        duplicates: list[str] = []
        for row_idx in range(2, ws.max_row + 1):
            para_id = _clean_cell(ws.cell(row_idx, header_index["para_id"]).value)
            if not para_id:
                continue
            if para_id in rows:
                duplicates.append(para_id)
                continue
            rows[para_id] = {
                "source_file": _clean_cell(ws.cell(row_idx, header_index["source_file"]).value),
                "para_id": para_id,
                "para_index": int(ws.cell(row_idx, header_index["para_index"]).value),
                "style": _clean_cell(ws.cell(row_idx, header_index["style"]).value),
                "CN": _clean_cell(ws.cell(row_idx, header_index["CN"]).value),
                "protected_tokens": _parse_json_cell(ws.cell(row_idx, header_index["protected_tokens"]).value, []),
                "term_hits": _parse_json_cell(ws.cell(row_idx, header_index["term_hits_json"]).value, []),
                "translations": {
                    header: _clean_cell(ws.cell(row_idx, header_index[header]).value)
                    for header, _ in target_languages
                },
            }
    finally:
        wb.close()
    if duplicates:
        raise ValueError(f"duplicate paragraph IDs: {duplicates}")
    return rows


def _expected_paragraphs(manifest: dict[str, Any]) -> dict[str, dict[str, Any]]:
    expected: dict[str, dict[str, Any]] = {}
    for doc in manifest.get("documents", []):
        for row in doc.get("paragraphs", []):
            expected[str(row["para_id"])] = row
    return expected


def _ordered_expected_rows(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    ordered: list[dict[str, Any]] = []
    for doc in manifest.get("documents", []):
        ordered.extend(doc.get("paragraphs", []))
    return ordered


def _validate_row_coverage(expected: dict[str, dict[str, Any]], rows: dict[str, dict[str, Any]]) -> None:
    expected_ids = set(expected)
    actual_ids = set(rows)
    missing = sorted(expected_ids - actual_ids)
    extra = sorted(actual_ids - expected_ids)
    if missing:
        raise ValueError(f"missing paragraph IDs: {missing}")
    if extra:
        raise ValueError(f"extra paragraph IDs: {extra}")


def _validate_input_drift(input_dir: Path, manifest: dict[str, Any]) -> None:
    for doc in manifest.get("documents", []):
        path = input_dir / str(doc["source_file"])
        if _sha256_file(path) != doc.get("sha256"):
            raise ValueError(f"input drift detected: {path.name}")


def _validate_all_translations(
    expected: dict[str, dict[str, Any]],
    rows: dict[str, dict[str, Any]],
    target_languages: list[tuple[str, str]],
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    for para_id, expected_row in expected.items():
        row = rows[para_id]
        source = str(expected_row["CN"])
        term_hits = expected_row.get("term_hits", [])
        protected_tokens = expected_row.get("protected_tokens", [])
        for header, code in target_languages:
            translation = row["translations"].get(header, "")
            issues.extend(_validate_translation(source, translation, protected_tokens, term_hits, header, code, row))
    return issues


def _resolve_language_pairs(
    languages: list[str] | None,
) -> list[tuple[str, str]]:
    if not languages:
        return list(TARGET_LANGUAGES)
    selected_headers = _resolve_language_headers(languages, valid_languages=SUPPORTED_LANGUAGES)
    language_by_header = {header: code for header, code in TARGET_LANGUAGES}
    return [(header, language_by_header[header]) for header in selected_headers]


def _resolve_language_headers(
    languages: list[str] | None,
    *,
    valid_languages: list[tuple[str, str]] | tuple[tuple[str, str], ...] = TARGET_LANGUAGES,
) -> list[str]:
    valid_headers = {header for header, _ in valid_languages}
    code_by_header = {header: code for header, code in valid_languages}
    valid_codes: dict[str, str] = {}
    for _, code in valid_languages:
        valid_codes.setdefault(code, CANONICAL_LANGUAGE_HEADER[code])
    if not languages:
        return [header for header, _ in valid_languages]

    selected: list[str] = []
    for lang in languages:
        raw = str(lang).strip()
        if raw in valid_headers:
            selected.append(CANONICAL_LANGUAGE_HEADER[code_by_header[raw]])
            continue
        code = raw.lower()
        if code in valid_codes:
            selected.append(valid_codes[code])
            continue
        normalized = raw.upper()
        if normalized in valid_headers:
            selected.append(CANONICAL_LANGUAGE_HEADER[code_by_header[normalized]])
            continue
        raise ValueError(f"unsupported language: {lang}")
    return selected


def _manifest_languages(manifest: dict[str, Any]) -> list[tuple[str, str]]:
    languages = []
    for item in manifest.get("languages", []):
        header = str(item.get("header", "")).strip().upper()
        code = str(item.get("code", "")).strip()
        if header and code:
            languages.append((header, code))
    return languages or list(TARGET_LANGUAGES)


def _infer_language_pairs_from_terms(pairs: list[tuple[Path, Path]]) -> list[tuple[str, str]]:
    selected: list[tuple[str, str]] = []
    seen: set[str] = set()
    for _, term_path in pairs:
        terms = load_announcement_terms(term_path)
        for spec in terms.languages:
            if spec.code not in seen:
                selected.append((spec.header, spec.code))
                seen.add(spec.code)
    if not selected:
        raise ValueError("no supported target language columns found in announcement terms")
    return selected


def _read_ai_response_rows(response_path: Path, expected_rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    if not response_path.exists():
        raise ValueError(f"AI response not found: {response_path}")

    rows: list[dict[str, str]] = []
    seen: set[str] = set()
    for line_no, line in enumerate(response_path.read_text(encoding="utf-8-sig").splitlines(), start=1):
        if not line.strip():
            continue
        try:
            row = json.loads(line)
        except json.JSONDecodeError as exc:
            raise ValueError(f"invalid JSON in {response_path.name}:{line_no}: {exc}") from exc
        if set(row) != {"para_id", "translation"}:
            raise ValueError(f"AI response row {line_no} must contain only para_id and translation")
        para_id = _clean_cell(row.get("para_id"))
        translation = _clean_cell(row.get("translation"))
        if not para_id:
            raise ValueError(f"AI response row {line_no} missing para_id")
        if para_id in seen:
            raise ValueError(f"duplicate AI response para_id: {para_id}")
        seen.add(para_id)
        rows.append({"para_id": para_id, "translation": translation})

    expected_ids = [str(row["para_id"]) for row in expected_rows]
    actual_ids = [row["para_id"] for row in rows]
    if actual_ids != expected_ids:
        missing = sorted(set(expected_ids) - set(actual_ids))
        extra = sorted(set(actual_ids) - set(expected_ids))
        if missing:
            raise ValueError(f"AI response missing paragraph IDs: {missing}")
        if extra:
            raise ValueError(f"AI response extra paragraph IDs: {extra}")
        raise ValueError("AI response paragraph order mismatch")
    return rows


def _write_ai_responses_to_workbook(
    workbook_path: Path,
    expected: dict[str, dict[str, Any]],
    translations_by_header: dict[str, dict[str, str]],
) -> None:
    wb = load_workbook(workbook_path)
    try:
        ws = wb["Translations"]
        headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
        header_index = {str(header): idx + 1 for idx, header in enumerate(headers)}
        row_by_para_id = {}
        for row_idx in range(2, ws.max_row + 1):
            para_id = _clean_cell(ws.cell(row_idx, header_index["para_id"]).value)
            if para_id:
                row_by_para_id[para_id] = row_idx
        missing_in_workbook = sorted(set(expected) - set(row_by_para_id))
        if missing_in_workbook:
            raise ValueError(f"translation workbook missing paragraph IDs: {missing_in_workbook}")
        for header, translations in translations_by_header.items():
            if header not in header_index:
                raise ValueError(f"translation workbook missing language column: {header}")
            for para_id, translation in translations.items():
                ws.cell(row_by_para_id[para_id], header_index[header]).value = translation
        wb.save(workbook_path)
    finally:
        wb.close()


def _validate_translation(
    source: str,
    translation: str,
    protected_tokens: list[str],
    term_hits: list[dict[str, Any]],
    lang_header: str,
    lang_code: str,
    row: dict[str, Any],
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    base = {
        "source_file": row["source_file"],
        "para_id": row["para_id"],
        "para_index": row["para_index"],
        "lang": lang_code,
        "source": source,
        "translation": translation,
    }
    if not translation.strip():
        issues.append({**base, "check_type": "empty_translation", "message": "Translation is empty"})
        return issues
    if _CJK_RE.search(translation):
        issues.append({**base, "check_type": "chinese_residue", "message": "Chinese residue found"})

    for token in protected_tokens:
        if token and token not in translation:
            issues.append({**base, "check_type": "protected_token_missing", "message": f"Missing protected token: {token}"})

    if _BRACKET_RE.search(source) and not _BRACKET_RE.search(translation):
        issues.append({**base, "check_type": "bracket_shape_missing", "message": "Source has bracketed content but translation has no brackets"})

    for hit in term_hits:
        target = str(hit.get("targets", {}).get(lang_header, "")).strip()
        if target and not _contains_term(translation, target):
            issues.append(
                {
                    **base,
                    "check_type": "term_missing",
                    "message": f"Missing term target for {hit.get('source')}: {target}",
                }
            )
    return issues


def _write_output_docx(
    input_dir: Path,
    manifest: dict[str, Any],
    rows: dict[str, dict[str, Any]],
    output_dir: Path,
    target_languages: list[tuple[str, str]],
) -> list[Path]:
    output_paths: list[Path] = []
    for doc_record in manifest.get("documents", []):
        source_name = str(doc_record["source_file"])
        source_path = input_dir / source_name
        paragraphs = doc_record.get("paragraphs", [])
        for header, code in target_languages:
            doc = Document(str(source_path))
            for para_record in paragraphs:
                para_id = str(para_record["para_id"])
                index = int(para_record["para_index"])
                translation = rows[para_id]["translations"][header]
                _replace_paragraph_text(doc.paragraphs[index], translation)
            output_path = output_dir / f"{Path(source_name).stem}_{code}.docx"
            doc.save(output_path)
            output_paths.append(output_path)
    return output_paths


def _write_qa_summary(path: Path, issues: list[dict[str, Any]], output_paths: list[Path]) -> None:
    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"
    summary.append(["metric", "value"])
    summary.append(["hard_blockers", len(issues)])
    summary.append(["output_docx_count", len(output_paths)])
    summary.append(["generated_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    details = wb.create_sheet("Issues")
    issue_headers = ["source_file", "para_id", "para_index", "lang", "check_type", "message", "source", "translation"]
    details.append(issue_headers)
    for issue in issues:
        details.append([issue.get(header, "") for header in issue_headers])
    outputs = wb.create_sheet("Outputs")
    outputs.append(["path"])
    for output_path in output_paths:
        outputs.append([str(output_path)])
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _load_manifest(work_dir: Path) -> dict[str, Any]:
    path = work_dir / MANIFEST_NAME
    if not path.exists():
        raise ValueError(f"manifest not found: {path}")
    return json.loads(path.read_text(encoding="utf-8"))


def _work_dir(input_dir: Path) -> Path:
    return input_dir / WORK_DIR_NAME / HARNESS_DIR_NAME


def _inspect_unsupported(path: Path, doc: Document) -> dict[str, int]:
    textboxes = 0
    hyperlinks = 0
    media = 0
    try:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            media = sum(1 for name in names if name.startswith("word/media/"))
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
            textboxes = document_xml.count("w:txbxContent")
            hyperlinks = document_xml.count("<w:hyperlink")
    except (KeyError, zipfile.BadZipFile):
        pass
    return {
        "tables": len(doc.tables),
        "media": media,
        "textboxes": textboxes,
        "hyperlinks": hyperlinks,
    }


def _paragraph_id(source_file: str, para_index: int, source: str) -> str:
    digest = hashlib.sha1(f"{source_file}\0{para_index}\0{source}".encode("utf-8")).hexdigest()[:10]
    return f"{Path(source_file).stem}:{para_index:04d}:{digest}"


def _protected_tokens(text: str) -> list[str]:
    tokens = []
    seen = set()
    for pattern in (_DATE_RE, _TIME_RANGE_RE, _BRACKET_TOKEN_RE, _NUMBER_RE):
        for token in pattern.findall(str(text)):
            if isinstance(token, tuple):
                token = next((part for part in token if part), "")
            token = str(token).strip()
            if not token or _CJK_RE.search(token) or token in seen:
                continue
            seen.add(token)
            tokens.append(token)
    return tokens


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _contains_term(translation: str, target: str) -> bool:
    normalized_translation = re.sub(r"\s+", " ", translation).casefold()
    normalized_target = re.sub(r"\s+", " ", target).casefold()
    return normalized_target in normalized_translation


def _language_code_for_header(header: str) -> str:
    return LANGUAGE_CODE_BY_HEADER.get(str(header or "").strip().upper(), "")


def _is_temp_file(path: Path) -> bool:
    return path.name.startswith("~$")


def _is_loose_announcement_terms_file(path: Path) -> bool:
    stem = path.stem.lower()
    return "announcement_terms" in stem or "术语译文交付表" in path.stem


def _is_generated_docx(path: Path) -> bool:
    stem = path.stem.lower()
    generated_suffixes = {f"_{code}" for _, code in SUPPORTED_LANGUAGES}
    generated_suffixes.update(f"_{header.lower()}" for header, _ in SUPPORTED_LANGUAGES)
    return any(stem.endswith(suffix) for suffix in generated_suffixes)


def _select_term_file_for_source(source_path: Path, term_files: list[Path]) -> Path:
    exact_prefix = f"{source_path.stem}_announcement_terms_"
    exact = [path for path in term_files if path.name.startswith(exact_prefix)]
    if exact:
        return sorted(exact, key=lambda item: item.name)[-1]
    if len(term_files) == 1:
        return term_files[0]
    raise ValueError(f"ambiguous announcement term workbook for {source_path.name}")


def _extract_date_stamp(path: Path) -> str:
    match = re.search(r"20\d{6}", path.name)
    return match.group(0) if match else datetime.now().strftime("%Y%m%d")


def _convert_txt_to_docx(source_path: Path, target_path: Path) -> None:
    text = source_path.read_text(encoding="utf-8-sig")
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(target_path)


def _clean_cell(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return "" if text.lower() == "nan" else text


def _parse_json_cell(value: Any, default: Any) -> Any:
    text = _clean_cell(value)
    if not text:
        return default
    return json.loads(text)


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.write_text(
        "\n".join(json.dumps(row, ensure_ascii=False) for row in rows) + ("\n" if rows else ""),
        encoding="utf-8",
    )
