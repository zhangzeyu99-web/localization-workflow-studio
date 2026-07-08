"""Announcement DOCX task inspection, staging, and workpack/manifest preparation.

Implementation module split out of utils/announcement_docx_harness.py. Import
these symbols through utils.announcement_docx_harness to keep the public
surface stable.
"""
from __future__ import annotations

import hashlib
import json
import re
import shutil
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook

from utils.announcement_docx_common import (
    FIXED_COLUMNS,
    HARNESS_DIR_NAME,
    MANIFEST_NAME,
    TRANSLATION_WORKBOOK_NAME,
    WORK_DIR_NAME,
    _CJK_RE,
    _is_generated_docx,
    _is_temp_file,
    _resolve_language_pairs,
    _sha256_file,
    _work_dir,
    _write_jsonl,
)
from utils.announcement_docx_terms import (
    _find_term_hits,
    _read_announcement_language_specs,
    load_announcement_terms,
)

_DATE_RE = re.compile(r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b")
_TIME_RANGE_RE = re.compile(r"\b\d{1,2}:\d{2}\s*(?:[-–~]\s*\d{1,2}:\d{2})?\b")
_NUMBER_RE = re.compile(r"\d+")
_BRACKET_TOKEN_RE = re.compile(r"\[[^\]]+\]|\([^)]+\)|\{[^}]+\}|【[^】]+】|（[^）]+）")


@dataclass(frozen=True)
class PreparedAnnouncementHarness:
    work_dir: Path
    translation_workbook: Path
    manifest_path: Path
    row_count: int
    doc_count: int


@dataclass(frozen=True)
class AnnouncementTaskInspection:
    input_dir: Path
    source_files: list[Path]
    term_files: list[Path]
    reference_files: list[Path]
    languages: list[tuple[str, str]]


@dataclass(frozen=True)
class StagedAnnouncementTask:
    staging_dir: Path
    source_files: list[Path]
    term_files: list[Path]
    languages: list[tuple[str, str]]


def prepare_announcement_docx_harness(
    input_dir: str | Path,
    *,
    languages: list[str] | None = None,
) -> PreparedAnnouncementHarness:
    """Create a structured workbook and per-language workpacks for DOCX announcements."""
    input_dir = Path(input_dir)
    work_dir = _work_dir(input_dir)
    if work_dir.exists():
        shutil.rmtree(work_dir)
    work_dir.mkdir(parents=True, exist_ok=True)
    pairs = discover_announcement_docx_pairs(input_dir)
    target_languages = (
        _resolve_language_pairs(languages)
        if languages
        else _infer_language_pairs_from_terms(pairs)
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Translations"
    headers = [*FIXED_COLUMNS, *[header for header, _ in target_languages]]
    ws.append(headers)

    manifest: dict[str, Any] = {
        "version": 1,
        "input_dir": str(input_dir),
        "languages": [{"header": header, "code": code} for header, code in target_languages],
        "documents": [],
    }
    workpack_rows: dict[str, list[dict[str, Any]]] = {code: [] for _, code in target_languages}
    row_count = 0

    for docx_path, term_path in pairs:
        terms = load_announcement_terms(term_path)
        doc = Document(str(docx_path))
        unsupported = _inspect_unsupported(docx_path, doc)
        doc_record = {
            "source_file": docx_path.name,
            "term_file": term_path.name,
            "sha256": _sha256_file(docx_path),
            "unsupported": unsupported,
            "paragraphs": [],
        }
        for index, paragraph in enumerate(doc.paragraphs):
            source = paragraph.text
            if not source or not source.strip():
                continue
            para_id = _paragraph_id(docx_path.name, index, source)
            style = paragraph.style.name if paragraph.style else ""
            term_hits = _find_term_hits(source, terms)
            protected_tokens = _protected_tokens(source)
            row = {
                "source_file": docx_path.name,
                "para_id": para_id,
                "para_index": index,
                "style": style,
                "CN": source,
                "protected_tokens": protected_tokens,
                "term_hits": term_hits,
            }
            doc_record["paragraphs"].append(row)
            ws.append(
                [
                    docx_path.name,
                    para_id,
                    index,
                    style,
                    source,
                    json.dumps(protected_tokens, ensure_ascii=False),
                    json.dumps(term_hits, ensure_ascii=False),
                    *["" for _ in target_languages],
                ]
            )
            for header, code in target_languages:
                workpack_rows[code].append(
                    {
                        "source_file": docx_path.name,
                        "para_id": para_id,
                        "para_index": index,
                        "source": source,
                        "term_hits": [
                            {"source": hit["source"], "target": hit["targets"].get(header, "")}
                            for hit in term_hits
                            if hit["targets"].get(header)
                        ],
                        "protected_tokens": protected_tokens,
                    }
                )
            row_count += 1
        manifest["documents"].append(doc_record)

    translation_workbook = work_dir / TRANSLATION_WORKBOOK_NAME
    wb.save(translation_workbook)
    wb.close()

    manifest_path = work_dir / MANIFEST_NAME
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    for _, code in target_languages:
        _write_jsonl(work_dir / f"workpack_{code}.jsonl", workpack_rows[code])

    return PreparedAnnouncementHarness(
        work_dir=work_dir,
        translation_workbook=translation_workbook,
        manifest_path=manifest_path,
        row_count=row_count,
        doc_count=len(pairs),
    )


def inspect_announcement_task_dir(input_dir: str | Path) -> AnnouncementTaskInspection:
    """Identify loose announcement sources, term delivery files, references, and target languages."""
    input_dir = Path(input_dir)
    source_files = [
        path
        for path in sorted(input_dir.iterdir(), key=lambda item: item.name)
        if path.is_file()
        and not _is_temp_file(path)
        and path.suffix.lower() in {".docx", ".txt"}
        and not _is_generated_docx(path)
    ]
    term_files: list[Path] = []
    reference_files: list[Path] = []
    language_pairs: list[tuple[str, str]] = []
    seen_languages: set[str] = set()
    for path in sorted(input_dir.glob("*.xlsx"), key=lambda item: item.name):
        if _is_temp_file(path):
            continue
        try:
            language_specs = _read_announcement_language_specs(path)
        except Exception:
            continue
        if not language_specs:
            continue
        if _is_loose_announcement_terms_file(path):
            term_files.append(path)
            for spec in language_specs:
                header = spec.header.strip().upper()
                code = spec.code
                if header not in seen_languages:
                    language_pairs.append((header, code))
                    seen_languages.add(header)
        else:
            reference_files.append(path)

    return AnnouncementTaskInspection(
        input_dir=input_dir,
        source_files=source_files,
        term_files=term_files,
        reference_files=reference_files,
        languages=language_pairs,
    )


def stage_announcement_task_dir(input_dir: str | Path) -> StagedAnnouncementTask:
    """Normalize a loose task folder into the strict DOCX+announcement_terms harness layout."""
    inspection = inspect_announcement_task_dir(input_dir)
    if not inspection.source_files:
        raise ValueError(f"no announcement source .docx/.txt files found in {inspection.input_dir}")
    if not inspection.term_files:
        raise ValueError(f"no announcement term delivery workbook found in {inspection.input_dir}")
    if not inspection.languages:
        raise ValueError(f"no supported target language columns found in {inspection.input_dir}")

    staging_dir = inspection.input_dir / WORK_DIR_NAME / HARNESS_DIR_NAME / "source_input"
    if staging_dir.exists():
        shutil.rmtree(staging_dir)
    staging_dir.mkdir(parents=True, exist_ok=True)

    staged_sources: list[Path] = []
    staged_terms: list[Path] = []
    for source_path in inspection.source_files:
        staged_source = staging_dir / f"{source_path.stem}.docx"
        if source_path.suffix.lower() == ".docx":
            shutil.copy2(source_path, staged_source)
        else:
            _convert_txt_to_docx(source_path, staged_source)
        staged_sources.append(staged_source)

        term_path = _select_term_file_for_source(source_path, inspection.term_files)
        term_date = _extract_date_stamp(term_path)
        staged_term = staging_dir / f"{staged_source.stem}_announcement_terms_{term_date}.xlsx"
        shutil.copy2(term_path, staged_term)
        staged_terms.append(staged_term)

    return StagedAnnouncementTask(
        staging_dir=staging_dir,
        source_files=staged_sources,
        term_files=staged_terms,
        languages=inspection.languages,
    )


def discover_announcement_docx_pairs(input_dir: str | Path) -> list[tuple[Path, Path]]:
    input_dir = Path(input_dir)
    term_files = sorted(
        path
        for path in input_dir.glob("*_announcement_terms_*.xlsx")
        if path.is_file() and not _is_temp_file(path)
    )
    docx_files = sorted(
        path
        for path in input_dir.glob("*.docx")
        if path.is_file() and not _is_temp_file(path) and not _is_generated_docx(path)
    )
    pairs: list[tuple[Path, Path]] = []
    for docx_path in docx_files:
        prefix = f"{docx_path.stem}_announcement_terms_"
        matches = [path for path in term_files if path.name.startswith(prefix)]
        if not matches:
            raise ValueError(f"announcement term workbook not found for {docx_path.name}")
        pairs.append((docx_path, sorted(matches, key=lambda item: item.name)[-1]))
    if not pairs:
        raise ValueError(f"no announcement DOCX files found in {input_dir}")
    return pairs


def _infer_language_pairs_from_terms(pairs: list[tuple[Path, Path]]) -> list[tuple[str, str]]:
    selected: list[tuple[str, str]] = []
    seen: set[str] = set()
    for _, term_path in pairs:
        terms = load_announcement_terms(term_path)
        for spec in terms.languages:
            if spec.code not in seen:
                selected.append((spec.header, spec.code))
                seen.add(spec.code)
    if not selected:
        raise ValueError("no supported target language columns found in announcement terms")
    return selected


def _inspect_unsupported(path: Path, doc: Document) -> dict[str, int]:
    textboxes = 0
    hyperlinks = 0
    media = 0
    try:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            media = sum(1 for name in names if name.startswith("word/media/"))
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
            textboxes = document_xml.count("w:txbxContent")
            hyperlinks = document_xml.count("<w:hyperlink")
    except (KeyError, zipfile.BadZipFile):
        pass
    return {
        "tables": len(doc.tables),
        "media": media,
        "textboxes": textboxes,
        "hyperlinks": hyperlinks,
    }


def _paragraph_id(source_file: str, para_index: int, source: str) -> str:
    digest = hashlib.sha1(f"{source_file}\0{para_index}\0{source}".encode("utf-8")).hexdigest()[:10]
    return f"{Path(source_file).stem}:{para_index:04d}:{digest}"


def _protected_tokens(text: str) -> list[str]:
    tokens = []
    seen = set()
    for pattern in (_DATE_RE, _TIME_RANGE_RE, _BRACKET_TOKEN_RE, _NUMBER_RE):
        for token in pattern.findall(str(text)):
            if isinstance(token, tuple):
                token = next((part for part in token if part), "")
            token = str(token).strip()
            if not token or _CJK_RE.search(token) or token in seen:
                continue
            seen.add(token)
            tokens.append(token)
    return tokens


def _is_loose_announcement_terms_file(path: Path) -> bool:
    stem = path.stem.lower()
    return "announcement_terms" in stem or "术语译文交付表" in path.stem


def _select_term_file_for_source(source_path: Path, term_files: list[Path]) -> Path:
    exact_prefix = f"{source_path.stem}_announcement_terms_"
    exact = [path for path in term_files if path.name.startswith(exact_prefix)]
    if exact:
        return sorted(exact, key=lambda item: item.name)[-1]
    if len(term_files) == 1:
        return term_files[0]
    raise ValueError(f"ambiguous announcement term workbook for {source_path.name}")


def _extract_date_stamp(path: Path) -> str:
    match = re.search(r"20\d{6}", path.name)
    return match.group(0) if match else datetime.now().strftime("%Y%m%d")


def _convert_txt_to_docx(source_path: Path, target_path: Path) -> None:
    text = source_path.read_text(encoding="utf-8-sig")
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(target_path)
