import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from utils.announcement_docx_harness import (
    TARGET_LANGUAGES,
    apply_announcement_translations,
    deliver_announcement_outputs,
    import_announcement_ai_responses,
    inspect_announcement_task_dir,
    load_announcement_terms,
    prepare_announcement_docx_harness,
    stage_announcement_task_dir,
)


SOURCE_TITLE = "\u670d\u52a1\u5668\u65f6\u95f4 2026 \u66f4\u65b0\u516c\u544a"
SOURCE_BODY = "\u6211\u4eec\u5c06\u4e8e\u670d\u52a1\u5668\u65f6\u95f4 2026/5/20 04:00-06:00 \u8fdb\u884c\u7ef4\u62a4\u3002"
SOURCE_SHORT = "\u8bad\u7ec3\u5e08\u8bf7\u7559\u610f"


def _write_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph(SOURCE_TITLE, style="Title")
    doc.add_paragraph(SOURCE_BODY)
    doc.add_paragraph(SOURCE_SHORT)
    doc.save(path)


def _write_terms(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Glossary"
    ws.append(["ID", "CN", "EN", "KR", "JP", "FR", "DE", "RU", "IT", "ES", "PT", "TK", "VI", "ID", "TH", "AR"])
    ws.append([
        "term_notice",
        "\u516c\u544a",
        "Notice",
        "공지",
        "Oshirase",
        "Annonce",
        "Ankündigung",
        "Объявление",
        "Annuncio",
        "Anuncio",
        "Anúncio",
        "Duyuru",
        "Thông báo",
        "Pengumuman",
        "ประกาศ",
        "إعلان",
    ])
    ws.append([
        "term_server_time_short",
        "\u670d\u52a1\u5668\u65f6\u95f4",
        "Server Time",
        "서버 시간",
        "Server Time",
        "Heure du serveur",
        "Serverzeit",
        "Время сервера",
        "Orario server",
        "Hora del servidor",
        "Horário de servidor",
        "Sunucu Saati",
        "Giờ máy chủ",
        "Waktu Server",
        "เวลาเซิร์ฟเวอร์",
        "وقت الخادم",
    ])
    ws.append([
        "term_server_time_long",
        "\u670d\u52a1\u5668\u65f6\u95f4 2026",
        "Server Time 2026",
        "서버 시간 2026",
        "Server Time 2026",
        "Heure du serveur 2026",
        "Serverzeit 2026",
        "Время сервера 2026",
        "Orario server 2026",
        "Hora del servidor 2026",
        "Horário de servidor 2026",
        "Sunucu Saati 2026",
        "Giờ máy chủ 2026",
        "Waktu Server 2026",
        "เวลาเซิร์ฟเวอร์ 2026",
        "وقت الخادم 2026",
    ])
    ws.append([
        "term_maintenance",
        "\u7ef4\u62a4",
        "Maintenance",
        "점검",
        "Maintenance",
        "Maintenance",
        "Wartung",
        "Обслуживание",
        "Manutenzione",
        "Mantenimiento",
        "Manutenção",
        "Bakım",
        "Bảo trì",
        "Maintenance",
        "ปรับปรุง",
        "صيانة",
    ])
    ws.append([
        "term_trainer",
        "\u8bad\u7ec3\u5e08",
        "Trainer",
        "트레이너",
        "Trainer",
        "Dresseur",
        "Trainer",
        "Тренер",
        "Allenatore",
        "Entrenador",
        "Treinador",
        "Eğitmen",
        "Huấn luyện viên",
        "Pelatih",
        "เทรนเนอร์",
        "مدرب",
    ])
    wb.save(path)


def _write_generic_noise_terms(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Glossary"
    ws.append(["ID", "CN", "EN"])
    ws.append(["generic_notice", "\u516c\u544a", "Notice"])
    ws.append(["generic_game", "\u6e38\u620f", "Game"])
    ws.append(["generic_use", "\u4f7f\u7528", "Use"])
    ws.append(["generic_issue", "\u95ee\u9898", "Issue"])
    ws.append(["generic_enter", "\u8fdb\u5165", "Enter"])
    ws.append(["generic_enter_game", "\u8fdb\u5165\u6e38\u620f", "Enter Game"])
    ws.append(["term_train", "\u65e0\u9650\u5217\u8f66", "Infinity Train"])
    ws.append(["term_train_event", "\u65e0\u9650\u5217\u8f66\u6d3b\u52a8", "Infinity Train Event"])
    ws.append(["term_artifact", "\u5723\u5668", "Artifact"])
    wb.save(path)


def _write_en_fr_terms(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Glossary"
    ws.append(["ID", "CN", "EN", "FR"])
    ws.append(["term_notice", "\u516c\u544a", "Notice", "Annonce"])
    ws.append(["term_trainer", "\u8bad\u7ec3\u5e08", "Trainer", "Dresseur"])
    wb.save(path)


def _write_kr_terms(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Glossary"
    ws.append(["ID", "CN", "KR"])
    ws.append(["term_hero", "\u82f1\u96c4", "\uc601\uc6c5"])
    ws.append(["term_awaken", "\u89c9\u9192", "\uac01\uc131"])
    wb.save(path)


def _write_terms_with_sentence_adaptations(path: Path) -> None:
    wb = Workbook()
    glossary = wb.active
    glossary.title = "Glossary"
    glossary.append(["ID", "CN", "EN", "FR"])
    glossary.append(["term_trainer", "\u8bad\u7ec3\u5e08", "Trainer", "Dresseur"])
    glossary.append(["term_maintenance", "\u7ef4\u62a4", "Maintenance", "Maintenance"])

    templates = wb.create_sheet("SentenceTemplates")
    templates.append(
        ["Priority", "MatchType", "ID", "AnnouncementCN", "OfficialCNTemplate", "EN", "FR"]
    )
    templates.append(
        [
            1,
            "official_exact",
            "sentence_exact",
            "\u7ef4\u62a4\u671f\u95f4\uff0c\u8bad\u7ec3\u5e08\u53ef\u9886\u53d62\u4efd\u5956\u52b1\u3002",
            "\u7ef4\u62a4\u671f\u95f4\uff0c\u8bad\u7ec3\u5e08\u53ef\u9886\u53d6<@1>\u4efd\u5956\u52b1\u3002",
            "During maintenance, coaches can claim <@1> rewards.",
            "Pendant la maintenance, les coaches peuvent obtenir <@1> recompenses.",
        ]
    )
    templates.append(
        [
            2,
            "official_similar",
            "sentence_similar",
            "\u8bad\u7ec3\u5e08",
            "\u8bad\u7ec3\u5e08\u8bf7\u7559\u610f\u6d3b\u52a8\u65f6\u95f4\u3002",
            "Trainers, please note the event schedule.",
            "Dresseurs, consultez les horaires de l'evenement.",
        ]
    )
    wb.save(path)


def _term_targets(term_hits_json: str, lang_header: str) -> str:
    hits = json.loads(term_hits_json)
    return " ".join(hit["targets"][lang_header] for hit in hits if hit["targets"].get(lang_header))


def _fill_translation_workbook(path: Path) -> None:
    wb = load_workbook(path)
    ws = wb["Translations"]
    headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
    fixed_headers = {
        "source_file",
        "para_id",
        "para_index",
        "style",
        "CN",
        "protected_tokens",
        "term_hits_json",
        "sentence_adaptations_json",
    }
    language_headers = [str(header) for header in headers if header not in fixed_headers]
    lang_cols = {header: headers.index(header) + 1 for header in language_headers}
    term_hits_col = headers.index("term_hits_json") + 1
    protected_tokens_col = headers.index("protected_tokens") + 1
    cn_col = headers.index("CN") + 1
    for row in range(2, ws.max_row + 1):
        source = str(ws.cell(row, cn_col).value or "")
        numbers = " ".join(__import__("re").findall(r"\d+", source))
        protected_tokens = " ".join(json.loads(str(ws.cell(row, protected_tokens_col).value or "[]")))
        for header in language_headers:
            terms = _term_targets(str(ws.cell(row, term_hits_col).value or "[]"), header)
            ws.cell(row, lang_cols[header]).value = f"{terms} {protected_tokens} {numbers}".strip() or f"{header} translated"
    wb.save(path)
    wb.close()


def _write_ai_response_files(work_dir: Path, response_dir: Path, languages=TARGET_LANGUAGES) -> None:
    response_dir.mkdir(parents=True, exist_ok=True)
    for _, code in languages:
        rows = []
        for line in (work_dir / f"workpack_{code}.jsonl").read_text(encoding="utf-8").splitlines():
            item = json.loads(line)
            parts = [term["target"] for term in item["term_hits"]]
            parts.extend(item["protected_tokens"])
            parts.append(f"{code} translated")
            translation = " ".join(dict.fromkeys(part for part in parts if part))
            if any(ch in item["source"] for ch in "()[]{}【】（）") and not any(ch in translation for ch in "()[]{}【】（）"):
                translation = f"[{translation}]"
            rows.append({"para_id": item["para_id"], "translation": translation})
        (response_dir / f"ai_response_{code}.jsonl").write_text(
            "\n".join(json.dumps(row, ensure_ascii=False) for row in rows) + "\n",
            encoding="utf-8",
        )


class AnnouncementDocxHarnessTests(unittest.TestCase):
    def test_load_terms_maps_duplicate_id_header_to_indonesian_by_position(self):
        with tempfile.TemporaryDirectory() as tmp:
            term_path = Path(tmp) / "sample_announcement_terms_20260526.xlsx"
            _write_terms(term_path)

            terms = load_announcement_terms(term_path)

            id_specs = [spec for spec in terms.languages if spec.header == "IDN"]
            self.assertEqual(len(id_specs), 1)
            self.assertEqual(id_specs[0].code, "idn")
            self.assertEqual(id_specs[0].column_index, 14)
            self.assertEqual(terms.by_language["IDN"]["\u8bad\u7ec3\u5e08"].target, "Pelatih")

    def test_prepare_uses_canonical_full_language_headers_and_alias_inputs(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "sample.docx")
            _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir, languages=["tk", "id", "ar"])
            manifest = json.loads(prepared.manifest_path.read_text(encoding="utf-8"))

            self.assertEqual(
                manifest["languages"],
                [{"header": "TR", "code": "tr"}, {"header": "IDN", "code": "idn"}, {"header": "AR", "code": "ar"}],
            )
            wb = load_workbook(prepared.translation_workbook, read_only=True, data_only=True)
            try:
                ws = wb["Translations"]
                headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
                self.assertEqual(headers[-3:], ["TR", "IDN", "AR"])
            finally:
                wb.close()

    def test_prepare_builds_translation_workbook_and_longest_term_hits(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "sample.docx")
            _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir)

            self.assertTrue(prepared.translation_workbook.exists())
            self.assertTrue(prepared.manifest_path.exists())
            self.assertTrue((prepared.work_dir / "workpack_en.jsonl").exists())

            wb = load_workbook(prepared.translation_workbook, read_only=True, data_only=True)
            try:
                ws = wb["Translations"]
                headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
                self.assertEqual(
                    headers,
                    [
                        "source_file",
                        "para_id",
                        "para_index",
                        "style",
                        "CN",
                        "protected_tokens",
                        "term_hits_json",
                        "sentence_adaptations_json",
                        "EN",
                        "KR",
                        "JP",
                        "FR",
                        "DE",
                        "RU",
                        "IT",
                        "ES",
                        "PT",
                        "TR",
                        "VI",
                        "IDN",
                        "TH",
                        "AR",
                    ],
                )
                first_hits = json.loads(ws.cell(2, headers.index("term_hits_json") + 1).value)
                self.assertIn("\u670d\u52a1\u5668\u65f6\u95f4", [hit["source"] for hit in first_hits])
                self.assertNotIn("\u670d\u52a1\u5668\u65f6\u95f4 2026", [hit["source"] for hit in first_hits])
                self.assertNotIn("\u516c\u544a", [hit["source"] for hit in first_hits])
                self.assertEqual(ws.max_row, 4)
            finally:
                wb.close()

    def test_load_terms_reads_optional_sentence_adaptation_sheet(self):
        with tempfile.TemporaryDirectory() as tmp:
            term_path = Path(tmp) / "sample_announcement_terms_20260526.xlsx"
            _write_terms_with_sentence_adaptations(term_path)

            terms = load_announcement_terms(term_path)

            self.assertEqual(len(terms.sentence_adaptations), 2)
            exact = terms.sentence_adaptations[0]
            self.assertEqual(exact.priority, 1)
            self.assertEqual(exact.match_type, "official_exact")
            self.assertEqual(exact.entry_id, "sentence_exact")
            self.assertEqual(exact.targets["EN"], "During maintenance, coaches can claim <@1> rewards.")

    def test_load_terms_rejects_invalid_sentence_adaptation_match_type(self):
        with tempfile.TemporaryDirectory() as tmp:
            term_path = Path(tmp) / "sample_announcement_terms_20260526.xlsx"
            _write_terms_with_sentence_adaptations(term_path)
            wb = load_workbook(term_path)
            wb["SentenceTemplates"].cell(2, 2).value = "copy_anyway"
            wb.save(term_path)
            wb.close()

            with self.assertRaisesRegex(ValueError, "unsupported MatchType"):
                load_announcement_terms(term_path)

    def test_inspect_reads_languages_from_glossary_when_sentence_sheet_is_active(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            (task_dir / "notice.txt").write_text("\u8bad\u7ec3\u5e08\u8bf7\u7559\u610f\n", encoding="utf-8")
            term_path = task_dir / "notice_announcement_terms_20260526.xlsx"
            _write_terms_with_sentence_adaptations(term_path)
            wb = load_workbook(term_path)
            wb.active = wb.sheetnames.index("SentenceTemplates")
            wb.save(term_path)
            wb.close()

            inspection = inspect_announcement_task_dir(task_dir)

            self.assertEqual(inspection.languages, [("EN", "en"), ("FR", "fr")])

    def test_prepare_retrieves_exact_template_with_dynamic_value_before_similar_reference(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u7ef4\u62a4\u671f\u95f4,\u8bad\u7ec3\u5e08\u53ef\u9886\u53d63\u4efd\u5956\u52b1\u3002")
            doc.save(task_dir / "sample.docx")
            _write_terms_with_sentence_adaptations(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir, languages=["en"])
            workpack = json.loads(
                (prepared.work_dir / "workpack_en.jsonl").read_text(encoding="utf-8").splitlines()[0]
            )

            self.assertEqual(
                [item["match_type"] for item in workpack["sentence_adaptations"]],
                ["official_exact", "official_similar"],
            )
            self.assertEqual(workpack["sentence_adaptations"][0]["target"], "During maintenance, coaches can claim <@1> rewards.")
            wb = load_workbook(prepared.translation_workbook, read_only=True, data_only=True)
            ws = wb["Translations"]
            headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
            adaptations = json.loads(ws.cell(2, headers.index("sentence_adaptations_json") + 1).value)
            wb.close()
            self.assertEqual([item["match_type"] for item in adaptations], ["official_exact", "official_similar"])

    def test_exact_sentence_adaptation_avoids_mechanical_term_rejection(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u7ef4\u62a4\u671f\u95f4\uff0c\u8bad\u7ec3\u5e08\u53ef\u9886\u53d62\u4efd\u5956\u52b1\u3002")
            doc.save(task_dir / "sample.docx")
            _write_terms_with_sentence_adaptations(task_dir / "sample_announcement_terms_20260526.xlsx")
            prepared = prepare_announcement_docx_harness(task_dir, languages=["en"])
            row = json.loads((prepared.work_dir / "workpack_en.jsonl").read_text(encoding="utf-8").strip())
            response_dir = task_dir / "responses"
            response_dir.mkdir()
            (response_dir / "ai_response_en.jsonl").write_text(
                json.dumps(
                    {
                        "para_id": row["para_id"],
                        "translation": "During maintenance, coaches can claim 2 rewards.",
                    }
                )
                + "\n",
                encoding="utf-8",
            )

            imported = import_announcement_ai_responses(
                task_dir,
                response_dir=response_dir,
                languages=["en"],
            )
            applied = apply_announcement_translations(task_dir, imported.translation_workbook)

            self.assertEqual(applied.hard_blockers, 0)
            qa = load_workbook(applied.qa_summary_path, read_only=True, data_only=True)
            summary = dict(qa["Summary"].iter_rows(min_row=2, values_only=True))
            qa.close()
            self.assertEqual(summary["sentence_adaptation_hit_rows"], 1)
            self.assertEqual(summary["official_exact_sentence_hits"], 1)
            self.assertEqual(summary["official_similar_sentence_hits"], 1)

    def test_prepare_filters_generic_terms_before_term_hits(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u516c\u544a\uff1a\u65e0\u9650\u5217\u8f66\u6d3b\u52a8\u5f00\u542f\uff0c\u8fdb\u5165\u6e38\u620f\u540e\u4f7f\u7528\u5723\u5668\u4fee\u590d\u95ee\u9898\u3002")
            doc.save(task_dir / "sample.docx")
            _write_generic_noise_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir, languages=["en"])
            rows = [
                json.loads(line)
                for line in (prepared.work_dir / "workpack_en.jsonl").read_text(encoding="utf-8").splitlines()
            ]
            targets = [hit["target"] for row in rows for hit in row["term_hits"]]

            self.assertEqual(targets, ["Infinity Train Event", "Artifact"])
            self.assertNotIn("Notice", targets)
            self.assertNotIn("Game", targets)
            self.assertNotIn("Use", targets)
            self.assertNotIn("Issue", targets)
            self.assertNotIn("Enter", targets)
            self.assertNotIn("Enter Game", targets)

    def test_prepare_infers_target_languages_from_term_columns(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "sample.docx")
            _write_en_fr_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir)

            wb = load_workbook(prepared.translation_workbook, read_only=True, data_only=True)
            ws = wb["Translations"]
            headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
            self.assertEqual(headers[-2:], ["EN", "FR"])
            self.assertNotIn("DE", headers)
            wb.close()
            self.assertTrue((prepared.work_dir / "workpack_en.jsonl").exists())
            self.assertTrue((prepared.work_dir / "workpack_fr.jsonl").exists())
            self.assertFalse((prepared.work_dir / "workpack_de.jsonl").exists())
            manifest = json.loads(prepared.manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(
                manifest["languages"],
                [{"header": "EN", "code": "en"}, {"header": "FR", "code": "fr"}],
            )

    def test_inspect_and_stage_loose_txt_task_infers_en_fr(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            source = task_dir / "\u5192\u9669\u516c\u544a.txt"
            source.write_text("\u516c\u544a\n\u8bad\u7ec3\u5e08\u8bf7\u7559\u610f\n", encoding="utf-8")
            _write_en_fr_terms(task_dir / "\u5192\u9669\u516c\u544a-\u672f\u8bed\u8bd1\u6587\u4ea4\u4ed8\u8868-20260526.xlsx")
            _write_en_fr_terms(task_dir / "Y \u8bed\u8a00\u8868-\u82f1\u6587EN.xlsx")

            inspection = inspect_announcement_task_dir(task_dir)
            staged = stage_announcement_task_dir(task_dir)
            prepared = prepare_announcement_docx_harness(staged.staging_dir)

            self.assertEqual([path.name for path in inspection.source_files], [source.name])
            self.assertEqual(inspection.languages, [("EN", "en"), ("FR", "fr")])
            self.assertEqual([path.name for path in inspection.reference_files], ["Y \u8bed\u8a00\u8868-\u82f1\u6587EN.xlsx"])
            self.assertTrue((staged.staging_dir / "\u5192\u9669\u516c\u544a.docx").exists())
            self.assertTrue((staged.staging_dir / "\u5192\u9669\u516c\u544a_announcement_terms_20260526.xlsx").exists())
            manifest = json.loads(prepared.manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(
                manifest["languages"],
                [{"header": "EN", "code": "en"}, {"header": "FR", "code": "fr"}],
            )

    def test_inspect_and_stage_loose_docx_task_infers_kr(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "notice.docx")
            _write_kr_terms(task_dir / "\u4e3b\u5bb0\u516c\u544a-\u672f\u8bed\u8bd1\u6587\u4ea4\u4ed8\u8868-20260526.xlsx")

            inspection = inspect_announcement_task_dir(task_dir)
            staged = stage_announcement_task_dir(task_dir)
            prepared = prepare_announcement_docx_harness(staged.staging_dir)

            self.assertEqual(inspection.languages, [("KR", "ko")])
            self.assertTrue((staged.staging_dir / "notice.docx").exists())
            self.assertTrue((staged.staging_dir / "notice_announcement_terms_20260526.xlsx").exists())
            manifest = json.loads(prepared.manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(manifest["languages"], [{"header": "KR", "code": "ko"}])

    def test_prepare_import_apply_supports_korean_term_column(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u82f1\u96c4\u89c9\u9192 2026/5/20")
            doc.save(task_dir / "sample.docx")
            _write_kr_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir)
            response_dir = task_dir / "ai_responses"
            _write_ai_response_files(prepared.work_dir, response_dir, languages=[("KR", "ko")])
            imported = import_announcement_ai_responses(task_dir, response_dir=response_dir)
            applied = apply_announcement_translations(task_dir, prepared.translation_workbook)

            self.assertEqual(imported.languages, ["KR"])
            self.assertEqual(applied.hard_blockers, 0)
            self.assertEqual(len(applied.output_docx_paths), 1)
            self.assertEqual(applied.output_docx_paths[0].name, "sample_ko.docx")

    def test_prepare_prefers_kr_jp_headers_when_language_codes_are_selected(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u82f1\u96c4\u89c9\u9192 2026/5/20")
            doc.save(task_dir / "sample.docx")
            _write_kr_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir, languages=["ko", "ja"])
            manifest = json.loads(prepared.manifest_path.read_text(encoding="utf-8"))

            self.assertEqual(manifest["languages"], [{"header": "KR", "code": "ko"}, {"header": "JP", "code": "ja"}])

    def test_prepare_does_not_match_numeric_terms_inside_longer_numbers(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("48\u5c0f\u65f6\u540e\u53ef\u786e\u8ba4")
            doc.save(task_dir / "sample.docx")
            wb = Workbook()
            ws = wb.active
            ws.append(["ID", "CN", "EN", "KR", "JP", "FR", "DE", "RU", "IT", "ES", "PT", "TK", "VI", "ID", "TH", "AR"])
            ws.append(["t1", "8\u5c0f\u65f6", "8H", "8시간", "8H", "8 heures", "8 Std.", "8 \u0447.", "8 ore", "8 horas", "8 horas", "8 Saat", "8 giờ", "8 jam", "8 \u0e0a\u0e21.", "8 ساعات"])
            wb.save(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir)

            wb = load_workbook(prepared.translation_workbook, read_only=True, data_only=True)
            ws = wb["Translations"]
            headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
            hits = json.loads(ws.cell(2, headers.index("term_hits_json") + 1).value)
            self.assertEqual(hits, [])
            wb.close()

    def test_apply_rejects_qa_blockers_before_writing_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "sample.docx")
            _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")
            prepared = prepare_announcement_docx_harness(task_dir)
            _fill_translation_workbook(prepared.translation_workbook)

            wb = load_workbook(prepared.translation_workbook)
            ws = wb["Translations"]
            headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
            ws.cell(2, headers.index("EN") + 1).value = "Maintenance \u4e2d\u6587"
            wb.save(prepared.translation_workbook)
            wb.close()

            with self.assertRaisesRegex(ValueError, "hard blockers"):
                apply_announcement_translations(task_dir, prepared.translation_workbook)

    def test_import_ai_responses_fills_workbook_before_apply(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "sample.docx")
            _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")
            prepared = prepare_announcement_docx_harness(task_dir)
            response_dir = task_dir / "ai_responses"
            _write_ai_response_files(prepared.work_dir, response_dir)

            imported = import_announcement_ai_responses(
                task_dir,
                response_dir=response_dir,
            )

            self.assertEqual(imported.row_count, 3)
            self.assertEqual(imported.languages, [header for header, _ in TARGET_LANGUAGES])
            applied = apply_announcement_translations(task_dir, prepared.translation_workbook)
            self.assertEqual(applied.hard_blockers, 0)

    def test_import_ai_responses_rejects_protocol_drift(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "sample.docx")
            _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")
            prepared = prepare_announcement_docx_harness(task_dir)
            response_dir = task_dir / "ai_responses"
            _write_ai_response_files(prepared.work_dir, response_dir)

            rows = [
                json.loads(line)
                for line in (response_dir / "ai_response_en.jsonl").read_text(encoding="utf-8").splitlines()
            ]
            rows[0], rows[1] = rows[1], rows[0]
            (response_dir / "ai_response_en.jsonl").write_text(
                "\n".join(json.dumps(row, ensure_ascii=False) for row in rows) + "\n",
                encoding="utf-8",
            )

            with self.assertRaisesRegex(ValueError, "order mismatch"):
                import_announcement_ai_responses(
                    task_dir,
                    response_dir=response_dir,
                    languages=["EN"],
                )

    def test_import_ai_responses_writes_issue_report_for_translation_qa_failure(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u8bad\u7ec3\u5e08\u8bf7\u7559\u610f")
            doc.save(task_dir / "sample.docx")
            _write_en_fr_terms(task_dir / "sample_announcement_terms_20260526.xlsx")
            prepared = prepare_announcement_docx_harness(task_dir, languages=["en"])
            workpack = json.loads((prepared.work_dir / "workpack_en.jsonl").read_text(encoding="utf-8").strip())
            response_dir = task_dir / "responses"
            response_dir.mkdir()
            (response_dir / "ai_response_en.jsonl").write_text(
                json.dumps({"para_id": workpack["para_id"], "translation": "Please note."}) + "\n",
                encoding="utf-8",
            )

            with self.assertRaisesRegex(ValueError, "qa_report="):
                import_announcement_ai_responses(
                    task_dir,
                    response_dir=response_dir,
                    languages=["en"],
                )

            report_path = prepared.work_dir / "ai_response_qa_en.json"
            report = json.loads(report_path.read_text(encoding="utf-8"))
            self.assertEqual(report["issue_count"], 1)
            self.assertEqual(report["issues"][0]["check_type"], "term_missing")

            (response_dir / "ai_response_en.jsonl").write_text(
                json.dumps({"para_id": workpack["para_id"], "translation": "Trainer, please note."}) + "\n",
                encoding="utf-8",
            )
            import_announcement_ai_responses(
                task_dir,
                response_dir=response_dir,
                languages=["en"],
            )
            self.assertFalse(report_path.exists())

    def test_prepare_does_not_require_chinese_bracket_text_as_protected_token(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u3010\u73a9\u6cd5\u66f4\u65b0\u3011")
            doc.add_paragraph("Token [VIP] 2026/5/20 04:00-06:00")
            doc.save(task_dir / "sample.docx")
            _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir)

            wb = load_workbook(prepared.translation_workbook, read_only=True, data_only=True)
            ws = wb["Translations"]
            headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
            protected_col = headers.index("protected_tokens") + 1
            self.assertEqual(json.loads(ws.cell(2, protected_col).value), [])
            self.assertIn("[VIP]", json.loads(ws.cell(3, protected_col).value))
            wb.close()

    def test_prepare_does_not_protect_numeric_month_when_chinese_date_is_localized(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("7\u670831\u65e5~8\u67085\u65e5\uff08\u5f00\u670d\u7b2c8\u5929\u53ca\u4ee5\u4e0a\uff09")
            doc.save(task_dir / "sample.docx")
            _write_en_fr_terms(task_dir / "sample_announcement_terms_20260526.xlsx")

            prepared = prepare_announcement_docx_harness(task_dir, languages=["en"])
            row = json.loads((prepared.work_dir / "workpack_en.jsonl").read_text(encoding="utf-8").strip())

            self.assertEqual(row["protected_tokens"], ["31", "5", "8"])

    def test_import_accepts_localized_bracket_shape_and_hyphenated_term(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            doc = Document()
            doc.add_paragraph("\u76ae\u80a4\u6280\u80fd\uff08Lv.1\uff09\uff1a\u7d2f\u8ba1\u5145\u503c")
            doc.save(task_dir / "sample.docx")
            wb = Workbook()
            ws = wb.active
            ws.append(["ID", "CN", "EN"])
            ws.append(["term_top_up", "\u5145\u503c", "Top Up"])
            wb.save(task_dir / "sample_announcement_terms_20260526.xlsx")
            prepared = prepare_announcement_docx_harness(task_dir, languages=["en"])
            row = json.loads((prepared.work_dir / "workpack_en.jsonl").read_text(encoding="utf-8").strip())
            response_dir = task_dir / "responses"
            response_dir.mkdir()
            (response_dir / "ai_response_en.jsonl").write_text(
                json.dumps(
                    {
                        "para_id": row["para_id"],
                        "translation": "Skin Skill (Lv.1): cumulative top-up",
                    }
                )
                + "\n",
                encoding="utf-8",
            )

            imported = import_announcement_ai_responses(
                task_dir,
                response_dir=response_dir,
                languages=["en"],
            )

            self.assertEqual(imported.row_count, 1)

    def test_apply_rejects_workbook_protocol_and_source_drift(self):
        cases = ("missing", "duplicate", "extra", "empty", "drift")
        for case in cases:
            with self.subTest(case=case), tempfile.TemporaryDirectory() as tmp:
                task_dir = Path(tmp)
                source_docx = task_dir / "sample.docx"
                _write_docx(source_docx)
                _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")
                prepared = prepare_announcement_docx_harness(task_dir)
                _fill_translation_workbook(prepared.translation_workbook)

                wb = load_workbook(prepared.translation_workbook)
                ws = wb["Translations"]
                headers = [ws.cell(1, col).value for col in range(1, ws.max_column + 1)]
                if case == "missing":
                    ws.delete_rows(2)
                elif case == "duplicate":
                    ws.append([ws.cell(2, col).value for col in range(1, ws.max_column + 1)])
                elif case == "extra":
                    row = [ws.cell(2, col).value for col in range(1, ws.max_column + 1)]
                    row[headers.index("para_id")] = "extra:0000"
                    ws.append(row)
                elif case == "empty":
                    ws.cell(2, headers.index("EN") + 1).value = ""
                wb.save(prepared.translation_workbook)
                wb.close()

                if case == "drift":
                    doc = Document(str(source_docx))
                    doc.add_paragraph("new paragraph")
                    doc.save(source_docx)

                expected = "hard blockers" if case == "empty" else case
                with self.assertRaisesRegex(ValueError, expected):
                    apply_announcement_translations(task_dir, prepared.translation_workbook)

    def test_apply_writes_docx_and_deliver_keeps_only_final_outputs(self):
        with tempfile.TemporaryDirectory() as tmp:
            task_dir = Path(tmp)
            _write_docx(task_dir / "sample.docx")
            _write_terms(task_dir / "sample_announcement_terms_20260526.xlsx")
            prepared = prepare_announcement_docx_harness(task_dir)
            _fill_translation_workbook(prepared.translation_workbook)

            applied = apply_announcement_translations(task_dir, prepared.translation_workbook)
            self.assertEqual(applied.hard_blockers, 0)
            self.assertEqual(len(applied.output_docx_paths), len(TARGET_LANGUAGES))
            en_doc = Document(str(applied.output_docx_paths[0]))
            self.assertEqual(len([p for p in en_doc.paragraphs if p.text.strip()]), 3)
            self.assertEqual(en_doc.paragraphs[0].style.name, "Title")

            delivered = deliver_announcement_outputs(task_dir, date_stamp="20260526")
            names = sorted(path.name for path in delivered.delivery_dir.iterdir())
            self.assertIn("QA\u6458\u8981.xlsx", names)
            self.assertEqual(sum(1 for name in names if name.endswith(".docx")), len(TARGET_LANGUAGES))
            self.assertFalse(any(name.endswith(".json") or name.endswith(".jsonl") for name in names))
            self.assertFalse((delivered.delivery_dir / "_work").exists())


if __name__ == "__main__":
    unittest.main()
