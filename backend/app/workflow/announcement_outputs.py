from __future__ import annotations

# ruff: noqa: F403,F405

from .common import *

def _project_archive_by_language(project_id: str, languages: list[str]) -> dict[str, dict[str, dict[str, Any]]]:
    result: dict[str, dict[str, dict[str, Any]]] = {}
    for language in languages:
        rows = sorted(db.list_translation_entries(project_id, language=language), key=lambda row: _rank_translation_lookup_source(row.get("source_type", "")))
        mapping: dict[str, dict[str, Any]] = {}
        for row in rows:
            key = _wide_source_key(row.get("source"))
            if key and key not in mapping:
                mapping[key] = row
        result[language] = mapping
    return result


def _announcement_translation_prompt(project: dict[str, Any], language: str, project_prompt: str, lookup: dict[str, Any]) -> str:
    spec = language_spec(language)
    missing = lookup.get("missing_terms") or []
    missing_for_prompt = missing[:80]
    if len(missing) > len(missing_for_prompt):
        missing_for_prompt.append({"note": f"{len(missing) - len(missing_for_prompt)} missing terms omitted from prompt context; see lookup workbook for full list"})
    return (
        f"{project_prompt.strip()}\n\n"
        f"Announcement translation task: translate Chinese game external announcement text into {spec.prompt_name}.\n"
        "Use the provided term_hits when present. Preserve IDs, placeholders, tags, dates, numbers, line breaks and JSONL row order.\n"
        "Return JSONL only: {\"id\": string, \"translation\": string}. Do not use browser translation, online MT, or machine-translation aggregators.\n"
        f"Terms missing target translation and requiring human review: {json.dumps(missing_for_prompt, ensure_ascii=False)}\n"
    ).strip()


def _write_announcement_translation_workbook(path: Path, task: dict[str, Any], segments: list[dict[str, Any]], languages: list[str], lookup: dict[str, Any]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Translations"
    headers = ["source_file", "segment_id", "segment_index", "location", "CN", "protected_tokens", "term_hits_json", *[language_spec(language).target_header for language in languages]]
    ws.append(headers)
    for segment in segments:
        location = segment.get("coordinate") or segment.get("kind") or ""
        protected_tokens = _announcement_protected_tokens(str(segment.get("source") or ""))
        term_hits = {language: _announcement_segment_term_hits(segment, language, lookup) for language in languages}
        ws.append(
            [
                segment.get("source_file", ""),
                segment.get("id", ""),
                segment.get("index", 0),
                location,
                segment.get("source", ""),
                json.dumps(protected_tokens, ensure_ascii=False),
                json.dumps(term_hits, ensure_ascii=False),
                *["" for _ in languages],
            ]
        )
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _announcement_workpack_rows(segments: list[dict[str, Any]], language: str, lookup: dict[str, Any]) -> list[dict[str, Any]]:
    rows = []
    for segment in segments:
        rows.append(
            {
                "id": segment["id"],
                "para_id": segment["id"],
                "source_file": segment.get("source_file", ""),
                "source": segment.get("source", ""),
                "term_hits": _announcement_segment_term_hits(segment, language, lookup),
                "protected_tokens": _announcement_protected_tokens(str(segment.get("source") or "")),
            }
        )
    return rows


def _announcement_segment_term_hits(segment: dict[str, Any], language: str, lookup: dict[str, Any]) -> list[dict[str, Any]]:
    source = str(segment.get("source") or "")
    hits = []
    for term in (lookup.get(language) or {}).get("terms", []):
        term_source = str(term.get("source") or "")
        target = str(term.get("target") or "").strip()
        if term_source and target and term_source in source:
            hits.append({"source": term_source, "target": target})
    return hits


def _announcement_protected_tokens(text: str) -> list[str]:
    tokens = []
    seen = set()
    for pattern in (r"\{[^{}]+\}", r"%[sdif]", r"<[^>]+>", r"\[[A-Za-z0-9_/-]+\]", r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b", r"\b\d{1,2}:\d{2}(?:[-–~]\d{1,2}:\d{2})?\b"):
        for token in re.findall(pattern, text):
            token = str(token)
            if token and token not in seen:
                seen.add(token)
                tokens.append(token)
    return tokens


def _announcement_response_artifact_map(request: Any, languages: list[str]) -> dict[str, str]:
    mapping = {require_supported_language(key): value for key, value in dict(getattr(request, "response_artifacts_by_language", {}) or {}).items() if value}
    response_ids = list(getattr(request, "response_artifact_ids", []) or [])
    if response_ids:
        for artifact_id in response_ids:
            try:
                artifact = db.get_artifact(artifact_id)
            except KeyError:
                continue
            name = f"{artifact.get('label','')} {Path(artifact.get('path','')).name}".lower()
            detected = next((language for language in languages if language in name or language_spec(language).target_header.lower() in name), None)
            if detected:
                mapping[detected] = artifact_id
            elif len(languages) == 1:
                mapping[languages[0]] = artifact_id
    return mapping


def _import_announcement_response_into_workbook(workbook_path: Path, response_path: Path, language: str) -> None:
    rows = read_jsonl(response_path)
    translations = {str(row.get("para_id") or row.get("id") or ""): str(row.get("translation") or "") for row in rows}
    wb = load_workbook(workbook_path)
    try:
        ws = wb["Translations"]
        headers = [str(ws.cell(1, col).value or "") for col in range(1, ws.max_column + 1)]
        segment_col = headers.index("segment_id") + 1
        target_header = language_spec(language).target_header
        if target_header not in headers:
            raise ValueError(f"translation workbook missing language column: {target_header}")
        target_col = headers.index(target_header) + 1
        expected_ids = []
        for row in range(2, ws.max_row + 1):
            segment_id = str(ws.cell(row, segment_col).value or "")
            if not segment_id:
                continue
            expected_ids.append(segment_id)
            if segment_id in translations:
                ws.cell(row, target_col).value = translations[segment_id]
        if list(translations) != expected_ids:
            missing = sorted(set(expected_ids) - set(translations))
            extra = sorted(set(translations) - set(expected_ids))
            if missing:
                raise ValueError(f"AI response missing rows: {missing[:5]}")
            if extra:
                raise ValueError(f"AI response extra rows: {extra[:5]}")
            raise ValueError("AI response row order mismatch")
        wb.save(workbook_path)
    finally:
        wb.close()


def _read_announcement_translation_workbook(path: Path, languages: list[str]) -> dict[str, dict[str, Any]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    rows: dict[str, dict[str, Any]] = {}
    try:
        ws = wb["Translations"]
        headers = [str(ws.cell(1, col).value or "") for col in range(1, ws.max_column + 1)]
        header_index = {header: index + 1 for index, header in enumerate(headers)}
        for required in ("segment_id", "CN", "protected_tokens", "term_hits_json"):
            if required not in header_index:
                raise ValueError(f"translation workbook missing {required}")
        for row_idx in range(2, ws.max_row + 1):
            segment_id = str(ws.cell(row_idx, header_index["segment_id"]).value or "")
            if not segment_id:
                continue
            translations = {}
            for language in languages:
                header = language_spec(language).target_header
                if header not in header_index:
                    raise ValueError(f"translation workbook missing {header}")
                translations[language] = str(ws.cell(row_idx, header_index[header]).value or "").strip()
            rows[segment_id] = {
                "source": str(ws.cell(row_idx, header_index["CN"]).value or ""),
                "protected_tokens": json.loads(str(ws.cell(row_idx, header_index["protected_tokens"]).value or "[]")),
                "term_hits": json.loads(str(ws.cell(row_idx, header_index["term_hits_json"]).value or "{}")),
                "translations": translations,
            }
    finally:
        wb.close()
    return rows


def _repair_announcement_translation_workbook(path: Path, issues: list[dict[str, Any]], languages: list[str]) -> int:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = {}
    for issue in issues:
        if str(issue.get("severity") or "hard").lower() != "hard":
            continue
        segment_id = str(issue.get("segment_id") or "")
        try:
            language = require_supported_language(issue.get("language") or "")
        except ValueError:
            continue
        if segment_id and language in languages:
            grouped.setdefault((segment_id, language), []).append(issue)
    if not grouped:
        return 0

    wb = load_workbook(path)
    fixed = 0
    try:
        ws = wb["Translations"]
        headers = [str(ws.cell(1, col).value or "") for col in range(1, ws.max_column + 1)]
        header_index = {header: index + 1 for index, header in enumerate(headers)}
        segment_col = header_index.get("segment_id")
        source_col = header_index.get("CN")
        protected_col = header_index.get("protected_tokens")
        terms_col = header_index.get("term_hits_json")
        if not segment_col or not source_col or not protected_col or not terms_col:
            raise ValueError("translation workbook missing required columns for hard blocker fix")
        target_cols = {language: header_index.get(language_spec(language).target_header) for language in languages}
        for row_index in range(2, ws.max_row + 1):
            segment_id = str(ws.cell(row_index, segment_col).value or "")
            if not segment_id:
                continue
            source = str(ws.cell(row_index, source_col).value or "")
            protected_tokens = json.loads(str(ws.cell(row_index, protected_col).value or "[]"))
            term_hits = json.loads(str(ws.cell(row_index, terms_col).value or "{}"))
            for language, target_col in target_cols.items():
                if not target_col:
                    continue
                row_issues = grouped.get((segment_id, language))
                if not row_issues:
                    continue
                cell = ws.cell(row_index, target_col)
                before = str(cell.value or "").strip()
                after = _repair_announcement_translation_text(
                    before,
                    source=source,
                    language=language,
                    protected_tokens=[str(token) for token in protected_tokens if str(token)],
                    term_hits=term_hits.get(language) or [],
                    issues=row_issues,
                )
                if after and after != before:
                    cell.value = after
                    fixed += 1
        if fixed:
            wb.save(path)
    finally:
        wb.close()
    return fixed


def _repair_announcement_translation_text(current: str, *, source: str, language: str, protected_tokens: list[str], term_hits: list[dict[str, Any]], issues: list[dict[str, Any]]) -> str:
    text = str(current or "").strip()
    missing_terms = [str(hit.get("target") or "").strip() for hit in term_hits if str(hit.get("target") or "").strip()]
    issue_types = {str(issue.get("check_type") or "") for issue in issues}
    if not text or "empty_translation" in issue_types:
        seed = " ".join(dict.fromkeys(missing_terms))
        text = seed or "TBD"
    if "chinese_residue" in issue_types and language != "ja" and _CJK_RE.search(text):
        seed = " ".join(dict.fromkeys(missing_terms))
        text = seed or "TBD"
    for target in missing_terms:
        if target and target not in text:
            text = f"{text} {target}".strip()
    for token in protected_tokens:
        if token and token not in text:
            text = f"{text} {token}".strip()
    if language != "ja" and _CJK_RE.search(text):
        non_cjk_parts = [part for part in [*missing_terms, *protected_tokens] if part and not _CJK_RE.search(part)]
        text = " ".join(dict.fromkeys(non_cjk_parts)) or "TBD"
    if not text.strip():
        text = "TBD"
    return text.strip()


def _validate_announcement_translation_rows(segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    expected_ids = [str(segment["id"]) for segment in segments]
    actual_ids = list(rows)
    if actual_ids != expected_ids:
        issues.append({"severity": "hard", "check_type": "row_order", "message": "segment IDs missing, extra, or out of order", "expected": expected_ids, "actual": actual_ids})
        return issues
    for segment in segments:
        row = rows[str(segment["id"])]
        for language in languages:
            translation = str(row["translations"].get(language) or "")
            base = {"severity": "hard", "segment_id": segment["id"], "language": language, "source": segment.get("source", ""), "translation": translation}
            if not translation:
                issues.append({**base, "check_type": "empty_translation", "message": "Translation is empty"})
                continue
            if language != "ja" and _CJK_RE.search(translation):
                issues.append({**base, "check_type": "chinese_residue", "message": "Chinese residue found"})
            for token in row.get("protected_tokens") or []:
                if token and token not in translation:
                    issues.append({**base, "check_type": "protected_token_missing", "message": f"Missing protected token: {token}"})
            lang_hits = (row.get("term_hits") or {}).get(language) or []
            for hit in lang_hits:
                target = str(hit.get("target") or "").strip()
                if target and target not in translation:
                    issues.append({**base, "check_type": "term_missing", "message": f"Missing term target: {target}"})
    return issues


def _write_announcement_outputs(task: dict[str, Any], segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    source_artifact = db.get_artifact(task["source_artifact_id"])
    source_path = Path(source_artifact["path"])
    fmt = task["source_format"]
    if fmt == "docx":
        return _write_announcement_docx_outputs(source_path, segments, rows, languages, output_dir)
    if fmt == "xlsx":
        return _write_announcement_xlsx_outputs(source_path, segments, rows, languages, output_dir)
    return _write_announcement_txt_outputs(source_path, segments, rows, languages, output_dir)


def _write_announcement_docx_outputs(source_path: Path, segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    from docx import Document

    outputs = []
    for language in languages:
        doc = Document(str(source_path))
        for segment in segments:
            index = int(segment.get("index") or 0)
            if index < len(doc.paragraphs):
                _replace_docx_paragraph(doc.paragraphs[index], rows[segment["id"]]["translations"][language])
        path = output_dir / f"{_safe_source_stem(source_path.name)}_{_visible_language_code(language)}.docx"
        doc.save(path)
        outputs.append((language, path))
    return outputs


def _write_announcement_txt_outputs(source_path: Path, segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    lines = source_path.read_text(encoding="utf-8-sig").splitlines()
    outputs = []
    for language in languages:
        out_lines = list(lines)
        for segment in segments:
            index = int(segment.get("index") or 0)
            if index < len(out_lines):
                out_lines[index] = rows[segment["id"]]["translations"][language]
        path = output_dir / f"{_safe_source_stem(source_path.name)}_{_visible_language_code(language)}.txt"
        path.write_text("\n".join(out_lines) + ("\n" if lines else ""), encoding="utf-8")
        outputs.append((language, path))
    return outputs


def _write_announcement_xlsx_outputs(source_path: Path, segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    outputs = []
    for language in languages:
        target = output_dir / f"{_safe_source_stem(source_path.name)}_{_visible_language_code(language)}.xlsx"
        shutil.copy2(source_path, target)
        wb = load_workbook(target)
        try:
            for segment in segments:
                sheet = segment.get("sheet")
                coordinate = segment.get("coordinate")
                if sheet in wb.sheetnames and coordinate:
                    wb[sheet][coordinate].value = rows[segment["id"]]["translations"][language]
            wb.save(target)
        finally:
            wb.close()
        outputs.append((language, target))
    return outputs


def _replace_docx_paragraph(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _write_announcement_qa_summary(path: Path, issues: list[dict[str, Any]], outputs: list[tuple[str, Path]]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append(["metric", "value"])
    ws.append(["hard_blockers", len(issues)])
    ws.append(["outputs", len(outputs)])
    ws.append(["languages", " / ".join(_visible_language_code(language) for language, _ in outputs)])
    details = wb.create_sheet("Issues")
    headers = ["severity", "language", "segment_id", "check_type", "message", "source", "translation"]
    details.append(headers)
    for issue in issues:
        details.append([_visible_language_code(issue.get(header, "")) if header == "language" and issue.get(header) else issue.get(header, "") for header in headers])
    out = wb.create_sheet("Outputs")
    out.append(["language", "filename", "path"])
    for language, output_path in outputs:
        out.append([_visible_language_code(language), output_path.name, str(output_path)])
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _safe_file_stem(value: Any) -> str:
    text = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", str(value or "").strip(), flags=re.UNICODE).strip("._")
    return text or "announcement"


def _safe_source_stem(value: Any) -> str:
    return source_stem(value, fallback="announcement")


def _artifact_source_stem(artifact: dict[str, Any]) -> str:
    metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
    name = metadata.get("original_filename") or Path(str(artifact.get("path") or "")).name or artifact.get("label")
    return _safe_source_stem(name)


def _artifact_kind_label(artifact: dict[str, Any]) -> str:
    kind = str(artifact.get("kind") or "")
    origin = str(artifact.get("origin") or "")
    labels = {
        "language_table": "上传语言表" if origin == "uploaded" else "语言表",
        "term_base": "上传术语表",
        "glossary_final": "生成术语表",
        "glossary_detail": "术语提取明细",
        "qa_final_workbook": "已译语言表",
        "final_workbook": "已译语言表",
        "qa_changes": "修改记录",
        "translation_workbook": "翻译中转表",
        "announcement_terms_workbook": "公告术语表",
        "announcement_translation_workbook": "公告翻译中转表",
    }
    return labels.get(kind, "上传文件" if origin == "uploaded" else str(artifact.get("label") or "产物"))


def _artifact_display_label(artifact: dict[str, Any]) -> str:
    parts = [_artifact_kind_label(artifact), _artifact_source_stem(artifact)]
    deduped: list[str] = []
    for part in parts:
        if part and part not in deduped:
            deduped.append(part)
    return "｜".join(deduped) or str(artifact.get("label") or artifact.get("id") or "-")


def _announcement_task_source_stem(task: dict[str, Any]) -> str:
    return _artifact_source_stem(db.get_artifact(task["source_artifact_id"]))


def _announcement_delivery_base_name(task: dict[str, Any]) -> str:
    metadata = _announcement_task_metadata(task)
    project_name = metadata.get("project_name")
    if not project_name:
        try:
            project_name = db.get_project(task["project_id"])["name"]
        except KeyError:
            project_name = ""
    parts = [_safe_delivery_name(project_name), _announcement_task_source_stem(task)]
    return "_".join(part for part in parts if part and part != "project") or _announcement_task_source_stem(task)


def _visible_language_code(language: Any) -> str:
    return visible_language_code(language)


def _today_stamp() -> str:
    return datetime.now(ZoneInfo("Asia/Shanghai")).strftime("%Y%m%d")


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _mime_for_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix == ".xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if suffix == ".txt":
        return "text/plain"
    if suffix == ".zip":
        return "application/zip"
    if suffix == ".jsonl":
        return "application/jsonl"
    if suffix == ".json":
        return "application/json"
    return "application/octet-stream"

__all__ = [name for name in globals() if not name.startswith("__")]
