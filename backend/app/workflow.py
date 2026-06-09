from __future__ import annotations

import asyncio
import csv
import hashlib
import importlib.util
import json
import os
import re
import shutil
import subprocess
import sys
import math
import time
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import ZipFile
import xml.etree.ElementTree as ET
from zoneinfo import ZoneInfo

from openpyxl import Workbook, load_workbook

from . import db
from .config import DATA_ROOT, GLOSSARY_ROOT, LOCALIZATION_ROOT, REAL_PROVIDERS, TEST_FAKE_PROVIDER, load_settings, normalize_provider_name, test_provider_enabled
from .delivery_naming import safe_delivery_name, source_stem
from .languages import ANNOUNCEMENT_LANGUAGE_ORDER, PROJECT_LANGUAGE_ORDER, alt_aliases, language_spec, normalize_language, require_supported_language, target_aliases, visible_language_code
from .providers import call_image_text, call_text, translate_batch
from .translation_batches import (
    AsyncTokenRateLimiter as _AsyncTokenRateLimiter,
    build_batch_manifest as _build_batch_manifest,
    cap_context_text as _cap_context_text,
    estimate_row_tokens as _estimate_row_tokens,
    estimate_text_tokens as _estimate_text_tokens,
    load_or_create_batch_manifest as _load_or_create_batch_manifest,
    manage_project_prompt_context as _manage_project_prompt_context,
    manifest_matches_rows as _manifest_matches_rows,
    project_context_summary as _project_context_summary,
    provider_retry_delay_seconds as _provider_retry_delay_seconds,
)

__all__ = [
    "_AsyncTokenRateLimiter",
    "_build_batch_manifest",
    "_cap_context_text",
    "_estimate_row_tokens",
    "_estimate_text_tokens",
    "_load_or_create_batch_manifest",
    "_manage_project_prompt_context",
    "_manifest_matches_rows",
    "_project_context_summary",
    "_provider_retry_delay_seconds",
]


HARNESS_SCHEMA_VERSION = 1
_GLOSSARY_EXTRACTOR_MODULE: Any | None = None
GLOBAL_HARNESS_CONTRACT: dict[str, Any] = {
    "source": "global_harness",
    "workpack": "translation_workpack.jsonl",
    "response_protocol": "jsonl:{id:int|str,translation:str}",
    "hard_gates": ["id", "placeholder", "tag", "newline", "input_fingerprint"],
    "qa_sources": ["workflow/localization/utils/quality_harness.py", "workflow/localization/fixtures/quality_regression.json"],
}

_CJK_RE = re.compile(r"[\u3400-\u9fff]")
RowId = int | str
LANGUAGE_ORDER = PROJECT_LANGUAGE_ORDER
AUTO_LANGUAGE_TARGET_ALIASES = {code: tuple(target_aliases(code)) for code in LANGUAGE_ORDER}
AUTO_LANGUAGE_ALT_ALIASES = {code: tuple(alt_aliases(code)) for code in LANGUAGE_ORDER}
_TARGET_DETECTION_ALIASES: dict[str, set[str]] = {
    "en": {"en", "english"},
    "ko": {"ko", "kr", "korean"},
    "ja": {"ja", "jp", "japanese"},
    "fr": {"fr", "fre", "french"},
    "de": {"de", "ger", "german"},
    "ru": {"ru", "rus", "russian"},
    "it": {"it", "ita", "italian"},
    "es": {"es", "spa", "spanish"},
    "pt": {"pt", "pt-br", "ptbr", "por", "portuguese"},
    "tr": {"tr", "tk", "tur", "turkish"},
    "idn": {"idn", "ind", "indonesian", "bahasa", "bahasa indonesia"},
    "th": {"th", "tha", "thai"},
    "ar": {"ar", "ara", "arabic"},
}
_STRUCTURAL_TARGET_HEADERS = {
    "id", "key", "编号", "序号",
    "cn", "zh", "source", "original", "chinese", "term", "原文", "中文", "术语",
    "category", "type", "分类", "类别", "类型",
    "note", "notes", "comment", "备注",
    "target", "translation", "译文",
}


def _looks_like_untranslated_seed(text: str, language: str) -> bool:
    value = str(text or "")
    if not _CJK_RE.search(value):
        return False
    if language == "ja":
        return False
    return True


def project_dir(project_id: str) -> Path:
    path = DATA_ROOT / "projects" / project_id
    path.mkdir(parents=True, exist_ok=True)
    for child in ("uploads", "profile", "glossary", "runs", "assets", "translations"):
        (path / child).mkdir(exist_ok=True)
    return path


def run_dir(run_id: str) -> Path:
    path = DATA_ROOT / "runs" / run_id
    path.mkdir(parents=True, exist_ok=True)
    return path


def default_project_harness(project: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema_version": HARNESS_SCHEMA_VERSION,
        "source": "project_harness",
        "project_id": project["id"],
        "project_name": project["name"],
        "project_metadata": {},
        "style_guidance": "",
        "target_audience": "",
        "tone": "",
        "forbidden_translations": [],
        "fixed_terms": [],
        "hard_rules": [],
        "soft_rules": [],
        "reference_examples": [],
        "manual_fixes": [],
        "qa_summary": {},
        "updated_at": "",
    }


def project_harness_path(project_id: str) -> Path:
    return project_dir(project_id) / "profile" / "project_harness.json"


def read_project_harness(project_id: str) -> dict[str, Any]:
    project = db.get_project(project_id)
    default = default_project_harness(project)
    path = project_harness_path(project_id)
    if not path.exists():
        return default
    payload = json.loads(path.read_text(encoding="utf-8"))
    merged = {**default, **payload}
    merged["project_id"] = project_id
    merged["project_name"] = project["name"]
    merged["schema_version"] = HARNESS_SCHEMA_VERSION
    return _sanitize_harness(merged)


def write_project_harness(project_id: str, updates: dict[str, Any]) -> dict[str, Any]:
    payload = read_project_harness(project_id)
    for key, value in updates.items():
        if value is not None:
            payload[key] = value
    payload["updated_at"] = db.now_iso()
    payload = _sanitize_harness(payload)
    path = project_harness_path(project_id)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


def harness_overview(project_id: str) -> dict[str, Any]:
    return {
        "global_harness": GLOBAL_HARNESS_CONTRACT,
        "project_harness": read_project_harness(project_id),
        "boundary": (
            "global_harness stores reusable workflow contracts and gates; "
            "project_harness stores this project's private requirements only."
        ),
    }


def _workbook_text_stats(path: Path) -> dict[str, int]:
    try:
        wb = load_workbook(path, read_only=True, data_only=False)
        ws = wb.active
        headers = [str(cell.value or "").strip() for cell in next(ws.iter_rows(min_row=1, max_row=1))]
        lowered = [header.lower() for header in headers]
        source_idx = next(
            (idx for idx, header in enumerate(lowered) if header in {"cn", "source", "原文", "中文"}),
            None,
        )
        target_idx = next(
            (idx for idx, header in enumerate(lowered) if header in {"en", "target", "translation", "译文", "英文"}),
            None,
        )
        source_rows = 0
        translated_rows = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            if source_idx is not None and row[source_idx] not in (None, ""):
                source_rows += 1
            if target_idx is not None and row[target_idx] not in (None, ""):
                translated_rows += 1
        wb.close()
        return {"source_rows": source_rows, "translated_rows": translated_rows}
    except Exception:
        return {"source_rows": 0, "translated_rows": 0}


def _language_assets_summary(project_id: str) -> str:
    candidates = [
        *db.list_artifacts(project_id=project_id, role="translation_workbook"),
        *db.list_artifacts(project_id=project_id, role="language_source"),
    ]
    best = {"source_rows": 0, "translated_rows": 0}
    for artifact in candidates:
        path = Path(str(artifact.get("path") or ""))
        if path.suffix.lower() != ".xlsx" or not path.exists():
            continue
        stats = _workbook_text_stats(path)
        if stats["source_rows"] > best["source_rows"]:
            best = stats
        elif stats["translated_rows"] > best["translated_rows"]:
            best["translated_rows"] = stats["translated_rows"]
    if best["source_rows"]:
        return f"{best['source_rows']} 条文本，已有英文 {best['translated_rows']} 条。"
    return "暂未统计语言表行数。"


def _project_material_labels(project_id: str) -> list[str]:
    labels: list[str] = []
    for role in ("glossary_source", "language_source", "translation_workbook"):
        for artifact in db.list_artifacts(project_id=project_id, role=role):
            label = str(artifact.get("label") or "").strip()
            if label and label not in labels:
                labels.append(label)
    return labels


def _is_warplane_project(project: dict[str, Any], intro: str, material_labels: list[str], asset_notes: list[str] | None = None) -> bool:
    text = " ".join(
        [
            str(project.get("name") or ""),
            str(project.get("type") or ""),
            str(project.get("description") or ""),
            intro,
            *material_labels,
            *(asset_notes or []),
        ]
    )
    return any(token in text for token in ("战机", "飞行射击", "导弹", "装备", "Warplane", "warplane"))


def _build_project_profile(project: dict[str, Any], intro: str, asset_notes: list[str], target_language: str = "en", material_packet: dict[str, Any] | None = None) -> dict[str, Any]:
    target_language = require_supported_language(target_language)
    spec = language_spec(target_language)
    material_labels = _project_material_labels(project["id"])
    description = project.get("description", "") or intro or "；".join(asset_notes[:3])
    is_warplane = _is_warplane_project(project, intro, material_labels, asset_notes)
    style_language = "英文" if target_language == "en" else spec.label
    if is_warplane:
        game_type = "科幻战机 / 飞行射击 / RPG养成"
        target_audience = "偏中重度、喜欢战机养成、战斗数值、装备强化和活动推进的移动端玩家。"
        content_scope = "战机、导弹、射击、弹幕等战斗内容；英雄、装备、技能、属性和战力成长；建造、升级、采集、生产等基地系统；角色剧情对话；活动、礼包和奖励。"
        translation_style = f"翻译为自然{style_language}；UI/玩法精简适配移动端；剧情自然、地道、通顺，保留角色语气；整体语气冷静、利落、偏科幻军事；战机、装备、导弹、技能和战斗数值要专业清晰，避免可爱化、生活化或过度口语化。"
        tone = "冷静、利落、偏科幻军事"
    else:
        game_type = project.get("type", "") or "游戏本地化项目"
        target_audience = "目标语区游戏玩家；以当前项目描述和参考素材为准。"
        content_scope = description or "UI、系统、任务、道具、活动和剧情文本。"
        translation_style = f"准确翻译为自然{style_language}；UI/按钮/任务短句清晰；剧情对话自然但不改设定；术语以项目术语表为准；保留变量、占位符、富文本标签、数字和换行。"
        tone = "自然、准确、游戏 UI 友好"
    seed = _display_profile_seed_from_packet(material_packet or {"materials": []})
    return {
        "project_name": project["name"],
        "project_type": project.get("type", ""),
        "description": description,
        "intro": intro,
        "asset_notes": asset_notes,
        "source_materials": material_labels,
        "game_type": game_type,
        "target_audience": target_audience,
        "content_scope": content_scope,
        "translation_style": translation_style,
        "language_assets": _language_assets_summary(project["id"]),
        "target_language": target_language,
        "target_language_label": spec.label,
        "target_language_name": spec.prompt_name,
        "tone": tone,
        "generated_date": db.now_iso()[:10],
        "analysis_source": "template",
        "analysis_warning": "未配置 API key，只生成模板草稿，未进行 AI 资料分析。",
        "display_game_type": seed.get("display_game_type") or game_type,
        "display_target_audience": seed.get("display_target_audience") or target_audience,
        "display_content_scope": seed.get("display_content_scope") or content_scope,
        "display_worldview": seed.get("display_worldview") or tone,
        "display_translation_style": seed.get("display_translation_style") or translation_style,
        "display_focus": seed.get("display_focus") or "",
        "display_key_terms": seed.get("display_key_terms") or "",
    }


def _display_profile_seed_from_packet(material_packet: dict[str, Any]) -> dict[str, str]:
    source_text = " ".join(str(material.get("excerpt") or "") for material in material_packet.get("materials", []) if isinstance(material, dict))

    def table_value(*labels: str) -> str:
        for label in labels:
            pattern = rf"\|\s*{re.escape(label)}(?:（[^|]*）)?\s*\|\s*([^|]+?)\s*\|"
            match = re.search(pattern, source_text)
            if match:
                return re.sub(r"\s+", " ", match.group(1)).strip()
        return ""

    term_match = re.search(r"术语保持一致[：:]\s*([^；;。]+)", source_text)
    return {
        "display_game_type": table_value("游戏类型"),
        "display_target_audience": table_value("目标用户"),
        "display_content_scope": table_value("内容构成", "内容范围"),
        "display_worldview": table_value("视觉与世界观", "世界观"),
        "display_translation_style": table_value("翻译风格", "风格要求"),
        "display_focus": table_value("重点注意", "注意事项"),
        "display_key_terms": re.sub(r"\s+", " ", term_match.group(1)).strip() if term_match else "",
    }


def _project_analysis_provider_prompt(project: dict[str, Any], intro: str, asset_notes: list[str], profile: dict[str, Any], material_packet: dict[str, Any]) -> str:
    return (
        "你正在为游戏本地化项目做真实资料分析，不是套模板。\n"
        "Return strict JSON only. No markdown fences or prose.\n"
        "必须只基于 evidence packet、项目 intro 和参考资料；资料没有支持的信息写入 missing_info，不要编造。\n"
        "display_* 字段必须用中文，给前端和译员看；执行字段可用目标语言或英文，给模型翻译/QA 使用。\n"
        "Required JSON fields: game_type, target_audience, content_scope, translation_style, tone, "
        "display_game_type, display_target_audience, display_content_scope, display_worldview, "
        "display_translation_style, display_focus, display_key_terms, confidence, missing_info.\n\n"
        f"Project name: {project.get('name', '')}\n"
        f"Project type: {project.get('type', '')}\n"
        f"Project description: {project.get('description', '')}\n"
        f"Target language: {profile.get('target_language_name', profile.get('target_language', 'en'))}\n"
        f"User intro:\n{intro}\n\n"
        f"Reference material notes:\n{json.dumps(asset_notes, ensure_ascii=False, indent=2)}\n\n"
        f"Evidence packet:\n{json.dumps({k: v for k, v in material_packet.items() if k != 'context'}, ensure_ascii=False, indent=2)}\n\n"
        f"Evidence context:\n{material_packet.get('context', '')}\n\n"
        f"Template baseline:\n{json.dumps({key: profile.get(key, '') for key in ('game_type', 'target_audience', 'content_scope', 'translation_style', 'tone')}, ensure_ascii=False, indent=2)}"
    )


def _apply_project_analysis_provider(project: dict[str, Any], intro: str, asset_notes: list[str], profile: dict[str, Any], material_packet: dict[str, Any]) -> dict[str, Any]:
    settings = load_settings()
    provider = normalize_provider_name(settings.get("provider"))
    if provider not in REAL_PROVIDERS or not settings.get("api_key"):
        return profile
    prompt = _cap_context_text(_project_analysis_provider_prompt(project, intro, asset_notes, profile, material_packet), 12000, "project analysis")
    payload = _parse_semantic_qa_payload(_call_semantic_provider(settings, prompt))
    if not isinstance(payload, dict):
        raise ValueError("project analysis provider must return a JSON object")
    updates: dict[str, Any] = {}
    for key in (
        "game_type", "target_audience", "content_scope", "translation_style", "tone",
        "display_game_type", "display_target_audience", "display_content_scope", "display_worldview",
        "display_translation_style", "display_focus", "display_key_terms", "confidence",
    ):
        value = str(payload.get(key) or "").strip()
        if value:
            updates[key] = value
    missing_info = payload.get("missing_info")
    if isinstance(missing_info, list):
        updates["missing_info"] = [str(item).strip() for item in missing_info if str(item).strip()]
    elif str(missing_info or "").strip():
        updates["missing_info"] = [str(missing_info).strip()]
    return {
        **profile,
        **updates,
        "analysis_source": "provider",
        "analysis_provider": provider,
        "analysis_model": str(settings.get("model") or ""),
        "analysis_warning": "",
    }


def _project_prompt_from_profile(profile: dict[str, Any]) -> str:
    spec = language_spec(profile.get("target_language") or "en")
    term_rule = (
        f"关键术语以随附术语表为准，{spec.target_header} 为标准译法，{spec.alt_header} 为项目中稳定出现的手动适配译法；"
        if spec.alt_header
        else f"关键术语以随附术语表为准，{spec.target_header} 为标准译法；"
    )
    return (
        f"你是一位资深游戏本地化译者，正在将《{profile['project_name']}》这款{profile['game_type']}游戏翻译为{spec.label}（{spec.prompt_name}）。\n"
        "译文需符合以下要求：\n"
        "1. 游戏内容/UI/玩法说明尽量精简，适配移动游戏按钮、弹窗、任务、道具和奖励说明；\n"
        "2. 剧情对话必须自然、地道、通顺，保留角色语气、冲突、幽默和情绪，不要逐字直译；\n"
        f"3. 项目内容范围：{profile['content_scope']}\n"
        f"4. 译文风格：{profile['translation_style']}\n"
        f"5. {term_rule}\n"
        f"6. 已有{spec.label}译文代表项目历史用法；如现有译法不自然，可以优化，但不要破坏已固定的系统术语；\n"
        "7. 保留所有游戏代码、变量、数字、换行、颜色标签、HTML/富文本标签和占位符，如 {0}、%s、<color> 等；\n"
        "8. 无法确认的专有名词或信息缺口用 [TBD] 标记，不要自行编造设定。\n"
        "输出协议：只返回 JSONL，每行包含 id 和 translation。"
    )



def _project_prompt_display_zh(profile: dict[str, Any]) -> str:
    spec = language_spec(profile.get("target_language") or "en")
    term_rule = (
        f"\u9879\u76ee\u672f\u8bed\u4ee5\u672f\u8bed\u8868\u4e3a\u51c6\uff1a{spec.target_header} \u662f\u6807\u51c6\u8bd1\u6cd5\uff0c{spec.alt_header} \u662f\u7a33\u5b9a\u51fa\u73b0\u7684\u624b\u52a8\u9002\u914d\u8bd1\u6cd5\u3002"
        if spec.alt_header
        else f"\u9879\u76ee\u672f\u8bed\u4ee5\u672f\u8bed\u8868\u4e3a\u51c6\uff1a{spec.target_header} \u662f\u6807\u51c6\u8bd1\u6cd5\u3002"
    )
    project_name = str(profile.get("project_name") or "\u5f53\u524d\u9879\u76ee").strip()
    source_text = " ".join(str(item or "") for item in profile.get("asset_notes") or [])

    def table_value(*labels: str) -> str:
        for label in labels:
            pattern = rf"\|\s*{re.escape(label)}(?:（[^|]*）)?\s*\|\s*([^|]+?)\s*\|"
            match = re.search(pattern, source_text)
            if match:
                return match.group(1).strip()
        return ""

    def first_text(*values: Any) -> str:
        for value in values:
            text = re.sub(r"\s+", " ", str(value or "")).strip()
            if text:
                return text
        return ""

    def sentence(label: str, value: str) -> str:
        value = value.strip().rstrip("。；;")
        return f"{label}：{value}。" if value else ""

    term_match = re.search(r"术语保持一致[：:]\s*([^；;。]+)", source_text)
    key_terms = term_match.group(1).strip() if term_match else ""
    game_type = first_text(profile.get("display_game_type"), table_value("游戏类型"), profile.get("game_type"))
    target_audience = first_text(profile.get("display_target_audience"), table_value("目标用户"), profile.get("target_audience"))
    content_scope = first_text(profile.get("display_content_scope"), table_value("内容构成", "内容范围"), profile.get("content_scope"))
    worldview = first_text(profile.get("display_worldview"), table_value("视觉与世界观", "世界观"), profile.get("tone"))
    style = first_text(profile.get("display_translation_style"), table_value("翻译风格", "风格要求"), profile.get("translation_style"))
    focus = first_text(profile.get("display_focus"), table_value("重点注意", "注意事项"))
    key_terms = first_text(profile.get("display_key_terms"), key_terms)
    lines = [
        f"\u4f60\u6b63\u5728\u5904\u7406\u300a{project_name}\u300b\u7684\u6e38\u620f\u672c\u5730\u5316\uff0c\u76ee\u6807\u8bed\u8a00\uff1a{spec.label}\u3002",
        sentence("项目定位", game_type),
        sentence("目标用户", target_audience),
        sentence("内容范围", content_scope),
        sentence("世界观/语气", worldview),
        sentence("风格要求", style),
        sentence("重点注意", focus),
        sentence("核心术语", key_terms),
        term_rule,
        f"\u5df2\u6709{spec.label}\u8bd1\u6587\u4ee3\u8868\u9879\u76ee\u5386\u53f2\u7528\u6cd5\uff1b\u5982\u9700\u4f18\u5316\uff0c\u4e0d\u80fd\u7834\u574f\u5df2\u56fa\u5b9a\u7684\u7cfb\u7edf\u672f\u8bed\u3002",
        "\u5fc5\u987b\u4fdd\u7559\u53d8\u91cf\u3001\u6570\u5b57\u3001\u6362\u884c\u3001\u989c\u8272\u6807\u7b7e\u3001HTML/\u5bcc\u6587\u672c\u6807\u7b7e\u548c\u5360\u4f4d\u7b26\uff0c\u4f8b\u5982 {0}\u3001%s\u3001<color>\u3002",
        "\u65e0\u6cd5\u786e\u8ba4\u7684\u4e13\u6709\u540d\u8bcd\u6216\u4fe1\u606f\u7f3a\u53e3\u7528 [TBD] \u6807\u8bb0\uff0c\u4e0d\u8981\u81ea\u884c\u7f16\u9020\u8bbe\u5b9a\u3002",
    ]
    return "\n".join(str(line).strip() for line in lines if str(line).strip())

def _project_brief_markdown(profile: dict[str, Any], prompt: str) -> str:
    game_type = str(profile.get("display_game_type") or profile.get("game_type") or "")
    target_audience = str(profile.get("display_target_audience") or profile.get("target_audience") or "")
    content_scope = str(profile.get("display_content_scope") or profile.get("content_scope") or "")
    translation_style = str(profile.get("display_translation_style") or profile.get("translation_style") or "")
    focus = str(profile.get("display_focus") or "")
    key_terms = str(profile.get("display_key_terms") or "")
    warning = str(profile.get("analysis_warning") or "")
    return (
        f"# {profile['project_name']} 翻译提示词与项目元信息\n\n"
        "## 🤖 AI 生成的专属翻译提示词\n\n"
        "```\n"
        f"{prompt}\n"
        "```\n\n"
        "## 📌 项目元信息\n\n"
        "| 项目 | 信息 |\n"
        "| --- | --- |\n"
        f"| 游戏类型 | {game_type} |\n"
        f"| 目标用户 | {target_audience} |\n"
        f"| 内容构成 | {content_scope} |\n"
        f"| 翻译风格 | {translation_style} |\n"
        f"| 重点注意 | {focus or '-'} |\n"
        f"| 核心术语 | {key_terms or '-'} |\n"
        f"| 语言资产 | {profile['language_assets']} |\n"
        f"| 分析状态 | {warning or profile.get('analysis_source', 'template')} |\n"
        f"| 生成日期 | {profile['generated_date']} |\n"
    )


def _project_analysis_report_markdown(profile: dict[str, Any], material_packet: dict[str, Any]) -> str:
    summary = material_packet.get("summary") or {}
    warnings = material_packet.get("warnings") or []
    materials = material_packet.get("materials") or []
    rows = [
        "# 项目资料 AI 分析报告",
        "",
        f"- 分析状态：{profile.get('analysis_source', 'template')}",
        f"- 模型：{profile.get('analysis_provider', '-')}/{profile.get('analysis_model', '-')}",
        f"- 资料总数：{summary.get('total', 0)}",
        f"- 已读取：{summary.get('parsed', 0)}",
        f"- 语言表候选：{summary.get('language_table_candidates', 0)}",
        f"- 图片资料：{summary.get('images', 0)}",
    ]
    if profile.get("analysis_warning"):
        rows.append(f"- 提示：{profile['analysis_warning']}")
    if warnings:
        rows.extend(["", "## 未完全分析的资料", *[f"- {warning}" for warning in warnings]])
    rows.append("")
    rows.append("## 资料读取明细")
    for material in materials:
        rows.append(f"- {material.get('label')}: {material.get('status')}；{material.get('note', '')}")
    return "\n".join(rows).strip() + "\n"


def _save_generated_project_harness(project: dict[str, Any], profile: dict[str, Any]) -> dict[str, Any]:
    current = read_project_harness(project["id"])
    metadata = dict(current.get("project_metadata") or {})
    metadata.update(
        {
            "game_type": profile["game_type"],
            "target_language": profile["target_language"],
            "content_scope": profile["content_scope"],
            "language_assets": profile["language_assets"],
            "source_materials": profile.get("source_materials", []),
            "generated_from": "project_analysis",
        }
    )
    return write_project_harness(
        project["id"],
        {
            "project_metadata": metadata,
            "style_guidance": profile["translation_style"],
            "target_audience": profile["target_audience"],
            "tone": profile["tone"],
        },
    )


def _profile_material_summary(material_packet: dict[str, Any]) -> dict[str, Any]:
    return {
        "summary": material_packet.get("summary") or {},
        "materials": [
            {
                "artifact_id": material.get("artifact_id"),
                "label": material.get("label"),
                "material_type": material.get("material_type"),
                "status": material.get("status"),
                "rows": material.get("rows"),
                "detected_languages": material.get("detected_languages") or [],
                "language_table_candidate": bool(material.get("language_table_candidate")),
                "warning": material.get("warning") or "",
            }
            for material in material_packet.get("materials", [])
            if isinstance(material, dict)
        ],
        "language_table_candidates": material_packet.get("language_table_candidates") or [],
        "warnings": material_packet.get("warnings") or [],
    }


def write_project_prompt(
    project: dict[str, Any],
    intro: str,
    asset_notes: list[str],
    target_language: str = "en",
    material_packet: dict[str, Any] | None = None,
) -> tuple[Path, Path, Path, Path, Path, str]:
    target_language = require_supported_language(target_language)
    root = project_dir(project["id"]) / "profile"
    root.mkdir(parents=True, exist_ok=True)
    material_packet = material_packet or build_project_material_packet(project["id"], [], load_settings(), run_visual_analysis=False)
    profile = _build_project_profile(project, intro, asset_notes, target_language=target_language, material_packet=material_packet)
    profile = _apply_project_analysis_provider(project, intro, asset_notes, profile, material_packet)
    profile["material_packet"] = _profile_material_summary(material_packet)
    profile["source_materials"] = [material.get("label") for material in material_packet.get("materials", []) if isinstance(material, dict)]
    prompt = _project_prompt_from_profile(profile)
    display_prompt = _project_prompt_display_zh(profile)
    _save_generated_project_harness(project, profile)
    profile_path = root / f"project_profile_{target_language}.json"
    prompt_path = root / f"translation_prompt_{target_language}.txt"
    brief_path = root / f"project_brief_{target_language}.md"
    packet_path = root / "project_material_packet.json"
    report_path = root / "project_analysis_report.md"
    profile_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    prompt_path.write_text(prompt, encoding="utf-8")
    brief_path.write_text(_project_brief_markdown(profile, prompt), encoding="utf-8")
    packet_path.write_text(json.dumps(material_packet, ensure_ascii=False, indent=2), encoding="utf-8")
    report_path.write_text(_project_analysis_report_markdown(profile, material_packet), encoding="utf-8")
    project_profile = dict(project.get("profile") or {})
    prompts = dict(project_profile.get("prompts_by_language") or {})
    prompts[target_language] = prompt
    display_prompts = dict(project_profile.get("display_prompts_by_language") or {})
    display_prompts[target_language] = display_prompt
    profiles = dict(project_profile.get("profiles_by_language") or {})
    profiles[target_language] = profile
    shared_profile_keys = (
        "game_type", "target_audience", "content_scope", "translation_style", "tone",
        "display_game_type", "display_target_audience", "display_content_scope", "display_worldview",
        "display_translation_style", "display_focus", "display_key_terms",
        "language_assets", "source_materials", "asset_notes", "material_packet", "analysis_source",
        "analysis_warning", "analysis_provider", "analysis_model", "confidence", "missing_info", "generated_date",
    )
    project_profile.update({key: profile[key] for key in shared_profile_keys if key in profile})
    project_profile["prompts_by_language"] = prompts
    project_profile["display_prompts_by_language"] = display_prompts
    project_profile["profiles_by_language"] = profiles
    updates: dict[str, Any] = {"profile": project_profile}
    if target_language == "en" or not str(project.get("prompt_text") or "").strip():
        updates["prompt_text"] = prompt
    db.update_project(project["id"], updates)
    return profile_path, prompt_path, brief_path, packet_path, report_path, prompt


def compile_project_harness_prompt(project: dict[str, Any], base_prompt: str, output_dir: Path) -> tuple[Path, Path, str, dict[str, Any]]:
    harness = read_project_harness(project["id"])
    parts = [base_prompt.strip()]
    project_parts = _project_harness_prompt_parts(harness)
    if project_parts:
        parts.append(
            "Project Harness (project-specific; apply only to this project, do not generalize):\n"
            + "\n".join(project_parts)
        )
    raw_compiled = "\n\n".join(part for part in parts if part)
    settings = load_settings()
    compiled = _manage_project_prompt_context(raw_compiled, settings)
    prompt_path = output_dir / "compiled_project_harness_prompt.txt"
    snapshot_path = output_dir / "project_harness_snapshot.json"
    snapshot = {
        "global_harness": GLOBAL_HARNESS_CONTRACT,
        "project_harness": harness,
        "summary": _harness_summary(harness),
        "context_budget": _project_context_summary(raw_compiled, settings),
    }
    prompt_path.write_text(compiled, encoding="utf-8")
    snapshot_path.write_text(json.dumps(snapshot, ensure_ascii=False, indent=2), encoding="utf-8")
    return prompt_path, snapshot_path, compiled, snapshot


def create_project_glossary_snapshot(project_id: str, run_id: str, output_dir: Path | None = None, language: str = "en") -> dict[str, Any]:
    language = require_supported_language(language)
    spec = language_spec(language)
    output = output_dir or run_dir(run_id) / "snapshots"
    output.mkdir(parents=True, exist_ok=True)
    path = output / ("project_glossary_snapshot.xlsx" if language == "en" else f"project_glossary_snapshot_{language}.xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Glossary"
    ws.append(["ID", "CN", spec.target_header, *(["EN2"] if spec.alt_header else []), "分类", "备注"])
    terms = db.list_glossary_terms(project_id, language=language)
    for term in reversed(terms):
        ws.append(_glossary_export_row(term, include_alt=bool(spec.alt_header)))
    wb.save(path)
    wb.close()
    return db.add_artifact(
        project_id,
        "Project glossary snapshot",
        path,
        "glossary_snapshot",
        run_id=run_id,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        origin="generated",
        metadata={"term_count": len(terms), "source": "project_glossary", "language": language},
    )


def create_prompt_and_harness_snapshots(project_id: str, run_id: str, output_dir: Path | None = None, language: str = "en") -> dict[str, Any]:
    language = require_supported_language(language)
    project = db.get_project(project_id)
    output = output_dir or run_dir(run_id) / "snapshots"
    output.mkdir(parents=True, exist_ok=True)
    profile = project.get("profile") or {}
    prompt_path = project_dir(project_id) / "profile" / f"translation_prompt_{language}.txt"
    if not prompt_path.exists():
        prompt_text = str((profile.get("prompts_by_language") or {}).get(language) or "")
        if prompt_text:
            prompt_path.write_text(prompt_text, encoding="utf-8")
        else:
            write_project_prompt(project, project.get("description", ""), [], target_language=language)
    base_prompt = prompt_path.read_text(encoding="utf-8")
    compiled_path, harness_path, compiled_prompt, harness_snapshot = compile_project_harness_prompt(project, base_prompt, output)
    prompt_artifact = db.add_artifact(
        project_id,
        "Prompt snapshot",
        compiled_path,
        "prompt_snapshot",
        run_id=run_id,
        mime="text/plain",
        origin="generated",
        metadata={"source": "project_prompt_and_harness", "language": language},
    )
    harness_artifact = db.add_artifact(
        project_id,
        "Project harness snapshot",
        harness_path,
        "project_harness_snapshot",
        run_id=run_id,
        mime="application/json",
        origin="generated",
        metadata={"source": "project_harness", "language": language},
    )
    return {
        "prompt": compiled_prompt,
        "prompt_artifact": prompt_artifact,
        "harness_artifact": harness_artifact,
        "harness_snapshot": harness_snapshot,
        "prompt_path": compiled_path,
        "harness_path": harness_path,
    }


def _quick_reference_excerpt(artifact: dict[str, Any], max_chars: int = 5000) -> str:
    path = Path(str(artifact.get("path") or ""))
    if not path.exists():
        return ""
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md", ".csv", ".json", ".jsonl"}:
            return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]
        if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
            wb = load_workbook(path, read_only=True, data_only=True)
            try:
                lines: list[str] = []
                remaining = max_chars
                for ws in wb.worksheets[:3]:
                    lines.append(f"[Sheet] {ws.title}")
                    for row in ws.iter_rows(max_row=16, values_only=True):
                        text = " | ".join(str(cell).strip() for cell in row if cell not in (None, ""))
                        if not text:
                            continue
                        lines.append(text)
                        remaining -= len(text)
                        if remaining <= 0:
                            break
                    if remaining <= 0:
                        break
                return "\n".join(lines)[:max_chars]
            finally:
                wb.close()
    except Exception:
        return "[reference read failed]"
    return f"[binary reference: {path.name}]"


def create_quick_reference_snapshot(project_id: str, run_id: str, reference_artifact_ids: list[str] | None, output_dir: Path | None = None) -> dict[str, Any] | None:
    ids = [str(item).strip() for item in (reference_artifact_ids or []) if str(item).strip()]
    if not ids:
        return None
    output = output_dir or run_dir(run_id) / "snapshots"
    output.mkdir(parents=True, exist_ok=True)
    rows: list[dict[str, Any]] = []
    context_parts = [
        "Temporary reference material for this quick task only.",
        "Use it as style/term/context guidance for this run. Do not treat it as permanent project memory unless the user imports it separately.",
    ]
    for artifact_id in ids:
        artifact = db.get_artifact(artifact_id)
        if artifact["project_id"] != project_id:
            raise KeyError(artifact_id)
        excerpt = _quick_reference_excerpt(artifact)
        item = {
            "id": artifact["id"],
            "label": artifact.get("label", ""),
            "kind": artifact.get("kind", ""),
            "role": artifact.get("role", ""),
            "origin": artifact.get("origin", ""),
            "original_filename": (artifact.get("metadata") or {}).get("original_filename", ""),
            "sha256": (artifact.get("metadata") or {}).get("sha256", ""),
            "size": artifact.get("size", 0),
            "excerpt": excerpt,
        }
        rows.append(item)
        context_parts.append(f"\nReference: {item['original_filename'] or item['label']} ({item['kind']})\n{excerpt}")
    snapshot_path = output / "quick_reference_snapshot.json"
    snapshot = {"source": "quick_task_reference", "reference_artifact_ids": ids, "references": rows}
    snapshot_path.write_text(json.dumps(snapshot, ensure_ascii=False, indent=2), encoding="utf-8")
    context = "\n".join(part for part in context_parts if part).strip()
    settings = load_settings()
    quick_ref_budget = max(500, int(settings.get("max_quick_reference_context_tokens") or 2000))
    managed_context = _cap_context_text(context, quick_ref_budget, "quick task references")
    context_summary = {
        "original_estimated_tokens": _estimate_text_tokens(context),
        "managed_estimated_tokens": _estimate_text_tokens(managed_context),
        "max_quick_reference_context_tokens": quick_ref_budget,
        "trimmed": managed_context != context,
    }
    artifact = db.add_artifact(
        project_id,
        "Quick task reference snapshot",
        snapshot_path,
        "quick_reference_snapshot",
        run_id=run_id,
        mime="application/json",
        origin="generated",
        metadata={"source": "quick_task_reference", "reference_artifact_ids": ids, "reference_count": len(rows), "context_budget": context_summary},
    )
    return {"artifact": artifact, "snapshot": snapshot, "context": managed_context, "context_budget": context_summary}


PROJECT_TEXT_MATERIAL_EXTENSIONS = {".md", ".markdown", ".txt"}
PROJECT_DOCX_MATERIAL_EXTENSIONS = {".docx"}
PROJECT_PDF_MATERIAL_EXTENSIONS = {".pdf"}
PROJECT_TABLE_MATERIAL_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
PROJECT_DELIMITED_MATERIAL_EXTENSIONS = {".csv", ".tsv"}
PROJECT_JSON_MATERIAL_EXTENSIONS = {".json", ".jsonl"}
PROJECT_IMAGE_MATERIAL_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
PROJECT_VIDEO_MATERIAL_EXTENSIONS = {".mp4", ".mov", ".mkv", ".avi"}
PROJECT_AUDIO_MATERIAL_EXTENSIONS = {".mp3", ".wav", ".m4a"}


def analyze_assets(artifact_ids: list[str], settings: dict[str, Any]) -> list[str]:
    packet = build_project_material_packet("", artifact_ids, settings, run_visual_analysis=False)
    return [str(material.get("note") or f"{material.get('label')}={material.get('status')}") for material in packet.get("materials", [])]


def _text_asset_note(path: Path) -> str:
    try:
        text = path.read_text(encoding="utf-8-sig", errors="replace")
    except OSError:
        return "text_material:read_failed"
    compact = re.sub(r"\s+", " ", text).strip()
    if not compact:
        return "text_material:empty"
    return f"text_material:{compact[:800]}"


def build_project_material_packet(
    project_id: str,
    artifact_ids: list[str],
    settings: dict[str, Any],
    *,
    run_visual_analysis: bool = True,
) -> dict[str, Any]:
    materials: list[dict[str, Any]] = []
    context_parts: list[str] = []
    warnings: list[str] = []
    language_table_candidates: list[dict[str, Any]] = []
    for artifact_id in artifact_ids:
        artifact = db.get_artifact(artifact_id)
        material = _parse_project_material_artifact(artifact, settings, run_visual_analysis=run_visual_analysis, project_id=project_id)
        materials.append(material)
        if material.get("excerpt"):
            context_parts.append(f"## {material.get('label')}\n{material.get('excerpt')}")
        if material.get("language_table_candidate"):
            language_table_candidates.append({
                "artifact_id": artifact_id,
                "label": material.get("label"),
                "languages": material.get("detected_languages", []),
                "rows": material.get("rows", 0),
            })
        if material.get("warning"):
            warnings.append(str(material["warning"]))
    parsed = [m for m in materials if str(m.get("status") or "").startswith(("parsed", "vision_analyzed"))]
    image_materials = [m for m in materials if m.get("material_type") == "image"]
    unsupported = [m for m in materials if str(m.get("status") or "").startswith(("archived_only", "unsupported"))]
    return {
        "version": 1,
        "source": "project_material_packet",
        "project_id": project_id,
        "artifact_ids": artifact_ids,
        "materials": materials,
        "context": _cap_context_text("\n\n".join(context_parts), 9000, "project material packet"),
        "summary": {
            "total": len(materials),
            "parsed": len(parsed),
            "unsupported": len(unsupported),
            "images": len(image_materials),
            "language_table_candidates": len(language_table_candidates),
            "warnings": len(warnings),
        },
        "language_table_candidates": language_table_candidates,
        "warnings": warnings,
    }


def _parse_project_material_artifact(artifact: dict[str, Any], settings: dict[str, Any], *, run_visual_analysis: bool, project_id: str = "") -> dict[str, Any]:
    path = Path(artifact["path"])
    suffix = path.suffix.lower()
    base = {
        "artifact_id": artifact.get("id"),
        "label": artifact.get("label") or path.name,
        "filename": path.name,
        "kind": artifact.get("kind", ""),
        "role": artifact.get("role", ""),
        "origin": artifact.get("origin", ""),
        "size": artifact.get("size", 0),
        "suffix": suffix,
        "status": "parsed",
        "note": "",
        "excerpt": "",
    }
    try:
        if suffix in PROJECT_TEXT_MATERIAL_EXTENSIONS:
            result = _parse_project_text_material(path)
        elif suffix in PROJECT_DOCX_MATERIAL_EXTENSIONS:
            result = _parse_project_docx_material(path)
        elif suffix in PROJECT_PDF_MATERIAL_EXTENSIONS:
            result = _parse_project_pdf_material(path, settings)
        elif suffix in PROJECT_TABLE_MATERIAL_EXTENSIONS:
            result = _parse_project_xlsx_material(path)
        elif suffix in PROJECT_DELIMITED_MATERIAL_EXTENSIONS:
            result = _parse_project_delimited_material(path)
        elif suffix in PROJECT_JSON_MATERIAL_EXTENSIONS:
            result = _parse_project_json_material(path)
        elif suffix in PROJECT_IMAGE_MATERIAL_EXTENSIONS:
            result = _parse_project_image_material(path, settings, run_visual_analysis=run_visual_analysis)
        elif suffix in PROJECT_VIDEO_MATERIAL_EXTENSIONS:
            result = _parse_project_video_material(path, settings, run_visual_analysis=run_visual_analysis, project_id=project_id, artifact_id=str(artifact.get("id") or "video"))
        elif suffix in PROJECT_AUDIO_MATERIAL_EXTENSIONS:
            result = {"material_type": "audio", "status": "archived_only:audio_analysis_not_supported", "warning": f"{path.name} 已归档，v1 未进行音频内容分析。", "excerpt": f"音频资料：{path.name}。v1 未分析内容。"}
        else:
            result = {"material_type": "unknown", "status": "archived_only:unknown_type", "warning": f"{path.name} 文件类型暂未分析。", "excerpt": f"未分析资料：{path.name}"}
    except Exception as exc:
        result = {"material_type": "unknown", "status": "parse_failed", "warning": f"{path.name} 解析失败：{user_facing_error(exc)}", "excerpt": ""}
    item = {**base, **result}
    item["note"] = _project_material_note(item)
    return item


def _project_material_note(material: dict[str, Any]) -> str:
    label = material.get("label") or material.get("filename") or "material"
    status = material.get("status") or "parsed"
    if material.get("language_table_candidate"):
        langs = "/".join(str(x) for x in material.get("detected_languages") or [])
        return f"{label}=language_table_candidate:{material.get('rows', 0)} rows:{langs}"
    if material.get("material_type") == "image":
        return f"{label}={status}:{str(material.get('visual_summary') or material.get('excerpt') or '')[:300]}"
    return f"{label}={status}:{str(material.get('excerpt') or '')[:300]}"


def _parse_project_text_material(path: Path) -> dict[str, Any]:
    text = _read_lookup_text_file(path)
    compact = re.sub(r"\s+", " ", text).strip()
    headings = re.findall(r"(?m)^#{1,4}\s+(.+)$", text)[:20]
    table_rows = re.findall(r"(?m)^\|.+\|$", text)[:40]
    excerpt = "\n".join(part for part in [
        "\n".join(f"# {heading}" for heading in headings[:8]),
        "\n".join(table_rows[:20]),
        compact[:3000],
    ] if part).strip()
    return {"material_type": "text", "status": "parsed_text", "chars": len(text), "headings": headings, "table_rows": table_rows[:20], "excerpt": excerpt}


def _parse_project_docx_material(path: Path) -> dict[str, Any]:
    with ZipFile(path) as archive:
        if "word/document.xml" not in archive.namelist():
            raise ValueError(f"Missing word/document.xml in DOCX file: {path}")
        root = ET.fromstring(archive.read("word/document.xml"))
    paragraphs: list[str] = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for paragraph in root.findall(".//w:p", ns):
        text = re.sub(r"\s+", " ", "".join(node.text or "" for node in paragraph.findall(".//w:t", ns))).strip()
        if text:
            paragraphs.append(text)
    return {"material_type": "docx", "status": "parsed_docx", "paragraphs": len(paragraphs), "excerpt": "\n".join(paragraphs[:120])[:5000]}


def _parse_project_pdf_material(path: Path, settings: dict[str, Any]) -> dict[str, Any]:
    if not (settings.get("multimodal", {}) or {}).get("pdf", True):
        return {"material_type": "pdf", "status": "archived_only:pdf_not_supported", "warning": f"{path.name} 已归档，当前设置未启用 PDF 文字提取。", "excerpt": ""}
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        return {"material_type": "pdf", "status": "archived_only:pdf_parser_missing", "warning": f"缺少 PDF 解析依赖：{exc}", "excerpt": ""}
    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages[:20]:
        text = re.sub(r"\s+", " ", page.extract_text() or "").strip()
        if text:
            chunks.append(text)
    return {"material_type": "pdf", "status": "parsed_pdf_text", "pages": len(reader.pages), "excerpt": "\n".join(chunks)[:5000]}


def _parse_project_xlsx_material(path: Path) -> dict[str, Any]:
    wb = load_workbook(path, read_only=True, data_only=True)
    sheets: list[dict[str, Any]] = []
    detected_languages: set[str] = set()
    total_rows = 0
    samples: list[dict[str, Any]] = []
    try:
        for ws in wb.worksheets[:12]:
            rows_iter = ws.iter_rows(values_only=True)
            first = next(rows_iter, None)
            headers = [str(value or "").strip() for value in (first or [])]
            normalized = _normalized_header_indices(headers)
            id_idx = _column_index(normalized, None, ["id", "key", "编号", "序号"], required=False)
            source_idx = _column_index(normalized, None, ["source", "original", "cn", "zh", "chinese", "原文", "中文", "简体中文"], required=False)
            reserved = {idx for idx in (id_idx, source_idx) if idx is not None}
            language_indices = _auto_language_indices(headers, reserved)
            detected_languages.update(language_indices.keys())
            row_count = 0
            for row in rows_iter:
                if any(str(cell or "").strip() for cell in row):
                    row_count += 1
                if len(samples) < 30 and source_idx is not None:
                    source = _value_at(row, source_idx)
                    if source:
                        item = {"sheet": ws.title, "source": source}
                        for code, (target_idx, _alt_idx) in language_indices.items():
                            target = _value_at(row, target_idx)
                            if target:
                                item[visible_language_code(code)] = target
                        samples.append(item)
            total_rows += row_count
            sheets.append({"name": ws.title, "headers": headers[:40], "rows": row_count, "source_column": headers[source_idx] if source_idx is not None and source_idx < len(headers) else "", "languages": sorted(visible_language_code(code) for code in language_indices)})
    finally:
        wb.close()
    language_table_candidate = total_rows > 0 and bool(detected_languages) and any(sheet.get("source_column") for sheet in sheets)
    excerpt = json.dumps({"sheets": sheets[:8], "samples": samples[:20]}, ensure_ascii=False, indent=2)
    return {
        "material_type": "table",
        "status": "parsed_xlsx",
        "rows": total_rows,
        "sheets": sheets,
        "samples": samples,
        "detected_languages": sorted(visible_language_code(code) for code in detected_languages),
        "language_table_candidate": language_table_candidate,
        "excerpt": excerpt[:7000],
    }


def _parse_project_delimited_material(path: Path) -> dict[str, Any]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    raw = _read_lookup_text_file(path)
    rows = list(csv.reader(raw.splitlines(), delimiter=delimiter))
    headers = [str(value or "").strip() for value in rows[0]] if rows else []
    samples = rows[1:31]
    return {"material_type": "table", "status": "parsed_delimited", "rows": max(0, len(rows) - 1), "headers": headers, "excerpt": json.dumps({"headers": headers, "samples": samples}, ensure_ascii=False)[:4000]}


def _parse_project_json_material(path: Path) -> dict[str, Any]:
    text = _read_lookup_text_file(path)
    try:
        payload = json.loads(text)
        excerpt = json.dumps(payload, ensure_ascii=False, indent=2)[:5000]
    except Exception:
        excerpt = text[:5000]
    return {"material_type": "json", "status": "parsed_json", "excerpt": excerpt}


def _parse_project_image_material(path: Path, settings: dict[str, Any], *, run_visual_analysis: bool) -> dict[str, Any]:
    if not run_visual_analysis:
        return {"material_type": "image", "status": "archived_only:image_not_requested", "excerpt": f"图片资料：{path.name}"}
    provider = normalize_provider_name(settings.get("provider"))
    if provider not in REAL_PROVIDERS or not settings.get("api_key"):
        return {"material_type": "image", "status": "archived_only:image_api_key_missing", "warning": f"{path.name} 已归档，未配置支持图片分析的 API。", "excerpt": f"图片资料：{path.name}；未配置 API，未做视觉分析。"}
    prompt = (
        "请分析这张游戏本地化项目资料图。返回严格 JSON："
        "{\"ui_type\":\"\", \"theme\":\"\", \"worldview\":\"\", \"characters_or_scene\":\"\", "
        "\"visible_text\":\"\", \"localization_style_hints\":\"\", \"confidence\":\"low|medium|high\"}。"
        "只基于图片可见信息，不要编造。"
    )
    try:
        raw = call_image_text(settings, prompt, path, system="Return strict JSON only.")
        payload = _parse_semantic_qa_payload(raw)
        if not isinstance(payload, dict):
            raise ValueError("image provider returned non-object")
        summary = "；".join(str(payload.get(key) or "").strip() for key in ("ui_type", "theme", "worldview", "characters_or_scene", "visible_text", "localization_style_hints") if str(payload.get(key) or "").strip())
        return {"material_type": "image", "status": "vision_analyzed", "visual_analysis": payload, "visual_summary": summary, "excerpt": summary[:2500]}
    except Exception as exc:
        return {"material_type": "image", "status": "vision_failed", "warning": f"{path.name} 图片分析失败：{user_facing_error(exc)}", "excerpt": f"图片资料：{path.name}；视觉分析失败。"}


def _extract_video_keyframes(path: Path, output_dir: Path, *, max_frames: int = 4) -> tuple[list[Path], dict[str, Any]]:
    try:
        import cv2  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"缺少视频抽帧依赖 OpenCV：{exc}") from exc
    output_dir.mkdir(parents=True, exist_ok=True)
    cap = cv2.VideoCapture(str(path))
    if not cap.isOpened():
        raise RuntimeError("视频无法打开，可能是格式不受支持或文件损坏")
    try:
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        fps = float(cap.get(cv2.CAP_PROP_FPS) or 0)
        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH) or 0)
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT) or 0)
        if frame_count > 0:
            ratios = [0.08, 0.32, 0.58, 0.84][:max_frames]
            indices = sorted({min(frame_count - 1, max(0, int(frame_count * ratio))) for ratio in ratios})
        else:
            indices = list(range(max_frames))
        frames: list[Path] = []
        for order, index in enumerate(indices, start=1):
            if frame_count > 0:
                cap.set(cv2.CAP_PROP_POS_FRAMES, index)
            ok, frame = cap.read()
            if not ok or frame is None:
                continue
            frame_path = output_dir / f"frame_{order:02d}.png"
            if cv2.imwrite(str(frame_path), frame):
                frames.append(frame_path)
        if not frames:
            raise RuntimeError("没有抽取到可用关键帧")
        return frames, {"frame_count": frame_count, "fps": fps, "width": width, "height": height, "frames": [frame.name for frame in frames]}
    finally:
        cap.release()


def _parse_project_video_material(path: Path, settings: dict[str, Any], *, run_visual_analysis: bool, project_id: str, artifact_id: str) -> dict[str, Any]:
    if not run_visual_analysis:
        return {"material_type": "video", "status": "archived_only:video_not_requested", "excerpt": f"视频资料：{path.name}"}
    provider = normalize_provider_name(settings.get("provider"))
    if provider not in REAL_PROVIDERS or not settings.get("api_key"):
        return {"material_type": "video", "status": "archived_only:video_api_key_missing", "warning": f"{path.name} 已归档，未配置支持视频关键帧分析的 API。", "excerpt": f"视频资料：{path.name}；未配置 API，未做画面分析。"}
    frame_root = (project_dir(project_id) / "profile" / "video_frames" / artifact_id) if project_id else (path.parent / ".lws_video_frames" / artifact_id)
    try:
        frames, frame_meta = _extract_video_keyframes(path, frame_root, max_frames=4)
    except Exception as exc:
        return {"material_type": "video", "status": "video_frame_extract_failed", "warning": f"{path.name} 抽帧失败：{user_facing_error(exc)}", "excerpt": f"视频资料：{path.name}；抽帧失败。"}
    prompt = (
        "请分析这个游戏本地化项目视频的关键帧。返回严格 JSON："
        "{\"ui_type\":\"\", \"theme\":\"\", \"worldview\":\"\", \"characters_or_scene\":\"\", "
        "\"visible_text\":\"\", \"localization_style_hints\":\"\", \"confidence\":\"low|medium|high\"}。"
        "只基于当前帧可见信息，不要编造；如果文字看不清，visible_text 写空。"
    )
    analyses: list[dict[str, Any]] = []
    warnings: list[str] = []
    for frame in frames:
        try:
            raw = call_image_text(settings, prompt, frame, system="Return strict JSON only.")
            payload = _parse_semantic_qa_payload(raw)
            if not isinstance(payload, dict):
                payload = {"summary": str(raw)[:1000]}
            payload["frame"] = frame.name
            analyses.append(payload)
        except Exception as exc:
            warnings.append(f"{frame.name}: {user_facing_error(exc)}")
    if not analyses:
        return {"material_type": "video", "status": "vision_failed_video", "warning": f"{path.name} 关键帧视觉分析失败：{'；'.join(warnings[:3])}", "frame_meta": frame_meta, "excerpt": f"视频资料：{path.name}；关键帧分析失败。"}
    summary_parts: list[str] = []
    for payload in analyses:
        summary = "；".join(str(payload.get(key) or "").strip() for key in ("ui_type", "theme", "worldview", "characters_or_scene", "visible_text", "localization_style_hints") if str(payload.get(key) or "").strip())
        if summary:
            summary_parts.append(f"{payload.get('frame')}: {summary}")
    warning = f"{path.name} 部分关键帧分析失败：{'；'.join(warnings[:3])}" if warnings else ""
    return {
        "material_type": "video",
        "status": "vision_analyzed_video",
        "frames_analyzed": len(analyses),
        "frame_meta": frame_meta,
        "visual_analysis": analyses,
        "warning": warning,
        "excerpt": "\n".join(summary_parts)[:4000],
    }


def _read_lookup_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _read_lookup_csv_text(path: Path) -> str:
    raw = _read_lookup_text_file(path)
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    rows = csv.reader(raw.splitlines(), delimiter=delimiter)
    cells: list[str] = []
    for row in rows:
        cells.extend(str(cell).strip() for cell in row if str(cell).strip())
    return "\n".join(cells)


def _read_lookup_xlsx_text(path: Path) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        cells: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells.extend(str(cell).strip() for cell in row if cell not in (None, ""))
        return "\n".join(cells)
    finally:
        wb.close()


def _read_lookup_docx_text(path: Path) -> str:
    with ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    root = ET.fromstring(xml)
    texts = [node.text for node in root.iter() if node.tag.endswith("}t") and node.text]
    return "\n".join(texts)


def _read_lookup_material_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown", ".json"}:
        return _read_lookup_text_file(path)
    if suffix in {".csv", ".tsv"}:
        return _read_lookup_csv_text(path)
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return _read_lookup_xlsx_text(path)
    if suffix == ".docx":
        return _read_lookup_docx_text(path)
    return ""


def _compact_lookup_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _count_lookup_hits(text: str, needle: str) -> tuple[int, int]:
    if not needle:
        return (0, -1)
    count = 0
    first = -1
    start = 0
    while True:
        index = text.find(needle, start)
        if index < 0:
            break
        if first < 0:
            first = index
        count += 1
        start = index + max(1, len(needle))
    return (count, first)


def _suppress_overlapping_lookup_hits(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    accepted: list[dict[str, Any]] = []
    spans: list[tuple[int, int]] = []
    for row in sorted(rows, key=lambda item: (int(item.get("first_position") or 0), -len(str(item.get("source") or "")), str(item.get("source") or ""))):
        start = int(row.get("first_position") or 0)
        end = start + len(str(row.get("source") or ""))
        if any(start < existing_end and end > existing_start for existing_start, existing_end in spans):
            continue
        accepted.append(row)
        spans.append((start, end))
    return accepted


def _rank_translation_lookup_source(source_type: str) -> int:
    priority = {
        "qa_passed": 0,
        "qa_final": 0,
        "manual": 1,
        "imported": 2,
        "archive": 2,
        "translation_archive": 2,
    }
    return priority.get(str(source_type or "").strip().lower(), 3)


def _lookup_terms(text: str, terms: list[dict[str, Any]], *, min_length: int, limit: int) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for term in terms:
        source = str(term.get("source") or "").strip()
        if len(source) < min_length or not (str(term.get("target") or "").strip() or str(term.get("target_alt") or "").strip()):
            continue
        hit_count, first_position = _count_lookup_hits(text, source)
        if not hit_count:
            continue
        rows.append(
            {
                "id": term.get("id"),
                "term_key": term.get("term_key", ""),
                "source": source,
                "target": term.get("target", ""),
                "target_alt": term.get("target_alt", ""),
                "language": term.get("language", "en"),
                "category": term.get("category", ""),
                "note": term.get("note", ""),
                "source_type": term.get("source_type", ""),
                "first_position": first_position,
                "hit_count": hit_count,
            }
        )
    return _suppress_overlapping_lookup_hits(rows)[:limit]


def _lookup_translation_entries(text: str, entries: list[dict[str, Any]], *, min_length: int, limit: int) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for entry in entries:
        source = str(entry.get("source") or "").strip()
        if len(source) < min_length or not (str(entry.get("target") or "").strip() or str(entry.get("target_alt") or "").strip()):
            continue
        hit_count, first_position = _count_lookup_hits(text, source)
        if not hit_count:
            continue
        rows.append(
            {
                "id": entry.get("id"),
                "entry_key": entry.get("entry_key", ""),
                "source": source,
                "target": entry.get("target", ""),
                "target_alt": entry.get("target_alt", ""),
                "language": entry.get("language", "en"),
                "sheet": entry.get("sheet", ""),
                "row_number": entry.get("row_number", 0),
                "note": entry.get("note", ""),
                "source_type": entry.get("source_type", ""),
                "source_artifact_id": entry.get("source_artifact_id", ""),
                "first_position": first_position,
                "hit_count": hit_count,
                "_priority": _rank_translation_lookup_source(str(entry.get("source_type") or "")),
            }
        )
    rows.sort(key=lambda item: (int(item.get("first_position") or 0), item.get("_priority", 3), -len(str(item.get("source") or "")), str(item.get("source") or "")))
    accepted = _suppress_overlapping_lookup_hits(rows)[:limit]
    for row in accepted:
        row.pop("_priority", None)
    return accepted


def _announcement_prompt_context(project: dict[str, Any], language: str, terms: list[dict[str, Any]], translations: list[dict[str, Any]]) -> str:
    spec = language_spec(language)
    lines = [
        "公告翻译检索上下文（不含正文译文）",
        f"项目：{project['name']}",
        f"目标语言：{spec.prompt_name} / {spec.target_header}",
        "用途：供下游长文本翻译工作流使用；必须优先遵守项目术语和 QA 通过译文参考。",
        "",
        "【命中的项目术语】",
    ]
    if terms:
        for term in terms:
            alt = f" / {term['target_alt']}" if term.get("target_alt") else ""
            note = f"（{term['note']}）" if term.get("note") else ""
            lines.append(f"- {term['source']} => {term['target']}{alt}{note}")
    else:
        lines.append("- 无命中；当前语言缺少术语约束或公告文本未命中现有术语。")
    lines.extend(["", "【命中的 QA/归档译文参考】"])
    if translations:
        for entry in translations:
            alt = f" / {entry['target_alt']}" if entry.get("target_alt") else ""
            meta = f"{entry.get('sheet') or 'Archive'}:{entry.get('row_number') or ''}".rstrip(":")
            lines.append(f"- {entry['source']} => {entry['target']}{alt}（{meta} / {entry.get('source_type') or 'archive'}）")
    else:
        lines.append("- 无命中；下游翻译需仅依赖项目提示词和通用语言质量要求。")
    lines.extend(
        [
            "",
            "【硬性要求】",
            "- 不修改变量、占位符、HTML/富文本标签、数字、专名和换行结构。",
            "- 术语命中时使用上述译法；同一概念在整篇公告中保持一致。",
            "- 本检索包不代表公告正文已翻译，正文翻译和 QA 由后续工作流完成。",
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _write_announcement_lookup_workbook(path: Path, project: dict[str, Any], language: str, text: str, terms: list[dict[str, Any]], translations: list[dict[str, Any]], prompt_context: str) -> None:
    spec = language_spec(language)
    wb = Workbook()
    overview = wb.active
    overview.title = "Overview"
    overview.append(["Field", "Value"])
    overview.append(["Project", project["name"]])
    overview.append(["Project ID", project["id"]])
    overview.append(["Language", language])
    overview.append(["Target header", spec.target_header])
    overview.append(["Text chars", len(text)])
    overview.append(["Matched terms", len(terms)])
    overview.append(["Matched translations", len(translations)])

    term_sheet = wb.create_sheet("MatchedTerms")
    term_sheet.append(["CN", spec.target_header, *([spec.alt_header] if spec.alt_header else []), "分类", "备注", "first_position", "hit_count", "term_id"])
    for term in terms:
        row = [term["source"], term["target"]]
        if spec.alt_header:
            row.append(term.get("target_alt", ""))
        row.extend([term.get("category", ""), term.get("note", ""), term["first_position"], term["hit_count"], term.get("id", "")])
        term_sheet.append(row)

    translation_sheet = wb.create_sheet("MatchedTranslations")
    translation_sheet.append(["ID", "CN", spec.target_header, *([spec.alt_header] if spec.alt_header else []), "source_type", "sheet", "row_number", "first_position", "hit_count", "entry_id"])
    for entry in translations:
        row = [entry.get("entry_key", ""), entry["source"], entry["target"]]
        if spec.alt_header:
            row.append(entry.get("target_alt", ""))
        row.extend([entry.get("source_type", ""), entry.get("sheet", ""), entry.get("row_number", 0), entry["first_position"], entry["hit_count"], entry.get("id", "")])
        translation_sheet.append(row)

    context_sheet = wb.create_sheet("PromptContext")
    context_sheet.append(["Prompt context"])
    for line in prompt_context.splitlines():
        context_sheet.append([line])
    wb.save(path)
    wb.close()


def run_announcement_lookup(project_id: str, request: Any) -> dict[str, Any]:
    project = db.get_project(project_id)
    language = require_supported_language(getattr(request, "language", "en") or "en")
    material_ids = list(getattr(request, "material_artifact_ids", []) or [])
    direct_text = str(getattr(request, "text", "") or "")
    if not material_ids and not direct_text.strip():
        raise ValueError("announcement lookup requires material_artifact_ids or text")

    text_parts: list[str] = []
    materials: list[dict[str, Any]] = []
    if direct_text.strip():
        text_parts.append(direct_text)
    for artifact_id in material_ids:
        artifact = db.get_artifact(artifact_id)
        if artifact["project_id"] != project_id:
            raise KeyError(artifact_id)
        path = Path(artifact["path"])
        material_text = _read_lookup_material_text(path) if path.exists() else ""
        text_parts.append(material_text)
        materials.append({"id": artifact["id"], "label": artifact.get("label", ""), "kind": artifact.get("kind", ""), "chars": len(material_text)})

    text = _compact_lookup_text("\n".join(text_parts))
    if not text:
        raise ValueError("announcement lookup text is empty")

    max_terms = max(0, min(int(getattr(request, "max_terms", 300) or 300), 1000))
    max_translation_rows = max(0, min(int(getattr(request, "max_translation_rows", 300) or 300), 1000))
    min_term_length = max(1, int(getattr(request, "min_term_length", 2) or 2))
    min_translation_length = max(1, int(getattr(request, "min_translation_length", 4) or 4))

    run = db.insert_run(
        project_id,
        kind="announcement_lookup",
        language=language,
        metadata={
            "request": {
                "material_artifact_ids": material_ids,
                "has_inline_text": bool(direct_text.strip()),
                "language": language,
                "include_glossary": bool(getattr(request, "include_glossary", True)),
                "include_translation_archive": bool(getattr(request, "include_translation_archive", True)),
            }
        },
    )
    db.add_event(run["id"], "announcement lookup started")
    try:
        glossary_rows = db.list_glossary_terms(project_id, language=language) if bool(getattr(request, "include_glossary", True)) else []
        archive_rows = db.list_translation_entries(project_id, language=language) if bool(getattr(request, "include_translation_archive", True)) else []
        matched_terms = _lookup_terms(text, glossary_rows, min_length=min_term_length, limit=max_terms)
        matched_translations = _lookup_translation_entries(text, archive_rows, min_length=min_translation_length, limit=max_translation_rows)
        text_fingerprint = hashlib.sha256(text.encode("utf-8")).hexdigest()
        constraint_status = "available" if matched_terms or matched_translations else "missing"
        summary = {
            "language": language,
            "text_chars": len(text),
            "text_fingerprint": text_fingerprint,
            "materials": len(materials),
            "matched_terms": len(matched_terms),
            "matched_translations": len(matched_translations),
            "constraint_status": constraint_status,
        }
        prompt_context = _announcement_prompt_context(project, language, matched_terms, matched_translations)
        manifest = {
            "kind": "announcement_lookup",
            "project_id": project_id,
            "project_name": project["name"],
            "language": language,
            "materials": materials,
            "text_fingerprint": text_fingerprint,
            "limits": {
                "min_term_length": min_term_length,
                "min_translation_length": min_translation_length,
                "max_terms": max_terms,
                "max_translation_rows": max_translation_rows,
            },
            "summary": summary,
            "matched_terms": matched_terms,
            "matched_translations": matched_translations,
        }

        output = run_dir(run["id"]) / "announcement_lookup"
        output.mkdir(parents=True, exist_ok=True)
        lang_code = _visible_language_code(language)
        workbook_path = output / f"announcement_lookup_{lang_code}.xlsx"
        manifest_path = output / f"announcement_lookup_manifest_{lang_code}.json"
        prompt_path = output / f"announcement_lookup_prompt_context_{lang_code}.txt"
        _write_announcement_lookup_workbook(workbook_path, project, language, text, matched_terms, matched_translations, prompt_context)
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        prompt_path.write_text(prompt_context, encoding="utf-8")

        artifact_metadata = {"language": language, "summary": summary, "text_fingerprint": text_fingerprint}
        artifacts = [
            db.add_artifact(project_id, f"Announcement lookup workbook ({lang_code})", workbook_path, "announcement_lookup_workbook", run_id=run["id"], mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata=artifact_metadata),
            db.add_artifact(project_id, f"Announcement lookup manifest ({lang_code})", manifest_path, "announcement_lookup_manifest", run_id=run["id"], mime="application/json", metadata=artifact_metadata),
            db.add_artifact(project_id, f"Announcement lookup prompt context ({lang_code})", prompt_path, "announcement_lookup_prompt_context", run_id=run["id"], mime="text/plain", metadata=artifact_metadata),
        ]
        db.add_event(run["id"], f"announcement lookup matched terms={len(matched_terms)} translations={len(matched_translations)}")
        run = db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "summary": summary, "manifest_path": str(manifest_path)})
        return {"run": run, "summary": summary, "artifacts": artifacts, "manifest": manifest}
    except Exception as exc:
        friendly = user_facing_error(exc)
        db.add_event(run["id"], f"announcement lookup failed: {friendly}", level="error")
        db.update_run(run["id"], status="failed", metadata={**run.get("metadata", {}), "error": friendly})
        raise


ANNOUNCEMENT_STEP = {
    "source": 1,
    "constraints": 2,
    "languages": 3,
    "terms": 4,
    "lookup": 5,
    "prepare": 6,
    "translate": 7,
    "apply": 8,
    "deliver": 9,
}


def list_announcement_tasks(project_id: str) -> list[dict[str, Any]]:
    db.get_project(project_id)
    return [_hydrate_announcement_task(task) for task in db.list_announcement_tasks(project_id)]


def get_announcement_task(task_id: str) -> dict[str, Any]:
    return _hydrate_announcement_task(db.get_announcement_task(task_id))


def cancel_announcement_task(task_id: str) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    metadata["canceled_at"] = db.now_iso()
    task = db.update_announcement_task(
        task_id,
        status="canceled",
        current_step=task.get("current_step") or 1,
        metadata=metadata,
    )
    for item in task.get("languages") or []:
        lang_meta = dict(item.get("metadata") or {})
        lang_meta["canceled_at"] = metadata["canceled_at"]
        db.upsert_announcement_task_language(
            task_id,
            task["project_id"],
            item["language"],
            status="canceled",
            current_step=item.get("current_step") or task.get("current_step") or 1,
            metadata=lang_meta,
        )
    return {"task": _hydrate_announcement_task(db.get_announcement_task(task_id))}


def cancel_announcement_translation_task(task_id: str) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    metadata["translation_cancel_requested_at"] = db.now_iso()
    metadata["reason"] = "announcement_translation_canceled"
    task = db.update_announcement_task(
        task_id,
        status="prepared",
        current_step=ANNOUNCEMENT_STEP["translate"],
        metadata=metadata,
    )
    for item in task.get("languages") or []:
        if item.get("status") in {"queued", "running"}:
            lang_meta = dict(item.get("metadata") or {})
            lang_meta["translation_cancel_requested_at"] = metadata["translation_cancel_requested_at"]
            db.upsert_announcement_task_language(
                task_id,
                task["project_id"],
                str(item["language"]),
                status="prepared",
                current_step=ANNOUNCEMENT_STEP["translate"],
                metadata=lang_meta,
            )
    return {"task": _hydrate_announcement_task(db.get_announcement_task(task_id))}


def create_announcement_task(project_id: str, request: Any) -> dict[str, Any]:
    project = db.get_project(project_id)
    source_artifact_id = str(getattr(request, "source_artifact_id", "") or "").strip()
    text = str(getattr(request, "text", "") or "")
    if not source_artifact_id and text.strip():
        source_artifact_id = _create_inline_announcement_source(project_id, text, getattr(request, "title", "") or "announcement")["id"]
    if not source_artifact_id:
        raise ValueError("announcement task requires source_artifact_id or text")
    source_artifact = db.get_artifact(source_artifact_id)
    if source_artifact["project_id"] != project_id:
        raise KeyError(source_artifact_id)
    source_format = _announcement_source_format(Path(source_artifact["path"]))
    if source_format not in {"docx", "txt", "xlsx"}:
        raise ValueError("announcement source must be DOCX, TXT, or XLSX")

    metadata = {
        "project_name": project["name"],
        "output_policy": str(getattr(request, "output_policy", "same_format") or "same_format"),
        "language_table_artifact_ids": list(getattr(request, "language_table_artifact_ids", []) or []),
        "constraint_artifact_ids": list(getattr(request, "constraint_artifact_ids", []) or []),
        "include_project_archive": bool(getattr(request, "include_project_archive", True)),
        "source": _announcement_source_manifest(source_artifact),
    }
    detected_languages = _detect_announcement_constraint_languages(project_id, metadata)
    requested_languages = _normalize_announcement_languages(getattr(request, "languages", []) or [], fallback=detected_languages)
    task = db.insert_announcement_task(
        project_id,
        {
            "title": str(getattr(request, "title", "") or source_artifact.get("label") or "公告翻译").strip(),
            "source_artifact_id": source_artifact_id,
            "source_format": source_format,
            "selected_languages": requested_languages,
            "status": "source_ready",
            "current_step": ANNOUNCEMENT_STEP["constraints"],
            "metadata": {**metadata, "detected_languages": detected_languages},
        },
    )
    return _hydrate_announcement_task(task)


def inspect_announcement_constraints(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    metadata = _merge_announcement_constraint_request(metadata, request)
    detected = _detect_announcement_constraint_languages(task["project_id"], metadata)
    selected = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=detected)
    metadata["detected_languages"] = detected
    metadata["language_constraints"] = _announcement_language_constraint_summary(task["project_id"], metadata, selected)
    confirmed = bool(getattr(request, "confirm_languages", False))
    next_step = ANNOUNCEMENT_STEP["terms"] if confirmed else ANNOUNCEMENT_STEP["languages"]
    next_status = "languages_ready" if confirmed else "constraints_ready" if detected or selected else "missing_constraints"
    task = db.update_announcement_task(
        task_id,
        status=next_status,
        current_step=next_step,
        selected_languages=selected,
        metadata=metadata,
    )
    for language in selected:
        db.upsert_announcement_task_language(task_id, task["project_id"], language, status=next_status, current_step=next_step)
    return {"task": _hydrate_announcement_task(task), "detected_languages": detected, "selected_languages": selected, "constraints": metadata["language_constraints"]}


def extract_announcement_terms(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    project_id = task["project_id"]
    metadata = _merge_announcement_constraint_request(_announcement_task_metadata(task), request)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("detected_languages") or [])
    source_text = _announcement_task_source_text(task)
    if not source_text:
        raise ValueError("announcement source text is empty")
    min_hit = max(1, int(getattr(request, "announcement_min_hit", 1) or 1))
    candidates = _announcement_constraint_rows(project_id, metadata, languages)
    rows = _select_announcement_constraint_rows(source_text, candidates, languages, min_hit=min_hit)

    run = db.insert_run(project_id, kind="announcement_terms", language=languages[0] if languages else "en", metadata={"task_id": task_id, "languages": languages})
    db.update_run(run["id"], status="running")
    output = run_dir(run["id"]) / "announcement_terms"
    output.mkdir(parents=True, exist_ok=True)
    base = _announcement_task_source_stem(task)
    stamp = _today_stamp()
    workbook_path = output / f"{base}_announcement_terms_{stamp}.xlsx"
    manifest_path = output / f"{base}_announcement_terms_manifest_{stamp}.json"
    validation_path = output / f"{base}_announcement_terms_validation_{stamp}.md"
    rows, ai_summary = _apply_announcement_ai_supplement(
        project_id=project_id,
        output_dir=output,
        base_name=base,
        source_text=source_text,
        rows=rows,
        candidates=candidates,
        languages=languages,
        request=request,
        project_name=db.get_project(project_id).get("name", ""),
    )
    _write_announcement_terms_workbook(workbook_path, rows, languages)
    summary = {"terms": len(rows), "languages": languages, "source_chars": len(source_text)}
    if ai_summary:
        summary["ai_supplement"] = {
            key: ai_summary[key]
            for key in ("enabled", "response_artifact_id", "provider", "provider_status", "provider_error", "term_count", "added_to_main", "project_name_translation_missing")
        }
    manifest = {"kind": "announcement_terms", "task_id": task_id, "project_id": project_id, "languages": languages, "summary": summary, "terms": rows}
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    validation_path.write_text(_announcement_terms_validation(summary, rows, languages), encoding="utf-8")
    artifacts = [
        db.add_artifact(project_id, "公告术语表", workbook_path, "announcement_terms_workbook", run_id=run["id"], mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata={"task_id": task_id, "languages": languages}),
        db.add_artifact(project_id, "公告术语 manifest", manifest_path, "announcement_terms_manifest", run_id=run["id"], mime="application/json", metadata={"task_id": task_id, "languages": languages}),
        db.add_artifact(project_id, "公告术语 validation", validation_path, "announcement_terms_validation", run_id=run["id"], mime="text/markdown", metadata={"task_id": task_id, "languages": languages}),
    ]
    if ai_summary:
        packet_artifact = db.add_artifact(project_id, "公告 AI 补充包", Path(ai_summary["packet_path"]), "announcement_ai_supplement_packet", run_id=run["id"], mime="application/json", metadata={"task_id": task_id, "languages": languages})
        response_artifact = None
        if ai_summary.get("response_path"):
            response_artifact = db.add_artifact(project_id, "公告 AI 补充响应", Path(ai_summary["response_path"]), "announcement_ai_supplement_response", run_id=run["id"], mime="application/json", metadata={"task_id": task_id, "languages": languages, "provider": ai_summary.get("provider", "")})
        report_artifact = db.add_artifact(project_id, "公告 AI 补充报告", Path(ai_summary["report_path"]), "announcement_ai_supplement_report", run_id=run["id"], mime="text/markdown", metadata={"task_id": task_id, "languages": languages})
        artifacts.extend([artifact for artifact in (packet_artifact, response_artifact, report_artifact) if artifact])
        summary["ai_supplement"]["packet_artifact_id"] = packet_artifact["id"]
        if response_artifact:
            summary["ai_supplement"]["response_artifact_id"] = response_artifact["id"]
        summary["ai_supplement"]["report_artifact_id"] = report_artifact["id"]
    metadata.update({"languages": languages, "terms": rows, "terms_artifact_id": artifacts[0]["id"], "terms_manifest_artifact_id": artifacts[1]["id"], "terms_validation_artifact_id": artifacts[2]["id"], "terms_summary": summary})
    if ai_summary:
        metadata["ai_supplement"] = summary["ai_supplement"]
    task = db.update_announcement_task(task_id, status="terms_ready", current_step=ANNOUNCEMENT_STEP["lookup"], selected_languages=languages, metadata=metadata)
    for language in languages:
        missing = sum(1 for row in rows if not str((row.get("translations") or {}).get(language) or "").strip())
        db.upsert_announcement_task_language(task_id, project_id, language, status="terms_ready", current_step=ANNOUNCEMENT_STEP["lookup"], metadata={"terms": len(rows), "missing_terms": missing})
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "summary": summary, "task_id": task_id})
    return {"task": _hydrate_announcement_task(task), "run": db.get_run(run["id"]), "summary": summary, "artifacts": artifacts, "manifest": manifest}


def import_announcement_terms(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    project_id = task["project_id"]
    metadata = _announcement_task_metadata(task)
    source_artifact_id = str(getattr(request, "terms_artifact_id", "") or "").strip()
    requested_languages = list(getattr(request, "languages", []) or [])
    raw_terms = list(getattr(request, "terms", []) or [])

    rows: list[dict[str, Any]] = []
    detected_languages: list[str] = []
    if source_artifact_id:
        artifact = db.get_artifact(source_artifact_id)
        if artifact["project_id"] != project_id:
            raise KeyError(source_artifact_id)
        detected_languages = _detect_language_columns(Path(artifact["path"]))
        languages = _normalize_announcement_languages(requested_languages, fallback=detected_languages or task.get("selected_languages") or metadata.get("languages") or [])
        rows = _read_language_table_rows(Path(artifact["path"]), languages)
        metadata["imported_terms_artifact_id"] = source_artifact_id
    else:
        rows = _normalize_announcement_terms_payload(raw_terms)
        detected_languages = _announcement_terms_languages(rows)
        languages = _normalize_announcement_languages(requested_languages, fallback=detected_languages or task.get("selected_languages") or metadata.get("languages") or [])
        rows = _filter_announcement_terms_languages(rows, languages)

    if not rows:
        raise ValueError("announcement terms are empty")
    if not languages:
        languages = _announcement_terms_languages(rows)
    if not languages:
        raise ValueError("announcement terms contain no target languages")

    return _save_announcement_terms(task_id, rows, languages, run_kind="announcement_terms_import")


def lookup_announcement_translations(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    project = db.get_project(task["project_id"])
    metadata = _merge_announcement_constraint_request(_announcement_task_metadata(task), request)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
    terms = list(metadata.get("terms") or [])
    if not terms:
        raise ValueError("extract terms before lookup")
    archive_by_language = _project_archive_by_language(task["project_id"], languages)
    lookup: dict[str, Any] = {}
    run = db.insert_run(task["project_id"], kind="announcement_translation_lookup", language=languages[0] if languages else "en", metadata={"task_id": task_id, "languages": languages})
    output = run_dir(run["id"]) / "announcement_lookup"
    output.mkdir(parents=True, exist_ok=True)
    artifacts: list[dict[str, Any]] = []
    for language in languages:
        rows = []
        for term in terms:
            source = str(term.get("source") or "").strip()
            archive_entry = archive_by_language.get(language, {}).get(_wide_source_key(source))
            table_target = str((term.get("translations") or {}).get(language) or "").strip()
            target = str((archive_entry or {}).get("target") or "").strip() or table_target
            rows.append({**term, "language": language, "target": target, "source_type": "qa_archive" if archive_entry else "language_table" if table_target else "missing"})
        missing = [row for row in rows if not str(row.get("target") or "").strip()]
        prompt_context = _announcement_prompt_context(
            project,
            language,
            [
                {"source": row.get("source", ""), "target": row.get("target", ""), "target_alt": "", "category": "", "note": f"{row.get('hit_count', 0)} hit(s)"}
                for row in rows
                if row.get("target")
            ],
            [],
        )
        context_path = output / f"{_announcement_task_source_stem(task)}_prompt_context_{_visible_language_code(language)}.txt"
        context_path.write_text(prompt_context, encoding="utf-8")
        artifacts.append(db.add_artifact(task["project_id"], f"公告 prompt context ({_visible_language_code(language)})", context_path, "announcement_lookup_prompt_context", run_id=run["id"], mime="text/plain", metadata={"task_id": task_id, "language": language}))
        lookup[language] = {"terms": rows, "missing_terms": [{"source": row.get("source", ""), "id": row.get("id", "")} for row in missing], "prompt_context_artifact_id": artifacts[-1]["id"]}
        db.upsert_announcement_task_language(task_id, task["project_id"], language, status="lookup_ready", current_step=ANNOUNCEMENT_STEP["prepare"], metadata={"terms": len(rows), "missing_terms": len(missing), "prompt_context_artifact_id": artifacts[-1]["id"]})
    summary = {"languages": languages, "terms": len(terms), "missing_terms": sum(len(lookup[language]["missing_terms"]) for language in languages)}
    manifest_path = output / "announcement_lookup_manifest.json"
    manifest = {"kind": "announcement_translation_lookup", "task_id": task_id, "project_id": task["project_id"], "summary": summary, "lookup": lookup}
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    artifacts.append(db.add_artifact(task["project_id"], "公告译文反查 manifest", manifest_path, "announcement_lookup_manifest", run_id=run["id"], mime="application/json", metadata={"task_id": task_id, "languages": languages}))
    metadata.update({"lookup": lookup, "lookup_manifest_artifact_id": artifacts[-1]["id"], "lookup_summary": summary})
    task = db.update_announcement_task(task_id, status="lookup_ready", current_step=ANNOUNCEMENT_STEP["prepare"], selected_languages=languages, metadata=metadata)
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "summary": summary})
    return {"task": _hydrate_announcement_task(task), "run": db.get_run(run["id"]), "summary": summary, "artifacts": artifacts, "manifest": manifest}


def prepare_announcement_translation(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    project = db.get_project(task["project_id"])
    metadata = _announcement_task_metadata(task)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
    if not languages:
        raise ValueError("select at least one target language")
    segments = _announcement_task_segments(task)
    if not segments:
        raise ValueError("announcement source contains no translatable text")
    lookup = metadata.get("lookup") or {}
    run = db.insert_run(task["project_id"], kind="announcement_prepare", language=languages[0], metadata={"task_id": task_id, "languages": languages})
    output = run_dir(run["id"]) / "announcement_prepare"
    output.mkdir(parents=True, exist_ok=True)
    source_stem = _announcement_task_source_stem(task)
    workbook_path = output / f"{source_stem}_announcement_translation_workbook.xlsx"
    manifest_path = output / "announcement_manifest.json"
    _write_announcement_translation_workbook(workbook_path, task, segments, languages, lookup)
    manifest = {"kind": "announcement_prepare", "task_id": task_id, "project_id": task["project_id"], "source_format": task["source_format"], "languages": languages, "segments": segments}
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    artifacts = [
        db.add_artifact(task["project_id"], "公告翻译中转表", workbook_path, "announcement_translation_workbook", run_id=run["id"], mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata={"task_id": task_id, "languages": languages}),
        db.add_artifact(task["project_id"], "公告翻译 manifest", manifest_path, "announcement_docx_manifest", run_id=run["id"], mime="application/json", metadata={"task_id": task_id, "languages": languages}),
    ]
    workpacks: dict[str, str] = {}
    prompts: dict[str, str] = {}
    for language in languages:
        prompt_snapshot = create_prompt_and_harness_snapshots(task["project_id"], run["id"], output / "snapshots" / language, language=language)
        raw_prompt = _announcement_translation_prompt(project, language, prompt_snapshot["prompt"], lookup.get(language, {}))
        prompt = _manage_project_prompt_context(raw_prompt, load_settings())
        lang_code = _visible_language_code(language)
        prompt_path = output / f"{source_stem}_prompt_{lang_code}.txt"
        prompt_path.write_text(prompt, encoding="utf-8")
        prompts[language] = db.add_artifact(task["project_id"], f"公告翻译提示词 ({lang_code})", prompt_path, "prompt_snapshot", run_id=run["id"], mime="text/plain", metadata={"task_id": task_id, "language": language})["id"]
        workpack_path = output / f"{source_stem}_workpack_{lang_code}.jsonl"
        write_jsonl(workpack_path, _announcement_workpack_rows(segments, language, lookup))
        workpack_artifact = db.add_artifact(task["project_id"], f"公告 workpack ({lang_code})", workpack_path, "announcement_workpack", run_id=run["id"], mime="application/jsonl", metadata={"task_id": task_id, "language": language})
        workpacks[language] = workpack_artifact["id"]
        artifacts.append(workpack_artifact)
        db.upsert_announcement_task_language(task_id, task["project_id"], language, status="prepared", current_step=ANNOUNCEMENT_STEP["translate"], metadata={"workpack_artifact_id": workpack_artifact["id"], "prompt_artifact_id": prompts[language], "translation_workbook_artifact_id": artifacts[0]["id"]})
    metadata.update({"segments": segments, "prepare_run_id": run["id"], "translation_workbook_artifact_id": artifacts[0]["id"], "manifest_artifact_id": artifacts[1]["id"], "workpack_artifact_ids": workpacks, "prompt_artifact_ids": prompts})
    task = db.update_announcement_task(task_id, status="prepared", current_step=ANNOUNCEMENT_STEP["translate"], selected_languages=languages, metadata=metadata)
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "summary": {"segments": len(segments), "languages": languages}})
    return {"task": _hydrate_announcement_task(task), "run": db.get_run(run["id"]), "summary": {"segments": len(segments), "languages": languages}, "artifacts": artifacts, "manifest": manifest}


def translate_announcement_task(task_id: str, request: Any, cancel_event: Any | None = None) -> dict[str, Any]:
    return asyncio.run(_translate_announcement_task(task_id, request, cancel_event=cancel_event))


async def _translate_announcement_task(task_id: str, request: Any, cancel_event: Any | None = None) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
    if not metadata.get("translation_workbook_artifact_id") or not metadata.get("workpack_artifact_ids"):
        raise ValueError("prepare announcement translation before AI translation")
    settings = load_settings()
    provider = normalize_provider_name(getattr(request, "provider", None) or settings.get("provider"))
    if getattr(request, "provider", None):
        settings["provider"] = provider
    if getattr(request, "protocol", None):
        settings["protocol"] = request.protocol
    if provider in REAL_PROVIDERS and not settings.get("api_key"):
        for language in languages:
            db.upsert_announcement_task_language(task_id, task["project_id"], language, status="awaiting_ai_response", current_step=ANNOUNCEMENT_STEP["translate"])
        return {"task": _hydrate_announcement_task(db.update_announcement_task(task_id, status="awaiting_ai_response", current_step=ANNOUNCEMENT_STEP["translate"], metadata=metadata)), "summary": {"status": "awaiting_ai_response", "reason": f"{provider} api_key is required; upload AI response instead"}, "artifacts": []}
    run = db.insert_run(task["project_id"], kind="announcement_translate", language=languages[0] if languages else "en", metadata={"task_id": task_id, "languages": languages, "provider": provider})
    metadata = {**metadata, "translate_run_id": run["id"]}
    db.update_announcement_task(task_id, status="running", current_step=ANNOUNCEMENT_STEP["translate"], metadata=metadata)
    output = run_dir(run["id"]) / "announcement_translate"
    output.mkdir(parents=True, exist_ok=True)
    artifacts: list[dict[str, Any]] = []
    response_artifacts: dict[str, str] = {}
    source_stem = _announcement_task_source_stem(task)
    for language in languages:
        lang_code = _visible_language_code(language)
        db.upsert_announcement_task_language(task_id, task["project_id"], language, status="running", current_step=ANNOUNCEMENT_STEP["translate"])
        workpack_artifact = db.get_artifact(metadata["workpack_artifact_ids"][language])
        rows = read_jsonl(Path(workpack_artifact["path"]))
        provider_rows = [{"id": row["id"], "source": row["source"], "term_hits": row.get("term_hits") or []} for row in rows]
        prompt = Path(db.get_artifact(metadata.get("prompt_artifact_ids", {}).get(language, ""))["path"]).read_text(encoding="utf-8") if metadata.get("prompt_artifact_ids", {}).get(language) else ""
        translated = await _translate_rows_with_orchestration(
            run_id=run["id"],
            rows=provider_rows,
            settings=settings,
            project_prompt=prompt,
            work_dir=output / language,
            batch_size=int(getattr(request, "batch_size", None) or settings.get("batch_size") or 90),
            language=language,
            cancel_event=cancel_event,
            confirm_api_budget=bool(getattr(request, "confirm_api_budget", False)),
        )
        if not translated and db.get_run(run["id"]).get("status") == "needs_input":
            task = db.update_announcement_task(task_id, status="needs_input", current_step=ANNOUNCEMENT_STEP["translate"], metadata=metadata)
            return {"task": _hydrate_announcement_task(task), "run": db.get_run(run["id"]), "summary": {"status": "needs_input", "reason": "api_budget_confirmation_required"}, "artifacts": artifacts}
        response_path = output / f"{source_stem}_ai_response_{lang_code}.jsonl"
        write_jsonl(response_path, [{"para_id": item["id"], "translation": item["translation"]} for item in translated])
        artifact = db.add_artifact(task["project_id"], f"公告 AI response ({lang_code})", response_path, "announcement_ai_response", run_id=run["id"], mime="application/jsonl", metadata={"task_id": task_id, "language": language, "provider": provider})
        response_artifacts[language] = artifact["id"]
        artifacts.append(artifact)
    import_result = import_announcement_ai_response(task_id, _SimpleRequest(languages=languages, response_artifacts_by_language=response_artifacts))
    metadata = _announcement_task_metadata(db.get_announcement_task(task_id))
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "response_artifact_ids": response_artifacts})
    return {"task": import_result["task"], "run": db.get_run(run["id"]), "summary": {"status": "translated", "languages": languages}, "artifacts": [*artifacts, *import_result.get("artifacts", [])]}


def import_announcement_ai_response(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
    workbook_artifact = db.get_artifact(metadata.get("translation_workbook_artifact_id", ""))
    response_map = _announcement_response_artifact_map(request, languages)
    if not response_map:
        raise ValueError("response_artifact_ids or response_artifacts_by_language is required")
    imported_languages: list[str] = []
    for language in languages:
        artifact_id = response_map.get(language)
        if not artifact_id:
            continue
        response_artifact = db.get_artifact(artifact_id)
        if response_artifact["project_id"] != task["project_id"]:
            raise KeyError(artifact_id)
        _import_announcement_response_into_workbook(Path(workbook_artifact["path"]), Path(response_artifact["path"]), language)
        imported_languages.append(language)
        db.upsert_announcement_task_language(task_id, task["project_id"], language, status="translated", current_step=ANNOUNCEMENT_STEP["apply"], metadata={"response_artifact_id": artifact_id, "translation_workbook_artifact_id": workbook_artifact["id"]})
    metadata.setdefault("response_artifact_ids", {}).update(response_map)
    task = db.update_announcement_task(task_id, status="translated", current_step=ANNOUNCEMENT_STEP["apply"], metadata=metadata)
    return {"task": _hydrate_announcement_task(task), "summary": {"languages": [language_spec(code).target_header for code in imported_languages], "imported": len(imported_languages)}, "artifacts": [workbook_artifact]}


def apply_announcement_task(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
    workbook_artifact_id = str(getattr(request, "translation_workbook_artifact_id", "") or metadata.get("translation_workbook_artifact_id") or "")
    if not workbook_artifact_id:
        raise ValueError("translation workbook is required")
    workbook_artifact = db.get_artifact(workbook_artifact_id)
    segments = metadata.get("segments") or _announcement_task_segments(task)
    rows = _read_announcement_translation_workbook(Path(workbook_artifact["path"]), languages)
    issues = _validate_announcement_translation_rows(segments, rows, languages)
    run = db.insert_run(task["project_id"], kind="announcement_apply", language=languages[0] if languages else "en", metadata={"task_id": task_id, "languages": languages})
    output = run_dir(run["id"]) / "announcement_apply"
    output.mkdir(parents=True, exist_ok=True)
    qa_path = output / "QA摘要.xlsx"
    if issues:
        _write_announcement_qa_summary(qa_path, issues, [])
        qa_artifact = db.add_artifact(task["project_id"], "公告 QA 摘要", qa_path, "announcement_qa_summary", run_id=run["id"], mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata={"task_id": task_id, "hard_blockers": len(issues)})
        metadata.update({"qa_summary_artifact_id": qa_artifact["id"], "hard_blockers": len(issues), "qa_issues": issues})
        task = db.update_announcement_task(task_id, status="qa_failed", current_step=ANNOUNCEMENT_STEP["apply"], metadata=metadata)
        db.update_run(run["id"], status="failed", metadata={**run.get("metadata", {}), "hard_blockers": len(issues)})
        raise ValueError(f"hard blockers: {len(issues)}")
    output_files = _write_announcement_outputs(task, segments, rows, languages, output / "outputs")
    _write_announcement_qa_summary(qa_path, [], output_files)
    artifacts = [db.add_artifact(task["project_id"], "公告 QA 摘要", qa_path, "announcement_qa_summary", run_id=run["id"], mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata={"task_id": task_id, "hard_blockers": 0})]
    output_artifacts: dict[str, str] = {}
    for language, path in output_files:
        artifact = db.add_artifact(task["project_id"], f"公告成品 ({_visible_language_code(language)})", path, "announcement_output_file", run_id=run["id"], mime=_mime_for_path(path), metadata={"task_id": task_id, "language": language})
        artifacts.append(artifact)
        output_artifacts[language] = artifact["id"]
        db.upsert_announcement_task_language(task_id, task["project_id"], language, status="applied", current_step=ANNOUNCEMENT_STEP["deliver"], metadata={"output_artifact_id": artifact["id"], "qa_summary_artifact_id": artifacts[0]["id"]})
    metadata.update({"qa_summary_artifact_id": artifacts[0]["id"], "output_artifact_ids": output_artifacts, "hard_blockers": 0})
    task = db.update_announcement_task(task_id, status="applied", current_step=ANNOUNCEMENT_STEP["deliver"], metadata=metadata)
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "outputs": output_artifacts})
    return {"task": _hydrate_announcement_task(task), "run": db.get_run(run["id"]), "summary": {"hard_blockers": 0, "languages": languages}, "artifacts": artifacts}


def fix_announcement_hard_blockers(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
    workbook_artifact_id = str(getattr(request, "translation_workbook_artifact_id", "") or metadata.get("translation_workbook_artifact_id") or "")
    if not workbook_artifact_id:
        raise ValueError("translation workbook is missing; prepare/import translations before fixing")
    source_artifact = db.get_artifact(workbook_artifact_id)
    source_path = Path(source_artifact["path"])
    if not source_path.exists():
        raise FileNotFoundError(str(source_path))
    issues = metadata.get("qa_issues") if isinstance(metadata.get("qa_issues"), list) else []
    if not issues:
        segments = metadata.get("segments") or _announcement_task_segments(task)
        rows = _read_announcement_translation_workbook(source_path, languages)
        issues = _validate_announcement_translation_rows(segments, rows, languages)
    if not issues:
        return {"task": _hydrate_announcement_task(task), "summary": {"fixed": 0, "remaining_hard_blockers": 0, "message": "no hard blockers"}, "artifacts": []}

    run = db.insert_run(task["project_id"], kind="announcement_fix", language=languages[0] if languages else "en", metadata={"task_id": task_id, "languages": languages, "source_artifact_id": workbook_artifact_id})
    output = run_dir(run["id"]) / "announcement_fix"
    output.mkdir(parents=True, exist_ok=True)
    fixed_path = output / f"{Path(source_path).stem}_hardblock_fixed.xlsx"
    shutil.copy2(source_path, fixed_path)
    fixed_count = _repair_announcement_translation_workbook(fixed_path, issues, languages)
    fixed_artifact = db.add_artifact(
        task["project_id"],
        "公告 Hard blocker 修复中转表",
        fixed_path,
        "announcement_translation_workbook",
        run_id=run["id"],
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        origin="generated",
        metadata={"task_id": task_id, "languages": languages, "source_artifact_id": workbook_artifact_id, "fixed_count": fixed_count},
    )
    metadata.update({
        "translation_workbook_artifact_id": fixed_artifact["id"],
        "hardblock_fix_artifact_id": fixed_artifact["id"],
        "hardblock_fix_count": fixed_count,
    })
    task = db.update_announcement_task(task_id, status="translated", current_step=ANNOUNCEMENT_STEP["apply"], metadata=metadata)
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "fixed_count": fixed_count, "fixed_artifact_id": fixed_artifact["id"]})

    apply_request = _SimpleRequest(languages=languages, translation_workbook_artifact_id=fixed_artifact["id"])
    try:
        applied = apply_announcement_task(task_id, apply_request)
        applied["summary"] = {**(applied.get("summary") or {}), "fixed": fixed_count, "remaining_hard_blockers": 0}
        applied["artifacts"] = [fixed_artifact, *(applied.get("artifacts") or [])]
        return applied
    except ValueError as exc:
        if "hard blockers" not in str(exc):
            raise
        current_task = _hydrate_announcement_task(db.get_announcement_task(task_id))
        remaining = int((current_task.get("metadata") or {}).get("hard_blockers") or 0)
        return {
            "task": current_task,
            "run": db.get_run(run["id"]),
            "summary": {"fixed": fixed_count, "remaining_hard_blockers": remaining, "message": "some hard blockers still need manual review"},
            "artifacts": [fixed_artifact],
        }


def deliver_announcement_task(task_id: str, request: Any) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    metadata = _announcement_task_metadata(task)
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
    force = bool(getattr(request, "force", False))
    hard_blockers = int(metadata.get("hard_blockers") or 0)
    if hard_blockers > 0 and not force:
        raise ValueError("hard blockers must be fixed before delivery")
    stamp = str(getattr(request, "date_stamp", "") or datetime.now().strftime("%Y%m%d"))
    existing_artifact = _find_existing_announcement_delivery(task, languages, stamp)
    if existing_artifact and not force:
        existing_run = None
        if existing_artifact.get("run_id"):
            try:
                existing_run = db.get_run(str(existing_artifact["run_id"]))
            except KeyError:
                existing_run = None
        existing_languages = _normalize_announcement_languages((existing_artifact.get("metadata") or {}).get("languages") or [], fallback=languages)
        metadata["delivery_artifact_id"] = existing_artifact["id"]
        task = db.update_announcement_task(task_id, status="delivered", current_step=ANNOUNCEMENT_STEP["deliver"], metadata=metadata)
        return {
            "task": _hydrate_announcement_task(task),
            "run": existing_run,
            "summary": {"languages": existing_languages or languages, "delivery_artifact_id": existing_artifact["id"], "reused": True, "date_stamp": stamp},
            "artifacts": [existing_artifact],
        }
    superseded_artifacts = _matching_announcement_delivery_artifacts(task, languages, stamp) if force else []
    run = db.insert_run(task["project_id"], kind="announcement_deliver", language=languages[0] if languages else "en", metadata={"task_id": task_id, "languages": languages})
    output = run_dir(run["id"]) / "announcement_delivery"
    output.mkdir(parents=True, exist_ok=True)
    output_artifact_ids = metadata.get("output_artifact_ids") or {}
    qa_artifact_id = metadata.get("qa_summary_artifact_id")
    if not output_artifact_ids:
        if not force:
            raise ValueError("apply announcement translations before delivery")
        generated = _force_materialize_announcement_outputs_for_delivery(task, metadata, languages, output, run["id"])
        output_artifact_ids = generated["output_artifact_ids"]
        qa_artifact_id = generated["qa_summary_artifact_id"]
        metadata.update(generated)
    zip_path = output / f"{_announcement_delivery_base_name(task)}_announcement_delivery_{stamp}.zip"
    with ZipFile(zip_path, "w") as archive:
        for language in languages:
            artifact_id = output_artifact_ids.get(language)
            if not artifact_id:
                continue
            artifact = db.get_artifact(artifact_id)
            archive.write(artifact["path"], f"{_visible_language_code(language)}/{Path(artifact['path']).name}")
        if qa_artifact_id:
            qa_artifact = db.get_artifact(qa_artifact_id)
            archive.write(qa_artifact["path"], "QA摘要.xlsx")
    artifact_metadata = {"task_id": task_id, "languages": languages, "date_stamp": stamp}
    if hard_blockers > 0 and force:
        artifact_metadata.update({"forced": True, "hard_blockers": hard_blockers})
    artifact = db.add_artifact(task["project_id"], "公告交付总包", zip_path, "announcement_delivery_package", run_id=run["id"], mime="application/zip", metadata=artifact_metadata)
    for old_artifact in superseded_artifacts:
        if old_artifact["id"] == artifact["id"]:
            continue
        db.update_artifact(old_artifact["id"], {"metadata": {**(old_artifact.get("metadata") or {}), "superseded": True, "superseded_by": artifact["id"], "superseded_at": datetime.now().isoformat(timespec="seconds")}})
    metadata["delivery_artifact_id"] = artifact["id"]
    task = db.update_announcement_task(task_id, status="delivered", current_step=ANNOUNCEMENT_STEP["deliver"], metadata=metadata)
    for language in languages:
        db.upsert_announcement_task_language(task_id, task["project_id"], language, status="delivered", current_step=ANNOUNCEMENT_STEP["deliver"])
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "delivery_artifact_id": artifact["id"], "forced": bool(hard_blockers > 0 and force), "hard_blockers": hard_blockers})
    return {"task": _hydrate_announcement_task(task), "run": db.get_run(run["id"]), "summary": {"languages": languages, "delivery_artifact_id": artifact["id"], "date_stamp": stamp, "forced": bool(hard_blockers > 0 and force), "hard_blockers": hard_blockers}, "artifacts": [artifact]}


def _force_materialize_announcement_outputs_for_delivery(task: dict[str, Any], metadata: dict[str, Any], languages: list[str], output: Path, run_id: str) -> dict[str, Any]:
    workbook_artifact_id = str(metadata.get("translation_workbook_artifact_id") or "")
    if not workbook_artifact_id:
        raise ValueError("translation workbook is missing; prepare/import translations before delivery")
    workbook_artifact = db.get_artifact(workbook_artifact_id)
    segments = metadata.get("segments") or _announcement_task_segments(task)
    rows = _read_announcement_translation_workbook(Path(workbook_artifact["path"]), languages)
    issues = metadata.get("qa_issues") if isinstance(metadata.get("qa_issues"), list) else _validate_announcement_translation_rows(segments, rows, languages)
    output_files = _write_announcement_outputs(task, segments, rows, languages, output / "forced_outputs")
    qa_path = output / "QA摘要.xlsx"
    _write_announcement_qa_summary(qa_path, issues, output_files)
    qa_artifact = db.add_artifact(task["project_id"], "公告 QA 摘要", qa_path, "announcement_qa_summary", run_id=run_id, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata={"task_id": task["id"], "hard_blockers": len(issues), "forced_delivery": True})
    output_artifact_ids: dict[str, str] = {}
    for language, path in output_files:
        artifact = db.add_artifact(task["project_id"], f"公告成品 ({_visible_language_code(language)})", path, "announcement_output_file", run_id=run_id, mime=_mime_for_path(path), metadata={"task_id": task["id"], "language": language, "forced_delivery": True})
        output_artifact_ids[language] = artifact["id"]
        db.upsert_announcement_task_language(task["id"], task["project_id"], language, status="applied_with_blockers", current_step=ANNOUNCEMENT_STEP["deliver"], metadata={"output_artifact_id": artifact["id"], "qa_summary_artifact_id": qa_artifact["id"], "forced_delivery": True})
    return {
        "qa_summary_artifact_id": qa_artifact["id"],
        "output_artifact_ids": output_artifact_ids,
        "hard_blockers": len(issues),
        "forced_delivery": True,
    }


def _find_existing_announcement_delivery(task: dict[str, Any], languages: list[str], date_stamp: str) -> dict[str, Any] | None:
    for artifact in _matching_announcement_delivery_artifacts(task, languages, date_stamp):
        if (artifact.get("metadata") or {}).get("superseded"):
            continue
        return artifact
    return None


def _matching_announcement_delivery_artifacts(task: dict[str, Any], languages: list[str], date_stamp: str) -> list[dict[str, Any]]:
    task_id = str(task.get("id") or "")
    expected_languages = set(_normalize_announcement_languages(languages, fallback=[]))
    matches: list[dict[str, Any]] = []
    for artifact in db.list_artifacts(project_id=task["project_id"], role="delivery", include_superseded=True):
        if artifact["kind"] not in {"announcement_delivery_package", "announcement_docx_delivery_package"}:
            continue
        metadata = artifact.get("metadata") or {}
        if str(metadata.get("task_id") or "") != task_id:
            continue
        artifact_languages = set(_normalize_announcement_languages(metadata.get("languages") or [], fallback=[]))
        if expected_languages and artifact_languages and artifact_languages != expected_languages:
            continue
        if _announcement_delivery_artifact_date(artifact) != date_stamp:
            continue
        if not Path(artifact["path"]).exists():
            continue
        matches.append(artifact)
    return matches


def _announcement_delivery_artifact_date(artifact: dict[str, Any]) -> str:
    metadata = artifact.get("metadata") or {}
    if metadata.get("date_stamp"):
        return str(metadata["date_stamp"])
    match = re.search(r"_announcement_delivery_(\d{8})\.zip$", Path(str(artifact.get("path") or "")).name)
    return match.group(1) if match else ""


def generate_announcement_terms_package(project_id: str, request: Any) -> dict[str, Any]:
    text = str(getattr(request, "text", "") or "")
    material_ids = list(getattr(request, "material_artifact_ids", []) or [])
    if not text.strip() and not material_ids:
        raise ValueError("announcement terms requires text or material_artifact_ids")
    if material_ids:
        text = "\n".join([text, *[_read_lookup_material_text(Path(db.get_artifact(artifact_id)["path"])) for artifact_id in material_ids]]).strip()
    if not text.strip():
        raise ValueError("announcement text is empty")
    languages = _normalize_announcement_languages(getattr(request, "languages", []) or [], fallback=["en"])
    metadata = {
        "language_table_artifact_ids": list(getattr(request, "language_table_artifact_ids", []) or []),
        "constraint_artifact_ids": [],
        "include_project_archive": False,
    }
    candidates = _announcement_constraint_rows(project_id, metadata, languages)
    rows = _select_announcement_constraint_rows(_compact_lookup_text(text), candidates, languages, min_hit=max(1, int(getattr(request, "announcement_min_hit", 1) or 1)))
    run = db.insert_run(project_id, kind="announcement_terms", language=languages[0] if languages else "en", metadata={"languages": languages})
    output = run_dir(run["id"]) / "announcement_terms"
    output.mkdir(parents=True, exist_ok=True)
    base = "announcement"
    if material_ids:
        try:
            base = _artifact_source_stem(db.get_artifact(material_ids[0]))
        except KeyError:
            base = "announcement"
    stamp = _today_stamp()
    workbook_path = output / f"{base}_announcement_terms_{stamp}.xlsx"
    manifest_path = output / f"{base}_announcement_terms_manifest_{stamp}.json"
    validation_path = output / f"{base}_announcement_terms_validation_{stamp}.md"
    rows, ai_summary = _apply_announcement_ai_supplement(
        project_id=project_id,
        output_dir=output,
        base_name=base,
        source_text=_compact_lookup_text(text),
        rows=rows,
        candidates=candidates,
        languages=languages,
        request=request,
        project_name=db.get_project(project_id).get("name", ""),
    )
    _write_announcement_terms_workbook(workbook_path, rows, languages)
    summary = {"terms": len(rows), "languages": languages, "source_chars": len(text)}
    if ai_summary:
        summary["ai_supplement"] = {
            key: ai_summary[key]
            for key in ("enabled", "response_artifact_id", "provider", "provider_status", "provider_error", "term_count", "added_to_main", "project_name_translation_missing")
        }
    manifest = {"kind": "announcement_terms", "project_id": project_id, "languages": languages, "summary": summary, "terms": rows}
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    validation_path.write_text(_announcement_terms_validation(summary, rows, languages), encoding="utf-8")
    artifacts = [
        db.add_artifact(project_id, "公告术语表", workbook_path, "announcement_terms_workbook", run_id=run["id"], mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata={"languages": languages}),
        db.add_artifact(project_id, "公告术语 validation", validation_path, "announcement_terms_validation", run_id=run["id"], mime="text/markdown", metadata={"languages": languages}),
        db.add_artifact(project_id, "公告术语 manifest", manifest_path, "announcement_terms_manifest", run_id=run["id"], mime="application/json", metadata={"languages": languages}),
    ]
    if ai_summary:
        packet_artifact = db.add_artifact(project_id, "公告 AI 补充包", Path(ai_summary["packet_path"]), "announcement_ai_supplement_packet", run_id=run["id"], mime="application/json", metadata={"languages": languages})
        response_artifact = None
        if ai_summary.get("response_path"):
            response_artifact = db.add_artifact(project_id, "公告 AI 补充响应", Path(ai_summary["response_path"]), "announcement_ai_supplement_response", run_id=run["id"], mime="application/json", metadata={"languages": languages, "provider": ai_summary.get("provider", "")})
        report_artifact = db.add_artifact(project_id, "公告 AI 补充报告", Path(ai_summary["report_path"]), "announcement_ai_supplement_report", run_id=run["id"], mime="text/markdown", metadata={"languages": languages})
        artifacts.extend([artifact for artifact in (packet_artifact, response_artifact, report_artifact) if artifact])
        summary["ai_supplement"]["packet_artifact_id"] = packet_artifact["id"]
        if response_artifact:
            summary["ai_supplement"]["response_artifact_id"] = response_artifact["id"]
        summary["ai_supplement"]["report_artifact_id"] = report_artifact["id"]
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "summary": summary})
    return {"run": db.get_run(run["id"]), "summary": summary, "artifacts": artifacts, "manifest": manifest}


def legacy_prepare_announcement_docx(project_id: str, request: Any) -> dict[str, Any]:
    if not list(getattr(request, "source_artifact_ids", []) or []):
        raise ValueError("source_artifact_ids is required")
    source_artifact_id = list(getattr(request, "source_artifact_ids", []))[0]
    create_request = _SimpleRequest(
        source_artifact_id=source_artifact_id,
        title="公告 DOCX",
        languages=list(getattr(request, "languages", []) or []),
        constraint_artifact_ids=[getattr(request, "terms_artifact_id")],
        language_table_artifact_ids=[getattr(request, "terms_artifact_id")],
        include_project_archive=False,
    )
    task = create_announcement_task(project_id, create_request)
    extract_announcement_terms(task["id"], _SimpleRequest(languages=create_request.languages, language_table_artifact_ids=create_request.language_table_artifact_ids, include_project_archive=False))
    lookup_announcement_translations(task["id"], _SimpleRequest(languages=create_request.languages, include_project_archive=False))
    prepared = prepare_announcement_translation(task["id"], _SimpleRequest(languages=create_request.languages))
    run = prepared["run"]
    db.update_run(run["id"], metadata={**run.get("metadata", {}), "task_id": task["id"], "legacy_prepare": True})
    artifacts = []
    for artifact in prepared["artifacts"]:
        if artifact["kind"] == "announcement_translation_workbook":
            artifact = db.update_artifact(artifact["id"], {"label": "Announcement DOCX translation workbook", "metadata": {**artifact.get("metadata", {}), "legacy_kind": "announcement_docx_translation_workbook"}})
            artifact["kind"] = "announcement_docx_translation_workbook"
        elif artifact["kind"] == "announcement_workpack":
            artifact = db.update_artifact(artifact["id"], {"label": artifact["label"].replace("公告", "Announcement DOCX"), "metadata": {**artifact.get("metadata", {}), "legacy_kind": "announcement_docx_workpack"}})
            artifact["kind"] = "announcement_docx_workpack"
        elif artifact["kind"] == "announcement_docx_manifest":
            pass
        artifacts.append(artifact)
    return {**prepared, "run": db.get_run(run["id"]), "artifacts": artifacts}


def legacy_import_announcement_docx_ai(project_id: str, request: Any) -> dict[str, Any]:
    run = db.get_run(getattr(request, "prepare_run_id"))
    task_id = str((run.get("metadata") or {}).get("task_id") or "")
    if not task_id:
        raise KeyError("task_id")
    result = import_announcement_ai_response(task_id, request)
    return {"summary": result["summary"], "task": result["task"], "artifacts": result["artifacts"]}


def legacy_apply_announcement_docx(project_id: str, request: Any) -> dict[str, Any]:
    run = db.get_run(getattr(request, "prepare_run_id"))
    task_id = str((run.get("metadata") or {}).get("task_id") or "")
    result = apply_announcement_task(task_id, request)
    legacy_artifacts = []
    for artifact in result["artifacts"]:
        if artifact["kind"] == "announcement_output_file" and str(artifact.get("path", "")).lower().endswith(".docx"):
            artifact["kind"] = "announcement_docx_output_docx"
        elif artifact["kind"] == "announcement_qa_summary":
            artifact["kind"] = "announcement_docx_qa_summary"
        legacy_artifacts.append(artifact)
    result["artifacts"] = legacy_artifacts
    return result


def legacy_deliver_announcement_docx(project_id: str, request: Any) -> dict[str, Any]:
    run = db.get_run(getattr(request, "prepare_run_id"))
    task_id = str((run.get("metadata") or {}).get("task_id") or "")
    result = deliver_announcement_task(task_id, request)
    for artifact in result["artifacts"]:
        if artifact["kind"] == "announcement_delivery_package":
            artifact["kind"] = "announcement_docx_delivery_package"
    return result


class _SimpleRequest:
    def __init__(self, **kwargs: Any) -> None:
        self.__dict__.update(kwargs)


def _hydrate_announcement_task(task: dict[str, Any]) -> dict[str, Any]:
    metadata = task.get("metadata") or {}
    translate_run_id = str(metadata.get("translate_run_id") or "")
    if translate_run_id:
        try:
            translate_run = db.get_run(translate_run_id)
            metadata = {
                **metadata,
                "translate_run_status": translate_run.get("status"),
                "translation_progress": (translate_run.get("metadata") or {}).get("translation_progress"),
                "api_budget_estimate": (translate_run.get("metadata") or {}).get("api_budget_estimate"),
            }
            task["metadata"] = metadata
        except KeyError:
            pass
    artifact_ids: set[str] = set()
    for key in (
        "terms_artifact_id",
        "terms_manifest_artifact_id",
        "terms_validation_artifact_id",
        "translation_workbook_artifact_id",
        "manifest_artifact_id",
        "qa_summary_artifact_id",
        "delivery_artifact_id",
    ):
        if metadata.get(key):
            artifact_ids.add(str(metadata[key]))
    for mapping_key in ("workpack_artifact_ids", "prompt_artifact_ids", "response_artifact_ids", "output_artifact_ids"):
        value = metadata.get(mapping_key)
        if isinstance(value, dict):
            artifact_ids.update(str(item) for item in value.values() if item)
    ai_supplement = metadata.get("ai_supplement")
    if isinstance(ai_supplement, dict):
        for key in ("packet_artifact_id", "report_artifact_id"):
            if ai_supplement.get(key):
                artifact_ids.add(str(ai_supplement[key]))
    artifacts: list[dict[str, Any]] = []
    for artifact_id in sorted(artifact_ids):
        try:
            artifacts.append(db.get_artifact(artifact_id))
        except KeyError:
            continue
    task["artifacts"] = artifacts
    return task


def _announcement_task_metadata(task: dict[str, Any]) -> dict[str, Any]:
    return dict(task.get("metadata") or {})


def _merge_announcement_constraint_request(metadata: dict[str, Any], request: Any) -> dict[str, Any]:
    merged = dict(metadata)
    for key in ("language_table_artifact_ids", "constraint_artifact_ids"):
        value = list(getattr(request, key, []) or [])
        if value:
            merged[key] = value
    if hasattr(request, "include_project_archive"):
        merged["include_project_archive"] = bool(getattr(request, "include_project_archive", True))
    return merged


def _create_inline_announcement_source(project_id: str, text: str, title: str) -> dict[str, Any]:
    output = project_dir(project_id) / "announcements" / "inline"
    output.mkdir(parents=True, exist_ok=True)
    path = output / f"{_safe_file_stem(title or 'announcement')}_{hashlib.sha1(text.encode('utf-8')).hexdigest()[:8]}.txt"
    path.write_text(text, encoding="utf-8")
    return db.add_artifact(project_id, "Inline announcement source", path, "asset", mime="text/plain", origin="uploaded")


def _announcement_source_format(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return "docx"
    if suffix in {".txt", ".md", ".markdown"}:
        return "txt"
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return "xlsx"
    return suffix.lstrip(".")


def _announcement_source_manifest(artifact: dict[str, Any]) -> dict[str, Any]:
    path = Path(artifact["path"])
    return {"artifact_id": artifact["id"], "label": artifact.get("label", ""), "path": str(path), "format": _announcement_source_format(path), "sha256": _file_sha256(path) if path.exists() else ""}


def _announcement_task_source_text(task: dict[str, Any]) -> str:
    artifact = db.get_artifact(task["source_artifact_id"])
    return _compact_lookup_text(_read_lookup_material_text(Path(artifact["path"])))


def _announcement_task_segments(task: dict[str, Any]) -> list[dict[str, Any]]:
    metadata = _announcement_task_metadata(task)
    if metadata.get("segments"):
        return list(metadata["segments"])
    artifact = db.get_artifact(task["source_artifact_id"])
    path = Path(artifact["path"])
    fmt = _announcement_source_format(path)
    if fmt == "docx":
        return _docx_announcement_segments(path)
    if fmt == "xlsx":
        return _xlsx_announcement_segments(path)
    return _txt_announcement_segments(path)


def _docx_announcement_segments(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    rows = []
    for index, paragraph in enumerate(doc.paragraphs):
        source = paragraph.text
        if not source.strip():
            continue
        rows.append({"id": _segment_id(path.name, index, source), "source_file": path.name, "index": index, "kind": "paragraph", "source": source, "style": paragraph.style.name if paragraph.style else ""})
    return rows


def _txt_announcement_segments(path: Path) -> list[dict[str, Any]]:
    text = path.read_text(encoding="utf-8-sig")
    rows = []
    for index, line in enumerate(text.splitlines()):
        if not line.strip():
            continue
        rows.append({"id": _segment_id(path.name, index, line), "source_file": path.name, "index": index, "kind": "line", "source": line})
    return rows


def _is_quick_text_path(path: Path) -> bool:
    return path.suffix.lower() in {".txt", ".md", ".markdown"}


def _quick_text_translation_rows(path: Path) -> list[dict[str, Any]]:
    return [
        {"id": segment["id"], "source": segment["source"], "index": segment["index"], "source_file": segment["source_file"]}
        for segment in _txt_announcement_segments(path)
    ]


def _write_quick_text_output(source_path: Path, translated_rows: list[dict[str, Any]], language: str, output_dir: Path) -> Path:
    translations = {str(row.get("id")): str(row.get("translation") or "") for row in translated_rows}
    segments = _txt_announcement_segments(source_path)
    by_index = {int(segment["index"]): translations.get(str(segment["id"]), segment["source"]) for segment in segments}
    raw_lines = source_path.read_text(encoding="utf-8-sig").splitlines(keepends=True)
    if not raw_lines and source_path.read_text(encoding="utf-8-sig").strip():
        raw_lines = [source_path.read_text(encoding="utf-8-sig")]
    output_parts: list[str] = []
    for index, raw in enumerate(raw_lines):
        if raw.endswith("\r\n"):
            content, newline = raw[:-2], "\r\n"
        elif raw.endswith("\n"):
            content, newline = raw[:-1], "\n"
        elif raw.endswith("\r"):
            content, newline = raw[:-1], "\r"
        else:
            content, newline = raw, ""
        output_parts.append((by_index.get(index, content) if content.strip() else content) + newline)
    suffix = source_path.suffix.lower() if source_path.suffix.lower() in {".txt", ".md", ".markdown"} else ".txt"
    output_path = output_dir / f"{_safe_source_stem(source_path.name)}_{_visible_language_code(language)}{suffix}"
    output_path.write_text("".join(output_parts), encoding="utf-8")
    return output_path


def _xlsx_announcement_segments(path: Path) -> list[dict[str, Any]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    rows: list[dict[str, Any]] = []
    try:
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    value = str(cell.value or "").strip()
                    if not value or not _CJK_RE.search(value):
                        continue
                    key = f"{ws.title}!{cell.coordinate}"
                    rows.append({"id": _segment_id(path.name, len(rows), f"{key}:{value}"), "source_file": path.name, "index": len(rows), "kind": "cell", "sheet": ws.title, "coordinate": cell.coordinate, "source": value})
    finally:
        wb.close()
    return rows


def _segment_id(source_file: str, index: int, source: str) -> str:
    digest_source = "\0".join([source_file, str(index), source])
    digest = hashlib.sha1(digest_source.encode("utf-8")).hexdigest()[:10]
    return f"{Path(source_file).stem}:{index:04d}:{digest}"


def _detect_announcement_constraint_languages(project_id: str, metadata: dict[str, Any]) -> list[str]:
    found: set[str] = set()
    for artifact_id in [*metadata.get("language_table_artifact_ids", []), *metadata.get("constraint_artifact_ids", [])]:
        try:
            artifact = db.get_artifact(artifact_id)
        except KeyError:
            continue
        if artifact["project_id"] != project_id:
            continue
        if _is_generated_announcement_terms_artifact(artifact):
            continue
        found.update(_detect_language_columns(Path(artifact["path"])))
    if metadata.get("include_project_archive", True):
        for language in ANNOUNCEMENT_LANGUAGE_ORDER:
            if db.list_translation_entries(project_id, language=language):
                found.add(language)
    return [language for language in ANNOUNCEMENT_LANGUAGE_ORDER if language in found]


def _announcement_language_constraint_summary(project_id: str, metadata: dict[str, Any], languages: list[str]) -> dict[str, Any]:
    table_rows = _language_table_rows_from_artifacts(project_id, [*metadata.get("language_table_artifact_ids", []), *metadata.get("constraint_artifact_ids", [])], languages)
    archive_counts = {language: len(db.list_translation_entries(project_id, language=language)) for language in languages}
    table_counts = {language: sum(1 for row in table_rows if row.get("translations", {}).get(language)) for language in languages}
    return {language: {"language_table": table_counts.get(language, 0), "qa_archive": archive_counts.get(language, 0)} for language in languages}


def _normalize_announcement_languages(raw: Any, fallback: list[str] | tuple[str, ...] = ()) -> list[str]:
    values = list(raw or []) or list(fallback or [])
    normalized: list[str] = []
    for value in values:
        try:
            code = require_supported_language(value)
        except ValueError:
            continue
        if code not in normalized:
            normalized.append(code)
    return [code for code in ANNOUNCEMENT_LANGUAGE_ORDER if code in normalized]


def _detect_language_columns(path: Path) -> list[str]:
    if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return []
    wb = load_workbook(path, read_only=True, data_only=True)
    found: set[str] = set()
    try:
        for ws in wb.worksheets:
            headers = [str(cell.value or "").strip() for cell in next(ws.iter_rows(min_row=1, max_row=1))]
            normalized = _header_index(headers, prefer_last=True)
            reserved_indices = set(_reserved_language_table_indices(headers))
            for code in ANNOUNCEMENT_LANGUAGE_ORDER:
                index = _language_column_index(normalized, code)
                if index is not None and index not in reserved_indices:
                    found.add(code)
    finally:
        wb.close()
    return [code for code in ANNOUNCEMENT_LANGUAGE_ORDER if code in found]


def _language_column_index(normalized_headers: dict[str, int], language: str) -> int | None:
    aliases = [alias.lower() for alias in target_aliases(language)]
    for alias in aliases:
        if alias in normalized_headers:
            return normalized_headers[alias]
    return None


def _header_index(headers: list[str], *, prefer_last: bool = False) -> dict[str, int]:
    normalized: dict[str, int] = {}
    for index, header in enumerate(headers):
        key = str(header or "").strip().lower()
        if not key:
            continue
        if prefer_last or key not in normalized:
            normalized[key] = index
    return normalized


def _reserved_language_table_indices(headers: list[str]) -> list[int]:
    normalized_first = _header_index(headers)
    indices = [
        _column_by_alias(normalized_first, ["id", "key", "编号", "序号"]),
        _column_by_alias(normalized_first, ["cn", "zh", "source", "chinese", "中文", "原文", "简体中文", "term", "术语"]),
    ]
    return [index for index in indices if index is not None]


def _column_by_alias(normalized_headers: dict[str, int], aliases: list[str]) -> int | None:
    for alias in aliases:
        key = alias.lower()
        if key in normalized_headers:
            return normalized_headers[key]
    return None


def _announcement_constraint_rows(project_id: str, metadata: dict[str, Any], languages: list[str]) -> list[dict[str, Any]]:
    rows = _language_table_rows_from_artifacts(project_id, [*metadata.get("language_table_artifact_ids", []), *metadata.get("constraint_artifact_ids", [])], languages)
    if metadata.get("include_project_archive", True):
        by_source: dict[str, dict[str, Any]] = {_wide_source_key(row.get("source")): row for row in rows if _wide_source_key(row.get("source"))}
        for language in languages:
            for entry in db.list_translation_entries(project_id, language=language):
                source = str(entry.get("source") or "").strip()
                if not source:
                    continue
                key = _wide_source_key(source)
                row = by_source.setdefault(key, {"id": entry.get("entry_key") or entry.get("id"), "source": source, "translations": {}, "sources": []})
                if str(entry.get("target") or "").strip():
                    current = row["translations"].get(language)
                    if not current or _rank_translation_lookup_source(entry.get("source_type", "")) < _rank_translation_lookup_source(row.get("source_type", "")):
                        row["translations"][language] = entry.get("target", "")
                        row["source_type"] = entry.get("source_type", "")
                row.setdefault("sources", []).append({"type": "qa_archive", "language": language, "entry_id": entry.get("id")})
        rows = list(by_source.values())
    return rows


def _language_table_rows_from_artifacts(project_id: str, artifact_ids: list[str], languages: list[str]) -> list[dict[str, Any]]:
    by_source: dict[str, dict[str, Any]] = {}
    scanned_artifacts = 0
    for artifact_id in artifact_ids:
        if not artifact_id:
            continue
        artifact = db.get_artifact(artifact_id)
        if artifact["project_id"] != project_id:
            raise KeyError(artifact_id)
        if _is_generated_announcement_terms_artifact(artifact):
            continue
        path = Path(artifact["path"])
        scanned_artifacts += 1
        if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
            raise ValueError(f"约束文件格式不正确：{path.name} 不是 XLSX 语言表。请上传完整语言表/术语交付表，不要把公告原文或 TXT 放在约束来源。")
        artifact_rows = _read_language_table_rows(path, languages)
        if not artifact_rows:
            visible = " / ".join(_visible_language_code(language) for language in languages) or "目标语言"
            raise ValueError(f"约束文件未识别到可反查词条：{path.name}。请检查表头是否包含 ID、CN/中文/原文，以及 {visible} 目标语言列。")
        for row in artifact_rows:
            key = _wide_source_key(row.get("source"))
            if not key:
                continue
            existing = by_source.setdefault(key, {"id": row.get("id", ""), "source": row.get("source", ""), "translations": {}, "sources": []})
            if row.get("id") and not existing.get("id"):
                existing["id"] = row["id"]
            for language, target in (row.get("translations") or {}).items():
                if target and not existing["translations"].get(language):
                    existing["translations"][language] = target
            existing["sources"].append({"type": "language_table", "artifact_id": artifact_id})
    if scanned_artifacts and not by_source:
        visible = " / ".join(_visible_language_code(language) for language in languages) or "目标语言"
        raise ValueError(f"约束文件未识别到可反查词条。请确认表头包含 ID、CN/中文/原文，以及 {visible} 目标语言列。")
    return list(by_source.values())


def _is_generated_announcement_terms_artifact(artifact: dict[str, Any]) -> bool:
    if artifact.get("kind") == "announcement_terms_workbook":
        return True
    text = " ".join(
        str(part or "")
        for part in (
            artifact.get("label"),
            artifact.get("path"),
            (artifact.get("metadata") or {}).get("original_filename"),
        )
    ).lower()
    if "announcement_terms" not in text and "公告术语" not in text:
        return False
    return _workbook_looks_like_announcement_terms(Path(str(artifact.get("path") or "")))


def _workbook_looks_like_announcement_terms(path: Path) -> bool:
    if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"} or not path.exists():
        return True
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        for ws in wb.worksheets:
            try:
                headers = [str(value or "").strip().lower() for value in next(ws.iter_rows(min_row=1, max_row=1, values_only=True))]
            except StopIteration:
                continue
            has_source = any(header in {"cn", "zh", "source", "chinese", "中文", "原文", "术语", "term"} for header in headers)
            has_hit_count = any(header in {"hit count", "hit_count", "hits", "命中次数"} or "命中" in header for header in headers)
            has_origin = any(header in {"来源", "origin", "source"} for header in headers)
            if has_source and has_hit_count and has_origin:
                return True
    finally:
        wb.close()
    return False


def _read_language_table_rows(path: Path, languages: list[str]) -> list[dict[str, Any]]:
    if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return []
    wb = load_workbook(path, read_only=True, data_only=True)
    rows: list[dict[str, Any]] = []
    try:
        for ws in wb.worksheets:
            iterator = ws.iter_rows(values_only=True)
            try:
                headers = [str(value or "").strip() for value in next(iterator)]
            except StopIteration:
                continue
            normalized_first = _header_index(headers)
            normalized_last = _header_index(headers, prefer_last=True)
            id_idx = _column_by_alias(normalized_first, ["id", "key", "编号", "序号"])
            source_idx = _column_by_alias(normalized_first, ["cn", "zh", "source", "chinese", "中文", "原文", "简体中文", "term", "术语"])
            if source_idx is None:
                continue
            reserved_indices = set(index for index in (id_idx, source_idx) if index is not None)
            lang_indices = {
                language: index if index not in reserved_indices else None
                for language in languages
                for index in [_language_column_index(normalized_last, language)]
            }
            if not any(index is not None for index in lang_indices.values()):
                continue
            for values in iterator:
                source = str(values[source_idx] or "").strip() if source_idx < len(values) else ""
                if not source:
                    continue
                translations = {}
                for language, index in lang_indices.items():
                    if index is not None and index < len(values):
                        value = str(values[index] or "").strip()
                        if value:
                            translations[language] = value
                rows.append({"id": str(values[id_idx] or "").strip() if id_idx is not None and id_idx < len(values) else "", "source": source, "translations": translations, "sheet": ws.title})
    finally:
        wb.close()
    return rows


def _select_announcement_constraint_rows(text: str, candidates: list[dict[str, Any]], languages: list[str], *, min_hit: int) -> list[dict[str, Any]]:
    selected: list[dict[str, Any]] = []
    for row in candidates:
        source = str(row.get("source") or "").strip()
        if len(source) < 2:
            continue
        hit_count, first_position = _count_lookup_hits(text, source)
        if hit_count < min_hit:
            continue
        selected.append({"id": row.get("id", ""), "source": source, "translations": {language: (row.get("translations") or {}).get(language, "") for language in languages}, "hit_count": hit_count, "first_position": first_position})
    selected.sort(key=lambda row: (int(row.get("first_position") or 0), -len(str(row.get("source") or "")), str(row.get("source") or "")))
    return _suppress_overlapping_lookup_hits(selected)


def _glossary_extractor_module() -> Any:
    global _GLOSSARY_EXTRACTOR_MODULE
    if _GLOSSARY_EXTRACTOR_MODULE is not None:
        return _GLOSSARY_EXTRACTOR_MODULE
    script_path = GLOSSARY_ROOT / "scripts" / "extract_glossary.py"
    spec = importlib.util.spec_from_file_location("lws_embedded_extract_glossary", script_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load glossary extractor: {script_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules.setdefault("lws_embedded_extract_glossary", module)
    spec.loader.exec_module(module)
    _GLOSSARY_EXTRACTOR_MODULE = module
    return module


def _announcement_ai_headers(languages: list[str]) -> list[str]:
    return ["ID", "CN", *[visible_language_code(language) for language in languages]]


def _announcement_term_to_ai_row(row: dict[str, Any], languages: list[str]) -> dict[str, object]:
    output: dict[str, object] = {
        "ID": str(row.get("id") or row.get("ID") or "").strip(),
        "CN": str(row.get("source") or row.get("CN") or "").strip(),
    }
    translations = row.get("translations") if isinstance(row.get("translations"), dict) else {}
    for language in languages:
        output[visible_language_code(language)] = str((translations or {}).get(language) or "").strip()
    return output


def _normalize_ai_supplement_response(response: dict[str, Any], languages: list[str]) -> dict[str, Any]:
    terms = response.get("supplement_terms")
    if not isinstance(terms, list):
        return {"supplement_terms": []}
    normalized_terms: list[dict[str, Any]] = []
    for term in terms:
        if not isinstance(term, dict):
            continue
        item = dict(term)
        translations = item.get("translations")
        if isinstance(translations, dict):
            normalized_translations = {str(key): value for key, value in translations.items()}
            lower_lookup = {str(key).strip().lower(): value for key, value in translations.items()}
            for language in languages:
                header = visible_language_code(language)
                if str(normalized_translations.get(header) or "").strip():
                    continue
                aliases = {language, header.lower(), *[alias.lower() for alias in target_aliases(language)]}
                for alias in aliases:
                    if alias in lower_lookup and str(lower_lookup[alias] or "").strip():
                        normalized_translations[header] = lower_lookup[alias]
                        break
            item["translations"] = normalized_translations
        normalized_terms.append(item)
    return {**response, "supplement_terms": normalized_terms}


def _announcement_ai_rows_to_terms(ai_rows: list[dict[str, object]], original_rows: list[dict[str, Any]], source_text: str, languages: list[str]) -> list[dict[str, Any]]:
    original_by_source = {_wide_source_key(row.get("source")): row for row in original_rows if _wide_source_key(row.get("source"))}
    output: list[dict[str, Any]] = []
    seen: set[str] = set()
    for ai_row in ai_rows:
        source = str(ai_row.get("CN") or "").strip()
        key = _wide_source_key(source)
        if not key or key in seen:
            continue
        seen.add(key)
        if key in original_by_source:
            output.append(original_by_source[key])
            continue
        hit_count, first_position = _count_lookup_hits(source_text, source)
        translations = {
            language: str(ai_row.get(visible_language_code(language)) or "").strip()
            for language in languages
        }
        output.append(
            {
                "id": str(ai_row.get("ID") or "").strip(),
                "source": source,
                "translations": translations,
                "hit_count": hit_count,
                "first_position": first_position,
                "source_type": "ai_supplement",
            }
        )
    return output


def _read_ai_supplement_response_artifact(project_id: str, artifact_id: str | None) -> dict[str, Any] | None:
    if not artifact_id:
        return None
    artifact = db.get_artifact(artifact_id)
    if artifact["project_id"] != project_id:
        raise KeyError(artifact_id)
    payload = json.loads(Path(artifact["path"]).read_text(encoding="utf-8-sig"))
    if not isinstance(payload, dict):
        raise ValueError("AI supplement response must be a JSON object")
    return payload


def _ai_supplement_provider_prompt(packet: dict[str, Any]) -> str:
    return (
        "You are auditing a game announcement glossary extraction result.\n"
        "Use only the supplied packet. Do not invent translations without evidence_rows.\n"
        "Return strict JSON only, no markdown fences, matching packet.response_schema.\n"
        "Only add terms that appear verbatim in announcement_text and are backed by evidence_ids.\n\n"
        f"Packet JSON:\n{json.dumps(packet, ensure_ascii=False, indent=2)}"
    )


def _call_ai_supplement_provider(settings: dict[str, Any], packet: dict[str, Any]) -> dict[str, Any]:
    text = _call_semantic_provider(settings, _ai_supplement_provider_prompt(packet))
    payload = _parse_semantic_qa_payload(text)
    if not isinstance(payload, dict):
        raise ValueError("AI supplement provider must return a JSON object")
    return payload


def _apply_announcement_ai_supplement(
    *,
    project_id: str,
    output_dir: Path,
    base_name: str,
    source_text: str,
    rows: list[dict[str, Any]],
    candidates: list[dict[str, Any]],
    languages: list[str],
    request: Any,
    project_name: str,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    if not bool(getattr(request, "ai_supplement", False)):
        return rows, {}
    extractor = _glossary_extractor_module()
    headers = _announcement_ai_headers(languages)
    packet_path = output_dir / f"{base_name}_ai_supplement_packet_{_today_stamp()}.json"
    report_path = output_dir / f"{base_name}_ai_supplement_report_{_today_stamp()}.md"
    matched_rows = [_announcement_term_to_ai_row(row, languages) for row in rows]
    candidate_rows = [_announcement_term_to_ai_row(row, languages) for row in candidates]
    packet = extractor.build_ai_supplement_packet(
        announcement_text=source_text,
        matched_rows=matched_rows,
        candidate_rows=candidate_rows,
        headers=headers,
        project_name=project_name,
    )
    packet_path.write_text(json.dumps(packet, ensure_ascii=False, indent=2), encoding="utf-8")
    response_artifact_id = str(getattr(request, "ai_supplement_response_artifact_id", "") or "")
    response_path: Path | None = None
    provider = ""
    provider_status = "not_configured"
    provider_error = ""
    response = _read_ai_supplement_response_artifact(project_id, response_artifact_id)
    if response is not None:
        provider = "uploaded"
        provider_status = "uploaded_response"
    elif not packet.get("evidence_rows"):
        provider_status = "no_evidence"
        response = {"supplement_terms": []}
    else:
        settings = load_settings()
        configured_provider = normalize_provider_name(settings.get("provider"))
        if configured_provider in REAL_PROVIDERS and settings.get("api_key"):
            provider = configured_provider
            try:
                response = _call_ai_supplement_provider(settings, packet)
                provider_status = "provider_response"
                response_path = output_dir / f"{base_name}_ai_supplement_response_{_today_stamp()}.json"
                response_path.write_text(json.dumps(response, ensure_ascii=False, indent=2), encoding="utf-8")
            except Exception as exc:
                provider_status = "provider_error"
                provider_error = user_facing_error(exc)
                response = {"supplement_terms": []}
        else:
            response = {"supplement_terms": []}
    response = _normalize_ai_supplement_response(response, languages)
    merged_ai_rows, report = extractor.apply_ai_supplement_response(
        announcement_rows=matched_rows,
        headers=headers,
        announcement_text=source_text,
        packet=packet,
        response=response,
        project_name=project_name,
    )
    report["provider"] = provider or provider_status
    if provider_error:
        report["provider_error"] = provider_error
    report_markdown = extractor.build_ai_supplement_report_markdown(
        report=report,
        packet_path=packet_path,
        response_path=Path(db.get_artifact(response_artifact_id)["path"]) if response_artifact_id else response_path,
        output_path=report_path,
    )
    report_path.write_text(report_markdown, encoding="utf-8")
    merged_rows = _announcement_ai_rows_to_terms(merged_ai_rows, rows, source_text, languages)
    report_terms = report.get("terms") if isinstance(report.get("terms"), list) else []
    ai_summary = {
        "enabled": True,
        "packet_path": str(packet_path),
        "report_path": str(report_path),
        "response_path": str(response_path or ""),
        "response_artifact_id": response_artifact_id,
        "provider": provider,
        "provider_status": provider_status,
        "provider_error": provider_error,
        "term_count": len(report_terms),
        "added_to_main": sum(1 for term in report_terms if isinstance(term, dict) and term.get("status") == "added_to_main"),
        "project_name_translation_missing": bool(report.get("project_name_translation_missing")),
        "report": report,
    }
    return merged_rows, ai_summary


def _save_announcement_terms(task_id: str, rows: list[dict[str, Any]], languages: list[str], *, run_kind: str) -> dict[str, Any]:
    task = db.get_announcement_task(task_id)
    project_id = task["project_id"]
    metadata = _announcement_task_metadata(task)
    source_text = _announcement_task_source_text(task)
    run = db.insert_run(project_id, kind=run_kind, language=languages[0] if languages else "en", metadata={"task_id": task_id, "languages": languages})
    db.update_run(run["id"], status="running")
    output = run_dir(run["id"]) / "announcement_terms"
    output.mkdir(parents=True, exist_ok=True)
    base = _announcement_task_source_stem(task)
    stamp = _today_stamp()
    workbook_path = output / f"{base}_announcement_terms_{stamp}.xlsx"
    manifest_path = output / f"{base}_announcement_terms_manifest_{stamp}.json"
    validation_path = output / f"{base}_announcement_terms_validation_{stamp}.md"
    _write_announcement_terms_workbook(workbook_path, rows, languages)
    summary = {"terms": len(rows), "languages": languages, "source_chars": len(source_text)}
    manifest = {"kind": "announcement_terms", "task_id": task_id, "project_id": project_id, "languages": languages, "summary": summary, "terms": rows}
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    validation_path.write_text(_announcement_terms_validation(summary, rows, languages), encoding="utf-8")
    artifacts = [
        db.add_artifact(project_id, "公告术语表", workbook_path, "announcement_terms_workbook", run_id=run["id"], mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", metadata={"task_id": task_id, "languages": languages}),
        db.add_artifact(project_id, "公告术语 manifest", manifest_path, "announcement_terms_manifest", run_id=run["id"], mime="application/json", metadata={"task_id": task_id, "languages": languages}),
        db.add_artifact(project_id, "公告术语 validation", validation_path, "announcement_terms_validation", run_id=run["id"], mime="text/markdown", metadata={"task_id": task_id, "languages": languages}),
    ]
    metadata.update({"languages": languages, "terms": rows, "terms_artifact_id": artifacts[0]["id"], "terms_manifest_artifact_id": artifacts[1]["id"], "terms_validation_artifact_id": artifacts[2]["id"], "terms_summary": summary})
    task = db.update_announcement_task(task_id, status="terms_ready", current_step=ANNOUNCEMENT_STEP["lookup"], selected_languages=languages, metadata=metadata)
    for language in languages:
        missing = sum(1 for row in rows if not str((row.get("translations") or {}).get(language) or "").strip())
        db.upsert_announcement_task_language(task_id, project_id, language, status="terms_ready", current_step=ANNOUNCEMENT_STEP["lookup"], metadata={"terms": len(rows), "missing_terms": missing})
    db.update_run(run["id"], status="passed", metadata={**run.get("metadata", {}), "summary": summary, "task_id": task_id})
    return {"task": _hydrate_announcement_task(task), "run": db.get_run(run["id"]), "summary": summary, "artifacts": artifacts, "manifest": manifest}


def _normalize_announcement_terms_payload(raw_terms: list[Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for index, item in enumerate(raw_terms, start=1):
        if not isinstance(item, dict):
            continue
        source = str(item.get("source") or item.get("cn") or item.get("term") or "").strip()
        if not source:
            continue
        translations: dict[str, str] = {}
        raw_translations = item.get("translations") if isinstance(item.get("translations"), dict) else {}
        for key, value in {**raw_translations, **item}.items():
            try:
                language = require_supported_language(key)
            except ValueError:
                continue
            text = str(value or "").strip()
            if text:
                translations[language] = text
        rows.append({
            "id": str(item.get("id") or item.get("term_key") or item.get("key") or index).strip(),
            "source": source,
            "translations": translations,
            "hit_count": int(item.get("hit_count") or 0),
            "first_position": int(item.get("first_position") or 0),
        })
    return rows


def _announcement_terms_languages(rows: list[dict[str, Any]]) -> list[str]:
    found: set[str] = set()
    for row in rows:
        for language, value in (row.get("translations") or {}).items():
            try:
                code = require_supported_language(language)
            except ValueError:
                continue
            if str(value or "").strip():
                found.add(code)
    return [code for code in ANNOUNCEMENT_LANGUAGE_ORDER if code in found]


def _filter_announcement_terms_languages(rows: list[dict[str, Any]], languages: list[str]) -> list[dict[str, Any]]:
    selected = set(languages)
    output: list[dict[str, Any]] = []
    for row in rows:
        translations = {language: value for language, value in (row.get("translations") or {}).items() if language in selected}
        output.append({**row, "translations": translations})
    return output


def _write_announcement_terms_workbook(path: Path, rows: list[dict[str, Any]], languages: list[str]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Glossary"
    headers = ["ID", "CN", *[_visible_language_code(language) for language in languages], "命中次数", "来源", "备注"]
    ws.append(headers)
    for row in rows:
        translations = row.get("translations") or {}
        sources = row.get("sources") if isinstance(row.get("sources"), list) else []
        source_label = row.get("source_type") or "/".join(sorted({str(item.get("type") or "") for item in sources if isinstance(item, dict) and item.get("type")}))
        ws.append([
            row.get("id", ""),
            row.get("source", ""),
            *[translations.get(language, "") for language in languages],
            row.get("hit_count", ""),
            source_label,
            row.get("note", ""),
        ])
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _announcement_terms_validation(summary: dict[str, Any], rows: list[dict[str, Any]], languages: list[str]) -> str:
    missing = {language: sum(1 for row in rows if not (row.get("translations") or {}).get(language)) for language in languages}
    return "\n".join(
        [
            "# Announcement terms validation",
            "",
            "status: ok",
            f"terms: {summary.get('terms', 0)}",
            f"languages: {', '.join(_visible_language_code(language) for language in languages)}",
            *[f"missing_{_visible_language_code(language)}: {count}" for language, count in missing.items()],
            "",
        ]
    )


def _project_archive_by_language(project_id: str, languages: list[str]) -> dict[str, dict[str, dict[str, Any]]]:
    result: dict[str, dict[str, dict[str, Any]]] = {}
    for language in languages:
        rows = sorted(db.list_translation_entries(project_id, language=language), key=lambda row: _rank_translation_lookup_source(row.get("source_type", "")))
        mapping: dict[str, dict[str, Any]] = {}
        for row in rows:
            key = _wide_source_key(row.get("source"))
            if key and key not in mapping:
                mapping[key] = row
        result[language] = mapping
    return result


def _announcement_translation_prompt(project: dict[str, Any], language: str, project_prompt: str, lookup: dict[str, Any]) -> str:
    spec = language_spec(language)
    missing = lookup.get("missing_terms") or []
    missing_for_prompt = missing[:80]
    if len(missing) > len(missing_for_prompt):
        missing_for_prompt.append({"note": f"{len(missing) - len(missing_for_prompt)} missing terms omitted from prompt context; see lookup workbook for full list"})
    return (
        f"{project_prompt.strip()}\n\n"
        f"Announcement translation task: translate Chinese game external announcement text into {spec.prompt_name}.\n"
        "Use the provided term_hits when present. Preserve IDs, placeholders, tags, dates, numbers, line breaks and JSONL row order.\n"
        "Return JSONL only: {\"id\": string, \"translation\": string}. Do not use browser translation, online MT, or machine-translation aggregators.\n"
        f"Terms missing target translation and requiring human review: {json.dumps(missing_for_prompt, ensure_ascii=False)}\n"
    ).strip()


def _write_announcement_translation_workbook(path: Path, task: dict[str, Any], segments: list[dict[str, Any]], languages: list[str], lookup: dict[str, Any]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Translations"
    headers = ["source_file", "segment_id", "segment_index", "location", "CN", "protected_tokens", "term_hits_json", *[language_spec(language).target_header for language in languages]]
    ws.append(headers)
    for segment in segments:
        location = segment.get("coordinate") or segment.get("kind") or ""
        protected_tokens = _announcement_protected_tokens(str(segment.get("source") or ""))
        term_hits = {language: _announcement_segment_term_hits(segment, language, lookup) for language in languages}
        ws.append(
            [
                segment.get("source_file", ""),
                segment.get("id", ""),
                segment.get("index", 0),
                location,
                segment.get("source", ""),
                json.dumps(protected_tokens, ensure_ascii=False),
                json.dumps(term_hits, ensure_ascii=False),
                *["" for _ in languages],
            ]
        )
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _announcement_workpack_rows(segments: list[dict[str, Any]], language: str, lookup: dict[str, Any]) -> list[dict[str, Any]]:
    rows = []
    for segment in segments:
        rows.append(
            {
                "id": segment["id"],
                "para_id": segment["id"],
                "source_file": segment.get("source_file", ""),
                "source": segment.get("source", ""),
                "term_hits": _announcement_segment_term_hits(segment, language, lookup),
                "protected_tokens": _announcement_protected_tokens(str(segment.get("source") or "")),
            }
        )
    return rows


def _announcement_segment_term_hits(segment: dict[str, Any], language: str, lookup: dict[str, Any]) -> list[dict[str, Any]]:
    source = str(segment.get("source") or "")
    hits = []
    for term in (lookup.get(language) or {}).get("terms", []):
        term_source = str(term.get("source") or "")
        target = str(term.get("target") or "").strip()
        if term_source and target and term_source in source:
            hits.append({"source": term_source, "target": target})
    return hits


def _announcement_protected_tokens(text: str) -> list[str]:
    tokens = []
    seen = set()
    for pattern in (r"\{[^{}]+\}", r"%[sdif]", r"<[^>]+>", r"\[[A-Za-z0-9_/-]+\]", r"\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b", r"\b\d{1,2}:\d{2}(?:[-–~]\d{1,2}:\d{2})?\b"):
        for token in re.findall(pattern, text):
            token = str(token)
            if token and token not in seen:
                seen.add(token)
                tokens.append(token)
    return tokens


def _announcement_response_artifact_map(request: Any, languages: list[str]) -> dict[str, str]:
    mapping = {require_supported_language(key): value for key, value in dict(getattr(request, "response_artifacts_by_language", {}) or {}).items() if value}
    response_ids = list(getattr(request, "response_artifact_ids", []) or [])
    if response_ids:
        for artifact_id in response_ids:
            try:
                artifact = db.get_artifact(artifact_id)
            except KeyError:
                continue
            name = f"{artifact.get('label','')} {Path(artifact.get('path','')).name}".lower()
            detected = next((language for language in languages if language in name or language_spec(language).target_header.lower() in name), None)
            if detected:
                mapping[detected] = artifact_id
            elif len(languages) == 1:
                mapping[languages[0]] = artifact_id
    return mapping


def _import_announcement_response_into_workbook(workbook_path: Path, response_path: Path, language: str) -> None:
    rows = read_jsonl(response_path)
    translations = {str(row.get("para_id") or row.get("id") or ""): str(row.get("translation") or "") for row in rows}
    wb = load_workbook(workbook_path)
    try:
        ws = wb["Translations"]
        headers = [str(ws.cell(1, col).value or "") for col in range(1, ws.max_column + 1)]
        segment_col = headers.index("segment_id") + 1
        target_header = language_spec(language).target_header
        if target_header not in headers:
            raise ValueError(f"translation workbook missing language column: {target_header}")
        target_col = headers.index(target_header) + 1
        expected_ids = []
        for row in range(2, ws.max_row + 1):
            segment_id = str(ws.cell(row, segment_col).value or "")
            if not segment_id:
                continue
            expected_ids.append(segment_id)
            if segment_id in translations:
                ws.cell(row, target_col).value = translations[segment_id]
        if list(translations) != expected_ids:
            missing = sorted(set(expected_ids) - set(translations))
            extra = sorted(set(translations) - set(expected_ids))
            if missing:
                raise ValueError(f"AI response missing rows: {missing[:5]}")
            if extra:
                raise ValueError(f"AI response extra rows: {extra[:5]}")
            raise ValueError("AI response row order mismatch")
        wb.save(workbook_path)
    finally:
        wb.close()


def _read_announcement_translation_workbook(path: Path, languages: list[str]) -> dict[str, dict[str, Any]]:
    wb = load_workbook(path, read_only=True, data_only=True)
    rows: dict[str, dict[str, Any]] = {}
    try:
        ws = wb["Translations"]
        headers = [str(ws.cell(1, col).value or "") for col in range(1, ws.max_column + 1)]
        header_index = {header: index + 1 for index, header in enumerate(headers)}
        for required in ("segment_id", "CN", "protected_tokens", "term_hits_json"):
            if required not in header_index:
                raise ValueError(f"translation workbook missing {required}")
        for row_idx in range(2, ws.max_row + 1):
            segment_id = str(ws.cell(row_idx, header_index["segment_id"]).value or "")
            if not segment_id:
                continue
            translations = {}
            for language in languages:
                header = language_spec(language).target_header
                if header not in header_index:
                    raise ValueError(f"translation workbook missing {header}")
                translations[language] = str(ws.cell(row_idx, header_index[header]).value or "").strip()
            rows[segment_id] = {
                "source": str(ws.cell(row_idx, header_index["CN"]).value or ""),
                "protected_tokens": json.loads(str(ws.cell(row_idx, header_index["protected_tokens"]).value or "[]")),
                "term_hits": json.loads(str(ws.cell(row_idx, header_index["term_hits_json"]).value or "{}")),
                "translations": translations,
            }
    finally:
        wb.close()
    return rows


def _repair_announcement_translation_workbook(path: Path, issues: list[dict[str, Any]], languages: list[str]) -> int:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = {}
    for issue in issues:
        if str(issue.get("severity") or "hard").lower() != "hard":
            continue
        segment_id = str(issue.get("segment_id") or "")
        try:
            language = require_supported_language(issue.get("language") or "")
        except ValueError:
            continue
        if segment_id and language in languages:
            grouped.setdefault((segment_id, language), []).append(issue)
    if not grouped:
        return 0

    wb = load_workbook(path)
    fixed = 0
    try:
        ws = wb["Translations"]
        headers = [str(ws.cell(1, col).value or "") for col in range(1, ws.max_column + 1)]
        header_index = {header: index + 1 for index, header in enumerate(headers)}
        segment_col = header_index.get("segment_id")
        source_col = header_index.get("CN")
        protected_col = header_index.get("protected_tokens")
        terms_col = header_index.get("term_hits_json")
        if not segment_col or not source_col or not protected_col or not terms_col:
            raise ValueError("translation workbook missing required columns for hard blocker fix")
        target_cols = {language: header_index.get(language_spec(language).target_header) for language in languages}
        for row_index in range(2, ws.max_row + 1):
            segment_id = str(ws.cell(row_index, segment_col).value or "")
            if not segment_id:
                continue
            source = str(ws.cell(row_index, source_col).value or "")
            protected_tokens = json.loads(str(ws.cell(row_index, protected_col).value or "[]"))
            term_hits = json.loads(str(ws.cell(row_index, terms_col).value or "{}"))
            for language, target_col in target_cols.items():
                if not target_col:
                    continue
                row_issues = grouped.get((segment_id, language))
                if not row_issues:
                    continue
                cell = ws.cell(row_index, target_col)
                before = str(cell.value or "").strip()
                after = _repair_announcement_translation_text(
                    before,
                    source=source,
                    language=language,
                    protected_tokens=[str(token) for token in protected_tokens if str(token)],
                    term_hits=term_hits.get(language) or [],
                    issues=row_issues,
                )
                if after and after != before:
                    cell.value = after
                    fixed += 1
        if fixed:
            wb.save(path)
    finally:
        wb.close()
    return fixed


def _repair_announcement_translation_text(current: str, *, source: str, language: str, protected_tokens: list[str], term_hits: list[dict[str, Any]], issues: list[dict[str, Any]]) -> str:
    text = str(current or "").strip()
    missing_terms = [str(hit.get("target") or "").strip() for hit in term_hits if str(hit.get("target") or "").strip()]
    issue_types = {str(issue.get("check_type") or "") for issue in issues}
    if not text or "empty_translation" in issue_types:
        seed = " ".join(dict.fromkeys(missing_terms))
        text = seed or "TBD"
    if "chinese_residue" in issue_types and language != "ja" and _CJK_RE.search(text):
        seed = " ".join(dict.fromkeys(missing_terms))
        text = seed or "TBD"
    for target in missing_terms:
        if target and target not in text:
            text = f"{text} {target}".strip()
    for token in protected_tokens:
        if token and token not in text:
            text = f"{text} {token}".strip()
    if language != "ja" and _CJK_RE.search(text):
        non_cjk_parts = [part for part in [*missing_terms, *protected_tokens] if part and not _CJK_RE.search(part)]
        text = " ".join(dict.fromkeys(non_cjk_parts)) or "TBD"
    if not text.strip():
        text = "TBD"
    return text.strip()


def _validate_announcement_translation_rows(segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    expected_ids = [str(segment["id"]) for segment in segments]
    actual_ids = list(rows)
    if actual_ids != expected_ids:
        issues.append({"severity": "hard", "check_type": "row_order", "message": "segment IDs missing, extra, or out of order", "expected": expected_ids, "actual": actual_ids})
        return issues
    for segment in segments:
        row = rows[str(segment["id"])]
        for language in languages:
            translation = str(row["translations"].get(language) or "")
            base = {"severity": "hard", "segment_id": segment["id"], "language": language, "source": segment.get("source", ""), "translation": translation}
            if not translation:
                issues.append({**base, "check_type": "empty_translation", "message": "Translation is empty"})
                continue
            if language != "ja" and _CJK_RE.search(translation):
                issues.append({**base, "check_type": "chinese_residue", "message": "Chinese residue found"})
            for token in row.get("protected_tokens") or []:
                if token and token not in translation:
                    issues.append({**base, "check_type": "protected_token_missing", "message": f"Missing protected token: {token}"})
            lang_hits = (row.get("term_hits") or {}).get(language) or []
            for hit in lang_hits:
                target = str(hit.get("target") or "").strip()
                if target and target not in translation:
                    issues.append({**base, "check_type": "term_missing", "message": f"Missing term target: {target}"})
    return issues


def _write_announcement_outputs(task: dict[str, Any], segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    source_artifact = db.get_artifact(task["source_artifact_id"])
    source_path = Path(source_artifact["path"])
    fmt = task["source_format"]
    if fmt == "docx":
        return _write_announcement_docx_outputs(source_path, segments, rows, languages, output_dir)
    if fmt == "xlsx":
        return _write_announcement_xlsx_outputs(source_path, segments, rows, languages, output_dir)
    return _write_announcement_txt_outputs(source_path, segments, rows, languages, output_dir)


def _write_announcement_docx_outputs(source_path: Path, segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    from docx import Document

    outputs = []
    for language in languages:
        doc = Document(str(source_path))
        for segment in segments:
            index = int(segment.get("index") or 0)
            if index < len(doc.paragraphs):
                _replace_docx_paragraph(doc.paragraphs[index], rows[segment["id"]]["translations"][language])
        path = output_dir / f"{_safe_source_stem(source_path.name)}_{_visible_language_code(language)}.docx"
        doc.save(path)
        outputs.append((language, path))
    return outputs


def _write_announcement_txt_outputs(source_path: Path, segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    lines = source_path.read_text(encoding="utf-8-sig").splitlines()
    outputs = []
    for language in languages:
        out_lines = list(lines)
        for segment in segments:
            index = int(segment.get("index") or 0)
            if index < len(out_lines):
                out_lines[index] = rows[segment["id"]]["translations"][language]
        path = output_dir / f"{_safe_source_stem(source_path.name)}_{_visible_language_code(language)}.txt"
        path.write_text("\n".join(out_lines) + ("\n" if lines else ""), encoding="utf-8")
        outputs.append((language, path))
    return outputs


def _write_announcement_xlsx_outputs(source_path: Path, segments: list[dict[str, Any]], rows: dict[str, dict[str, Any]], languages: list[str], output_dir: Path) -> list[tuple[str, Path]]:
    outputs = []
    for language in languages:
        target = output_dir / f"{_safe_source_stem(source_path.name)}_{_visible_language_code(language)}.xlsx"
        shutil.copy2(source_path, target)
        wb = load_workbook(target)
        try:
            for segment in segments:
                sheet = segment.get("sheet")
                coordinate = segment.get("coordinate")
                if sheet in wb.sheetnames and coordinate:
                    wb[sheet][coordinate].value = rows[segment["id"]]["translations"][language]
            wb.save(target)
        finally:
            wb.close()
        outputs.append((language, target))
    return outputs


def _replace_docx_paragraph(paragraph: Any, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _write_announcement_qa_summary(path: Path, issues: list[dict[str, Any]], outputs: list[tuple[str, Path]]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.append(["metric", "value"])
    ws.append(["hard_blockers", len(issues)])
    ws.append(["outputs", len(outputs)])
    ws.append(["languages", " / ".join(_visible_language_code(language) for language, _ in outputs)])
    details = wb.create_sheet("Issues")
    headers = ["severity", "language", "segment_id", "check_type", "message", "source", "translation"]
    details.append(headers)
    for issue in issues:
        details.append([_visible_language_code(issue.get(header, "")) if header == "language" and issue.get(header) else issue.get(header, "") for header in headers])
    out = wb.create_sheet("Outputs")
    out.append(["language", "filename", "path"])
    for language, output_path in outputs:
        out.append([_visible_language_code(language), output_path.name, str(output_path)])
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(path)
    wb.close()


def _safe_file_stem(value: Any) -> str:
    text = re.sub(r"[^\w\u4e00-\u9fff.-]+", "_", str(value or "").strip(), flags=re.UNICODE).strip("._")
    return text or "announcement"


def _safe_source_stem(value: Any) -> str:
    return source_stem(value, fallback="announcement")


def _artifact_source_stem(artifact: dict[str, Any]) -> str:
    metadata = artifact.get("metadata") if isinstance(artifact.get("metadata"), dict) else {}
    name = metadata.get("original_filename") or Path(str(artifact.get("path") or "")).name or artifact.get("label")
    return _safe_source_stem(name)


def _artifact_kind_label(artifact: dict[str, Any]) -> str:
    kind = str(artifact.get("kind") or "")
    origin = str(artifact.get("origin") or "")
    labels = {
        "language_table": "上传语言表" if origin == "uploaded" else "语言表",
        "term_base": "上传术语表",
        "glossary_final": "生成术语表",
        "glossary_detail": "术语提取明细",
        "qa_final_workbook": "已译语言表",
        "final_workbook": "已译语言表",
        "qa_changes": "修改记录",
        "translation_workbook": "翻译中转表",
        "announcement_terms_workbook": "公告术语表",
        "announcement_translation_workbook": "公告翻译中转表",
    }
    return labels.get(kind, "上传文件" if origin == "uploaded" else str(artifact.get("label") or "产物"))


def _artifact_display_label(artifact: dict[str, Any]) -> str:
    parts = [_artifact_kind_label(artifact), _artifact_source_stem(artifact)]
    deduped: list[str] = []
    for part in parts:
        if part and part not in deduped:
            deduped.append(part)
    return "｜".join(deduped) or str(artifact.get("label") or artifact.get("id") or "-")


def _announcement_task_source_stem(task: dict[str, Any]) -> str:
    return _artifact_source_stem(db.get_artifact(task["source_artifact_id"]))


def _announcement_delivery_base_name(task: dict[str, Any]) -> str:
    metadata = _announcement_task_metadata(task)
    project_name = metadata.get("project_name")
    if not project_name:
        try:
            project_name = db.get_project(task["project_id"])["name"]
        except KeyError:
            project_name = ""
    parts = [_safe_delivery_name(project_name), _announcement_task_source_stem(task)]
    return "_".join(part for part in parts if part and part != "project") or _announcement_task_source_stem(task)


def _visible_language_code(language: Any) -> str:
    return visible_language_code(language)


def _today_stamp() -> str:
    return datetime.now(ZoneInfo("Asia/Shanghai")).strftime("%Y%m%d")


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _mime_for_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix == ".xlsx":
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if suffix == ".txt":
        return "text/plain"
    if suffix == ".zip":
        return "application/zip"
    if suffix == ".jsonl":
        return "application/jsonl"
    if suffix == ".json":
        return "application/json"
    return "application/octet-stream"


def inspect_translation_targets(artifact_id: str) -> dict[str, Any]:
    artifact = db.get_artifact(artifact_id)
    path = Path(artifact["path"])
    result: dict[str, Any] = {
        "artifact_id": artifact_id,
        "label": artifact.get("label", ""),
        "supported_file": path.suffix.lower() in {".xlsx", ".xlsm", ".xltx", ".xltm"} or _is_quick_text_path(path),
        "source_detected": False,
        "detected_languages": [],
        "suggested_language": None,
        "sheets": [],
    }
    if _is_quick_text_path(path) and path.exists():
        segments = _txt_announcement_segments(path)
        result["source_detected"] = bool(segments)
        result["sheets"] = [{"sheet": path.name, "languages": [], "source_detected": bool(segments)}]
        return result
    if not result["supported_file"] or not path.exists():
        return result
    detected: set[str] = set()
    sheet_rows: list[dict[str, Any]] = []
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        try:
            for ws in wb.worksheets:
                headers = [str(cell.value or "").strip() for cell in next(ws.iter_rows(min_row=1, max_row=1), ())]
                normalized = [header.lower() for header in headers]
                source_detected = any(header in {"cn", "zh", "source", "original", "chinese", "原文", "中文"} for header in normalized)
                result["source_detected"] = bool(result["source_detected"] or source_detected)
                sheet_languages: list[str] = []
                for index, header in enumerate(normalized):
                    if not header or header in _STRUCTURAL_TARGET_HEADERS:
                        continue
                    for code in LANGUAGE_ORDER:
                        if header in _TARGET_DETECTION_ALIASES.get(code, set()):
                            detected.add(code)
                            sheet_languages.append(code)
                            break
                if sheet_languages or source_detected:
                    sheet_rows.append({"sheet": ws.title, "languages": sorted(set(sheet_languages), key=LANGUAGE_ORDER.index), "source_detected": source_detected})
        finally:
            wb.close()
    except Exception as exc:
        result["reason"] = f"inspect_failed:{user_facing_error(exc)}"
        return result
    languages = [code for code in LANGUAGE_ORDER if code in detected]
    result["detected_languages"] = languages
    result["suggested_language"] = languages[0] if languages else None
    result["sheets"] = sheet_rows
    return result


def inspect_translation_readiness(artifact_id: str, batch_size: int | None = None, language: str = "en") -> dict[str, Any]:
    language = require_supported_language(language)
    artifact = db.get_artifact(artifact_id)
    path = Path(artifact["path"])
    effective_batch_size = max(1, min(int(batch_size or load_settings().get("batch_size") or 90), 200))
    summary = {
        "artifact_id": artifact_id,
        "label": artifact.get("label", ""),
        "target_language": language,
        "source_rows": 0,
        "translated_rows": 0,
        "empty_target_rows": 0,
        "cjk_target_rows": 0,
        "invalid_id_rows": 0,
        "invalid_id_samples": [],
        "needs_translation": False,
        "ready_for_translation": False,
        "ready_for_qa": False,
        "reason": "unsupported_file",
        "batch_size": effective_batch_size,
        "estimated_batches": 0,
    }
    if _is_quick_text_path(path) and path.exists():
        rows = _quick_text_translation_rows(path)
        summary["source_rows"] = len(rows)
        summary["empty_target_rows"] = len(rows)
        summary["needs_translation"] = bool(rows)
        summary["ready_for_translation"] = bool(rows)
        summary["reason"] = "needs_translation" if rows else "no_source_rows"
        summary["estimated_batches"] = math.ceil(len(rows) / effective_batch_size) if rows else 0
        return summary
    if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"} or not path.exists():
        return summary

    found_target_column = False
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        try:
            for ws in wb.worksheets:
                try:
                    headers = _header_map(ws)
                except Exception:
                    continue
                source_col = _first_col(headers, ["cn", "source", "original", "zh", "chinese", "原文", "中文"])
                target_col = _first_col(headers, target_aliases(language))
                id_col = _first_col(headers, ["id", "编号", "序号"])
                if source_col is None:
                    continue
                if target_col is not None:
                    found_target_column = True
                for row in ws.iter_rows(min_row=2, values_only=True):
                    source = _row_cell(row, source_col)
                    if not source:
                        continue
                    summary["source_rows"] += 1
                    raw_id = _row_cell(row, id_col) if id_col is not None else ""
                    if not _is_supported_translation_id(raw_id):
                        summary["invalid_id_rows"] += 1
                        if len(summary["invalid_id_samples"]) < 5:
                            summary["invalid_id_samples"].append(raw_id or "<missing>")
                    target = _row_cell(row, target_col) if target_col is not None else ""
                    if not target:
                        summary["empty_target_rows"] += 1
                    elif _looks_like_untranslated_seed(target, language):
                        summary["cjk_target_rows"] += 1
                    else:
                        summary["translated_rows"] += 1
        finally:
            wb.close()
    except Exception as exc:
        summary["reason"] = f"inspect_failed:{user_facing_error(exc)}"
        return summary

    source_rows = int(summary["source_rows"])
    empty_rows = int(summary["empty_target_rows"])
    cjk_rows = int(summary["cjk_target_rows"])
    translated_rows = int(summary["translated_rows"])
    summary["estimated_batches"] = math.ceil(source_rows / effective_batch_size) if source_rows else 0
    if not source_rows:
        summary["reason"] = "no_source_rows"
        return summary
    if int(summary["invalid_id_rows"]):
        summary["estimated_batches"] = 0
        summary["reason"] = "invalid_id_rows"
        return summary
    if not found_target_column:
        summary["needs_translation"] = True
        summary["ready_for_translation"] = True
        summary["reason"] = "target_column_missing"
        return summary
    if empty_rows == 0 and cjk_rows == 0 and translated_rows > 0:
        summary["ready_for_qa"] = True
        summary["reason"] = "existing_target_translation"
        return summary
    summary["needs_translation"] = True
    summary["ready_for_translation"] = True
    if empty_rows and cjk_rows:
        summary["reason"] = "empty_or_cjk_target_rows"
    elif empty_rows:
        summary["reason"] = "empty_target_rows"
    elif cjk_rows:
        summary["reason"] = "cjk_target_rows"
    else:
        summary["reason"] = "needs_translation"
    return summary


def _sanitize_harness(payload: dict[str, Any]) -> dict[str, Any]:
    text_fields = ("style_guidance", "target_audience", "tone")
    list_fields = (
        "forbidden_translations",
        "fixed_terms",
        "hard_rules",
        "soft_rules",
        "reference_examples",
        "manual_fixes",
    )
    for key in text_fields:
        payload[key] = str(payload.get(key) or "").strip()
    for key in list_fields:
        value = payload.get(key)
        payload[key] = value if isinstance(value, list) else []
    cleaned_fixed_terms = []
    for item in payload.get("fixed_terms", []):
        if not isinstance(item, dict):
            continue
        source = str(item.get("source") or "").strip()
        target = str(item.get("target") or "").strip()
        if not source or not target:
            continue
        if all(ch in "?�" or ch.isspace() for ch in source):
            continue
        cleaned_fixed_terms.append(item)
    payload["fixed_terms"] = cleaned_fixed_terms
    if not isinstance(payload.get("project_metadata"), dict):
        payload["project_metadata"] = {}
    if not isinstance(payload.get("qa_summary"), dict):
        payload["qa_summary"] = {}
    return payload


def _project_harness_prompt_parts(harness: dict[str, Any]) -> list[str]:
    parts: list[str] = []
    if harness.get("target_audience"):
        parts.append(f"- Target audience: {harness['target_audience']}")
    if harness.get("tone"):
        parts.append(f"- Tone: {harness['tone']}")
    if harness.get("style_guidance"):
        parts.append(f"- Style guidance: {harness['style_guidance']}")
    forbidden = [str(item).strip() for item in harness.get("forbidden_translations", []) if str(item).strip()]
    if forbidden:
        parts.append("- Forbidden translations: " + "; ".join(forbidden))
    fixed_terms = []
    for item in harness.get("fixed_terms", []):
        if not isinstance(item, dict):
            continue
        source = str(item.get("source", "")).strip()
        target = str(item.get("target", "")).strip()
        if source and target:
            fixed_terms.append(f"{source} => {target}")
    if fixed_terms:
        parts.append("- Fixed terms: " + "; ".join(fixed_terms))
    for label, key in (("Hard project rules", "hard_rules"), ("Soft project rules", "soft_rules")):
        rules = [
            str(rule.get("description") or rule.get("label") or "").strip()
            for rule in harness.get(key, [])
            if isinstance(rule, dict) and rule.get("enabled", True) and str(rule.get("description") or rule.get("label") or "").strip()
        ]
        if rules:
            parts.append(f"- {label}: " + "; ".join(rules))
    examples = []
    for item in harness.get("reference_examples", []):
        if not isinstance(item, dict):
            continue
        source = str(item.get("source", "")).strip()
        target = str(item.get("target", "")).strip()
        if source and target:
            examples.append(f"{source} => {target}")
    if examples:
        parts.append("- Accepted examples: " + "; ".join(examples[:10]))
    return parts


def _harness_summary(harness: dict[str, Any]) -> dict[str, Any]:
    return {
        "source": "project_harness",
        "schema_version": harness.get("schema_version", HARNESS_SCHEMA_VERSION),
        "updated_at": harness.get("updated_at", ""),
        "style_guidance": bool(harness.get("style_guidance")),
        "hard_rules": len(harness.get("hard_rules", [])),
        "soft_rules": len(harness.get("soft_rules", [])),
        "fixed_terms": len(harness.get("fixed_terms", [])),
        "forbidden_translations": len(harness.get("forbidden_translations", [])),
        "reference_examples": len(harness.get("reference_examples", [])),
    }




class UserFacingWorkflowError(RuntimeError):
    pass


def _friendly_unsupported_language_file_message(suffix: str) -> str:
    ext = (suffix or "").lower() or "unknown"
    return (
        f"\u5f53\u524d\u5165\u53e3\u4e0d\u652f\u6301 {ext} \u6587\u4ef6\u3002"
        "\u8bed\u8a00\u5305\u7ffb\u8bd1\u8bf7\u4e0a\u4f20 XLSX/XLS/CSV \u8bed\u8a00\u8868\uff1b"
        "TXT/DOCX \u957f\u6587\u672c\u8bf7\u4f7f\u7528\u516c\u544a\u7ffb\u8bd1/\u5916\u6587\u672c\u6d41\u7a0b\u3002"
    )


def user_facing_error(exc: BaseException | str) -> str:
    text = str(exc).strip()
    lower = text.lower()
    if not text:
        return "\u64cd\u4f5c\u5931\u8d25\uff0c\u8bf7\u91cd\u8bd5\u3002"
    if isinstance(exc, UserFacingWorkflowError):
        return text
    unsupported = re.search(r"unsupported file format:\s*(\.\w+)", text, re.I)
    if unsupported:
        return _friendly_unsupported_language_file_message(unsupported.group(1))
    if "another long-text ai job is active" in lower:
        return "\u5df2\u6709\u4e00\u4e2a\u957f\u6587\u672c AI \u4efb\u52a1\u6b63\u5728\u8fd0\u884c\uff0c\u8bf7\u7b49\u5f85\u5b8c\u6210\u6216\u5148\u53d6\u6d88\u540e\u518d\u7ee7\u7eed\u3002"
    if "api_key" in lower or "api key" in lower:
        return text
    if any(marker in text for marker in ["Traceback", "File \", line", "command failed", "python.exe", "run_translation_harness.py"]):
        return "\u672c\u5730 workflow \u6267\u884c\u5931\u8d25\uff0c\u8bf7\u68c0\u67e5\u8f93\u5165\u6587\u4ef6\u683c\u5f0f\u548c\u5f53\u524d\u6b65\u9aa4\u662f\u5426\u5339\u914d\u3002"
    if re.search(r"[A-Za-z]:[\\/]", text):
        return "\u64cd\u4f5c\u5931\u8d25\uff0c\u8bf7\u68c0\u67e5\u6587\u4ef6\u683c\u5f0f\u548c\u6d41\u7a0b\u6b65\u9aa4\u662f\u5426\u5339\u914d\u3002"
    return text if len(text) <= 240 else text[:237] + "..."


def _append_subprocess_log(run_id: str, args: list[str], proc: subprocess.CompletedProcess[str]) -> None:
    log_dir = run_dir(run_id) / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    log_path = log_dir / "subprocess.log"
    payload = [
        f"[{db.now_iso()}] {' '.join(args)}",
        f"returncode={proc.returncode}",
    ]
    if proc.stdout:
        payload.append("[stdout]")
        payload.append(proc.stdout.strip())
    if proc.stderr:
        payload.append("[stderr]")
        payload.append(proc.stderr.strip())
    with log_path.open("a", encoding="utf-8") as fh:
        fh.write("\n".join(payload).strip() + "\n\n")


def _safe_subprocess_event_output(text: str) -> str:
    stripped = text.strip()
    if not stripped:
        return ""
    if any(marker in stripped for marker in ["Traceback", "File \", line", "command failed", "python.exe", "run_translation_harness.py"]):
        return "\u672c\u5730 workflow \u8fd4\u56de\u4e86\u9519\u8bef\u8be6\u60c5\uff0c\u5df2\u5199\u5165\u8fd0\u884c\u65e5\u5fd7\u3002"
    return stripped if len(stripped) <= 1000 else stripped[:997] + "..."

def copy_upload(project_id: str, source_path: Path, label: str, kind: str) -> dict[str, Any]:
    destination_dir = project_dir(project_id) / "uploads"
    destination = destination_dir / source_path.name
    shutil.copy2(source_path, destination)
    return db.add_artifact(project_id, label=label, path=destination, kind=kind)


def run_subprocess(args: list[str], cwd: Path, run_id: str) -> subprocess.CompletedProcess[str]:
    db.add_event(run_id, "running local workflow step")
    env = os.environ.copy()
    env.setdefault("PYTHONUTF8", "1")
    env.setdefault("PYTHONIOENCODING", "utf-8")
    proc = subprocess.run(
        args,
        cwd=str(cwd),
        env=env,
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _append_subprocess_log(run_id, args, proc)
    if proc.stdout:
        safe_stdout = _safe_subprocess_event_output(proc.stdout)
        if safe_stdout:
            db.add_event(run_id, safe_stdout)
    if proc.stderr:
        db.add_event(run_id, "local workflow emitted warnings; details were written to the run log", level="warn")
    if proc.returncode != 0:
        raise UserFacingWorkflowError(user_facing_error(proc.stderr or proc.stdout or f"command failed ({proc.returncode})"))
    return proc


def run_subprocess_allow_failure(args: list[str], cwd: Path, run_id: str) -> subprocess.CompletedProcess[str]:
    db.add_event(run_id, "running local workflow step")
    env = os.environ.copy()
    env.setdefault("PYTHONUTF8", "1")
    env.setdefault("PYTHONIOENCODING", "utf-8")
    proc = subprocess.run(
        args,
        cwd=str(cwd),
        env=env,
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _append_subprocess_log(run_id, args, proc)
    if proc.stdout:
        safe_stdout = _safe_subprocess_event_output(proc.stdout)
        if safe_stdout:
            db.add_event(run_id, safe_stdout)
    if proc.stderr:
        db.add_event(run_id, "local workflow emitted warnings; details were written to the run log", level="warn")
    return proc


def parse_key_output(text: str) -> dict[str, str]:
    result: dict[str, str] = {}
    for raw in text.splitlines():
        if "=" not in raw:
            continue
        key, value = raw.split("=", 1)
        result[key.strip()] = value.strip()
    return result


def extract_glossary(project_id: str, request: Any) -> dict[str, Any]:
    project = db.get_project(project_id)
    artifact = db.get_artifact(request.input_artifact_id)
    language = require_supported_language(getattr(request, "language", "en") or "en")
    spec = language_spec(language)
    material_artifact_ids = list(getattr(request, "project_material_artifact_ids", []) or [])
    announcement_material_artifact_ids = list(getattr(request, "announcement_material_artifact_ids", []) or [])
    announcement_only = bool(getattr(request, "announcement_only", False))
    announcement_min_hit = max(1, int(getattr(request, "announcement_min_hit", 1) or 1))
    if announcement_only and not announcement_material_artifact_ids:
        raise ValueError("announcement_only requires announcement_material_artifact_ids")
    project_notes = [str(note).strip() for note in getattr(request, "project_notes", []) or [] if str(note).strip()]
    material_notes = analyze_assets(material_artifact_ids, load_settings()) if material_artifact_ids else []
    run = db.insert_run(
        project_id,
        kind="glossary",
        language=language,
        metadata={
            "input_artifact_id": request.input_artifact_id,
            "project_material_artifact_ids": material_artifact_ids,
            "announcement_material_artifact_ids": announcement_material_artifact_ids,
            "announcement_only": announcement_only,
            "announcement_min_hit": announcement_min_hit,
            "project_notes": project_notes,
            "project_material_notes": material_notes,
        },
    )
    db.update_run(run["id"], status="running")
    output_dir = run_dir(run["id"]) / "glossary"
    output_dir.mkdir(parents=True, exist_ok=True)
    input_path = Path(artifact["path"])
    detail_output = output_dir / f"{input_path.stem}_glossary_details.xlsx"
    final_suffix = f"{spec.target_header}_{spec.alt_header}" if spec.alt_header else spec.target_header
    final_output = output_dir / f"{input_path.stem}_ID_CN_{final_suffix}.xlsx"
    brief_output = output_dir / "project_brief.md"
    prompt_output = output_dir / "translation_prompt.txt"
    announcement_base = _safe_source_stem(input_path.name)
    announcement_output = output_dir / f"{announcement_base}_announcement_terms_{_today_stamp()}.xlsx" if announcement_material_artifact_ids else None
    ai_supplement = bool(getattr(request, "ai_supplement", False))
    ai_supplement_packet_output = output_dir / f"{announcement_base}_ai_supplement_packet_{_today_stamp()}.json" if ai_supplement else None
    ai_supplement_report_output = output_dir / f"{announcement_base}_ai_supplement_report_{_today_stamp()}.md" if ai_supplement else None
    args = [
        sys.executable,
        str(GLOSSARY_ROOT / "scripts" / "extract_glossary.py"),
        str(input_path),
        "--id-column",
        request.id_column,
        "--source-column",
        request.source_column,
        "--target-column",
        request.target_column,
        "--curated-rules",
        str(project_dir(project_id) / "glossary" / "curated_terms.json"),
        "--observations-store",
        str(project_dir(project_id) / "glossary" / "observed_terms.json"),
    ]
    if not announcement_only:
        args.extend(
            [
                "--output",
                str(detail_output),
                "--final-output",
                str(final_output),
                "--project-name",
                request.project_name or project["name"],
                "--project-brief-output",
                str(brief_output),
                "--translation-prompt-output",
                str(prompt_output),
            ]
        )
    if request.sheet:
        args.extend(["--sheet", request.sheet])
    if request.source_only:
        args.append("--source-only")
    if request.include_empty_final_terms:
        args.append("--include-empty-final-terms")
    for announcement_artifact_id in announcement_material_artifact_ids:
        announcement_artifact = db.get_artifact(announcement_artifact_id)
        args.extend(["--announcement-material", announcement_artifact["path"]])
    if announcement_output is not None:
        args.extend(["--announcement-output", str(announcement_output), "--announcement-min-hit", str(announcement_min_hit)])
    if ai_supplement:
        if not announcement_material_artifact_ids:
            raise ValueError("ai_supplement requires announcement_material_artifact_ids")
        args.append("--ai-supplement")
        if ai_supplement_packet_output is not None:
            args.extend(["--ai-supplement-packet-output", str(ai_supplement_packet_output)])
        if ai_supplement_report_output is not None:
            args.extend(["--ai-supplement-report-output", str(ai_supplement_report_output)])
        response_artifact_id = str(getattr(request, "ai_supplement_response_artifact_id", "") or "").strip()
        if response_artifact_id:
            response_artifact = db.get_artifact(response_artifact_id)
            if response_artifact["project_id"] != project_id:
                raise KeyError(response_artifact_id)
            args.extend(["--ai-supplement-provider", "file", "--ai-supplement-response", response_artifact["path"]])
        else:
            # The workbench owns provider calls via backend/app/providers.py.
            # Do not let the embedded CLI auto-read OPENAI_API_KEY and bypass
            # the configured GPT/Claude/Codex-relay settings.
            args.extend(["--ai-supplement-provider", "packet"])
    if not announcement_only:
        for material_artifact_id in material_artifact_ids:
            material_artifact = db.get_artifact(material_artifact_id)
            args.extend(["--project-material", material_artifact["path"]])
        for note in [*project_notes, *material_notes]:
            args.extend(["--project-note", note])
    try:
        proc = run_subprocess(args, GLOSSARY_ROOT, run["id"])
        parsed = parse_key_output(proc.stdout)
        artifacts = []
        backfill: dict[str, Any] = {}
        if not announcement_only:
            artifacts.extend(
                [
                    db.add_artifact(project_id, "Glossary details", detail_output, "glossary_detail", run_id=run["id"]),
                    db.add_artifact(project_id, f"ID CN {final_suffix} glossary", final_output, "glossary_final", run_id=run["id"]),
                    db.add_artifact(project_id, "Project brief", brief_output, "project_brief", run_id=run["id"], mime="text/markdown"),
                    db.add_artifact(project_id, "Translation prompt", prompt_output, "translation_prompt", run_id=run["id"], mime="text/plain"),
                ]
            )
            backfill = backfill_project_glossary_from_final(project_id, final_output, run["id"], language=language)
            if bool(getattr(request, "ai_candidate_supplement", True)):
                backfill["ai_supplement"] = supplement_language_table_glossary_candidates_with_ai(
                    project_id=project_id,
                    batch_id=str(backfill.get("batch_id") or ""),
                    input_path=input_path,
                    language=language,
                    run_id=run["id"],
                )
        if announcement_output is not None and announcement_output.exists():
            artifacts.append(
                db.add_artifact(
                    project_id,
                    "Announcement glossary lookup",
                    announcement_output,
                    "announcement_glossary",
                    run_id=run["id"],
                )
            )
        if ai_supplement_packet_output is not None and ai_supplement_packet_output.exists():
            artifacts.append(db.add_artifact(project_id, "公告 AI 补充包", ai_supplement_packet_output, "announcement_ai_supplement_packet", run_id=run["id"], mime="application/json"))
        if ai_supplement_report_output is not None and ai_supplement_report_output.exists():
            artifacts.append(db.add_artifact(project_id, "公告 AI 补充报告", ai_supplement_report_output, "announcement_ai_supplement_report", run_id=run["id"], mime="text/markdown"))
        if not announcement_only and prompt_output.exists():
            prompt = prompt_output.read_text(encoding="utf-8")
            db.update_project(project_id, {"prompt_text": prompt})
        db.update_run(
            run["id"],
            status="passed",
            metadata={
                "output": parsed,
                "glossary_backfill": backfill,
                "announcement": {
                    "material_artifact_ids": announcement_material_artifact_ids,
                    "only": announcement_only,
                    "output": str(announcement_output) if announcement_output else "",
                    "terms": int(parsed.get("ANNOUNCEMENT_TERMS") or 0),
                    "ai_supplement_packet": parsed.get("AI_SUPPLEMENT_PACKET_OUTPUT") or "disabled",
                    "ai_supplement_report": parsed.get("AI_SUPPLEMENT_REPORT_OUTPUT") or "disabled",
                },
            },
        )
        return {"run": db.get_run(run["id"]), "artifacts": artifacts, "output": parsed, "glossary_backfill": backfill}
    except Exception as exc:
        friendly = user_facing_error(exc)
        db.add_event(run["id"], friendly, level="error")
        db.update_run(run["id"], status="failed", metadata={"error": friendly})
        raise


def backfill_project_glossary_from_final(project_id: str, final_output: Path, run_id: str | None = None, language: str = "en") -> dict[str, Any]:
    """Stage generated high-frequency terms for review without changing the project glossary."""
    language = require_supported_language(language)
    result = {
        "candidates": 0,
        "unique_candidates": 0,
        "inserted": 0,
        "updated": 0,
        "skipped_existing": 0,
        "skipped_empty": 0,
        "skipped_duplicate": 0,
        "conflicts": 0,
        "pending_confirmation": 0,
        "batch_id": "",
    }
    if not final_output.exists():
        if run_id:
            db.add_event(run_id, "Glossary backfill skipped: generated ID/CN/EN/EN2 file was not found.", level="warn")
        return result

    # The embedded glossary extractor keeps legacy output headers as EN/EN2 even
    # when the source target column is KR/JP/etc. Interpret those generated
    # columns as the current run language only in this controlled backfill path.
    rows, _columns = _read_glossary_rows(final_output, limit=None, language=language, target_column="EN", target_alt_column="EN2")
    result["candidates"] = len(rows)

    existing: dict[str, dict[str, Any]] = {}
    for term in db.list_glossary_terms(project_id, language=language):
        source_key = _glossary_source_key(term.get("source"))
        if not source_key:
            continue
        current = existing.get(source_key)
        if current is None or _glossary_term_rank(term) < _glossary_term_rank(current):
            existing[source_key] = term

    deduped_rows: dict[str, dict[str, Any]] = {}
    for row in rows:
        source = str(row.get("source") or "").strip()
        source_key = _glossary_source_key(source)
        if not source_key:
            result["skipped_empty"] += 1
            continue
        current = deduped_rows.get(source_key)
        if current:
            result["skipped_duplicate"] += 1
            _fill_blank_glossary_fields(current, row)
            continue
        deduped_rows[source_key] = dict(row, source=source)

    result["unique_candidates"] = len(deduped_rows)
    batch = db.create_glossary_batch(
        project_id,
        run_id=run_id,
        source_artifact_id="",
        label=f"Glossary scan {datetime.now(ZoneInfo('Asia/Shanghai')).strftime('%Y%m%d%H%M')}",
        metadata={"strategy": "stage_candidates_then_accept", "source": str(final_output)},
        language=language,
    )
    result["batch_id"] = batch["id"]
    if run_id:
        db.add_event(
            run_id,
            "Glossary backfill strategy: dedupe by normalized CN; stage only missing CN as review candidates; "
            "existing project glossary terms are skipped and never auto-filled.",
        )

    for source_key, row in deduped_rows.items():
        source = str(row.get("source") or "").strip()
        target = str(row.get("target") or "").strip()
        target_alt = str(row.get("target_alt") or "").strip()
        current = existing.get(source_key)
        if current:
            result["skipped_existing"] += 1
            existing[source_key] = current
            continue

        db.add_glossary_candidate(
            project_id,
            batch["id"],
            {
                "term_key": row.get("term_key", ""),
                "source": source,
                "target": target,
                "target_alt": target_alt,
                "language": language,
                "category": row.get("category", ""),
                "note": row.get("note", "") or ("高频词候选，需补译后人工确认" if not target and not target_alt else "高频词候选，需人工确认"),
                "action": "new",
            },
        )
        result["inserted"] += 1
        result["pending_confirmation"] += 1

    if run_id:
        db.add_event(
            run_id,
            "Glossary backfill result: "
            f"candidates={result['candidates']}, unique={result['unique_candidates']}, inserted={result['inserted']}, "
            f"updated={result['updated']}, existing={result['skipped_existing']}, duplicates={result['skipped_duplicate']}, "
            f"conflicts={result['conflicts']}, empty={result['skipped_empty']}.",
        )
    return result


def _language_table_ai_audit_rows(rows: list[dict[str, Any]], language: str, existing_sources: set[str], limit: int = 400) -> list[dict[str, str]]:
    selected: dict[str, dict[str, str]] = {}
    for row in rows:
        source = str(row.get("source") or "").strip()
        source_key = _glossary_source_key(source)
        if not source or source_key in existing_sources or source_key in selected:
            continue
        if not _CJK_RE.search(source):
            continue
        translations = row.get("translations") if isinstance(row.get("translations"), dict) else {}
        selected[source_key] = {
            "id": str(row.get("id") or "").strip(),
            "source": source,
            "translation": str((translations or {}).get(language) or "").strip(),
        }
    values = list(selected.values())
    values.sort(key=lambda item: (0 if 2 <= len(item["source"]) <= 24 else 1, len(item["source"]), item["source"]))
    return values[:limit]


def _glossary_ai_supplement_prompt(
    *,
    project: dict[str, Any],
    language: str,
    candidates: list[dict[str, Any]],
    audit_rows: list[dict[str, str]],
) -> str:
    spec = language_spec(language)
    profile = project.get("profile") or {}
    prompt_text = str((profile.get("prompts_by_language") or {}).get(language) or project.get("prompt_text") or "").strip()
    prompt_text = _manage_project_prompt_context(prompt_text, load_settings())
    candidate_sources = [
        {"cn": item.get("source"), "translation": item.get("target")}
        for item in candidates[:500]
        if str(item.get("source") or "").strip()
    ]
    packet = {
        "task": "language_table_glossary_ai_supplement",
        "target_language": spec.prompt_name,
        "visible_language": spec.target_header,
        "project": {
            "name": project.get("name", ""),
            "type": project.get("type", ""),
            "profile": {
                key: profile.get(key)
                for key in ("game_type", "target_audience", "content_scope", "translation_style", "tone")
                if profile.get(key)
            },
            "prompt": prompt_text,
        },
        "existing_candidates": candidate_sources,
        "source_rows_for_audit": audit_rows,
        "response_schema": {
            "supplement_terms": [
                {
                    "cn": "必须逐字出现在 source_rows_for_audit.source 中的中文术语",
                    "translation": f"{spec.target_header} 建议译文；不确定则留空",
                    "confidence": "medium|high",
                    "reason": "为什么本地规则漏掉且值得人工审核",
                    "evidence_ids": ["source_rows_for_audit.id"],
                }
            ]
        },
    }
    return (
        "你在做游戏语言表术语候选的漏词审计。本地脚本已经先按规则扫描了一批候选，"
        "你只补充明显漏掉的游戏专名、系统名、玩法名、道具/角色/活动名或关键 UI 术语。\n"
        "硬规则：\n"
        "1. 只能从 source_rows_for_audit 里补，cn 必须逐字出现在某条 source 中。\n"
        "2. 不要补整句、长句、纯标点、普通动词、泛泛说明词。\n"
        "3. 不要重复 existing_candidates 里已有的 CN。\n"
        "4. 有同一行译文证据时给 translation；不确定就留空，交给人工。\n"
        "5. 只返回 confidence 为 medium 或 high 的项，最多 30 条。\n"
        "6. 返回严格 JSON 对象，不要 markdown。\n\n"
        f"Packet JSON:\n{json.dumps(packet, ensure_ascii=False, indent=2)}"
    )


def _normalize_glossary_ai_supplement_terms(payload: dict[str, Any]) -> list[dict[str, Any]]:
    terms = payload.get("supplement_terms")
    if not isinstance(terms, list):
        return []
    return [item for item in terms if isinstance(item, dict)]


def supplement_language_table_glossary_candidates_with_ai(
    *,
    project_id: str,
    batch_id: str,
    input_path: Path,
    language: str,
    run_id: str | None = None,
) -> dict[str, Any]:
    language = require_supported_language(language)
    result: dict[str, Any] = {
        "status": "skipped",
        "added": 0,
        "reviewed_rows": 0,
        "provider": "",
        "reason": "",
    }
    if not batch_id:
        result["reason"] = "no_candidate_batch"
        return result

    settings = load_settings()
    provider = normalize_provider_name(settings.get("provider"))
    result["provider"] = provider
    if provider == TEST_FAKE_PROVIDER:
        result["reason"] = "test_provider"
        return result
    if provider in REAL_PROVIDERS and not str(settings.get("api_key") or "").strip():
        result["reason"] = "api_key_missing"
        if run_id:
            db.add_event(run_id, "AI glossary supplement skipped: API key is not configured.", level="warn")
        return result

    project = db.get_project(project_id)
    existing_sources = {
        _glossary_source_key(item.get("source"))
        for item in db.list_glossary_terms(project_id, language=language)
        if _glossary_source_key(item.get("source"))
    }
    batch_candidates = db.list_glossary_candidates(project_id, batch_id=batch_id, language=language)
    existing_sources.update(
        _glossary_source_key(item.get("source"))
        for item in batch_candidates
        if _glossary_source_key(item.get("source"))
    )
    audit_rows = _language_table_ai_audit_rows(_read_language_table_rows(input_path, [language]), language, existing_sources)
    result["reviewed_rows"] = len(audit_rows)
    if not audit_rows:
        result["reason"] = "no_audit_rows"
        return result

    prompt = _glossary_ai_supplement_prompt(
        project=project,
        language=language,
        candidates=batch_candidates,
        audit_rows=audit_rows,
    )
    try:
        response_text = call_text(settings, prompt, system="Return strict JSON only.")
        payload = _parse_semantic_qa_payload(response_text)
    except Exception as exc:
        result["status"] = "provider_error"
        result["reason"] = user_facing_error(exc)
        if run_id:
            db.add_event(run_id, f"AI glossary supplement failed: {result['reason']}", level="warn")
        return result

    evidence_by_id = {str(item.get("id") or "").strip(): item for item in audit_rows if str(item.get("id") or "").strip()}
    evidence_sources = [item["source"] for item in audit_rows]
    added = 0
    skipped = 0
    for term in _normalize_glossary_ai_supplement_terms(payload):
        cn = str(term.get("cn") or "").strip()
        source_key = _glossary_source_key(cn)
        confidence = str(term.get("confidence") or "").strip().lower()
        if not cn or source_key in existing_sources or confidence not in {"medium", "high"}:
            skipped += 1
            continue
        evidence_ids = term.get("evidence_ids") if isinstance(term.get("evidence_ids"), list) else []
        evidence_rows = [evidence_by_id.get(str(item).strip()) for item in evidence_ids]
        evidence_rows = [item for item in evidence_rows if item]
        if evidence_rows:
            appears = any(cn in str(item.get("source") or "") for item in evidence_rows)
        else:
            appears = any(cn in source for source in evidence_sources)
        if not appears:
            skipped += 1
            continue
        translation = str(term.get("translation") or "").strip()
        note_bits = ["AI 漏词补充候选，需人工确认"]
        if confidence:
            note_bits.append(f"置信度 {confidence}")
        reason = str(term.get("reason") or "").strip()
        if reason:
            note_bits.append(reason)
        db.add_glossary_candidate(
            project_id,
            batch_id,
            {
                "term_key": str((evidence_rows[0] or {}).get("id") or "") if evidence_rows else "",
                "source": cn,
                "target": translation,
                "target_alt": "",
                "language": language,
                "category": "",
                "note": "；".join(note_bits),
                "action": "new",
                "translation_status": "suggested" if translation else "needs_translation",
                "translation_source": "ai_supplement" if translation else "none",
                "metadata": {
                    "ai_supplement": {
                        "provider": provider,
                        "model": settings.get("model") or "",
                        "confidence": confidence,
                        "reason": reason,
                        "evidence_ids": evidence_ids,
                    }
                },
            },
        )
        existing_sources.add(source_key)
        added += 1

    result.update({"status": "passed", "added": added, "skipped": skipped, "reason": ""})
    if run_id:
        db.add_event(run_id, f"AI glossary supplement added {added} candidates, skipped {skipped}.")
    return result


async def translate_missing_glossary_candidates(project_id: str, batch_id: str) -> dict[str, Any]:
    project = db.get_project(project_id)
    batch = db.get_glossary_batch(batch_id)
    if batch["project_id"] != project_id:
        raise KeyError(batch_id)
    language = require_supported_language(batch.get("language") or "en")
    settings = load_settings()
    provider = normalize_provider_name(settings.get("provider"))
    if provider in REAL_PROVIDERS and not settings.get("api_key"):
        raise ValueError(f"{provider} api_key is required to translate glossary candidates")

    pending = db.list_glossary_candidates(project_id, batch_id=batch_id, status="pending", language=language)
    missing = [candidate for candidate in pending if not str(candidate.get("target") or "").strip()]
    if not missing:
        return {
            "batch": batch,
            "translated_count": 0,
            "skipped_count": len(pending),
            "candidates": db.list_glossary_candidates(project_id, batch_id=batch_id, language=language),
        }

    rows: list[dict[str, Any]] = []
    id_to_candidate: dict[int, dict[str, Any]] = {}
    for index, candidate in enumerate(missing, start=1):
        rows.append(
            {
                "id": index,
                "source": str(candidate.get("source") or ""),
                "text_type": "glossary_term_candidate",
                "term_key": str(candidate.get("term_key") or ""),
                "note": str(candidate.get("note") or ""),
            }
        )
        id_to_candidate[index] = candidate

    prompt = _glossary_candidate_translation_prompt(project, rows, language=language)
    try:
        items = await translate_batch(rows, settings, prompt)
    except Exception as exc:
        raise ValueError(f"glossary candidate translation failed: {user_facing_error(exc)}") from exc
    translated_count = 0
    for item in items:
        candidate = id_to_candidate.get(int(item.id))
        if candidate is None:
            continue
        translation = str(item.translation or "").strip()
        if not translation:
            continue
        metadata = dict(candidate.get("metadata") or {})
        metadata["model"] = {
            "provider": provider,
            "model": settings.get("model") or "",
            "preset": settings.get("preset") or "",
            "translated_at": db.now_iso(),
        }
        db.update_glossary_candidate(
            candidate["id"],
            {
                "target": translation,
                "translation_status": "suggested",
                "translation_source": "model",
                "metadata": metadata,
            },
        )
        translated_count += 1

    if batch.get("run_id"):
        db.add_event(batch["run_id"], f"Glossary candidate translation filled {translated_count} missing {language.upper()} values.")
    return {
        "batch": db.get_glossary_batch(batch_id),
        "translated_count": translated_count,
        "skipped_count": len(pending) - len(missing),
        "candidates": db.list_glossary_candidates(project_id, batch_id=batch_id, language=language),
    }


def translate_missing_glossary_candidates_sync(project_id: str, batch_id: str) -> dict[str, Any]:
    return asyncio.run(translate_missing_glossary_candidates(project_id, batch_id))


def _glossary_candidate_translation_prompt(project: dict[str, Any], rows: list[dict[str, Any]], language: str = "en") -> str:
    language = require_supported_language(language)
    spec = language_spec(language)
    profile = project.get("profile") or {}
    prompt_text = str((profile.get("prompts_by_language") or {}).get(language) or project.get("prompt_text") or "").strip()
    prompt_text = _manage_project_prompt_context(prompt_text, load_settings())
    profile_summary = {
        key: profile.get(key)
        for key in ("game_type", "target_audience", "content_scope", "translation_style", "tone", "language_assets")
        if profile.get(key)
    }
    existing_terms = [
        {"source": term.get("source"), "target": term.get("target"), "target_alt": term.get("target_alt")}
        for term in db.list_glossary_terms(project["id"], language=language)[:200]
        if str(term.get("source") or "").strip() and str(term.get("target") or "").strip()
    ]
    term_instruction = (
        f"Translate only the {spec.target_header} term. Do not create {spec.alt_header}, notes, categories, explanations, or markdown. "
        if spec.alt_header
        else f"Translate only the {spec.target_header} term. Do not create notes, categories, explanations, or markdown. "
    )
    return (
        f"Translate short game glossary term candidates from Chinese to {spec.prompt_name}. "
        "Return JSONL only; each line must be {\"id\": number, \"translation\": string}. "
        f"{term_instruction}"
        "Keep UI terms concise and consistent with existing glossary.\n\n"
        f"Project: {project.get('name', '')}\n"
        f"Profile: {json.dumps(profile_summary, ensure_ascii=False)}\n"
        f"Project prompt:\n{prompt_text}\n\n"
        f"Existing glossary examples:\n{json.dumps(existing_terms, ensure_ascii=False)}\n\n"
        f"Candidates:\n" + "\n".join(json.dumps(row, ensure_ascii=False) for row in rows)
    )


def _glossary_source_key(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").strip()).casefold()


def _glossary_term_rank(term: dict[str, Any]) -> tuple[int, int, int]:
    has_translation = bool(str(term.get("target") or "").strip() or str(term.get("target_alt") or "").strip())
    confirmed = bool(term.get("confirmed"))
    curated_source = str(term.get("source_type") or "") in {"manual", "imported", "curated"}
    return (0 if confirmed else 1, 0 if has_translation else 1, 0 if curated_source else 1)


def _fill_blank_glossary_fields(base: dict[str, Any], incoming: dict[str, Any]) -> None:
    for field in ("target", "target_alt", "category", "note"):
        if not str(base.get(field) or "").strip() and str(incoming.get(field) or "").strip():
            base[field] = incoming.get(field, "")


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.write_text("\n".join(json.dumps(row, ensure_ascii=False) for row in rows) + "\n", encoding="utf-8")


COMPLETE_LANGUAGE_TABLE_GLOSSARY_IMPORT_MESSAGE = (
    "这是完整语言表，不是项目术语表。请在「语言表/待翻译内容」上传后，"
    "到 STEP5「高频词扫描 / 从完整语言表扫描高频术语候选」生成候选术语；"
    "候选经人工确认后才会进入项目术语库。"
)
COMPLETE_LANGUAGE_TABLE_PROJECT_MATERIAL_MESSAGE = (
    "这个文件看起来是完整语言表，请上传到 STEP4「语言表」。"
    "它不会作为项目资料参与术语提取。"
)
_LARGE_LANGUAGE_TABLE_ROW_THRESHOLD = 1000


def is_complete_language_table_for_glossary_import(path: Path, sheet: str | None = None, row_threshold: int = _LARGE_LANGUAGE_TABLE_ROW_THRESHOLD) -> bool:
    if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return False
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[sheet] if sheet else wb[wb.sheetnames[0]]
        header_row = next(ws.iter_rows(min_row=1, max_row=1), None)
        if header_row is None:
            return False
        headers = [str(cell.value or "").strip() for cell in header_row]
        normalized = _normalized_header_indices(headers)
        term_key_idx = _column_index(normalized, None, ["id", "key", "编号", "序号"], required=False)
        source_idx = _column_index(normalized, None, ["source", "original", "cn", "zh", "chinese", "原文", "中文", "简体中文"], required=False)
        if term_key_idx is None or source_idx is None:
            return False
        reserved = {term_key_idx, source_idx}
        language_indices = _auto_language_indices(headers, reserved)
        if not language_indices:
            return False
        source_rows = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            if _value_at(row, source_idx):
                source_rows += 1
                if source_rows > row_threshold:
                    return True
        return False
    finally:
        wb.close()


def guard_complete_language_table_for_glossary_import(path: Path, sheet: str | None = None) -> None:
    if is_complete_language_table_for_glossary_import(path, sheet=sheet):
        raise ValueError(COMPLETE_LANGUAGE_TABLE_GLOSSARY_IMPORT_MESSAGE)


def guard_complete_language_table_for_project_material(path: Path, sheet: str | None = None) -> None:
    if is_complete_language_table_for_glossary_import(path, sheet=sheet):
        raise ValueError(COMPLETE_LANGUAGE_TABLE_PROJECT_MATERIAL_MESSAGE)


def preview_glossary_import(project_id: str, request: Any, import_all: bool = False) -> dict[str, Any]:
    project = db.get_project(project_id)
    _ = project
    artifact = db.get_artifact(request.artifact_id)
    path = Path(artifact["path"])
    guard_complete_language_table_for_glossary_import(path, sheet=getattr(request, "sheet", None))
    language = require_supported_language(getattr(request, "language", "en") or "en")
    auto_languages = bool(getattr(request, "auto_languages", True))
    if auto_languages and not getattr(request, "target_column", None) and not getattr(request, "target_alt_column", None):
        rows, columns, languages = _read_multilingual_glossary_rows(
            path,
            sheet=getattr(request, "sheet", None),
            term_key_column=getattr(request, "term_key_column", None),
            source_column=getattr(request, "source_column", None),
            category_column=getattr(request, "category_column", None),
            note_column=getattr(request, "note_column", None),
            limit=None if import_all else int(getattr(request, "limit", 100) or 100),
        )
        if languages:
            return {"artifact": artifact, "columns": columns, "rows": rows, "total_rows": len(rows), "language": "auto", "languages": languages}
    rows, columns = _read_glossary_rows(
        path,
        sheet=getattr(request, "sheet", None),
        term_key_column=getattr(request, "term_key_column", None),
        source_column=getattr(request, "source_column", None),
        target_column=getattr(request, "target_column", None),
        target_alt_column=getattr(request, "target_alt_column", None),
        category_column=getattr(request, "category_column", None),
        note_column=getattr(request, "note_column", None),
        language=language,
        limit=None if import_all else int(getattr(request, "limit", 100) or 100),
    )
    return {"artifact": artifact, "columns": columns, "rows": rows, "total_rows": len(rows), "language": language}


def import_glossary(project_id: str, request: Any) -> dict[str, Any]:
    preview = preview_glossary_import(project_id, request, import_all=True)
    language = preview["language"]
    payloads = []
    for row in preview["rows"]:
        if not row.get("source"):
            continue
        payloads.append(
            {
                "term_key": row.get("term_key", ""),
                "source": row.get("source", ""),
                "target": row.get("target", ""),
                "target_alt": row.get("target_alt", ""),
                "language": row.get("language") or language,
                "category": row.get("category", ""),
                "note": row.get("note", ""),
                "source_type": "imported",
                "confirmed": True,
            }
        )
    imported = db.upsert_glossary_terms_bulk(project_id, payloads)
    return {"imported_count": len(imported), "terms": imported, "preview": preview, "languages": preview.get("languages") or ([language] if language != "auto" else [])}


def export_glossary(project_id: str, fmt: str, language: str | None = None) -> dict[str, Any] | Path:
    project = db.get_project(project_id)
    language = require_supported_language(language or "en") if language else None
    terms = db.list_glossary_terms(project_id, language=language)
    if fmt == "json":
        return {
            "project_id": project_id,
            "language": language,
            "terms": [dict(zip(("term_key", "source", "target", "target_alt", "category", "note"), _glossary_export_row(term))) | {"language": term.get("language", "en")} for term in terms],
        }
    output_dir = project_dir(project_id) / "glossary" / "exports"
    output_dir.mkdir(parents=True, exist_ok=True)
    suffix = _export_language_suffix(language)
    if language:
        columns = ["ID", "CN", _visible_language_code(language), *(["EN2"] if language == "en" else []), "分类", "备注"]
        rows = [_glossary_export_row(term, include_alt=language == "en") for term in terms]
    else:
        wide = list_glossary_wide(project_id)
        languages = list(wide.get("languages") or [])
        columns = ["ID", "CN", *_wide_language_columns(languages), "分类", "备注"]
        rows = _glossary_wide_export_rows(wide, languages)
    if fmt == "csv":
        path = output_dir / _export_filename(project, "glossary", suffix, "csv")
        with path.open("w", encoding="utf-8-sig", newline="") as fh:
            writer = csv.writer(fh)
            writer.writerow(columns)
            writer.writerows(rows)
        return path
    path = output_dir / _export_filename(project, "glossary", suffix, "xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Glossary"
    ws.append(columns)
    for row in rows:
        ws.append(row)
    wb.save(path)
    wb.close()
    return path


def _export_language_suffix(language: str | None) -> str:
    return _visible_language_code(language) if language else "ALL"


def _export_filename(project: dict[str, Any], kind: str, suffix: str, ext: str) -> str:
    return f"{_safe_delivery_name(project['name'])}_{kind}_{suffix}_{_today_stamp()}.{ext}"


def _glossary_export_row(term: dict[str, Any], *, include_alt: bool = True) -> list[Any]:
    row = [
        term.get("term_key", ""),
        term.get("source", ""),
        term.get("target", ""),
    ]
    if include_alt:
        row.append(term.get("target_alt", ""))
    row.extend([term.get("category", ""), term.get("note", "")])
    return row


def _wide_language_columns(languages: list[str]) -> list[str]:
    columns: list[str] = []
    for code in languages:
        columns.append(_visible_language_code(code))
        if code == "en":
            columns.append("EN2")
    return columns


def _glossary_wide_export_rows(wide: dict[str, Any], languages: list[str]) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for row in wide["rows"]:
        translations = row.get("translations") or {}
        values = [row.get("term_key", ""), row.get("source", "")]
        for code in languages:
            entry = translations.get(code) or {}
            values.append(entry.get("target", ""))
            if code == "en":
                values.append(entry.get("target_alt", ""))
        values.extend([row.get("category", ""), row.get("note", "")])
        rows.append(values)
    return rows


def import_translation_archive(project_id: str, request: Any, source_type: str = "imported") -> dict[str, Any]:
    language = require_supported_language(getattr(request, "language", "en") or "en")
    artifact = db.get_artifact(request.artifact_id)
    if artifact["project_id"] != project_id:
        raise KeyError("artifact")
    if bool(getattr(request, "auto_languages", True)) and not getattr(request, "target_column", None) and not getattr(request, "target_alt_column", None):
        rows = _read_multilingual_translation_rows(
            Path(artifact["path"]),
            sheet=getattr(request, "sheet", None),
            id_column=getattr(request, "id_column", None),
            source_column=getattr(request, "source_column", None),
            note_column=getattr(request, "note_column", None),
            source_artifact_id=artifact["id"],
            source_type=source_type,
        )
        if not rows:
            rows = _read_translation_rows(
                Path(artifact["path"]),
                sheet=getattr(request, "sheet", None),
                id_column=getattr(request, "id_column", None),
                source_column=getattr(request, "source_column", None),
                target_column=getattr(request, "target_column", None),
                target_alt_column=getattr(request, "target_alt_column", None),
                note_column=getattr(request, "note_column", None),
                language=language,
                source_artifact_id=artifact["id"],
                source_type=source_type,
            )
    else:
        rows = _read_translation_rows(
            Path(artifact["path"]),
            sheet=getattr(request, "sheet", None),
            id_column=getattr(request, "id_column", None),
            source_column=getattr(request, "source_column", None),
            target_column=getattr(request, "target_column", None),
            target_alt_column=getattr(request, "target_alt_column", None),
            note_column=getattr(request, "note_column", None),
            language=language,
            source_artifact_id=artifact["id"],
            source_type=source_type,
        )
    imported = db.upsert_translation_entries_bulk(project_id, [row for row in rows if row.get("source") or row.get("target")])
    languages = [code for code in LANGUAGE_ORDER if any(row.get("language") == code for row in rows)]
    return {"project_id": project_id, "artifact_id": artifact["id"], "imported_count": len(imported), "entries": imported, "languages": languages or [language]}


def archive_translation_artifact(project_id: str, artifact_id: str, language: str = "en", source_type: str = "qa_passed") -> dict[str, Any]:
    class Request:
        pass

    request = Request()
    request.artifact_id = artifact_id
    request.language = language
    request.sheet = None
    request.id_column = None
    request.source_column = None
    request.target_column = None
    request.target_alt_column = None
    request.note_column = None
    return import_translation_archive(project_id, request, source_type=source_type)


def export_translation_archive(project_id: str, fmt: str, language: str | None = None) -> dict[str, Any] | Path:
    project = db.get_project(project_id)
    language = require_supported_language(language or "en") if language else None
    entries = db.list_translation_entries(project_id, language=language)
    if fmt == "json":
        return {"project_id": project_id, "language": language, "entries": [_translation_export_payload(entry) for entry in entries]}
    output_dir = project_dir(project_id) / "translations" / "exports"
    output_dir.mkdir(parents=True, exist_ok=True)
    suffix = _export_language_suffix(language)
    if language:
        columns = ["ID", "CN", _visible_language_code(language), *(["EN2"] if language == "en" else []), "备注"]
        rows = [_translation_export_row(entry, include_alt=language == "en") for entry in entries]
    else:
        wide = list_translation_archive_wide(project_id)
        languages = list(wide.get("languages") or [])
        columns = ["ID", "CN", *_wide_language_columns(languages), "备注"]
        rows = _translation_wide_export_rows(wide, languages)
    if fmt == "csv":
        path = output_dir / _export_filename(project, "translations", suffix, "csv")
        with path.open("w", encoding="utf-8-sig", newline="") as fh:
            writer = csv.writer(fh)
            writer.writerow(columns)
            writer.writerows(rows)
        return path
    path = output_dir / _export_filename(project, "translations", suffix, "xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = "Translations"
    ws.append(columns)
    for row in rows:
        ws.append(row)
    wb.save(path)
    wb.close()
    return path


def list_glossary_wide(project_id: str) -> dict[str, Any]:
    db.get_project(project_id)
    rows = _wide_rows(
        db.list_glossary_terms(project_id),
        key_field="term_key",
        shared_fields=("term_key", "category", "note"),
    )
    return {"project_id": project_id, **rows}


def list_translation_archive_wide(project_id: str) -> dict[str, Any]:
    db.get_project(project_id)
    rows = _wide_rows(
        db.list_translation_entries(project_id),
        key_field="entry_key",
        shared_fields=("entry_key", "note"),
    )
    return {"project_id": project_id, **rows}


def _wide_rows(items: list[dict[str, Any]], *, key_field: str, shared_fields: tuple[str, ...]) -> dict[str, Any]:
    grouped: dict[str, list[dict[str, Any]]] = {}
    for item in items:
        source_key = _wide_source_key(item.get("source"))
        if not source_key:
            continue
        grouped.setdefault(source_key, []).append(item)

    wide_rows: list[dict[str, Any]] = []
    coverage: dict[str, int] = {}
    for source_key, group in grouped.items():
        translations: dict[str, dict[str, Any]] = {}
        for code in LANGUAGE_ORDER:
            candidates = [item for item in group if normalize_language(item.get("language") or "en") == code and (str(item.get("target") or "").strip() or str(item.get("target_alt") or "").strip())]
            if not candidates:
                continue
            selected = sorted(candidates, key=lambda item: str(item.get("updated_at") or ""), reverse=True)[0]
            payload = {
                "id": selected.get("id", ""),
                "language": code,
                "target": selected.get("target", ""),
                "target_alt": selected.get("target_alt", ""),
            }
            translations[code] = payload
            coverage[code] = coverage.get(code, 0) + 1
        shared = {field: _first_non_blank(group, field) for field in shared_fields}
        wide_rows.append(
            {
                "source_key": source_key,
                "source": _first_non_blank(group, "source"),
                **shared,
                "translations": translations,
                "languages": [code for code in LANGUAGE_ORDER if code in translations],
                "conflicts": _wide_conflicts(group, ("source", *shared_fields)),
            }
        )

    languages = [code for code in LANGUAGE_ORDER if coverage.get(code, 0) > 0]
    wide_rows.sort(key=lambda row: (str(row.get("source") or ""), str(row.get(key_field) or "")))
    return {"languages": languages, "coverage": {code: coverage[code] for code in languages}, "row_count": len(wide_rows), "rows": wide_rows}


def _wide_source_key(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").strip()).casefold()


def _first_non_blank(rows: list[dict[str, Any]], field: str) -> str:
    for row in rows:
        value = str(row.get(field) or "").strip()
        if value and value != "-":
            return value
    return ""


def _wide_conflicts(rows: list[dict[str, Any]], fields: tuple[str, ...]) -> list[dict[str, Any]]:
    conflicts: list[dict[str, Any]] = []
    for field in fields:
        values: list[str] = []
        for row in rows:
            value = str(row.get(field) or "").strip()
            if value and value != "-" and value not in values:
                values.append(value)
        if len(values) > 1:
            conflicts.append({"field": field, "values": values})
    return conflicts


def _normalized_header_indices(headers: list[str]) -> dict[str, int]:
    normalized: dict[str, int] = {}
    for index, header in enumerate(headers):
        key = str(header or "").strip().lower()
        if key and key not in normalized:
            normalized[key] = index
    return normalized


def _auto_language_indices(headers: list[str], reserved_indices: set[int] | None = None) -> dict[str, tuple[int, int | None]]:
    reserved = reserved_indices or set()
    normalized_headers: dict[str, int] = {}
    for index, header in enumerate(headers):
        if index in reserved:
            continue
        key = str(header or "").strip().lower()
        if key:
            normalized_headers[key] = index
    detected: dict[str, tuple[int, int | None]] = {}
    for code in LANGUAGE_ORDER:
        target_idx = _column_index(normalized_headers, None, list(AUTO_LANGUAGE_TARGET_ALIASES[code]), required=False)
        if target_idx is None:
            continue
        alt_idx = None
        if code == "en":
            alt_idx = _column_index(normalized_headers, None, list(AUTO_LANGUAGE_ALT_ALIASES["en"]), required=False)
        detected[code] = (target_idx, alt_idx)
    return detected


def _read_multilingual_glossary_rows(
    path: Path,
    sheet: str | None = None,
    term_key_column: str | None = None,
    source_column: str | None = None,
    category_column: str | None = None,
    note_column: str | None = None,
    limit: int | None = 100,
) -> tuple[list[dict[str, Any]], dict[str, Any], list[str]]:
    if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return [], {}, []
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[sheet] if sheet else wb[wb.sheetnames[0]]
        headers = [str(cell.value or "").strip() for cell in next(ws.iter_rows(min_row=1, max_row=1))]
        normalized = _normalized_header_indices(headers)
        term_key_idx = _column_index(normalized, term_key_column, ["id", "key", "编号", "序号"], required=False)
        source_idx = _column_index(normalized, source_column, ["source", "original", "cn", "zh", "chinese", "term", "原文", "中文", "术语"])
        category_idx = _column_index(normalized, category_column, ["category", "type", "分类", "类别", "类型"], required=False)
        note_idx = _column_index(normalized, note_column, ["note", "notes", "comment", "备注"], required=False)
        reserved = {index for index in (term_key_idx, source_idx, category_idx, note_idx) if index is not None}
        language_indices = _auto_language_indices(headers, reserved)
        if not language_indices:
            return [], {}, []
        rows: list[dict[str, Any]] = []
        source_rows = 0
        for row in ws.iter_rows(min_row=2, values_only=True):
            source = _value_at(row, source_idx)
            if not source:
                continue
            source_rows += 1
            if limit is not None and source_rows > limit:
                break
            for code, (target_idx, alt_idx) in language_indices.items():
                target = _value_at(row, target_idx)
                target_alt = _value_at(row, alt_idx) if code == "en" else ""
                if not target and not target_alt:
                    continue
                rows.append(
                    {
                        "term_key": _value_at(row, term_key_idx) if term_key_idx is not None else "",
                        "source": source,
                        "target": target,
                        "target_alt": target_alt,
                        "language": code,
                        "category": _value_at(row, category_idx) if category_idx is not None else "",
                        "note": _value_at(row, note_idx) if note_idx is not None else "",
                    }
                )
        return rows, {
            "term_key": headers[term_key_idx] if term_key_idx is not None else "",
            "source": headers[source_idx],
            "languages": {code: {"target": headers[target_idx], "target_alt": headers[alt_idx] if alt_idx is not None else ""} for code, (target_idx, alt_idx) in language_indices.items()},
            "category": headers[category_idx] if category_idx is not None else "",
            "note": headers[note_idx] if note_idx is not None else "",
        }, [code for code in LANGUAGE_ORDER if code in language_indices and any(row.get("language") == code for row in rows)]
    finally:
        wb.close()


def _read_multilingual_translation_rows(
    path: Path,
    sheet: str | None = None,
    id_column: str | None = None,
    source_column: str | None = None,
    note_column: str | None = None,
    source_artifact_id: str = "",
    source_type: str = "imported",
) -> list[dict[str, Any]]:
    if path.suffix.lower() not in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        return []
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[sheet] if sheet else wb[wb.sheetnames[0]]
        headers = [str(cell.value or "").strip() for cell in next(ws.iter_rows(min_row=1, max_row=1))]
        normalized = _normalized_header_indices(headers)
        id_idx = _column_index(normalized, id_column, ["id", "key", "编号", "序号"], required=False)
        source_idx = _column_index(normalized, source_column, ["source", "original", "cn", "zh", "chinese", "原文", "中文"])
        note_idx = _column_index(normalized, note_column, ["note", "notes", "comment", "备注"], required=False)
        reserved = {index for index in (id_idx, source_idx, note_idx) if index is not None}
        language_indices = _auto_language_indices(headers, reserved)
        if not language_indices:
            return []
        rows: list[dict[str, Any]] = []
        for row_index, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            source = _value_at(row, source_idx)
            if not source:
                continue
            for code, (target_idx, alt_idx) in language_indices.items():
                target = _value_at(row, target_idx)
                target_alt = _value_at(row, alt_idx) if code == "en" else ""
                if not target and not target_alt:
                    continue
                rows.append(
                    {
                        "entry_key": _value_at(row, id_idx) if id_idx is not None else "",
                        "source": source,
                        "target": target,
                        "target_alt": target_alt,
                        "language": code,
                        "sheet": ws.title,
                        "row_number": row_index,
                        "note": _value_at(row, note_idx) if note_idx is not None else "",
                        "source_type": source_type,
                        "source_artifact_id": source_artifact_id,
                    }
                )
        return rows
    finally:
        wb.close()


def _read_translation_rows(
    path: Path,
    sheet: str | None = None,
    id_column: str | None = None,
    source_column: str | None = None,
    target_column: str | None = None,
    target_alt_column: str | None = None,
    note_column: str | None = None,
    language: str = "en",
    source_artifact_id: str = "",
    source_type: str = "imported",
) -> list[dict[str, Any]]:
    language = require_supported_language(language)
    if path.suffix.lower() == ".json":
        payload = json.loads(path.read_text(encoding="utf-8"))
        raw_rows = payload.get("entries") if isinstance(payload, dict) else payload
        rows = []
        for row in (raw_rows or []):
            if not isinstance(row, dict):
                continue
            normalized = {str(key or "").strip().lower(): value for key, value in row.items()}
            rows.append(_translation_row_from_mapping(normalized, int(row.get("row_number") or 0), str(row.get("sheet") or "").strip(), language, source_artifact_id, source_type))
        return rows
    if path.suffix.lower() == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as fh:
            reader = csv.DictReader(fh)
            rows = []
            for index, row in enumerate(reader, start=2):
                normalized = {str(key or "").strip().lower(): value for key, value in row.items()}
                rows.append(_translation_row_from_mapping(normalized, index, "", language, source_artifact_id, source_type))
            return rows

    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[sheet] if sheet else wb[wb.sheetnames[0]]
        headers = [str(cell.value or "").strip() for cell in next(ws.iter_rows(min_row=1, max_row=1))]
        normalized = {header.lower(): index for index, header in enumerate(headers) if header}
        id_idx = _column_index(normalized, id_column, ["id", "key", "编号", "序号"], required=False)
        source_idx = _column_index(normalized, source_column, ["source", "original", "cn", "zh", "chinese", "原文", "中文"])
        target_idx = _column_index(normalized, target_column, target_aliases(language))
        target_alt_idx = _column_index(normalized, target_alt_column, alt_aliases(language), required=False)
        note_idx = _column_index(normalized, note_column, ["note", "notes", "comment", "备注"], required=False)
        rows = []
        for row_index, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            source = _value_at(row, source_idx)
            target = _value_at(row, target_idx)
            if not source and not target:
                continue
            rows.append(
                {
                    "entry_key": _value_at(row, id_idx) if id_idx is not None else "",
                    "source": source,
                    "target": target,
                    "target_alt": _value_at(row, target_alt_idx) if target_alt_idx is not None else "",
                    "language": language,
                    "sheet": ws.title,
                    "row_number": row_index,
                    "note": _value_at(row, note_idx) if note_idx is not None else "",
                    "source_type": source_type,
                    "source_artifact_id": source_artifact_id,
                }
            )
        return rows
    finally:
        wb.close()


def _translation_row_from_mapping(
    row: dict[str, Any],
    row_number: int,
    sheet: str,
    language: str,
    source_artifact_id: str,
    source_type: str,
) -> dict[str, Any]:
    language = require_supported_language(language)
    def pick(*names: str) -> str:
        for name in names:
            value = row.get(name.lower())
            if value not in (None, ""):
                return str(value).strip()
        return ""

    return {
        "entry_key": pick("id", "key", "entry_key", "编号", "序号"),
        "source": pick("cn", "source", "original", "原文", "中文"),
        "target": pick("target", *target_aliases(language)),
        "target_alt": pick("target_alt", *alt_aliases(language)),
        "language": language,
        "sheet": sheet,
        "row_number": row_number,
        "note": pick("note", "notes", "comment", "备注"),
        "source_type": source_type,
        "source_artifact_id": source_artifact_id,
    }


def _translation_export_payload(entry: dict[str, Any]) -> dict[str, Any]:
    return {
        "entry_key": entry.get("entry_key", ""),
        "source": entry.get("source", ""),
        "target": entry.get("target", ""),
        "target_alt": entry.get("target_alt", ""),
        "language": entry.get("language", "en"),
        "note": entry.get("note", ""),
    }


def _translation_export_row(entry: dict[str, Any], *, include_alt: bool = True) -> list[Any]:
    row = [
        entry.get("entry_key", ""),
        entry.get("source", ""),
        entry.get("target", ""),
    ]
    if include_alt:
        row.append(entry.get("target_alt", ""))
    row.append(entry.get("note", ""))
    return row


def _translation_wide_export_rows(wide: dict[str, Any], languages: list[str]) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for row in wide["rows"]:
        translations = row.get("translations") or {}
        values = [row.get("entry_key", ""), row.get("source", "")]
        for code in languages:
            entry = translations.get(code) or {}
            values.append(entry.get("target", ""))
            if code == "en":
                values.append(entry.get("target_alt", ""))
        values.append(row.get("note", ""))
        rows.append(values)
    return rows


def list_project_deliverables(project_id: str) -> list[dict[str, Any]]:
    project = db.get_project(project_id)
    deliverables: list[dict[str, Any]] = []
    for run in db.list_runs(project_id):
        if run["kind"] not in {"translation", "qa"} or run["status"] != "passed":
            continue
        final_artifact = _deliverable_final_artifact(run)
        if not final_artifact or not Path(final_artifact["path"]).exists():
            continue
        deliverables.append(_deliverable_summary(project, run, final_artifact))
    deliverables.extend(_announcement_deliverable_summaries(project))
    return deliverables


def _announcement_deliverable_summaries(project: dict[str, Any]) -> list[dict[str, Any]]:
    summaries: list[dict[str, Any]] = []
    for task in db.list_announcement_tasks(project["id"]):
        if task.get("status") != "delivered":
            continue
        metadata = task.get("metadata") or {}
        delivery_artifact_id = str(metadata.get("delivery_artifact_id") or "")
        if not delivery_artifact_id:
            continue
        try:
            package_artifact = db.get_artifact(delivery_artifact_id)
        except KeyError:
            continue
        if (package_artifact.get("metadata") or {}).get("superseded"):
            continue
        if not Path(package_artifact["path"]).exists():
            continue
        languages = _normalize_announcement_languages(task.get("selected_languages") or [], fallback=metadata.get("languages") or [])
        output_files = []
        output_artifact_ids = metadata.get("output_artifact_ids") if isinstance(metadata.get("output_artifact_ids"), dict) else {}
        for language in languages:
            artifact_id = str(output_artifact_ids.get(language) or "")
            if not artifact_id:
                continue
            try:
                artifact = db.get_artifact(artifact_id)
            except KeyError:
                continue
            if Path(artifact["path"]).exists():
                output_files.append(_artifact_delivery_file(f"output_{_visible_language_code(language)}", artifact))
        qa_file = None
        qa_artifact_id = str(metadata.get("qa_summary_artifact_id") or "")
        if qa_artifact_id:
            try:
                qa_artifact = db.get_artifact(qa_artifact_id)
                if Path(qa_artifact["path"]).exists():
                    qa_file = _artifact_delivery_file("qa_summary", qa_artifact)
            except KeyError:
                qa_file = None
        package_file = _artifact_delivery_file("package", package_artifact)
        language_label = " / ".join(_visible_language_code(language) for language in languages) or "-"
        source_rows = int(metadata.get("segment_count") or metadata.get("terms_count") or metadata.get("term_count") or 0)
        return_label = task.get("title") or _announcement_task_source_stem(task)
        summaries.append(
            {
                "run_id": package_artifact.get("run_id") or task["id"],
                "task_code": "ANN",
                "task_id": _short_run_id(task["id"]),
                "task_label": f"ANN-{_short_run_id(task['id'])}",
                "task_type": "公告任务",
                "language": language_label,
                "created_at": task.get("created_at", ""),
                "updated_at": task.get("updated_at", ""),
                "status": "delivered",
                "processed_rows": source_rows,
                "source_rows": source_rows,
                "translated_rows": source_rows,
                "provider": "-",
                "model": "-",
                "input_label": return_label,
                "qa_status": "passed",
                "qa_hard_errors": int(metadata.get("hard_blockers") or 0),
                "qa_soft_warnings": 0,
                "files": {
                    "package": package_file,
                    "qa_summary": qa_file,
                    "outputs": output_files,
                },
                "source_artifacts": {
                    "announcement_delivery_package": package_artifact["id"],
                    "announcement_outputs": [item.get("artifact_id") for item in output_files if item.get("artifact_id")],
                    "announcement_qa_summary": qa_artifact_id,
                },
            }
        )
    return summaries


def build_delivery_package(project_id: str, run_id: str | None = None) -> dict[str, Any]:
    project = db.get_project(project_id)
    deliverables = list_project_deliverables(project_id)
    if not deliverables:
        raise ValueError("QA 未通过，暂无最终交付 workbook")
    selected = deliverables[0]
    if run_id:
        selected = next((item for item in deliverables if item["run_id"] == run_id), None)
        if not selected:
            raise ValueError("指定任务未通过 QA，暂无最终交付")
    run = db.get_run(selected["run_id"])
    final_source = _deliverable_final_artifact(run)
    if not final_source or not Path(final_source["path"]).exists():
        raise ValueError("暂无最终交付文件")
    changes_source = _run_artifact(run["id"], "qa_changes")

    output_dir = project_dir(project_id) / "delivery"
    output_dir.mkdir(parents=True, exist_ok=True)
    if final_source["kind"] == "final_text":
        final_path = _delivery_final_output_path(project, run, final_source)
        shutil.copy2(final_source["path"], final_path)
        summary = _deliverable_summary(project, run, final_source)
        summary["files"] = {"final": _delivery_file("final", final_path)}
        return {"project_id": project_id, "project_name": project["name"], "deliverable": summary, "files": list(summary["files"].values())}

    final_path, changes_path = _delivery_output_paths(project, run)
    shutil.copy2(final_source["path"], final_path)
    _normalize_delivery_workbook_headers(final_path, run.get("language") or "en")
    if changes_source and Path(changes_source["path"]).exists():
        shutil.copy2(changes_source["path"], changes_path)
    else:
        empty_changes = write_qa_changes_report(output_dir, [])
        empty_changes.replace(changes_path)

    summary = _deliverable_summary(project, run, final_source)
    summary["files"] = {
        "final": _delivery_file("final", final_path),
        "changes": _delivery_file("changes", changes_path),
    }
    return {"project_id": project_id, "project_name": project["name"], "deliverable": summary, "files": list(summary["files"].values())}


def _deliverable_summary(project: dict[str, Any], run: dict[str, Any], final_artifact: dict[str, Any]) -> dict[str, Any]:
    changes_artifact = _run_artifact(run["id"], "qa_changes")
    final_path, changes_path = _delivery_output_paths(project, run)
    if final_artifact["kind"] == "final_text":
        final_path = _delivery_final_output_path(project, run, final_artifact)
    task_code, task_run_id = _effective_task_identity(run)
    metadata = run.get("metadata", {})
    quality_summary = metadata.get("quality_summary") or {}
    provider, model = _deliverable_provider_model(metadata, quality_summary)
    input_label = _input_artifact_label(run, run["project_id"])
    processed = (
        {"processed_rows": int(metadata.get("translated_rows") or metadata.get("source_rows") or 0), "source_rows": int(metadata.get("source_rows") or 0), "translated_rows": int(metadata.get("translated_rows") or 0)}
        if final_artifact["kind"] == "final_text"
        else _workbook_processed_rows(Path(final_artifact["path"]))
    )
    files = {"final": _delivery_file("final", final_path) if final_path.exists() else _expected_delivery_file("final", final_path)}
    if final_artifact["kind"] != "final_text":
        files["changes"] = _delivery_file("changes", changes_path) if changes_path.exists() else _expected_delivery_file("changes", changes_path)
    return {
        "run_id": run["id"],
        "task_code": task_code,
        "task_id": _short_run_id(task_run_id),
        "task_label": f"{task_code}-{_short_run_id(task_run_id)}",
        "task_type": _task_type_label(task_code),
        "language": _visible_language_code(run.get("language") or "en"),
        "created_at": run.get("created_at", ""),
        "updated_at": run.get("updated_at", ""),
        "status": run.get("status", ""),
        "processed_rows": processed["processed_rows"] or int(metadata.get("translated_rows") or 0),
        "source_rows": processed["source_rows"],
        "translated_rows": processed["translated_rows"],
        "provider": provider,
        "model": model,
        "input_label": input_label,
        "qa_status": "passed" if quality_summary.get("passed", run.get("status") == "passed") else "failed",
        "qa_hard_errors": int(quality_summary.get("hard_errors") or 0),
        "qa_soft_warnings": _soft_warning_count(quality_summary),
        "files": files,
        "source_artifacts": {
            "qa_final_workbook": final_artifact["id"] if final_artifact["kind"] == "qa_final_workbook" else "",
            "final_text": final_artifact["id"] if final_artifact["kind"] == "final_text" else "",
            "qa_changes": changes_artifact["id"] if changes_artifact else "",
        },
    }


def _deliverable_final_artifact(run: dict[str, Any]) -> dict[str, Any] | None:
    return _run_artifact(run["id"], "qa_final_workbook") or _run_artifact(run["id"], "final_text")


def _effective_task_identity(run: dict[str, Any], seen: set[str] | None = None) -> tuple[str, str]:
    seen = seen or set()
    if run["id"] in seen:
        return _fallback_task_code(run), run["id"]
    seen.add(run["id"])
    metadata = run.get("metadata", {})
    source_run_id = metadata.get("manual_fix_source_run_id") or metadata.get("model_fix_source_run_id") or metadata.get("source_run_id")
    if source_run_id:
        try:
            source_run = db.get_run(str(source_run_id))
            if source_run["project_id"] == run["project_id"]:
                source_code, source_id = _effective_task_identity(source_run, seen)
                if run["kind"] == "qa" and source_run["kind"] in {"translation", "qa"}:
                    return source_code, source_id
        except KeyError:
            pass
    task_code = str(metadata.get("task_code") or "").upper()
    if task_code not in {"A", "T", "QA"}:
        task_code = _fallback_task_code(run)
    return task_code, run["id"]


def _fallback_task_code(run: dict[str, Any]) -> str:
    if run["kind"] == "translation":
        return "T"
    if run["kind"] == "qa":
        return "QA"
    return str(run["kind"] or "TASK").upper()


def _task_type_label(task_code: str) -> str:
    return {"A": "完整工作流", "T": "翻译任务", "QA": "校对任务"}.get(task_code, task_code)


def _short_run_id(run_id: str) -> str:
    return str(run_id).removeprefix("run_")[:6]


def _delivery_output_paths(project: dict[str, Any], run: dict[str, Any]) -> tuple[Path, Path]:
    output_dir = project_dir(project["id"]) / "delivery"
    output_dir.mkdir(parents=True, exist_ok=True)
    task_code, task_run_id = _effective_task_identity(run)
    timestamp = _delivery_timestamp(run.get("created_at", ""))
    language = _visible_language_code(run.get("language") or "en")
    prefix = f"{_safe_delivery_name(project['name'])}_{language}_{timestamp}_{task_code}-{_short_run_id(task_run_id)}"
    return output_dir / f"{prefix}_final.xlsx", output_dir / f"{prefix}_changes.xlsx"


def _delivery_final_output_path(project: dict[str, Any], run: dict[str, Any], source_artifact: dict[str, Any]) -> Path:
    output_dir = project_dir(project["id"]) / "delivery"
    output_dir.mkdir(parents=True, exist_ok=True)
    task_code, task_run_id = _effective_task_identity(run)
    timestamp = _delivery_timestamp(run.get("created_at", ""))
    language = _visible_language_code(run.get("language") or "en")
    suffix = Path(str(source_artifact.get("path") or "")).suffix.lower() or ".txt"
    prefix = f"{_safe_delivery_name(project['name'])}_{language}_{timestamp}_{task_code}-{_short_run_id(task_run_id)}"
    return output_dir / f"{prefix}_final{suffix}"


def _normalize_delivery_workbook_headers(path: Path, language: Any) -> None:
    code = require_supported_language(language or "en")
    target = _visible_language_code(code)
    aliases = {alias.strip().lower() for alias in target_aliases(code)}
    if not aliases:
        return
    wb = load_workbook(path)
    changed = False
    try:
        for ws in wb.worksheets:
            for cell in ws[1]:
                value = str(cell.value or "").strip()
                if value and value.lower() in aliases and value != target:
                    cell.value = target
                    changed = True
        if changed:
            wb.save(path)
    finally:
        wb.close()


def _delivery_timestamp(value: str) -> str:
    try:
        parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        if parsed.tzinfo is None:
            return parsed.strftime("%Y%m%d%H%M")
        return parsed.astimezone(ZoneInfo("Asia/Shanghai")).strftime("%Y%m%d%H%M")
    except Exception:
        return datetime.now(ZoneInfo("Asia/Shanghai")).strftime("%Y%m%d%H%M")


def _run_artifact(run_id: str, kind: str) -> dict[str, Any] | None:
    artifacts = [artifact for artifact in db.list_artifacts(run_id=run_id) if artifact["kind"] == kind]
    return artifacts[0] if artifacts else None


def _input_artifact_label(run: dict[str, Any], project_id: str, seen: set[str] | None = None) -> str:
    seen = seen or set()
    run_id = str(run.get("id") or "")
    if run_id:
        if run_id in seen:
            return "-"
        seen.add(run_id)
    metadata = run.get("metadata") or {}
    source_run_id = metadata.get("manual_fix_source_run_id") or metadata.get("model_fix_source_run_id") or metadata.get("source_run_id")
    if source_run_id:
        try:
            source_run = db.get_run(str(source_run_id))
            if source_run.get("project_id") == project_id:
                source_label = _input_artifact_label(source_run, project_id, seen)
                if source_label and source_label != "-":
                    return source_label
        except KeyError:
            pass
    input_artifacts = metadata.get("input_artifacts") if isinstance(metadata.get("input_artifacts"), dict) else {}
    candidates = [
        input_artifacts.get("source_workbook"),
        metadata.get("input_artifact_id"),
        input_artifacts.get("translation_workbook"),
    ]
    for artifact_id in candidates:
        if not artifact_id:
            continue
        try:
            artifact = db.get_artifact(str(artifact_id))
            if artifact["project_id"] == project_id:
                return _artifact_display_label(artifact)
        except KeyError:
            continue
    return "-"


def _workbook_processed_rows(path: Path) -> dict[str, int]:
    stats = {"source_rows": 0, "translated_rows": 0, "processed_rows": 0}
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        try:
            for ws in wb.worksheets:
                header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), ())
                headers = {
                    str(value).strip().lower(): index
                    for index, value in enumerate(header_row, start=1)
                    if value is not None and str(value).strip()
                }
                source_col = _first_col(headers, ["cn", "source", "original", "原文", "中文"])
                target_col = _first_col(headers, ["en", "target", "translation", "译文", "英文"])
                if source_col is None or target_col is None:
                    continue
                for row in ws.iter_rows(min_row=2, values_only=True):
                    has_source = bool(_row_cell(row, source_col))
                    has_target = bool(_row_cell(row, target_col))
                    if has_source:
                        stats["source_rows"] += 1
                    if has_target:
                        stats["translated_rows"] += 1
                    if has_source and has_target:
                        stats["processed_rows"] += 1
        finally:
            wb.close()
    except Exception:
        return stats
    return stats


def _soft_warning_count(summary: dict[str, Any]) -> int:
    total = 0
    for key in ("global_harness_quality", "project_harness_quality", "semantic_qa"):
        payload = summary.get(key) if isinstance(summary.get(key), dict) else {}
        total += int(payload.get("soft_warnings") or payload.get("warnings") or 0)
    return total


def _deliverable_provider_model(metadata: dict[str, Any], quality_summary: dict[str, Any]) -> tuple[str, str]:
    model_info = metadata.get("model") if isinstance(metadata.get("model"), dict) else {}
    if model_info.get("provider"):
        return str(model_info.get("provider") or "-"), str(model_info.get("model") or "-")
    semantic_qa = metadata.get("semantic_qa") if isinstance(metadata.get("semantic_qa"), dict) else quality_summary.get("semantic_qa", {})
    if not isinstance(semantic_qa, dict):
        return "-", "-"
    status = str(semantic_qa.get("status") or "")
    if status == "skipped_no_key":
        return "rules-only", "-"
    return str(semantic_qa.get("provider") or "-"), str(semantic_qa.get("model") or "-")


def _safe_delivery_name(name: str) -> str:
    return safe_delivery_name(name)


def _delivery_file(kind: str, path: Path) -> dict[str, str]:
    return {"kind": kind, "filename": path.name, "path": str(path)}


def _artifact_delivery_file(kind: str, artifact: dict[str, Any]) -> dict[str, str]:
    path = Path(str(artifact.get("path") or ""))
    return {
        "kind": kind,
        "filename": path.name,
        "path": str(path),
        "artifact_id": str(artifact.get("id") or ""),
        "download_url": f"/api/artifacts/{artifact['id']}/download",
    }


def _expected_delivery_file(kind: str, path: Path) -> dict[str, str]:
    return {"kind": kind, "filename": path.name, "path": ""}


def _write_empty_workbook(path: Path, headers: list[str], note: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Delivery"
    ws.append(headers)
    ws.append(["", note, ""])
    wb.save(path)
    wb.close()


def _latest_artifact(project_id: str, role: str | None = None, kind: str | None = None) -> dict[str, Any] | None:
    artifacts = db.list_artifacts(project_id=project_id, role=role)
    if kind:
        artifacts = [artifact for artifact in artifacts if artifact["kind"] == kind]
    return artifacts[0] if artifacts else None


def _read_glossary_rows(
    path: Path,
    sheet: str | None = None,
    term_key_column: str | None = None,
    source_column: str | None = None,
    target_column: str | None = None,
    target_alt_column: str | None = None,
    category_column: str | None = None,
    note_column: str | None = None,
    language: str = "en",
    limit: int | None = 100,
) -> tuple[list[dict[str, Any]], dict[str, str]]:
    language = require_supported_language(language)
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[sheet] if sheet else wb[wb.sheetnames[0]]
        headers = [str(cell.value or "").strip() for cell in next(ws.iter_rows(min_row=1, max_row=1))]
        normalized = {header.lower(): index for index, header in enumerate(headers) if header}
        term_key_idx = _column_index(normalized, term_key_column, ["id", "key", "编号", "序号"], required=False)
        source_idx = _column_index(normalized, source_column, ["source", "original", "cn", "zh", "chinese", "term", "原文", "中文", "术语"])
        target_idx = _column_index(normalized, target_column, target_aliases(language))
        target_alt_idx = _column_index(normalized, target_alt_column, alt_aliases(language), required=False)
        category_idx = _column_index(normalized, category_column, ["category", "type", "分类", "类别", "类型"], required=False)
        note_idx = _column_index(normalized, note_column, ["note", "notes", "comment", "备注"], required=False)
        rows = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if limit is not None and len(rows) >= limit:
                break
            source = _value_at(row, source_idx)
            target = _value_at(row, target_idx)
            if not source and not target:
                continue
            rows.append(
                {
                    "term_key": _value_at(row, term_key_idx) if term_key_idx is not None else "",
                    "source": source,
                    "target": target,
                    "target_alt": _value_at(row, target_alt_idx) if target_alt_idx is not None else "",
                    "category": _value_at(row, category_idx) if category_idx is not None else "",
                    "note": _value_at(row, note_idx) if note_idx is not None else "",
                }
            )
        return rows, {
            "term_key": headers[term_key_idx] if term_key_idx is not None else "",
            "source": headers[source_idx],
            "target": headers[target_idx],
            "target_alt": headers[target_alt_idx] if target_alt_idx is not None else "",
            "category": headers[category_idx] if category_idx is not None else "",
            "note": headers[note_idx] if note_idx is not None else "",
        }
    finally:
        wb.close()


def _column_index(normalized_headers: dict[str, int], explicit: str | None, candidates: list[str], required: bool = True) -> int | None:
    if explicit:
        hit = normalized_headers.get(explicit.strip().lower())
        if hit is not None:
            return hit
        if required:
            raise KeyError(f"column not found: {explicit}")
    for candidate in candidates:
        hit = normalized_headers.get(candidate.lower())
        if hit is not None:
            return hit
    if required:
        raise KeyError(f"none of columns found: {', '.join(candidates)}")
    return None


def _value_at(row: tuple[Any, ...], index: int | None) -> str:
    if index is None or index < 0 or index >= len(row):
        return ""
    value = row[index]
    return "" if value is None else str(value).strip()


def run_project_harness_qa(final_workbook: Path, harness: dict[str, Any], language: str = "en") -> dict[str, Any]:
    language = require_supported_language(language)
    issues: list[dict[str, Any]] = []
    if not final_workbook.exists():
        return {
            "source": "project_harness",
            "passed": False,
            "hard_errors": 1,
            "soft_warnings": 0,
            "issues": [{"severity": "hard", "message": "final workbook missing", "rule_source": "project_harness"}],
        }

    forbidden = [str(item).strip() for item in harness.get("forbidden_translations", []) if str(item).strip()]
    fixed_terms = [item for item in harness.get("fixed_terms", []) if isinstance(item, dict)]
    hard_rules = [item for item in harness.get("hard_rules", []) if isinstance(item, dict) and item.get("enabled", True)]

    wb = load_workbook(final_workbook, data_only=False)
    try:
        for ws in wb.worksheets:
            headers = _header_map(ws)
            source_col = _first_col(headers, ["cn", "source", "original", "原文", "中文"])
            target_col = _first_col(headers, target_aliases(language))
            if target_col is None:
                continue
            for row_index in range(2, ws.max_row + 1):
                source = _cell_text(ws.cell(row_index, source_col).value) if source_col else ""
                target = _cell_text(ws.cell(row_index, target_col).value)
                if not target:
                    continue
                for phrase in forbidden:
                    if phrase in target:
                        issues.append(
                            _project_issue(
                                ws.title,
                                row_index,
                                "forbidden_translation",
                                f"Translation contains forbidden phrase: {phrase}",
                                target,
                            )
                        )
                for term in fixed_terms:
                    source_term = str(term.get("source", "")).strip()
                    target_term = str(term.get("target", "")).strip()
                    if source_term and target_term and source_term in source and target_term not in target:
                        issues.append(
                            _project_issue(
                                ws.title,
                                row_index,
                                "fixed_term_missing",
                                f"Source term '{source_term}' must use '{target_term}'",
                                target,
                            )
                        )
                for rule in hard_rules:
                    pattern = str(rule.get("pattern", "")).strip()
                    if not pattern:
                        continue
                    try:
                        matched = re.search(pattern, target) is not None
                    except re.error as exc:
                        issues.append(
                            _project_issue(
                                ws.title,
                                row_index,
                                "invalid_project_rule",
                                f"Invalid project hard-rule pattern '{pattern}': {exc}",
                                target,
                                severity="soft",
                            )
                        )
                        continue
                    if matched:
                        message = str(rule.get("description") or rule.get("label") or f"Project rule matched: {pattern}")
                        issues.append(_project_issue(ws.title, row_index, "project_hard_rule", message, target))
    finally:
        wb.close()

    hard_errors = len([issue for issue in issues if issue["severity"] == "hard"])
    soft_warnings = len([issue for issue in issues if issue["severity"] == "soft"])
    return {
        "source": "project_harness",
        "passed": hard_errors == 0,
        "hard_errors": hard_errors,
        "soft_warnings": soft_warnings,
        "issues": issues[:100],
        "active_overlay": _harness_summary(harness),
    }


def list_improvements(project_id: str) -> list[dict[str, Any]]:
    path = project_dir(project_id) / "profile" / "improvement_suggestions.json"
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def create_improvement_review(run_id: str) -> dict[str, Any]:
    run = db.get_run(run_id)
    project_id = run["project_id"]
    metadata = run.get("metadata", {})
    suggestions = list_improvements(project_id)
    quality = metadata.get("quality", {})
    project_quality = metadata.get("project_harness_quality", {})
    if project_quality.get("hard_errors"):
        suggestions.append(
            _improvement_item(
                "project_harness",
                run_id,
                "Review project-specific hard rules and manual fixes",
                "Project harness QA produced hard errors; update this project's harness only after human review.",
            )
        )
    if quality and not quality.get("passed", True):
        suggestions.append(
            _improvement_item(
                "studio_integration",
                run_id,
                "Review reusable QA adapter coverage",
                "Global quality gate failed; inspect whether Studio needs better reporting or retry controls.",
            )
        )
    suggestions.append(
        _improvement_item(
            "upstream_backfeed",
            run_id,
            "Prepare upstream backfeed candidate",
            "If this run exposed a reusable gap, create a human-reviewed issue or PR against the source workflow repo.",
        )
    )
    path = project_dir(project_id) / "profile" / "improvement_suggestions.json"
    path.write_text(json.dumps(suggestions, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"project_id": project_id, "run_id": run_id, "suggestions": suggestions}


def list_quality_issues(run_id: str) -> dict[str, Any]:
    run = db.get_run(run_id)
    metadata = run.get("metadata", {})
    summary = metadata.get("quality_summary", {})
    issues: list[dict[str, Any]] = []
    for source_key, payload in (
        ("global_harness", metadata.get("quality") or summary.get("global_harness_quality") or {}),
        ("project_harness", metadata.get("project_harness_quality") or summary.get("project_harness_quality") or {}),
        ("semantic_qa", metadata.get("semantic_qa") or summary.get("semantic_qa") or {}),
    ):
        issues.extend(_normalize_quality_issues(source_key, payload))
    hard_errors = len([issue for issue in issues if issue["severity"] == "hard"])
    return {
        "run_id": run_id,
        "project_id": run["project_id"],
        "status": run["status"],
        "hard_errors": hard_errors,
        "issues": issues,
    }


def apply_manual_fixes(run_id: str, request: Any) -> dict[str, Any]:
    run = db.get_run(run_id)
    project_id = run["project_id"]
    fixes = [fix.model_dump() if hasattr(fix, "model_dump") else dict(fix) for fix in getattr(request, "fixes", [])]
    if not fixes:
        raise ValueError("manual fixes are required")

    source_artifact = _workbook_artifact_for_quality_run(run)
    source_path = Path(source_artifact["path"])
    if not source_path.exists():
        raise FileNotFoundError(str(source_path))

    output_dir = run_dir(run_id) / "manual_fixes"
    output_dir.mkdir(parents=True, exist_ok=True)
    fixed_path = output_dir / f"{source_path.stem}_manual_fixed.xlsx"
    shutil.copy2(source_path, fixed_path)

    applied = _apply_workbook_fixes(fixed_path, fixes, run_id)
    fixed_artifact = db.add_artifact(
        project_id,
        "Manual fixed workbook",
        fixed_path,
        "manual_fixed_workbook",
        run_id=run_id,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        origin="manual",
        metadata={"source_run_id": run_id, "source_artifact_id": source_artifact["id"], "manual_fix_count": len(applied)},
    )
    harness = read_project_harness(project_id)
    write_project_harness(
        project_id,
        {
            "manual_fixes": [*harness.get("manual_fixes", []), *applied],
            "qa_summary": {
                **harness.get("qa_summary", {}),
                "last_manual_fix_run": run_id,
                "last_manual_fix_artifact": fixed_artifact["id"],
            },
        },
    )
    _append_improvement_items(
        project_id,
        [
            _improvement_item(
                "project_harness",
                run_id,
                "Review manual fixes for project harness reuse",
                "Manual QA fixes were applied; review whether they should become project-specific rules, fixed terms, or reference examples.",
            )
        ],
    )

    result: dict[str, Any] = {
        "source_run": run,
        "fixed_artifact": fixed_artifact,
        "manual_fixes": applied,
        "qa_result": None,
    }
    if getattr(request, "rerun_qa", True):
        qa_run = db.insert_run(
            project_id,
            kind="qa",
            language=run.get("language", "en"),
            metadata={
                "input_artifact_id": fixed_artifact["id"],
                "manual_fix_source_run_id": run_id,
                "manual_fix_source_artifact_id": source_artifact["id"],
                "manual_fix_count": len(applied),
                "manual_fixes": applied,
            },
        )
        result["qa_result"] = run_qa_sync(qa_run["id"])
    return result


def apply_model_fixes(run_id: str, request: Any) -> dict[str, Any]:
    run = db.get_run(run_id)
    project = db.get_project(run["project_id"])
    settings = load_settings()
    provider = normalize_provider_name(settings.get("provider"))
    if provider not in REAL_PROVIDERS or not settings.get("api_key"):
        raise ValueError("模型修复需要配置 GPT / Claude / GPT 中转站 API key，不能在未配置真实 API 时生成可交付修复。")

    max_issues = max(1, min(int(getattr(request, "max_issues", 80) or 80), 200))
    issue_payload = list_quality_issues(run_id)
    issues = [
        issue
        for issue in issue_payload.get("issues", [])
        if issue.get("sheet") and int(issue.get("row") or 0) > 1 and issue.get("severity") in {"hard", "soft"}
    ][:max_issues]
    if not issues:
        raise ValueError("没有可交给模型修复的行级 QA 问题。")

    source_artifact = _workbook_artifact_for_quality_run(run)
    source_path = Path(source_artifact["path"])
    if not source_path.exists():
        raise FileNotFoundError(str(source_path))
    rows = [_model_fix_row_context(source_path, issue) for issue in issues]
    prompt = _model_fix_prompt(project, run, rows)
    text = _call_semantic_provider(settings, prompt)
    payload = _parse_semantic_qa_payload(text)
    fixes = _normalize_model_fixes(payload, rows)
    if not fixes:
        raise ValueError("模型没有返回可应用的修复。")

    output_dir = run_dir(run_id) / "model_fixes"
    output_dir.mkdir(parents=True, exist_ok=True)
    fixed_path = output_dir / f"{source_path.stem}_model_fixed.xlsx"
    shutil.copy2(source_path, fixed_path)
    applied = _apply_workbook_fixes(fixed_path, fixes, run_id)
    fixed_artifact = db.add_artifact(
        project["id"],
        "Model fixed workbook",
        fixed_path,
        "manual_fixed_workbook",
        run_id=run_id,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        origin="provider",
        metadata={
            "source_run_id": run_id,
            "source_artifact_id": source_artifact["id"],
            "model_fix_count": len(applied),
            "provider": provider,
            "model": settings.get("model") or "",
        },
    )
    _append_improvement_items(
        project["id"],
        [
            _improvement_item(
                "project_harness",
                run_id,
                "Review model fixes for reusable project rules",
                "Model QA fixes were applied; review whether repeated fixes should become project terms, fixed names, or project-specific rules.",
            )
        ],
    )

    result: dict[str, Any] = {
        "source_run": run,
        "fixed_artifact": fixed_artifact,
        "model_fixes": applied,
        "qa_result": None,
    }
    if getattr(request, "rerun_qa", True):
        qa_run = db.insert_run(
            project["id"],
            kind="qa",
            language=run.get("language", "en"),
            metadata={
                "input_artifact_id": fixed_artifact["id"],
                "model_fix_source_run_id": run_id,
                "model_fix_source_artifact_id": source_artifact["id"],
                "model_fix_count": len(applied),
                "manual_fixes": applied,
                "task_origin": "model_fix_continuation",
                "task_code": (run.get("metadata") or {}).get("task_code"),
            },
        )
        result["qa_result"] = run_qa_sync(qa_run["id"])
    return result


def create_semantic_qa_context(run_id: str) -> dict[str, Any]:
    run = db.get_run(run_id)
    project = db.get_project(run["project_id"])
    output_dir = run_dir(run_id) / "semantic_qa"
    output_dir.mkdir(parents=True, exist_ok=True)
    report = {
        "status": "needs_model_review",
        "message": "Semantic QA context is prepared; model review is not auto-marked as passed.",
        "run_id": run_id,
        "project_id": project["id"],
        "sources": {
            "global_harness": GLOBAL_HARNESS_CONTRACT,
            "project_harness": read_project_harness(project["id"]),
        },
        "run_quality": {
            "global": run.get("metadata", {}).get("quality", {}),
            "project_harness": run.get("metadata", {}).get("project_harness_quality", {}),
        },
    }
    path = output_dir / "semantic_qa_context.json"
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    artifact = db.add_artifact(project["id"], "Semantic QA context", path, "semantic_qa_context", run_id=run_id, mime="application/json")
    return {"run": db.get_run(run_id), "artifact": artifact, "semantic_qa": report}


def run_qa_sync(run_id: str) -> dict[str, Any]:
    run = db.get_run(run_id)
    project = db.get_project(run["project_id"])
    metadata = run.get("metadata", {})
    input_artifact_id = metadata.get("input_artifact_id")
    if not input_artifact_id:
        raise KeyError("input_artifact_id")
    workbook_artifact = db.get_artifact(input_artifact_id)
    workbook_path = Path(workbook_artifact["path"])
    db.update_run(run_id, status="running")

    output_dir = run_dir(run_id) / "qa"
    output_dir.mkdir(parents=True, exist_ok=True)
    language = require_supported_language(run.get("language") or "en")
    glossary_snapshot = create_project_glossary_snapshot(project["id"], run_id, output_dir / "snapshots", language=language)
    snapshots = create_prompt_and_harness_snapshots(project["id"], run_id, output_dir / "snapshots", language=language)
    reference_snapshot = create_quick_reference_snapshot(project["id"], run_id, metadata.get("reference_artifact_ids"), output_dir / "snapshots")
    qa_result = run_localization_qa(
        project=project,
        run_id=run_id,
        workbook_path=workbook_path,
        output_dir=output_dir,
        glossary_snapshot=glossary_snapshot,
        harness_snapshot=snapshots["harness_snapshot"],
        workbook_artifact=workbook_artifact,
        run_metadata=metadata,
        manual_fixes=metadata.get("manual_fixes") or [],
        language=language,
    )
    input_artifacts = {
        "translation_workbook": workbook_artifact["id"],
        "glossary_snapshot": glossary_snapshot["id"],
        "prompt_snapshot": snapshots["prompt_artifact"]["id"],
        "harness_snapshot": snapshots["harness_artifact"]["id"],
    }
    if reference_snapshot:
        input_artifacts["quick_reference_snapshot"] = reference_snapshot["artifact"]["id"]
    if qa_result.get("qa_final_artifact"):
        input_artifacts["qa_final_workbook"] = qa_result["qa_final_artifact"]["id"]
    status = "passed" if qa_result["quality_summary"]["passed"] else "failed"
    archive_result = None
    if status == "passed" and qa_result.get("qa_final_artifact"):
        archive_result = archive_translation_artifact(
            project["id"],
            qa_result["qa_final_artifact"]["id"],
            language=run.get("language") or "en",
            source_type="qa_passed",
        )
    db.update_run(
        run_id,
        status=status,
        metadata={
            **metadata,
            "task_origin": metadata.get("task_origin") or "direct_import",
            "input_artifacts": input_artifacts,
            "quality": qa_result["quality"],
            "project_harness_quality": qa_result["project_harness_quality"],
            "semantic_qa": qa_result["semantic_qa"],
            "quality_summary": qa_result["quality_summary"],
            "translation_archive": archive_result,
        },
    )
    return {"run": db.get_run(run_id), "artifacts": qa_result["artifacts"], "quality_summary": qa_result["quality_summary"]}


def run_localization_qa(
    project: dict[str, Any],
    run_id: str,
    workbook_path: Path,
    output_dir: Path,
    glossary_snapshot: dict[str, Any],
    harness_snapshot: dict[str, Any],
    workbook_artifact: dict[str, Any] | None = None,
    run_metadata: dict[str, Any] | None = None,
    manual_fixes: list[dict[str, Any]] | None = None,
    language: str = "en",
) -> dict[str, Any]:
    language = require_supported_language(language)
    output_dir.mkdir(parents=True, exist_ok=True)
    machine_dir = output_dir / "machine_review"
    machine_dir.mkdir(parents=True, exist_ok=True)
    review_args = [
        sys.executable,
        str(LOCALIZATION_ROOT / "process_language.py"),
        "--input",
        str(workbook_path),
        "--lang",
        language,
        "--output-dir",
        str(machine_dir),
        "--auto-fix",
        "--term-base",
        glossary_snapshot["path"],
    ]
    run_subprocess(review_args, LOCALIZATION_ROOT, run_id)
    qa_workbook = machine_dir / f"result_{language}.xlsx"
    qa_report = machine_dir / f"report_{language}.xlsx"
    _normalize_review_workbook_sheet_names(qa_workbook, workbook_path)
    quality_args = [
        sys.executable,
        str(LOCALIZATION_ROOT / "scripts" / "run_quality_harness.py"),
        "--workbook",
        str(qa_workbook),
        "--term-base",
        glossary_snapshot["path"],
        "--lang",
        language,
        "--json",
    ]
    if language == "en":
        quality_args.insert(2, str(LOCALIZATION_ROOT / "fixtures" / "quality_regression.json"))
    quality = _run_quality_json(quality_args, run_id)
    project_harness_quality = run_project_harness_qa(qa_workbook, harness_snapshot["project_harness"], language=language)
    semantic_qa = run_semantic_qa_report(run_id, project["id"], qa_workbook, quality, project_harness_quality, language=language)
    hard_errors = _hard_error_count(quality) + int(project_harness_quality.get("hard_errors", 0)) + int(semantic_qa.get("hard_errors", 0))
    passed = hard_errors == 0
    summary = {
        "version": 1,
        "run_id": run_id,
        "project_id": project["id"],
        "passed": passed,
        "hard_errors": hard_errors,
        "sources": {
            "translation_workbook": (workbook_artifact or {}).get("id", ""),
            "qa_workbook": str(qa_workbook),
            "glossary_snapshot": glossary_snapshot["id"],
            "global_harness": GLOBAL_HARNESS_CONTRACT,
            "project_harness": "project_harness_snapshot.json",
            "semantic_qa": semantic_qa.get("status", ""),
        },
        "global_harness_quality": quality,
        "project_harness_quality": project_harness_quality,
        "semantic_qa": semantic_qa,
    }
    metadata = run_metadata or {}
    if metadata.get("manual_fix_source_run_id"):
        summary["sources"]["manual_fix_source_run"] = metadata["manual_fix_source_run_id"]
    if metadata.get("model_fix_source_run_id"):
        summary["sources"]["model_fix_source_run"] = metadata["model_fix_source_run_id"]
    summary_path = output_dir / "quality_summary.json"
    summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    auto_fixes = _collect_workbook_translation_changes(workbook_path, qa_workbook)
    changes_path = write_qa_changes_report(output_dir, manual_fixes or [], auto_fixes)
    artifacts = [
        db.add_artifact(
            project["id"],
            "QA reviewed workbook",
            qa_workbook,
            "qa_result",
            run_id=run_id,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
        db.add_artifact(
            project["id"],
            "QA report",
            qa_report,
            "qa_report",
            run_id=run_id,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
        db.add_artifact(project["id"], "Quality summary", summary_path, "quality_summary", run_id=run_id, mime="application/json"),
        db.add_artifact(
            project["id"],
            "QA changes",
            changes_path,
            "qa_changes",
            run_id=run_id,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
    ]
    qa_final_artifact = None
    if passed:
        qa_final_artifact = db.add_artifact(
            project["id"],
            "QA final workbook",
            qa_workbook,
            "qa_final_workbook",
            run_id=run_id,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            origin="generated",
            metadata={"language": language, "source_workbook": str(workbook_path), "glossary_snapshot": glossary_snapshot["id"]},
        )
        artifacts.append(qa_final_artifact)
    return {
        "artifacts": artifacts,
        "qa_final_artifact": qa_final_artifact,
        "quality": quality,
        "project_harness_quality": project_harness_quality,
        "semantic_qa": semantic_qa,
        "quality_summary": summary,
        "qa_workbook": qa_workbook,
    }


def write_qa_changes_report(output_dir: Path, manual_fixes: list[dict[str, Any]], auto_fixes: list[dict[str, Any]] | None = None) -> Path:
    path = output_dir / "qa_changes.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "QA Changes"
    ws.append(["工作表", "行号", "问题ID", "修改前", "修改后", "规则来源", "备注"])
    if manual_fixes:
        for fix in manual_fixes:
            ws.append(
                [
                    fix.get("sheet", ""),
                    fix.get("row", ""),
                    fix.get("issue_id", ""),
                    fix.get("previous_translation", ""),
                    fix.get("translation", ""),
                    fix.get("rule_source", "manual_fix"),
                    fix.get("note", ""),
                ]
            )
    for fix in auto_fixes or []:
        ws.append(
            [
                fix.get("sheet", ""),
                fix.get("row", ""),
                fix.get("issue_id", "auto_fix"),
                fix.get("previous_translation", ""),
                fix.get("translation", ""),
                "localization_auto_fix",
                fix.get("note", ""),
            ]
        )
    if not manual_fixes and not auto_fixes:
        ws.append(["", "", "", "", "", "qa", "未应用修改"])
    wb.save(path)
    wb.close()
    return path


def _run_quality_json(args: list[str], run_id: str) -> dict[str, Any]:
    proc = run_subprocess_allow_failure(args, LOCALIZATION_ROOT, run_id)
    if not proc.stdout.strip():
        raise RuntimeError(f"quality harness returned no JSON: {proc.stderr}")
    return json.loads(proc.stdout)


def _hard_error_count(quality: dict[str, Any]) -> int:
    if quality.get("passed"):
        return 0
    hard_issues = [
        issue
        for issue in quality.get("issues", [])
        if str(issue.get("severity", "hard")).lower() not in {"warning", "soft", "info"}
    ]
    issue_count = len(hard_issues) + len(quality.get("failures", []))
    if issue_count:
        return issue_count
    return 1 if not quality.get("issues") and not quality.get("failures") else 0


def run_semantic_qa_report(
    run_id: str,
    project_id: str,
    workbook_path: Path,
    quality: dict[str, Any],
    project_quality: dict[str, Any],
    language: str = "en",
) -> dict[str, Any]:
    language = require_supported_language(language)
    spec = language_spec(language)
    settings = load_settings()
    provider = normalize_provider_name(settings.get("provider"))
    model = str(settings.get("model") or "")
    issue_context = {
        "global_issue_count": len(quality.get("issues", [])),
        "project_harness_issue_count": len(project_quality.get("issues", [])),
        "sample_issues": (quality.get("issues", []) or [])[:20],
        "project_issues": (project_quality.get("issues", []) or [])[:20],
    }
    base = {
        "source": "semantic_qa",
        "provider": provider,
        "model": model,
        "prompt_context": {"run_id": run_id, "project_id": project_id, "language": language, **issue_context},
        "issues": [],
        "soft_warnings": 0,
    }
    if provider not in REAL_PROVIDERS or not settings.get("api_key"):
        return {**base, "status": "skipped_no_key", "passed": True, "hard_errors": 0}

    prompt = (
        f"You are doing semantic QA for a {spec.prompt_name} game localization workbook. "
        "Review the machine QA context and return strict JSON only: "
        "{\"passed\": boolean, \"issues\": [{\"severity\":\"hard|soft\", \"message\":\"...\", \"sheet\":\"\", \"row\":0}]}.\n"
        f"Workbook: {workbook_path.name}\n"
        f"Context:\n{json.dumps(issue_context, ensure_ascii=False)}"
    )
    try:
        text = _call_semantic_provider(settings, prompt)
        payload = _parse_semantic_qa_payload(text)
        issues = payload.get("issues", []) if isinstance(payload.get("issues"), list) else []
        hard_errors = len([issue for issue in issues if str(issue.get("severity", "hard")).lower() == "hard"])
        return {
            **base,
            "status": "model_reviewed",
            "passed": bool(payload.get("passed", hard_errors == 0)) and hard_errors == 0,
            "hard_errors": hard_errors,
            "soft_warnings": len([issue for issue in issues if str(issue.get("severity", "")).lower() == "soft"]),
            "issues": issues,
        }
    except Exception as exc:
        return {
            **base,
            "status": "provider_error",
            "passed": False,
            "hard_errors": 1,
            "issues": [{"severity": "hard", "message": f"Semantic QA provider failed: {user_facing_error(exc)}", "sheet": "", "row": 0}],
        }


def _call_semantic_provider(settings: dict[str, Any], prompt: str) -> str:
    return call_text(settings, prompt, system="Return strict JSON only.")


def _parse_semantic_qa_payload(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        decoder = json.JSONDecoder()
        start = cleaned.find("{")
        while start != -1:
            try:
                payload, _ = decoder.raw_decode(cleaned[start:])
                if isinstance(payload, dict):
                    return payload
            except json.JSONDecodeError:
                start = cleaned.find("{", start + 1)
                continue
        raise


def _collect_workbook_translation_changes(before_path: Path, after_path: Path) -> list[dict[str, Any]]:
    if not before_path.exists() or not after_path.exists():
        return []
    before_wb = load_workbook(before_path, read_only=True, data_only=True)
    after_wb = load_workbook(after_path, read_only=True, data_only=True)
    changes: list[dict[str, Any]] = []
    try:
        before_ws = before_wb[before_wb.sheetnames[0]]
        after_ws = after_wb[before_ws.title] if before_ws.title in after_wb.sheetnames else after_wb[after_wb.sheetnames[0]]
        before_headers = _header_map(before_ws)
        after_headers = _header_map(after_ws)
        before_id_col = _first_col(before_headers, ["id", "key", "编号", "序号"])
        after_id_col = _first_col(after_headers, ["id", "key", "编号", "序号"])
        before_target_col = _first_col(before_headers, ["en", "translation", "target", "译文", "英文"])
        after_target_col = _first_col(after_headers, ["en", "translation", "target", "译文", "英文"])
        if before_id_col is None or after_id_col is None or before_target_col is None or after_target_col is None:
            return []
        after_by_id: dict[str, tuple[int, str]] = {}
        for row_index, row in enumerate(after_ws.iter_rows(min_row=2, values_only=True), start=2):
            row_id = _row_cell(row, after_id_col)
            if row_id:
                after_by_id[row_id] = (row_index, _row_cell(row, after_target_col))
        for row_index, row in enumerate(before_ws.iter_rows(min_row=2, values_only=True), start=2):
            row_id = _row_cell(row, before_id_col)
            if not row_id or row_id not in after_by_id:
                continue
            after_row, after_text = after_by_id[row_id]
            before_text = _row_cell(row, before_target_col)
            if before_text != after_text:
                changes.append(
                    {
                        "sheet": after_ws.title,
                        "row": after_row,
                        "issue_id": "auto_fix",
                        "previous_translation": before_text,
                        "translation": after_text,
                        "note": f"localization workflow auto-fix for ID {row_id}",
                    }
                )
    finally:
        before_wb.close()
        after_wb.close()
    return changes


def _normalize_review_workbook_sheet_names(review_path: Path, source_path: Path) -> None:
    if not review_path.exists() or not source_path.exists():
        return
    source_wb = load_workbook(source_path, read_only=True, data_only=True)
    review_wb = load_workbook(review_path)
    try:
        source_title = source_wb.sheetnames[0] if source_wb.sheetnames else ""
        if source_title and review_wb.sheetnames:
            first = review_wb[review_wb.sheetnames[0]]
            if first.title != source_title and source_title not in review_wb.sheetnames:
                first.title = source_title
                review_wb.save(review_path)
    finally:
        source_wb.close()
        review_wb.close()


def _row_cell(row: tuple[Any, ...], column: int) -> str:
    if column < 1 or column > len(row):
        return ""
    value = row[column - 1]
    return "" if value is None else str(value).strip()


def _normalize_translation_id(value: Any) -> RowId | None:
    if value is None:
        return None
    if isinstance(value, bool):
        text = str(value).strip()
        return text or None
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(value) if value.is_integer() else str(value).strip()
    text = str(value).strip()
    if not text:
        return None
    if re.fullmatch(r"-?(0|[1-9]\d*)", text):
        return int(text)
    return text


def _is_supported_translation_id(value: Any) -> bool:
    return _normalize_translation_id(value) is not None


def _normalize_quality_issues(source_key: str, payload: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for index, issue in enumerate(payload.get("issues", []) or []):
        severity = str(issue.get("severity") or "hard").lower()
        if severity in {"warning", "soft", "info"}:
            severity = "soft" if severity != "info" else "info"
        else:
            severity = "hard"
        rows.append(
            {
                "id": str(issue.get("id") or f"{source_key}:{index}:{issue.get('sheet', '')}:{issue.get('row', '')}:{issue.get('check_type', '')}"),
                "source": source_key,
                "rule_source": issue.get("rule_source") or issue.get("source") or source_key,
                "severity": severity,
                "sheet": issue.get("sheet") or "",
                "row": int(issue.get("row") or 0),
                "check_type": issue.get("check_type") or issue.get("type") or issue.get("code") or "quality_issue",
                "message": issue.get("message") or issue.get("detail") or "",
                "current_translation": issue.get("translation") or issue.get("target") or issue.get("actual") or "",
            }
        )
    for index, failure in enumerate(payload.get("failures", []) or []):
        rows.append(
            {
                "id": str(failure.get("id") or f"{source_key}:failure:{index}"),
                "source": source_key,
                "rule_source": source_key,
                "severity": "hard",
                "sheet": failure.get("sheet") or "",
                "row": int(failure.get("row") or 0),
                "check_type": "fixture_failure",
                "message": failure.get("message") or f"Fixture failure: {failure.get('id', index)}",
                "current_translation": failure.get("actual") or "",
            }
        )
    return rows


def _workbook_artifact_for_quality_run(run: dict[str, Any]) -> dict[str, Any]:
    metadata = run.get("metadata", {})
    input_artifact_id = metadata.get("input_artifacts", {}).get("translation_workbook") or metadata.get("input_artifact_id")
    if input_artifact_id:
        artifact = db.get_artifact(input_artifact_id)
        if artifact["role"] in {"translation_workbook", "translation_draft", "language_source", "quick_input"}:
            return artifact
    artifacts = db.list_artifacts(run_id=run["id"], role="translation_workbook") or db.list_artifacts(run_id=run["id"], role="language_source")
    if artifacts:
        return artifacts[0]
    raise KeyError("translation workbook artifact not found")


def _model_fix_row_context(path: Path, issue: dict[str, Any]) -> dict[str, Any]:
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        sheet_name = str(issue.get("sheet") or wb.sheetnames[0])
        requested_ws = wb[sheet_name] if sheet_name in wb.sheetnames else None
        issue_record_id = issue.get("id") or issue.get("record_id") or ""
        row_index = int(issue.get("row") or 0)
        resolved = _resolve_workbook_row_for_issue(wb, requested_ws, row_index, issue_record_id)
        ws, row_index = resolved if resolved else (requested_ws or wb[wb.sheetnames[0]], row_index)
        header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), ())
        headers = {
            str(value).strip().lower(): index
            for index, value in enumerate(header_row, start=1)
            if value is not None and str(value).strip()
        }
        source_col = _first_col(headers, ["cn", "source", "original", "原文", "中文"])
        target_col = _first_col(headers, ["en", "target", "translation", "译文", "英文"])
        id_col = _first_col(headers, ["id", "key", "编号", "序号"])
        row_values = next(ws.iter_rows(min_row=row_index, max_row=row_index, values_only=True), ())
        return {
            "issue_id": issue.get("id", ""),
            "sheet": ws.title,
            "row": row_index,
            "record_id": _row_cell(row_values, id_col) if id_col else "",
            "source_text": _row_cell(row_values, source_col) if source_col else "",
            "current_translation": _row_cell(row_values, target_col) if target_col else issue.get("current_translation", ""),
            "severity": issue.get("severity", "hard"),
            "check_type": issue.get("check_type", ""),
            "message": issue.get("message", ""),
            "rule_source": issue.get("rule_source") or issue.get("source") or "qa",
        }
    finally:
        wb.close()


def _resolve_workbook_row_for_issue(wb: Any, requested_ws: Any | None, row_index: int, record_id: Any) -> tuple[Any, int] | None:
    normalized_record_id = _normalize_translation_id(record_id)
    if requested_ws is not None and row_index >= 2:
        headers = _header_map(requested_ws)
        id_col = _first_col(headers, ["id", "key", "编号", "序号"])
        if id_col is None or normalized_record_id is None:
            return requested_ws, row_index
        current_id = _normalize_translation_id(requested_ws.cell(row_index, id_col).value)
        if current_id == normalized_record_id:
            return requested_ws, row_index
    if normalized_record_id is None:
        return (requested_ws, row_index) if requested_ws is not None and row_index >= 2 else None
    for ws in wb.worksheets:
        headers = _header_map(ws)
        id_col = _first_col(headers, ["id", "key", "编号", "序号"])
        if id_col is None:
            continue
        for candidate_row in range(2, ws.max_row + 1):
            if _normalize_translation_id(ws.cell(candidate_row, id_col).value) == normalized_record_id:
                return ws, candidate_row
    return (requested_ws, row_index) if requested_ws is not None and row_index >= 2 else None


def _model_fix_prompt(project: dict[str, Any], run: dict[str, Any], rows: list[dict[str, Any]]) -> str:
    language = require_supported_language(run.get("language") or "en")
    profile = project.get("profile") or {}
    prompt = str((profile.get("prompts_by_language") or {}).get(language) or project.get("prompt_text") or "").strip()
    prompt = _manage_project_prompt_context(prompt, load_settings())
    harness = read_project_harness(project["id"])
    return (
        "你是游戏本地化 QA 修复模型。请根据项目提示词、项目规则、术语要求和 QA 问题，"
        "只修复译文，不改原文，不解释过程。必须保留变量、数字、HTML/BBCode 标签、换行和占位符。"
        "如果无法确定，保留原译文并在 note 写明需要人工确认。\n\n"
        "返回严格 JSON：{\"fixes\":[{\"issue_id\":\"...\",\"record_id\":\"...\",\"sheet\":\"...\",\"row\":2,\"translation\":\"...\",\"note\":\"...\"}]}。"
        "必须优先沿用待修复行里的 issue_id 和 record_id；sheet/row 仅作辅助定位。\n"
        f"项目：{project.get('name','')}\n"
        f"任务：{run.get('id','')}\n"
        f"项目提示词：\n{prompt}\n\n"
        f"项目规则：\n{json.dumps(harness, ensure_ascii=False)}\n\n"
        f"待修复行：\n{json.dumps(rows, ensure_ascii=False, indent=2)}"
    )


def _normalize_model_fixes(payload: dict[str, Any], rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    fixes_by_issue = {str(row["issue_id"]): row for row in rows}
    fixes_by_position = {(str(row["sheet"]), int(row["row"])): row for row in rows}
    fixes: list[dict[str, Any]] = []
    seen: set[tuple[str, int]] = set()
    for item in payload.get("fixes", []) if isinstance(payload.get("fixes"), list) else []:
        issue_id = str(item.get("issue_id") or "")
        sheet = str(item.get("sheet") or "")
        row_index = int(item.get("row") or 0)
        source = fixes_by_issue.get(issue_id) or fixes_by_position.get((sheet, row_index))
        if not source:
            continue
        key = (str(source["sheet"]), int(source["row"]))
        if key in seen:
            continue
        translation = str(item.get("translation") or "").strip()
        if not translation:
            continue
        seen.add(key)
        fixes.append(
            {
                "issue_id": source["issue_id"],
                "sheet": source["sheet"],
                "row": source["row"],
                "record_id": source.get("record_id", ""),
                "source_text": source.get("source_text", ""),
                "translation": translation,
                "note": str(item.get("note") or f"model_fix:{source['check_type']}").strip(),
                "rule_source": "model_fix",
            }
        )
    return fixes


def _apply_workbook_fixes(path: Path, fixes: list[dict[str, Any]], source_run_id: str) -> list[dict[str, Any]]:
    wb = load_workbook(path)
    applied: list[dict[str, Any]] = []
    try:
        for fix in fixes:
            row_index = int(fix.get("row") or 0)
            if row_index < 2:
                raise ValueError(f"invalid workbook row: {row_index}")
            sheet_name = str(fix.get("sheet") or wb.sheetnames[0]).strip()
            if sheet_name not in wb.sheetnames:
                resolved = _resolve_workbook_row_for_issue(wb, None, row_index, fix.get("record_id"))
                if not resolved:
                    raise KeyError(f"sheet not found: {sheet_name}")
                ws, row_index = resolved
                sheet_name = ws.title
            else:
                ws = wb[sheet_name]
                resolved = _resolve_workbook_row_for_issue(wb, ws, row_index, fix.get("record_id"))
                if resolved:
                    ws, row_index = resolved
                    sheet_name = ws.title
            target_col = _first_col(_header_map(ws), ["en", "translation", "target", "译文", "英文"])
            if target_col is None:
                raise KeyError(f"target column not found in sheet: {sheet_name}")
            cell = ws.cell(row_index, target_col)
            previous = _cell_text(cell.value)
            translation = str(fix.get("translation") or "").strip()
            cell.value = translation
            applied.append(
                {
                    "id": db.new_id("fix"),
                    "source_run_id": source_run_id,
                    "issue_id": fix.get("issue_id") or "",
                    "sheet": sheet_name,
                    "row": row_index,
                    "column": target_col,
                    "previous_translation": previous,
                    "translation": translation,
                    "note": str(fix.get("note") or "").strip(),
                    "rule_source": str(fix.get("rule_source") or "manual_fix").strip() or "manual_fix",
                    "applied_at": db.now_iso(),
                }
            )
        wb.save(path)
    finally:
        wb.close()
    return applied


def _append_improvement_items(project_id: str, items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    suggestions = list_improvements(project_id)
    suggestions.extend(items)
    path = project_dir(project_id) / "profile" / "improvement_suggestions.json"
    path.write_text(json.dumps(suggestions, ensure_ascii=False, indent=2), encoding="utf-8")
    return suggestions


def _header_map(ws: Any) -> dict[str, int]:
    result: dict[str, int] = {}
    if getattr(ws, "max_row", 0) < 1 or getattr(ws, "max_column", 0) < 1:
        return result
    try:
        header_row = ws[1]
    except IndexError:
        return result
    for cell in header_row:
        if cell.value is None:
            continue
        result[str(cell.value).strip().lower()] = int(cell.column)
    return result


def _first_col(headers: dict[str, int], names: list[str]) -> int | None:
    for name in names:
        hit = headers.get(name.lower())
        if hit is not None:
            return hit
    return None


def _cell_text(value: Any) -> str:
    return "" if value is None else str(value)


def _project_issue(
    sheet: str,
    row: int,
    check_type: str,
    message: str,
    translation: str,
    severity: str = "hard",
) -> dict[str, Any]:
    return {
        "source": "project_harness",
        "rule_source": "project_harness",
        "severity": severity,
        "sheet": sheet,
        "row": row,
        "check_type": check_type,
        "message": message,
        "translation": translation,
    }


def _improvement_item(category: str, run_id: str, title: str, detail: str) -> dict[str, Any]:
    return {
        "id": db.new_id("imp"),
        "category": category,
        "run_id": run_id,
        "title": title,
        "detail": detail,
        "status": "pending_review",
        "created_at": db.now_iso(),
    }


def _completed_batch_rows(path: Path, batch: list[dict[str, Any]]) -> list[dict[str, Any]] | None:
    if not path.exists():
        return None
    try:
        rows = read_jsonl(path)
    except Exception:
        return None
    expected_ids = [_normalize_translation_id(row.get("id")) for row in batch]
    actual_ids = [_normalize_translation_id(row.get("id")) for row in rows if "id" in row]
    if actual_ids != expected_ids:
        return None
    if any("translation" not in row or not str(row.get("translation") or "").strip() for row in rows):
        return None
    return rows


def _write_batch_error(path: Path, batch_index: int, attempt: int, exc: Exception) -> None:
    path.write_text(
        json.dumps(
            {
                "batch_index": batch_index,
                "attempt": attempt,
                "error": user_facing_error(exc),
                "created_at": db.now_iso(),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def _translation_progress(
    *,
    total_rows: int,
    total_batches: int,
    completed_batches: int,
    completed_rows: int,
    batch_size: int,
    started_at: float,
    current_batch: int | None = None,
    failed_batch: int | None = None,
) -> dict[str, Any]:
    elapsed_seconds = max(0.0, time.monotonic() - started_at)
    average_batch_seconds = elapsed_seconds / completed_batches if completed_batches else None
    remaining_batches = max(0, total_batches - completed_batches)
    eta_seconds = int(average_batch_seconds * remaining_batches) if average_batch_seconds is not None else None
    percent = round((completed_batches / total_batches) * 100, 2) if total_batches else 100.0
    return {
        "total_rows": total_rows,
        "completed_rows": completed_rows,
        "total_batches": total_batches,
        "completed_batches": completed_batches,
        "remaining_batches": remaining_batches,
        "current_batch": current_batch,
        "failed_batch": failed_batch,
        "batch_size": batch_size,
        "percent": percent,
        "elapsed_seconds": int(elapsed_seconds),
        "average_batch_seconds": round(average_batch_seconds, 2) if average_batch_seconds is not None else None,
        "eta_seconds": eta_seconds,
    }


def _update_translation_progress(run_id: str, progress: dict[str, Any], status: str = "running") -> None:
    current = db.get_run(run_id)
    db.update_run(run_id, status=status, metadata={**current.get("metadata", {}), "translation_progress": progress})


def _terminal_translation_progress(progress: Any, status: str) -> Any:
    if not isinstance(progress, dict):
        return progress
    normalized = dict(progress)
    normalized["current_batch"] = None
    if str(normalized.get("lease_status") or "").lower() == "running":
        normalized["lease_status"] = status
    return normalized


def _translation_cancel_path(work_dir: Path) -> Path:
    return work_dir / "cancel.requested"


def _cancel_requested(run_id: str, work_dir: Path, cancel_event: Any | None = None) -> bool:
    if cancel_event is not None and getattr(cancel_event, "is_set", lambda: False)():
        return True
    if _translation_cancel_path(work_dir).exists():
        return True
    try:
        return db.get_run(run_id).get("status") == "canceled"
    except KeyError:
        return True


def _structural_tokens(text: str) -> list[str]:
    patterns = [
        r"\{[^{}]+\}",
        r"%[sdif]",
        r"##\d+",
        r"\[(?!/?color\b)(?:[A-Za-z]+\d+|\d+)\]",
        r"\[[a-zA-Z]+[^\]]*\]",
        r"\[/[a-zA-Z]+\]",
        r"<[^>]+>",
        r"&[A-Za-z][A-Za-z0-9]+;",
    ]
    hits: list[str] = []
    for pattern in patterns:
        hits.extend(re.findall(pattern, str(text or "")))
    return hits


def _validate_translated_batch(batch: list[dict[str, Any]], rows: list[dict[str, Any]], language: str) -> list[dict[str, Any]]:
    expected_ids = [_normalize_translation_id(row.get("id")) for row in batch]
    actual_ids = [_normalize_translation_id(row.get("id")) for row in rows if "id" in row]
    if actual_ids != expected_ids:
        raise ValueError(f"response IDs mismatch: expected={expected_ids[:8]}, actual={actual_ids[:8]}")
    if len(set(map(str, actual_ids))) != len(actual_ids):
        raise ValueError("response contains duplicate IDs")
    validated: list[dict[str, Any]] = []
    for source_row, row in zip(batch, rows):
        translation = str(row.get("translation") or "")
        if not translation.strip():
            raise ValueError(f"row {source_row.get('id')} returned empty translation")
        source = str(source_row.get("source") or "")
        missing_tokens = [token for token in _structural_tokens(source) if token not in translation]
        if missing_tokens:
            raise ValueError(f"row {source_row.get('id')} lost structural token(s): {missing_tokens[:5]}")
        if source.count("\n") != translation.count("\n"):
            raise ValueError(f"row {source_row.get('id')} changed actual newline count")
        if source.count("\\n") != translation.count("\\n"):
            raise ValueError(f"row {source_row.get('id')} changed escaped newline count")
        if language in {"en", "ko"} and _looks_like_untranslated_seed(translation, language):
            raise ValueError(f"row {source_row.get('id')} still contains obvious Chinese text")
        validated.append({"id": _normalize_translation_id(row.get("id")), "translation": translation})
    return validated


def _manifest_progress(
    manifest: dict[str, Any],
    *,
    batch_size: int,
    started_at: float,
    current_batch: int | None = None,
    failed_batch: int | None = None,
    rate_limit_wait_seconds: float | None = None,
) -> dict[str, Any]:
    batches = manifest.get("batches") or []
    completed_batches = [batch for batch in batches if batch.get("status") == "passed"]
    completed_rows = sum(int(batch.get("row_count") or 0) for batch in completed_batches)
    progress = _translation_progress(
        total_rows=int(manifest.get("total_rows") or 0),
        total_batches=len(batches),
        completed_batches=len(completed_batches),
        completed_rows=completed_rows,
        batch_size=batch_size,
        started_at=started_at,
        current_batch=current_batch,
        failed_batch=failed_batch,
    )
    progress.update(
        {
            "max_concurrent_batches": int(manifest.get("max_concurrent_batches") or 1),
            "estimated_total_input_tokens": int(manifest.get("estimated_total_input_tokens") or 0),
            "rate_limit_wait_seconds": round(rate_limit_wait_seconds, 2) if rate_limit_wait_seconds else 0,
            "fingerprint": str(manifest.get("input_fingerprint") or ""),
            "lease_status": (db.get_job_lease("long_text") or {}).get("status", ""),
            "invalidated_reason": str(manifest.get("invalidated_reason") or ""),
        }
    )
    return progress


async def _translate_rows_with_orchestration(
    *,
    run_id: str,
    rows: list[dict[str, Any]],
    settings: dict[str, Any],
    project_prompt: str,
    work_dir: Path,
    batch_size: int,
    language: str,
    cancel_event: Any | None = None,
    confirm_api_budget: bool = False,
) -> list[dict[str, Any]]:
    batch_size = max(1, min(int(batch_size or settings.get("batch_size") or 90), 200))
    provider_prompt = _manage_project_prompt_context(project_prompt, settings)
    context_summary = _project_context_summary(project_prompt, settings)
    cancel_path = _translation_cancel_path(work_dir)
    if not (cancel_event is not None and getattr(cancel_event, "is_set", lambda: False)()) and cancel_path.exists():
        cancel_path.unlink()
    manifest_path = work_dir / "batch_manifest.json"
    batches_dir = work_dir / f"batches_{batch_size}"
    manifest = _load_or_create_batch_manifest(manifest_path, rows, project_prompt, settings, batch_size, language)
    batches_dir.mkdir(parents=True, exist_ok=True)
    manifest["project_context"] = context_summary
    manifest["max_concurrent_batches"] = max(1, min(int(settings.get("max_concurrent_batches") or 2), 4))
    manifest["updated_at"] = db.now_iso()
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    if context_summary.get("trimmed"):
        db.add_event(
            run_id,
            "project context trimmed before provider call: "
            f"{context_summary.get('original_estimated_tokens')} -> {context_summary.get('managed_estimated_tokens')} estimated tokens",
            level="warning",
        )

    budget_warning_tokens = int(settings.get("api_budget_warning_tokens") or 1000000)
    estimated_total = int(manifest.get("estimated_total_input_tokens") or 0)
    if estimated_total > budget_warning_tokens and not confirm_api_budget:
        current = db.get_run(run_id)
        db.update_run(
            run_id,
            status="needs_input",
            metadata={
                **current.get("metadata", {}),
                "reason": "api_budget_confirmation_required",
                "api_budget_estimate": {
                    "estimated_input_tokens": estimated_total,
                    "warning_tokens": budget_warning_tokens,
                    "estimated_batches": len(manifest.get("batches") or []),
                },
                "translation_progress": _manifest_progress(manifest, batch_size=batch_size, started_at=time.monotonic()),
            },
        )
        db.add_event(run_id, f"translation paused for API budget confirmation: estimated_input_tokens={estimated_total}, warning={budget_warning_tokens}", level="warning")
        return []

    started_at = time.monotonic()
    limiter = _AsyncTokenRateLimiter(
        int(settings.get("max_requests_per_minute") or 12),
        int(settings.get("max_estimated_tokens_per_minute") or 120000),
    )
    max_attempts = max(1, min(int(settings.get("max_batch_attempts") or 3), 5))
    concurrency = max(1, min(int(settings.get("max_concurrent_batches") or 2), 4))
    manifest_lock = asyncio.Lock()
    failure: Exception | None = None

    def batch_rows(batch_meta: dict[str, Any]) -> list[dict[str, Any]]:
        start = int(batch_meta.get("start") or 0)
        count = int(batch_meta.get("row_count") or 0)
        return rows[start : start + count]

    async def persist_manifest(current_batch: int | None = None, failed_batch: int | None = None, status: str = "running", rate_wait: float | None = None) -> None:
        async with manifest_lock:
            manifest["updated_at"] = db.now_iso()
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
            _update_translation_progress(
                run_id,
                _manifest_progress(manifest, batch_size=batch_size, started_at=started_at, current_batch=current_batch, failed_batch=failed_batch, rate_limit_wait_seconds=rate_wait),
                status=status,
            )

    async def process_batch(batch_meta: dict[str, Any]) -> None:
        nonlocal failure
        if failure is not None:
            return
        batch_index = int(batch_meta["batch_index"])
        batch = batch_rows(batch_meta)
        batch_path = batches_dir / f"batch_{batch_index:05d}.jsonl"
        request_path = batches_dir / f"batch_{batch_index:05d}.request.jsonl"
        raw_response_path = batches_dir / f"batch_{batch_index:05d}.raw_response.jsonl"
        error_path = batches_dir / f"batch_{batch_index:05d}.error.json"
        if not request_path.exists():
            write_jsonl(request_path, batch)
        batch_meta["request_path"] = str(request_path)
        completed = _completed_batch_rows(batch_path, batch)
        if completed is not None:
            batch_meta.update({"status": "passed", "response_path": str(batch_path), "error_path": "", "updated_at": db.now_iso()})
            db.add_event(run_id, f"resume: batch {batch_index}/{len(manifest.get('batches') or [])} already completed; rows={len(completed)}")
            await persist_manifest(current_batch=batch_index)
            return
        for attempt in range(int(batch_meta.get("attempts") or 0) + 1, max_attempts + 1):
            if _cancel_requested(run_id, work_dir, cancel_event):
                batch_meta.update({"status": "canceled", "attempts": attempt - 1, "updated_at": db.now_iso()})
                await persist_manifest(current_batch=batch_index, status="canceled")
                raise RuntimeError("translation canceled")
            batch_meta.update({"status": "running", "attempts": attempt, "updated_at": db.now_iso()})
            await persist_manifest(current_batch=batch_index)
            wait_seconds = await limiter.acquire(int(batch_meta.get("estimated_input_tokens") or 1))
            if wait_seconds:
                db.add_event(run_id, f"rate limit wait before batch {batch_index}: {round(wait_seconds, 2)}s")
                await persist_manifest(current_batch=batch_index, rate_wait=wait_seconds)
            db.add_event(run_id, f"translating batch {batch_index}/{len(manifest.get('batches') or [])}: rows={len(batch)}, attempt={attempt}/{max_attempts}")
            try:
                prompt = provider_prompt
                if attempt > 1:
                    prompt = f"{provider_prompt}\n\nRepair request: previous output for this batch failed local validation. Return the full corrected batch only, preserving IDs, order, placeholders, tags, entities, and newlines."
                items = await translate_batch(batch, settings, prompt)
                batch_output = [{"id": item.id, "translation": item.translation} for item in items]
                write_jsonl(raw_response_path, batch_output)
                batch_meta["raw_response_path"] = str(raw_response_path)
                validated = _validate_translated_batch(batch, batch_output, language)
                write_jsonl(batch_path, validated)
                if error_path.exists():
                    error_path.unlink()
                batch_meta.update({"status": "passed", "response_path": str(batch_path), "error_path": "", "updated_at": db.now_iso()})
                db.add_event(run_id, f"batch {batch_index}/{len(manifest.get('batches') or [])} completed and persisted: rows={len(validated)}")
                await persist_manifest(current_batch=batch_index)
                return
            except Exception as exc:
                _write_batch_error(error_path, batch_index, attempt, exc)
                batch_meta.update({"status": "failed", "error_path": str(error_path), "updated_at": db.now_iso()})
                db.add_event(run_id, f"batch {batch_index}/{len(manifest.get('batches') or [])} failed attempt {attempt}/{max_attempts}: {user_facing_error(exc)}", level="warning")
                await persist_manifest(current_batch=batch_index, failed_batch=batch_index, status="running" if attempt < max_attempts else "failed")
                if attempt >= max_attempts:
                    failure = exc
                    raise
                delay = _provider_retry_delay_seconds(exc, attempt)
                db.add_event(run_id, f"batch {batch_index}/{len(manifest.get('batches') or [])} retry backoff: {round(delay, 2)}s")
                await asyncio.sleep(delay)

    queue: asyncio.Queue[dict[str, Any]] = asyncio.Queue()
    for item in manifest.get("batches") or []:
        completed = _completed_batch_rows(Path(item.get("response_path") or ""), batch_rows(item)) if item.get("status") == "passed" else None
        if completed is not None:
            db.add_event(run_id, f"resume: batch {int(item.get('batch_index') or 0)}/{len(manifest.get('batches') or [])} already completed; rows={len(completed)}")
            continue
        item["status"] = "pending"
        await queue.put(item)

    if queue.empty():
        await persist_manifest(status="running")
    else:
        await persist_manifest()

    async def worker() -> None:
        nonlocal failure
        while failure is None:
            if _cancel_requested(run_id, work_dir, cancel_event):
                failure = RuntimeError("translation canceled")
                return
            try:
                item = queue.get_nowait()
            except asyncio.QueueEmpty:
                return
            try:
                await process_batch(item)
            finally:
                queue.task_done()

    workers = [asyncio.create_task(worker()) for _ in range(concurrency)]
    try:
        await asyncio.gather(*workers)
    finally:
        for worker_task in workers:
            if not worker_task.done():
                worker_task.cancel()
    if failure is not None:
        if str(failure) == "translation canceled":
            db.add_event(run_id, "translation canceled")
            _update_translation_progress(run_id, _manifest_progress(manifest, batch_size=batch_size, started_at=started_at), status="canceled")
        raise failure

    translated_rows: list[dict[str, Any]] = []
    for item in sorted(manifest.get("batches") or [], key=lambda value: int(value.get("batch_index") or 0)):
        translated_rows.extend(read_jsonl(Path(item["response_path"])))
    await persist_manifest(status="running")
    return translated_rows


async def _translate_quick_text_run(
    *,
    run: dict[str, Any],
    input_artifact: dict[str, Any],
    settings: dict[str, Any],
    batch_size: int,
    language: str,
    readiness: dict[str, Any],
    request: Any,
    cancel_event: Any | None = None,
) -> dict[str, Any]:
    run_id = run["id"]
    project = db.get_project(run["project_id"])
    metadata = run.get("metadata", {})
    input_path = Path(input_artifact["path"])
    rows = _quick_text_translation_rows(input_path)
    if not rows:
        reason = "TXT 文件没有检测到可翻译文本。"
        db.update_run(run_id, status="needs_input", metadata={**metadata, "reason": reason, "translation_readiness": readiness})
        db.add_event(run_id, reason)
        return {"run": db.get_run(run_id), "artifacts": [], "quality": None, "translation_readiness": readiness}

    db.update_run(run_id, status="running")
    db.add_event(run_id, f"quick TXT translation preflight: source_lines={len(rows)}, batch_size={batch_size}, estimated_batches={readiness.get('estimated_batches') or '-'}")
    work_dir = run_dir(run_id) / "quick_text_translation"
    work_dir.mkdir(parents=True, exist_ok=True)
    snapshot_dir = work_dir / "snapshots"
    glossary_snapshot = create_project_glossary_snapshot(project["id"], run_id, snapshot_dir, language=language)
    snapshots = create_prompt_and_harness_snapshots(project["id"], run_id, snapshot_dir, language=language)
    reference_snapshot = create_quick_reference_snapshot(project["id"], run_id, metadata.get("reference_artifact_ids"), snapshot_dir)
    prompt = snapshots["prompt"]
    prompt_snapshot = snapshots["prompt_artifact"]
    harness_snapshot_artifact = snapshots["harness_artifact"]
    if reference_snapshot and reference_snapshot.get("context"):
        prompt = _manage_project_prompt_context(f"{prompt}\n\nQuick Task References:\n{reference_snapshot['context']}", settings)
    prompt = _manage_project_prompt_context(
        f"{prompt}\n\n快速 TXT 任务：逐行翻译 source 字段，保持每个 id 的顺序和结构。只返回 JSONL，每行包含 id 和 translation。",
        settings,
    )

    workpack_path = work_dir / "quick_text_workpack.jsonl"
    write_jsonl(workpack_path, rows)
    workpack_artifact = db.add_artifact(project["id"], "快速 TXT workpack", workpack_path, "translation_workpack", run_id=run_id, mime="application/jsonl", metadata={"language": language, "source_artifact_id": input_artifact["id"]})
    translated_rows = await _translate_rows_with_orchestration(
        run_id=run_id,
        rows=rows,
        settings=settings,
        project_prompt=prompt,
        work_dir=work_dir,
        batch_size=batch_size,
        language=language,
        cancel_event=cancel_event,
        confirm_api_budget=bool(getattr(request, "confirm_api_budget", False)),
    )
    if not translated_rows and db.get_run(run_id).get("status") == "needs_input":
        return {"run": db.get_run(run_id), "artifacts": [workpack_artifact], "quality": None, "translation_readiness": readiness}

    response_path = work_dir / "translation_response.jsonl"
    write_jsonl(response_path, translated_rows)
    response_artifact = db.add_artifact(project["id"], "快速 TXT translation response", response_path, "translation_response", run_id=run_id, mime="application/jsonl", metadata={"language": language, "source_artifact_id": input_artifact["id"]})
    output_path = _write_quick_text_output(input_path, translated_rows, language, work_dir)
    final_artifact = db.add_artifact(project["id"], "快速 TXT 最终译文", output_path, "final_text", run_id=run_id, mime="text/plain", role="delivery", origin="generated", metadata={"language": language, "source_artifact_id": input_artifact["id"]})
    manifest_path = work_dir / "quick_text_translation_manifest.json"
    manifest = {
        "kind": "quick_text_translation",
        "run_id": run_id,
        "project_id": project["id"],
        "language": language,
        "source_artifact_id": input_artifact["id"],
        "source_rows": len(rows),
        "final_artifact_id": final_artifact["id"],
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    manifest_artifact = db.add_artifact(project["id"], "快速 TXT translation manifest", manifest_path, "translation_manifest", run_id=run_id, mime="application/json", metadata={"language": language})
    input_artifacts = {
        "source_text": input_artifact["id"],
        "final_text": final_artifact["id"],
        "translation_workpack": workpack_artifact["id"],
        "translation_response": response_artifact["id"],
        "prompt_snapshot": prompt_snapshot["id"],
        "harness_snapshot": harness_snapshot_artifact["id"],
        "glossary_snapshot": glossary_snapshot["id"],
    }
    if reference_snapshot:
        input_artifacts["quick_reference_snapshot"] = reference_snapshot["artifact"]["id"]
    quality_summary = {"passed": True, "hard_errors": 0, "soft_warnings": 0, "rows": len(rows), "format": input_path.suffix.lower().lstrip(".") or "txt"}
    final_metadata = db.get_run(run_id).get("metadata", {})
    db.update_run(
        run_id,
        status="passed",
        metadata={
            **final_metadata,
            "task_origin": metadata.get("task_origin") or "quick_task",
            "input_artifacts": input_artifacts,
            "quality_summary": quality_summary,
            "quality": {"passed": True, "issues": []},
            "translation_readiness": readiness,
            "translated_rows": len(rows),
            "source_rows": len(rows),
            "output_format": input_path.suffix.lower().lstrip(".") or "txt",
        },
    )
    db.add_event(run_id, f"quick TXT translation finished: rows={len(rows)}, output={output_path.name}")
    return {
        "run": db.get_run(run_id),
        "artifacts": [final_artifact, response_artifact, workpack_artifact, manifest_artifact, glossary_snapshot, prompt_snapshot, harness_snapshot_artifact],
        "quality": {"passed": True, "issues": []},
        "translation_readiness": readiness,
    }


async def translate_run(run_id: str, request: Any, cancel_event: Any | None = None) -> dict[str, Any]:
    run = db.get_run(run_id)
    language = require_supported_language(run.get("language") or "en")
    project = db.get_project(run["project_id"])
    metadata = run.get("metadata", {})
    input_artifact = db.get_artifact(metadata["input_artifact_id"])
    settings = load_settings()
    if request.provider and str(request.provider).strip() == TEST_FAKE_PROVIDER and not test_provider_enabled():
        reason = "测试 provider 未启用；正式任务请使用已配置的 GPT / Claude API。"
        db.update_run(run_id, status="needs_input", metadata={**metadata, "reason": reason})
        db.add_event(run_id, reason)
        return {"run": db.get_run(run_id), "artifacts": [], "quality": None}
    if request.provider:
        settings["provider"] = normalize_provider_name(request.provider)
    if request.protocol:
        settings["protocol"] = request.protocol
    if getattr(request, "preset", None):
        settings["preset"] = request.preset
    batch_size = int(request.batch_size or metadata.get("batch_size") or settings.get("batch_size") or 90)
    batch_size = max(1, min(batch_size, 200))
    readiness = inspect_translation_readiness(input_artifact["id"], batch_size=batch_size, language=language)
    if _is_quick_text_path(Path(input_artifact["path"])) and metadata.get("task_origin") != "quick_task":
        reason = _friendly_unsupported_language_file_message(Path(input_artifact["path"]).suffix)
        db.update_run(
            run_id,
            status="needs_input",
            metadata={
                **metadata,
                "reason": reason,
                "translation_readiness": readiness,
            },
        )
        db.add_event(run_id, reason)
        return {"run": db.get_run(run_id), "artifacts": [], "quality": None, "translation_readiness": readiness}
    if readiness.get("reason") == "unsupported_file":
        reason = _friendly_unsupported_language_file_message(Path(input_artifact["path"]).suffix)
        db.update_run(
            run_id,
            status="needs_input",
            metadata={
                **metadata,
                "reason": reason,
                "translation_readiness": readiness,
            },
        )
        db.add_event(run_id, reason)
        return {"run": db.get_run(run_id), "artifacts": [], "quality": None, "translation_readiness": readiness}
    if readiness.get("reason") == "invalid_id_rows":
        reason = "language table ID column must be present and non-empty before translation or QA"
        db.update_run(
            run_id,
            status="needs_input",
            metadata={
                **metadata,
                "reason": reason,
                "translation_readiness": readiness,
            },
        )
        db.add_event(run_id, f"translation skipped: {reason}")
        return {"run": db.get_run(run_id), "artifacts": [], "quality": None, "translation_readiness": readiness}
    if readiness.get("ready_for_qa"):
        db.update_run(
            run_id,
            status="needs_input",
            metadata={
                **metadata,
                "reason": "input already contains target translations; run QA instead",
                "translation_readiness": readiness,
            },
        )
        db.add_event(run_id, "translation skipped: input already contains target translations; run QA instead")
        return {"run": db.get_run(run_id), "artifacts": [], "quality": None, "translation_readiness": readiness}
    effective_provider = normalize_provider_name(settings.get("provider"))
    if effective_provider in REAL_PROVIDERS and not settings.get("api_key"):
        db.update_run(
            run_id,
            status="needs_input",
            metadata={**metadata, "reason": f"{effective_provider} api_key is required for formal translation"},
        )
        return {"run": db.get_run(run_id), "artifacts": [], "quality": None}
    if metadata.get("task_origin") == "quick_task" and _is_quick_text_path(Path(input_artifact["path"])):
        return await _translate_quick_text_run(
            run=run,
            input_artifact=input_artifact,
            settings=settings,
            batch_size=batch_size,
            language=language,
            readiness=readiness,
            request=request,
            cancel_event=cancel_event,
        )

    db.update_run(run_id, status="running")
    db.add_event(
        run_id,
        f"translation preflight: source_rows={readiness['source_rows']}, translated_rows={readiness['translated_rows']}, "
        f"empty_target_rows={readiness['empty_target_rows']}, cjk_target_rows={readiness['cjk_target_rows']}, "
        f"batch_size={batch_size}, estimated_batches={readiness['estimated_batches']}",
    )
    work_dir = run_dir(run_id) / "translation"
    work_dir.mkdir(parents=True, exist_ok=True)
    snapshot_dir = work_dir / "snapshots"
    language = require_supported_language(run.get("language") or "en")
    glossary_snapshot = create_project_glossary_snapshot(project["id"], run_id, snapshot_dir, language=language)
    snapshots = create_prompt_and_harness_snapshots(project["id"], run_id, snapshot_dir, language=language)
    reference_snapshot = create_quick_reference_snapshot(project["id"], run_id, metadata.get("reference_artifact_ids"), snapshot_dir)
    prompt = snapshots["prompt"]
    prompt_snapshot = snapshots["prompt_artifact"]
    harness_snapshot_artifact = snapshots["harness_artifact"]
    harness_snapshot = snapshots["harness_snapshot"]
    prompt_path = snapshots["prompt_path"]
    if reference_snapshot and reference_snapshot.get("context"):
        raw_prompt = f"{prompt}\n\nQuick Task References:\n{reference_snapshot['context']}"
        settings = load_settings()
        prompt = _manage_project_prompt_context(raw_prompt, settings)
        prompt_path = snapshot_dir / "compiled_project_harness_prompt_with_quick_refs.txt"
        prompt_path.write_text(prompt, encoding="utf-8")
        prompt_snapshot = db.add_artifact(
            project["id"],
            "Prompt snapshot with quick references",
            prompt_path,
            "prompt_snapshot",
            run_id=run_id,
            mime="text/plain",
            origin="generated",
            metadata={
                "source": "project_prompt_harness_and_quick_references",
                "language": language,
                "reference_artifact_ids": metadata.get("reference_artifact_ids") or [],
                "context_budget": _project_context_summary(raw_prompt, settings),
            },
        )

    prepare_args = [
        sys.executable,
        str(LOCALIZATION_ROOT / "scripts" / "run_translation_harness.py"),
        "--input",
        input_artifact["path"],
        "--lang",
        language,
        "--output-dir",
        str(work_dir),
        "--style-hint-file",
        str(prompt_path),
        "--term-base",
        glossary_snapshot["path"],
    ]
    try:
        db.add_event(run_id, "preparing translation workpack with localization workflow")
        run_subprocess(prepare_args, LOCALIZATION_ROOT, run_id)
        workpack_path = work_dir / "translation_workpack.jsonl"
        rows = read_jsonl(workpack_path)
        manifest_preview = _load_or_create_batch_manifest(work_dir / "batch_manifest.json", rows, prompt, settings, batch_size, language)
        db.add_event(run_id, f"workpack prepared: rows={len(rows)}, dynamic_batches={len(manifest_preview.get('batches') or [])}, concurrency={settings.get('max_concurrent_batches')}")
        translated_rows = await _translate_rows_with_orchestration(
            run_id=run_id,
            rows=rows,
            settings=settings,
            project_prompt=prompt,
            work_dir=work_dir,
            batch_size=batch_size,
            language=language,
            cancel_event=cancel_event,
            confirm_api_budget=bool(getattr(request, "confirm_api_budget", False)),
        )
        if not translated_rows and db.get_run(run_id).get("status") == "needs_input":
            return {"run": db.get_run(run_id), "artifacts": [], "quality": None, "translation_readiness": readiness}
        response_path = work_dir / "translation_response.jsonl"
        write_jsonl(response_path, translated_rows)
        db.add_artifact(project["id"], "Translation response JSONL", response_path, "translation_response", run_id=run_id, mime="application/jsonl")

        db.add_event(run_id, "applying translation response and running strict harness validation")
        apply_args = [
            sys.executable,
            str(LOCALIZATION_ROOT / "scripts" / "run_translation_harness.py"),
            "--input",
            input_artifact["path"],
            "--lang",
            language,
            "--output-dir",
            str(work_dir),
            "--response",
            str(response_path),
            "--term-base",
            glossary_snapshot["path"],
        ]
        apply_proc = run_subprocess(apply_args, LOCALIZATION_ROOT, run_id)
        parsed = parse_key_output(apply_proc.stdout)
        raw_workbook = Path(parsed.get("final_workbook", ""))
        raw_artifact = db.add_artifact(
            project["id"],
            "Raw translated workbook",
            raw_workbook,
            "raw_translated_workbook",
            run_id=run_id,
            origin="generated",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            metadata={"language": language, "source_workbook": input_artifact["id"]},
        )
        db.add_event(run_id, "running localization QA gate after translation")
        qa_result = run_localization_qa(
            project=project,
            run_id=run_id,
            workbook_path=raw_workbook,
            output_dir=work_dir / "qa",
            glossary_snapshot=glossary_snapshot,
            harness_snapshot=harness_snapshot,
            workbook_artifact=raw_artifact,
            run_metadata=metadata,
            language=language,
        )
        status = "passed" if qa_result["quality_summary"]["passed"] else "failed"
        artifacts = [
            raw_artifact,
            db.add_artifact(project["id"], "Translation manifest", work_dir / "translation_manifest.json", "translation_manifest", run_id=run_id, mime="application/json"),
            glossary_snapshot,
            prompt_snapshot,
            harness_snapshot_artifact,
            *qa_result["artifacts"],
        ]
        input_artifacts = {
            "source_workbook": input_artifact["id"],
            "raw_translated_workbook": raw_artifact["id"],
            "translation_workbook": raw_artifact["id"],
            "glossary_snapshot": glossary_snapshot["id"],
            "prompt_snapshot": prompt_snapshot["id"],
            "harness_snapshot": harness_snapshot_artifact["id"],
        }
        if reference_snapshot:
            input_artifacts["quick_reference_snapshot"] = reference_snapshot["artifact"]["id"]
        if qa_result.get("qa_final_artifact"):
            input_artifacts["qa_final_workbook"] = qa_result["qa_final_artifact"]["id"]
            input_artifacts["translation_workbook"] = qa_result["qa_final_artifact"]["id"]
        archive_result = None
        if status == "passed" and qa_result.get("qa_final_artifact"):
            archive_result = archive_translation_artifact(
                project["id"],
                qa_result["qa_final_artifact"]["id"],
                language=run.get("language") or "en",
                source_type="qa_passed",
            )
            db.add_event(run_id, f"translation archive updated: rows={archive_result['imported_count']}")
        db.add_event(run_id, f"translation run finished: status={status}")
        final_metadata = db.get_run(run_id).get("metadata", {})
        final_progress = _terminal_translation_progress(final_metadata.get("translation_progress"), status)
        db.update_run(
            run_id,
            status=status,
            metadata={
                **final_metadata,
                "translation_progress": final_progress,
                "task_origin": metadata.get("task_origin") or "translation_run",
                "input_artifacts": input_artifacts,
                "quality": qa_result["quality"],
                "project_harness_quality": qa_result["project_harness_quality"],
                "semantic_qa": qa_result["semantic_qa"],
                "quality_summary": qa_result["quality_summary"],
                "harness": harness_snapshot["summary"],
                "model": {
                    "provider": settings.get("provider"),
                    "protocol": settings.get("protocol"),
                    "preset": settings.get("preset"),
                    "model": settings.get("model"),
                    "reasoning_effort": settings.get("reasoning_effort"),
                },
                "batch_size": batch_size,
                "translation_readiness": readiness,
                "translation_archive": archive_result,
            },
        )
        return {
            "run": db.get_run(run_id),
            "artifacts": artifacts,
            "quality": qa_result["quality"],
            "project_harness_quality": qa_result["project_harness_quality"],
            "quality_summary": qa_result["quality_summary"],
        }
    except Exception as exc:
        friendly = user_facing_error(exc)
        db.add_event(run_id, friendly, level="error")
        failed_metadata = db.get_run(run_id).get("metadata", {})
        status = "canceled" if str(exc) == "translation canceled" else "failed"
        db.update_run(run_id, status=status, metadata={**failed_metadata, "translation_progress": _terminal_translation_progress(failed_metadata.get("translation_progress"), status), "error": friendly})
        if isinstance(exc, UserFacingWorkflowError):
            raise
        raise UserFacingWorkflowError(friendly) from exc


def run_translate_sync(run_id: str, request: Any, cancel_event: Any | None = None) -> dict[str, Any]:
    return asyncio.run(translate_run(run_id, request, cancel_event=cancel_event))


def cancel_translation_run(run_id: str) -> dict[str, Any]:
    run = db.get_run(run_id)
    work_dir = run_dir(run_id) / "translation"
    work_dir.mkdir(parents=True, exist_ok=True)
    _translation_cancel_path(work_dir).write_text(db.now_iso(), encoding="utf-8")
    db.cancel_job_lease("long_text", run_id)
    metadata = run.get("metadata", {})
    db.update_run(run_id, status="canceled", metadata={**metadata, "cancel_requested_at": db.now_iso()})
    db.add_event(run_id, "translation cancel requested")
    return db.get_run(run_id)


def translation_run_progress(run_id: str) -> dict[str, Any]:
    run = db.get_run(run_id)
    metadata = run.get("metadata", {})
    progress = metadata.get("translation_progress")
    if run.get("status") in {"passed", "failed", "needs_input", "canceled"}:
        progress = _terminal_translation_progress(progress, str(run.get("status") or ""))
    return {
        "run": run,
        "progress": progress,
        "api_budget_estimate": metadata.get("api_budget_estimate"),
        "reason": metadata.get("reason"),
    }


def translation_batch_file(run_id: str, batch_index: int, kind: str) -> Path:
    if batch_index < 1:
        raise ValueError("batch_index must be positive")
    if kind not in {"request", "response", "raw-response", "error"}:
        raise ValueError("batch file kind must be request, response, raw-response, or error")
    run = db.get_run(run_id)
    metadata = run.get("metadata") or {}
    progress = metadata.get("translation_progress") or {}
    batch_size = int(progress.get("batch_size") or metadata.get("batch_size") or 90)
    suffix = {"request": ".request.jsonl", "response": ".jsonl", "raw-response": ".raw_response.jsonl", "error": ".error.json"}[kind]
    path = run_dir(run_id) / "translation" / f"batches_{batch_size}" / f"batch_{batch_index:05d}{suffix}"
    if not path.exists():
        raise KeyError(str(path))
    return path


def reconcile_interrupted_background_jobs() -> dict[str, int]:
    db.mark_running_job_leases_interrupted()
    translation_runs = 0
    announcement_tasks = 0
    for run in db.list_runs():
        if run.get("kind") == "translation" and run.get("status") in {"queued", "running"}:
            metadata = dict(run.get("metadata") or {})
            metadata["reason"] = "background_job_interrupted"
            metadata["interrupted_at"] = db.now_iso()
            db.update_run(run["id"], status="needs_input", metadata=metadata)
            db.add_event(run["id"], "background translation job was interrupted; resume from saved batches")
            translation_runs += 1
    for project in db.list_projects():
        for task in db.list_announcement_tasks(project["id"]):
            if task.get("status") in {"queued", "running"}:
                metadata = dict(task.get("metadata") or {})
                metadata["reason"] = "background_job_interrupted"
                metadata["interrupted_at"] = db.now_iso()
                db.update_announcement_task(task["id"], status="needs_input", current_step=ANNOUNCEMENT_STEP["translate"], metadata=metadata)
                for item in task.get("languages") or []:
                    if item.get("status") in {"queued", "running"}:
                        lang_meta = dict(item.get("metadata") or {})
                        lang_meta["reason"] = "background_job_interrupted"
                        db.upsert_announcement_task_language(
                            task["id"],
                            task["project_id"],
                            str(item["language"]),
                            status="prepared",
                            current_step=ANNOUNCEMENT_STEP["translate"],
                            metadata=lang_meta,
                        )
                announcement_tasks += 1
    return {"translation_runs": translation_runs, "announcement_tasks": announcement_tasks}
